import os
import pytest
from docx import Document
from lxml import etree
import latex2mathml.converter
import mathml2omml

def test_latex_to_mathml_to_omml():
    # 1. LaTeX to MathML
    latex = r"E = m c^2"
    mathml = latex2mathml.converter.convert(latex)
    assert "<math" in mathml
    assert "E" in mathml
    assert "mc" in mathml or "c" in mathml
    
    # 2. MathML to OMML
    omml = mathml2omml.convert(mathml)
    assert "<m:oMath" in omml
    assert "<m:t>E</m:t>" in omml or "E" in omml
    
    # Ensure namespaces parse cleanly in lxml
    if "xmlns:m=" not in omml:
        omml = omml.replace("<m:oMath", '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"', 1)
    
    omml_el = etree.fromstring(omml)
    assert omml_el.tag.endswith("oMath")

def test_embed_omml_in_docx():
    latex = r"\frac{a}{b}"
    mathml = latex2mathml.converter.convert(latex)
    omml = mathml2omml.convert(mathml)
    if "xmlns:m=" not in omml:
        omml = omml.replace("<m:oMath", '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"', 1)
    
    # Create empty docx document
    doc = Document()
    p = doc.add_paragraph("An equation here: ")
    
    # Parse OMML and append to paragraph element
    omml_el = etree.fromstring(omml)
    p._element.append(omml_el)
    
    # Add trailing text run
    p.add_run(" trailing text.")
    
    # Save document
    test_docx_path = "tests/temp_math_test.docx"
    doc.save(test_docx_path)
    
    try:
        assert os.path.exists(test_docx_path)
        # Reload and check structure
        reloaded_doc = Document(test_docx_path)
        reloaded_p = reloaded_doc.paragraphs[0]
        # Check that child element exists
        p_children_tags = [etree.QName(c.tag).localname for c in list(reloaded_p._element)]
        assert "oMath" in p_children_tags
    finally:
        if os.path.exists(test_docx_path):
            os.remove(test_docx_path)
