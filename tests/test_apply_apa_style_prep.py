"""Tests for `apply_apa_style_prep` — specifically the switch from a
style-name skip (`REF-U`) to a paragraph-identity skip (whatever the
bibliography detector claimed).

The rule is: no stylistic label (`REF-U`, `Reference-Alphabetical`, or any
other) should be used as a hard condition to skip or apply `cite_bib`.
Skipping happens iff the bibliography detection logic identified the
paragraph as a reference entry.
"""
from __future__ import annotations

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from app.processing.legacy.validation_core import apply_apa_style_prep


def _ensure_styles(doc, *names):
    for name in names:
        try:
            doc.styles[name]
        except KeyError:
            doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)


def _run_styles(paragraph):
    """Return the set of character-style names in use on `paragraph`'s runs."""
    out = set()
    for r in paragraph.runs:
        if r.style and r.style.name:
            out.add(r.style.name)
    return out


def test_reference_alphabetical_style_is_not_a_skip_condition(tmp_path):
    """A body paragraph carrying the `Reference-Alphabetical` style must
    still receive `cite_bib` on its in-text citations. The old code path
    hard-skipped on `REF-U`; the new path skips only paragraphs the
    bibliography detector recognised as reference entries — style names
    like `Reference-Alphabetical` never gate `cite_bib`."""
    doc = Document()
    _ensure_styles(doc, "Reference-Alphabetical")

    # Body first — the cite_bib loop stops at the bibliography boundary.
    body = doc.add_paragraph(
        "As widgets rotate (Smith, 2022), tolerances shift.",
        style="Reference-Alphabetical",
    )

    # Bibliography via explicit <ref-open>/<ref-close>; REF-U/Reference-
    # Alphabetical are not what tells the detector where the bib is.
    doc.add_paragraph("<ref-open>")
    doc.add_paragraph("Smith, J. (2022). Widget dynamics. Journal of Things, 12, 1-10.")
    doc.add_paragraph("<ref-close>")

    apply_apa_style_prep(doc)

    assert "cite_bib" in _run_styles(body), _run_styles(body)


def test_ref_u_outside_bibliography_no_longer_hard_skipped(tmp_path):
    """A body paragraph accidentally styled `REF-U` but living OUTSIDE the
    detected bibliography section is now eligible for `cite_bib`. The
    detector only claims paragraphs it actually identified as reference
    entries, so a stray REF-U body paragraph doesn't slip through
    untagged."""
    doc = Document()
    _ensure_styles(doc, "REF-U")

    # Body paragraph carrying REF-U style, placed BEFORE the bibliography.
    # The old code would have hard-skipped it on the style name; the new
    # code lets it through because the detector only claims the entries
    # inside <ref-open>/<ref-close>.
    body = doc.add_paragraph(
        "Rotation shifts tolerances (Smith, 2022).",
        style="REF-U",
    )

    doc.add_paragraph("<ref-open>")
    doc.add_paragraph("Smith, J. (2022). Widget dynamics. Journal of Things, 12, 1-10.")
    doc.add_paragraph("<ref-close>")

    apply_apa_style_prep(doc)

    assert "cite_bib" in _run_styles(body), _run_styles(body)


def test_ref_u_fallback_bibliography_entries_still_skipped(tmp_path):
    """When no <ref-open>/<ref-close> markers are present, the fallback
    treats REF-U paragraphs as bibliography entries. Those specific
    paragraph elements go into the detector's `bib_para_elements` set and
    must be skipped during `cite_bib` application — otherwise the
    parenthetical year inside an entry (`Smith, J. (2022).`) would be
    mistaken for a citation and tagged."""
    doc = Document()
    _ensure_styles(doc, "REF-U")

    body = doc.add_paragraph("Widgets rotate (Smith, 2022).")
    bib = doc.add_paragraph(
        "Smith, J. (2022). Widget dynamics. Journal of Things, 12, 1-10.",
        style="REF-U",
    )

    apply_apa_style_prep(doc)

    # Body citation IS tagged.
    assert "cite_bib" in _run_styles(body), _run_styles(body)
    # Bibliography entry (identified by fallback) is NOT.
    assert "cite_bib" not in _run_styles(bib), _run_styles(bib)


def test_bibliography_inside_markers_is_skipped_regardless_of_style(tmp_path):
    """Paragraphs inside <ref-open>/<ref-close> are bibliography entries by
    identity, whatever style they carry. Their parenthetical year must not
    be treated as an in-text citation."""
    doc = Document()

    body = doc.add_paragraph("As shown (Smith, 2022), it works.")
    doc.add_paragraph("<ref-open>")
    # Entry uses the built-in Normal style, NOT REF-U — the old code would
    # have failed to skip it; the new code skips by paragraph identity.
    bib = doc.add_paragraph(
        "Smith, J. (2022). Widget dynamics. Journal of Things, 12, 1-10."
    )
    doc.add_paragraph("<ref-close>")

    apply_apa_style_prep(doc)

    assert "cite_bib" in _run_styles(body)
    assert "cite_bib" not in _run_styles(bib), _run_styles(bib)
