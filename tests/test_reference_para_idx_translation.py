"""Test the body-level → all-w:p index translation used by the APA reference
review pipeline.

The bug this locks in: CitationProcessor numbers paragraphs body-level
(skipping paragraphs inside tables), while `docx_to_xhtml_runs` — the
XHTML the WYSIWYG editor consumes — numbers every `<w:p>` in document
order (table-cell paragraphs included). Without translation, a reference
entry's body-level `para_idx` maps onto whatever table-cell paragraph
happens to sit at the same numeric position in the editor. The frontend
then stamps REF{n} bookmarks on that table cell.

The translator returned by `_build_body_to_allp_translator` converts a
body-level index into the equivalent all-w:p index so downstream frontend
lookups land on the correct paragraph.
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

from app.domains.review.service import _build_body_to_allp_translator


def _build_doc_with_table_before_refs(path: Path) -> Document:
    """A doc shape that reproduces the bug: several body paragraphs, then a
    table (a few rows/cells with paragraph content), then the references."""
    doc = Document()

    # Body paragraphs 0-2.
    doc.add_paragraph("Intro paragraph one.")
    doc.add_paragraph("Intro paragraph two.")
    doc.add_paragraph("Intro paragraph three.")

    # Table with 3 rows × 2 cols. Each cell holds a paragraph → 6 extra
    # <w:p> elements that count in the all-w:p scheme but NOT in the
    # body-level scheme.
    table = doc.add_table(rows=3, cols=2)
    for row_i, row in enumerate(table.rows):
        for col_i, cell in enumerate(row.cells):
            cell.text = f"cell-{row_i}-{col_i}"

    # Body paragraph 3.
    doc.add_paragraph("Post-table body paragraph.")

    # References section.
    doc.add_paragraph("References")
    doc.add_paragraph(
        "Smith, J. (2022). Widget dynamics. Journal of Things, 12, 1-10."
    )
    doc.add_paragraph(
        "Jones, A. (2021). Gadget analysis. Journal of Stuff, 5, 20-30."
    )

    doc.save(path)
    return Document(path)


def test_translator_maps_body_indices_past_a_table(tmp_path):
    """Body index for a paragraph AFTER the table must translate to a
    larger all-w:p index (larger by the number of table-cell paragraphs)."""
    doc_path = tmp_path / "sample.docx"
    doc = _build_doc_with_table_before_refs(doc_path)
    translate = _build_body_to_allp_translator(doc)

    body_paragraphs = list(doc.paragraphs)
    all_p_elements = list(doc.element.body.iter(qn("w:p")))

    # Sanity: the two schemes diverge because of the table.
    assert len(all_p_elements) == len(body_paragraphs) + 6, (
        len(all_p_elements), len(body_paragraphs)
    )

    # Every body paragraph must translate to the position of its element
    # in the all-w:p iteration.
    for body_idx, para in enumerate(body_paragraphs):
        expected = all_p_elements.index(para._element)
        assert translate(body_idx) == expected, (
            f"body_idx={body_idx} translated to {translate(body_idx)}, expected {expected}"
        )

    # And specifically, the first post-table body paragraph — this is the
    # case that used to send REF{n} bookmarks into table cells. Body index
    # is 3 (four body paragraphs before), all-w:p index is 3 + 6 = 9.
    post_table_body_idx = 3
    post_table_para = body_paragraphs[post_table_body_idx]
    assert post_table_para.text == "Post-table body paragraph."
    assert translate(post_table_body_idx) == all_p_elements.index(post_table_para._element)
    assert translate(post_table_body_idx) == 9


def test_translator_passes_through_out_of_range_and_none(tmp_path):
    """`None` becomes -1; negative and virtual-encoded indices (values >=
    len(doc.paragraphs) that CitationProcessor tags to table/textbox
    citations) are returned unchanged so we don't mis-translate them."""
    doc_path = tmp_path / "sample.docx"
    doc = _build_doc_with_table_before_refs(doc_path)
    translate = _build_body_to_allp_translator(doc)

    assert translate(None) == -1
    assert translate(-1) == -1
    # Virtual index from CitationProcessor's table-cell scheme
    # (`_para_offset + tbl_idx * 1000 + cell_counter`). Left alone.
    big = len(list(doc.paragraphs)) + 5000
    assert translate(big) == big


def test_translator_matches_docx_to_xhtml_body_p_map(tmp_path):
    """The translator's output must be exactly the index that
    `docx_to_xhtml_runs.body_p_map` would emit as `data-para-idx` on the
    same paragraph — that's the whole point of the translation, so the
    frontend's `paraByIdx.get(para_idx)` lookup succeeds."""
    doc_path = tmp_path / "sample.docx"
    doc = _build_doc_with_table_before_refs(doc_path)
    translate = _build_body_to_allp_translator(doc)

    # Replicate docx_to_xhtml_runs.body_p_map construction (line 1099-1101).
    body_p_map = {
        p_elem: i
        for i, p_elem in enumerate(doc.element.body.iter(qn("w:p")))
    }

    for body_idx, para in enumerate(doc.paragraphs):
        translated = translate(body_idx)
        xhtml_idx = body_p_map[para._element]
        assert translated == xhtml_idx, (
            f"body_idx={body_idx}: translator gave {translated}, "
            f"docx_to_xhtml_runs would emit {xhtml_idx}"
        )
