"""Unit tests for bookmark-based paragraph & run tracking and in-place XML delta patching."""

import os
import tempfile
from pathlib import Path
from docx import Document
from docx.oxml.ns import qn

from app.processing.docx_to_xhtml_runs import DocxToXhtmlRunsEngine
from app.processing.xhtml_to_docx_delta import XhtmlToDocxDeltaEngine, _find_para_by_bookmark, _find_run_by_bookmark


def test_dual_bookmark_generation_and_patching():
    # 1. Create a mock DOCX file
    doc = Document()
    p1 = doc.add_paragraph()
    r1 = p1.add_run("first run")
    r2 = p1.add_run(" and second run")
    
    p2 = doc.add_paragraph("Paragraph 2 content")
    
    # Save mock file
    fd, docx_path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    try:
        doc.save(docx_path)
        
        # 2. Convert to HTML and assign bookmarks
        exporter = DocxToXhtmlRunsEngine()
        html_content = exporter.convert(docx_path)
        
        # Verify HTML contains paragraph and run bookmarks
        assert "p_bm_" in html_content
        assert "r_bm_" in html_content
        assert 'data-bookmark="p_bm_' in html_content
        assert 'data-bookmark="r_bm_' in html_content
        
        # 3. Reload DOCX and verify bookmarks exist inside the file XML
        doc_with_bm = Document(docx_path)
        p1_bm = None
        for child in doc_with_bm.paragraphs[0]._p:
            if child.tag == qn("w:bookmarkStart") and child.get(qn("w:name")).startswith("p_bm_"):
                p1_bm = child.get(qn("w:name"))
                break
        assert p1_bm is not None, "Paragraph bookmark was not persisted in DOCX!"
        
        # Find run bookmark
        r1_bm = None
        for child in doc_with_bm.paragraphs[0]._p:
            if child.tag == qn("w:bookmarkStart") and child.get(qn("w:name")).startswith("r_bm_"):
                r1_bm = child.get(qn("w:name"))
                break
        assert r1_bm is not None, "Run bookmark was not persisted in DOCX!"
        
        # 4. Modify HTML content (e.g. bolding the first run)
        import lxml.html
        root = lxml.html.fromstring(html_content)
        
        # Find the span matching the first run bookmark
        span_r1 = root.xpath(f"//span[@data-bookmark='{r1_bm}']")[0]
        # Wrap it with bold tag
        span_r1.tag = "strong"
        span_r1.set("style", "font-weight: bold;")
        
        # Write modified HTML to temp file
        fd_h, html_path = tempfile.mkstemp(suffix=".html")
        os.close(fd_h)
        try:
            with open(html_path, "w", encoding="utf-8") as f:
                f.write(lxml.html.tostring(root, encoding="utf-8").decode("utf-8"))
            
            # 5. Apply delta patch back to DOCX
            importer = XhtmlToDocxDeltaEngine()
            importer.convert(html_path, docx_path, username="Test Editor")
            
            # 6. Verify that only the targeted run was modified in-place and formatting was preserved
            doc_patched = Document(docx_path)
            patched_para = doc_patched.paragraphs[0]
            
            # Find the run via bookmark in patched document
            patched_run = _find_run_by_bookmark(patched_para, r1_bm)
            assert patched_run is not None, "Patched run could not be resolved!"
            assert patched_run.text == "first run", "Run text was corrupted!"
            assert patched_run.bold is True, "Run bold formatting was not applied in-place!"
            
            # Verify that the sibling runs and structures are intact
            assert len(patched_para.runs) >= 2
            
        finally:
            if os.path.exists(html_path):
                os.remove(html_path)
                
    finally:
        if os.path.exists(docx_path):
            os.remove(docx_path)


def test_inplace_same_name_versioning_and_xhtml_cascade_delete(db_session, user_factory, file_record_factory):
    # 1. Create editor user
    editor = user_factory("editor_test", role_names=("Editor",))

    # 2. Create mock file record and processed copy
    record, processed = file_record_factory(create_processed=True, paragraphs=["Original first", "Original second"])
    assert processed is not None
    processed_id = processed.id
    processed_path = processed.path
    processed_filename = processed.filename
    original_version = processed.version

    # Ensure XHTML directory exists and we write an XHTML content
    from app.domains.review.service import _get_xhtml_path
    xhtml_path = _get_xhtml_path(processed_path)
    os.makedirs(os.path.dirname(xhtml_path), exist_ok=True)
    with open(xhtml_path, "w", encoding="utf-8") as f:
        f.write("<html><body><p>Mock Content</p></body></html>")
    
    assert os.path.exists(processed_path)
    assert os.path.exists(xhtml_path)

    # 3. Perform a delta patch with save_xhtml_delta_and_convert
    from app.domains.review.service import save_xhtml_delta_and_convert
    import logging
    logger = logging.getLogger("test_logger")

    # Generate a run-anchored XHTML content for save
    from app.processing.docx_to_xhtml_runs import DocxToXhtmlRunsEngine
    exporter = DocxToXhtmlRunsEngine()
    html_content = exporter.convert(processed_path)

    # Run the delta save engine
    result = save_xhtml_delta_and_convert(
        db_session,
        file_id=processed_id,
        html_content=html_content,
        username=editor.username,
        logger=logger
    )

    # 4. Verify in-place same name version updates!
    assert result["status"] == "ok"
    assert result["file_id"] == processed_id

    # Retrieve processed record from database again
    from app.models import File, FileVersion
    db_session.expire_all()
    updated_processed = db_session.query(File).filter(File.id == processed_id).first()

    assert updated_processed is not None
    assert updated_processed.filename == processed_filename, "Filename must remain exactly the same!"
    assert updated_processed.path == processed_path, "Path must remain exactly the same!"
    assert updated_processed.version == original_version + 1, "Version number must be incremented by 1!"

    # Verify that a FileVersion archive entry was successfully created
    archived_versions = db_session.query(FileVersion).filter(FileVersion.file_id == processed_id).all()
    assert len(archived_versions) == 1
    assert archived_versions[0].version_num == original_version
    assert os.path.exists(archived_versions[0].path), "Archived copy must exist physically on disk!"

    # 5. Verify XHTML cascade deletion!
    from app.services.file_service import delete_file_and_capture_context
    delete_file_and_capture_context(db_session, file_id=processed_id)

    # Verify DB row is gone
    deleted_row = db_session.query(File).filter(File.id == processed_id).first()
    assert deleted_row is None

    # Verify both DOCX and XHTML files are deleted from the disk location
    assert not os.path.exists(processed_path), "DOCX file was not deleted!"
    assert not os.path.exists(xhtml_path), "Associated XHTML file was not deleted!"


def test_build_bookmark_para_index_deep():
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from app.processing.xhtml_to_docx_delta import _build_bookmark_para_index

    doc = Document()
    # Add a block SDT
    sdt = OxmlElement('w:sdt')
    sdtContent = OxmlElement('w:sdtContent')
    sdt.append(sdtContent)
    
    # Add a paragraph inside SDT with a bookmark
    p_inside = OxmlElement('w:p')
    bm_start = OxmlElement('w:bookmarkStart')
    bm_start.set(qn('w:name'), 'test_bm_inside_sdt')
    bm_start.set(qn('w:id'), '100')
    p_inside.append(bm_start)
    sdtContent.append(p_inside)
    
    doc.element.body.append(sdt)
    
    # Build index and verify
    index = _build_bookmark_para_index(doc)
    assert 'test_bm_inside_sdt' in index
    assert index['test_bm_inside_sdt']._p == p_inside


def test_patch_paragraph_runs_removes_sdt():
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import lxml.html
    from app.processing.xhtml_to_docx_delta import XhtmlToDocxDeltaEngine

    doc = Document()
    p = doc.add_paragraph()
    
    # Add inline SDT
    sdt = OxmlElement('w:sdt')
    sdtContent = OxmlElement('w:sdtContent')
    sdt.append(sdtContent)
    r = OxmlElement('w:r')
    t = OxmlElement('w:t')
    t.text = "Citation"
    r.append(t)
    sdtContent.append(r)
    p._p.append(sdt)
    
    # Verify the sdt is inside the paragraph first
    assert p._p.find(qn('w:sdt')) is not None
    
    # Patch this paragraph with simple HTML content
    engine = XhtmlToDocxDeltaEngine()
    html_el = lxml.html.fromstring("<p>New Content</p>")
    engine._patch_paragraph_runs(p, html_el, doc, "Test User")
    
    # Verify that the sdt wrapper is completely removed from paragraph element
    assert p._p.find(qn('w:sdt')) is None
    # Verify the new run with "New Content" text is present
    assert p.text == "New Content"


def test_patch_paragraph_runs_creates_inline_sdt():
    from docx import Document
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    import lxml.html
    from app.processing.xhtml_to_docx_delta import XhtmlToDocxDeltaEngine

    doc = Document()
    p = doc.add_paragraph()
    
    # Patch this paragraph with HTML content containing sdt-inline class span
    engine = XhtmlToDocxDeltaEngine()
    html_el = lxml.html.fromstring(
        '<p>Before <span class="sdt-inline" data-alias="FigureRef" data-tag="FigureRef">Figure 1.2</span> After</p>'
    )
    engine._patch_paragraph_runs(p, html_el, doc, "Test User")
    
    # Verify that the w:sdt wrapper is created in the paragraph element
    sdt_node = p._p.find(qn('w:sdt'))
    assert sdt_node is not None, "Inline w:sdt element was not recreated!"
    
    # Verify alias and tag are correct
    sdtPr = sdt_node.find(qn('w:sdtPr'))
    alias = sdtPr.find(qn('w:alias')).get(qn('w:val'))
    tag = sdtPr.find(qn('w:tag')).get(qn('w:val'))
    assert alias == "FigureRef"
    assert tag == "FigureRef"
    
    # Verify text inside the sdtContent
    sdtContent = sdt_node.find(qn('w:sdtContent'))
    text = "".join(t.text or "" for t in sdtContent.findall(f".//{qn('w:t')}"))
    assert text == "Figure 1.2"


if __name__ == "__main__":
    import pytest
    pytest.main([__file__])


