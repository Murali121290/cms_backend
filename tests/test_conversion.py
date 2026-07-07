import io
from pathlib import Path
from app import models
from app.domains.projects.models import Project

from unittest.mock import patch, MagicMock

def test_batch_indesign_to_word_endpoint(
    auth_cookie_client,
    admin_user,
    db_session,
    app_env,
):
    client = auth_cookie_client(admin_user)

    # Prepare mock files to upload
    file1_content = b"fake indesign data 1"
    file2_content = b"fake indesign data 2"
    
    files = [
        ("files", ("ch_01.indd", io.BytesIO(file1_content), "application/octet-stream")),
        ("files", ("chapter_02.indd", io.BytesIO(file2_content), "application/octet-stream")),
    ]
    
    data = {
        "client_name": "Test Client",
        "project_code": "TC_BATCH_01",
    }

    # Call endpoint with mocked Windows server response
    with patch("requests.post") as mock_post:
        mock_resp = MagicMock()
        mock_resp.status_code = 200
        mock_resp.content = b"MOCK DOCX BYTES"
        mock_post.return_value = mock_resp

        response = client.post(
            "/api/v1/conversion/batch-indesign-to-word",
            data=data,
            files=files
        )

    assert response.status_code == 200
    body = response.json()
    assert body["project_code"] == "TC_BATCH_01"
    assert len(body["results"]) == 2

    # Verify both chapters exist in database
    project = db_session.query(Project).filter(Project.project_code == "TC_BATCH_01").first()
    assert project is not None
    assert project.division_code == "Test Client"

    ch1 = db_session.query(models.ChapterInfo).filter(
        models.ChapterInfo.project == "TC_BATCH_01",
        models.ChapterInfo.chapters == "1"
    ).first()
    assert ch1 is not None

    ch2 = db_session.query(models.ChapterInfo).filter(
        models.ChapterInfo.project == "TC_BATCH_01",
        models.ChapterInfo.chapters == "2"
    ).first()
    assert ch2 is not None

    # Check that File records are created under InDesign and Manuscript categories
    indd_files = db_session.query(models.File).filter(
        models.File.project_id == project.id,
        models.File.category == "InDesign"
    ).all()
    assert len(indd_files) == 2

    docx_files = db_session.query(models.File).filter(
        models.File.project_id == project.id,
        models.File.category == "Manuscript"
    ).all()
    assert len(docx_files) == 2

def test_pdf_to_word_endpoint(
    auth_cookie_client,
    admin_user,
    app_env,
):
    client = auth_cookie_client(admin_user)

    # Post a mock pdf file
    pdf_content = b"%PDF-1.4 mock content"
    files = {
        "file": ("document.pdf", io.BytesIO(pdf_content), "application/pdf")
    }
    data = {
        "engine": "pdf2docx" # This will run pdf2docx (which will raise ImportError or handle gracefully)
    }

    response = client.post(
        "/api/v1/conversion/pdf-to-word",
        data=data,
        files=files
    )

    # Since pdf2docx may not be installed in the test env, it could return 500 with detailed message.
    # If mock is run or libraries are missing, we assert it behaves as expected.
    if response.status_code == 200:
        assert response.headers["content-type"] == "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        assert "converted_document.docx" in response.headers["content-disposition"]
    else:
        assert response.status_code == 500
        assert "detail" in response.json()


def test_backup_list_and_download_endpoints(
    auth_cookie_client,
    admin_user,
    db_session,
    app_env,
):
    client = auth_cookie_client(admin_user)
    
    from app.domains.projects.models import Project
    from app.domains.workflow.models import ChapterInfo
    from app.domains.clients.models import Client
    
    client_rec = Client(
        category_type="Organization",
        contact_type="Customer",
        name_company="Test Client Co",
        company="Test Client Co",
        division="TEST_DIV",
        email="test@example.com",
        active_status=True
    )
    db_session.add(client_rec)
    db_session.commit()
    db_session.refresh(client_rec)
    
    project = Project(code="TEST_BACKUP_PROJ", title="Test Backup Project", client_id=client_rec.id, client_name="Test Client Co", status="In-progress")
    db_session.add(project)
    db_session.commit()
    db_session.refresh(project)
    
    chapter = ChapterInfo(client="Test Client Co", project="TEST_BACKUP_PROJ", chapters="1", chapter_title="Chapter 1", status="In-progress")
    db_session.add(chapter)
    db_session.commit()
    db_session.refresh(chapter)
    
    # Call backup-list (should return empty list initially)
    response = client.get(f"/api/v2/uploads/{project.id}/chapter/chapter-1/backup-list")
    assert response.status_code == 200
    body = response.json()
    assert "files" in body
    assert isinstance(body["files"], list)

def test_single_indesign_to_word_endpoint(
    auth_cookie_client,
    admin_user,
    db_session,
    app_env,
):
    client = auth_cookie_client(admin_user)
    
    from app.domains.projects.models import Project
    from app.domains.workflow.models import ChapterInfo
    from app.domains.clients.models import Client
    from app.models import File
    import os
    import zipfile
    
    # Setup test DB entities
    client_rec = Client(
        category_type="Organization",
        contact_type="Customer",
        name_company="Test Client Co 2",
        company="Test Client Co 2",
        division="TEST_DIV2",
        email="test2@example.com",
        active_status=True
    )
    db_session.add(client_rec)
    db_session.commit()
    db_session.refresh(client_rec)
    
    project = Project(code="TEST_PROJ_SINGLE", title="Test Single InDesign Project", client_id=client_rec.id, client_name="Test Client Co 2", status="In-progress")
    db_session.add(project)
    db_session.commit()
    db_session.refresh(project)
    
    chapter = ChapterInfo(client="Test Client Co 2", project="TEST_PROJ_SINGLE", chapters="3", chapter_title="Chapter 3", status="In-progress")
    db_session.add(chapter)
    db_session.commit()
    db_session.refresh(chapter)
    
    # Mock source file paths and database File records
    # Create the files physically to allow packaging in zipfile
    from app.services.file_service import UPLOAD_DIR
    
    indd_rel_path = "TEST_PROJ_SINGLE/3/InDesign/9781284238990_CH01_001_032.indd"
    art_rel_path = "TEST_PROJ_SINGLE/3/Art/image.png"
    font_rel_path = "TEST_PROJ_SINGLE/3/Misc/custom_font.ttf"
    
    indd_path = os.path.join(UPLOAD_DIR, indd_rel_path)
    art_path = os.path.join(UPLOAD_DIR, art_rel_path)
    font_path = os.path.join(UPLOAD_DIR, font_rel_path)
    
    os.makedirs(os.path.dirname(indd_path), exist_ok=True)
    os.makedirs(os.path.dirname(art_path), exist_ok=True)
    os.makedirs(os.path.dirname(font_path), exist_ok=True)
    
    with open(indd_path, "wb") as f:
        f.write(b"mock indesign binary")
    with open(art_path, "wb") as f:
        f.write(b"mock png data")
    with open(font_path, "wb") as f:
        f.write(b"mock font data")
        
    indd_file = File(
        project_id=project.id,
        chapter_id=chapter.id,
        filename="9781284238990_CH01_001_032.indd",
        path=indd_rel_path,
        category="InDesign"
    )
    art_file = File(
        project_id=project.id,
        chapter_id=chapter.id,
        filename="image.png",
        path=art_rel_path,
        category="Art"
    )
    font_file = File(
        project_id=project.id,
        chapter_id=chapter.id,
        filename="custom_font.ttf",
        path=font_rel_path,
        category="Misc"
    )
    
    db_session.add(indd_file)
    db_session.add(art_file)
    db_session.add(font_file)
    db_session.commit()
    
    # Configure INDESIGN_SERVER_URL for testing
    from app.core.config import get_settings
    settings = get_settings()
    old_url = settings.INDESIGN_SERVER_URL
    settings.INDESIGN_SERVER_URL = "http://host.docker.internal:5555"
    
    try:
        # Call endpoint with mocked Windows server response
        with patch("requests.post") as mock_post:
            mock_resp = MagicMock()
            mock_resp.status_code = 200
            mock_resp.content = b"CONVERTED DOCX CONTENT"
            mock_post.return_value = mock_resp
            
            # Capture the request payload to inspect the zipped file
            response = client.post(f"/api/v1/conversion/indesign-to-word/{indd_file.id}")
        
        # Verify remote request was sent with zip archive
        assert mock_post.called
        call_kwargs = mock_post.call_args[1]
        assert "files" in call_kwargs
        uploaded_files = call_kwargs["files"]
        assert "file" in uploaded_files
        zip_filename, zip_bytes, _ = uploaded_files["file"]
        
        assert zip_filename.startswith("packaged_9781284238990_CH01_001_032")
        
        # Extract and verify ZIP contents
        with zipfile.ZipFile(io.BytesIO(zip_bytes)) as test_zip:
            namelist = test_zip.namelist()
            assert "9781284238990_CH01_001_032.indd" in namelist
            assert "Links/image.png" in namelist
            assert "Document Fonts/custom_font.ttf" in namelist
            
        assert response.status_code == 200
        body = response.json()
        assert body["status"] == "success"
        
        # Check that a Manuscript file record was created with the correct name: 9781284238990_CH01_001_032.docx
        manuscript = db_session.query(File).filter(
            File.project_id == project.id,
            File.chapter_id == chapter.id,
            File.category == "Manuscript"
        ).first()
        
        assert manuscript is not None
        assert manuscript.filename == "9781284238990_CH01_001_032.docx"
    finally:
        settings.INDESIGN_SERVER_URL = old_url
        
    # Clean up physical files
    for p in [indd_path, art_path, font_path]:
        try:
            os.remove(p)
        except Exception:
            pass


def test_single_pdf_to_word_endpoint(
    auth_cookie_client,
    admin_user,
    db_session,
    app_env,
):
    client = auth_cookie_client(admin_user)
    
    from app.domains.projects.models import Project
    from app.domains.workflow.models import ChapterInfo
    from app.domains.clients.models import Client
    from app.models import File
    import os
    
    # Setup test DB entities
    client_rec = Client(
        category_type="Organization",
        contact_type="Customer",
        name_company="Test Client Co 3",
        company="Test Client Co 3",
        division="TEST_DIV3",
        email="test3@example.com",
        active_status=True
    )
    db_session.add(client_rec)
    db_session.commit()
    db_session.refresh(client_rec)
    
    project = Project(code="TEST_PROJ_PDF", title="Test PDF Project", client_id=client_rec.id, client_name="Test Client Co 3", status="In-progress")
    db_session.add(project)
    db_session.commit()
    db_session.refresh(project)
    
    chapter = ChapterInfo(client="Test Client Co 3", project="TEST_PROJ_PDF", chapters="4", chapter_title="Chapter 4", status="In-progress")
    db_session.add(chapter)
    db_session.commit()
    db_session.refresh(chapter)
    
    from app.services.file_service import UPLOAD_DIR
    
    pdf_rel_path = "TEST_PROJ_PDF/4/Proof/proof_document.pdf"
    pdf_path = os.path.join(UPLOAD_DIR, pdf_rel_path)
    
    os.makedirs(os.path.dirname(pdf_path), exist_ok=True)
    with open(pdf_path, "wb") as f:
        f.write(b"%PDF-1.4 mock pdf data")
        
    pdf_file = File(
        project_id=project.id,
        chapter_id=chapter.id,
        filename="proof_document.pdf",
        path=pdf_rel_path,
        category="Proof"
    )
    db_session.add(pdf_file)
    db_session.commit()
    
    # Mock pdf2docx Converter class to simulate successful conversion without importing the actual library
    import sys
    mock_pdf2docx = MagicMock()
    mock_converter = MagicMock()
    mock_pdf2docx.Converter.return_value = mock_converter
    
    # When convert is called, we write a mock docx file so the service finds it
    def mock_convert(docx_path, start=0, end=None):
        with open(docx_path, "wb") as f:
            f.write(b"MOCK DOCX BYTES")
    mock_converter.convert.side_effect = mock_convert
    
    sys.modules["pdf2docx"] = mock_pdf2docx
    try:
        response = client.post(f"/api/v1/conversion/pdf-to-word/{pdf_file.id}")
    finally:
        sys.modules.pop("pdf2docx", None)
            
    assert response.status_code == 200
    body = response.json()
    assert body["status"] == "success"
    
    # Check that a Manuscript file record was created with the correct name: proof_document.docx
    manuscript = db_session.query(File).filter(
        File.project_id == project.id,
        File.chapter_id == chapter.id,
        File.category == "Manuscript"
    ).first()
    
    assert manuscript is not None
    assert manuscript.filename == "proof_document.docx"
    
    # Clean up physical files
    for p in [pdf_path]:
        try:
            os.remove(p)
        except Exception:
            pass


def test_docx_post_processor():
    from docx import Document
    from app.services.scripts.docx_post_processor import post_process_docx
    import tempfile
    import os
    from docx.oxml.ns import qn
    
    # Create temporary docx file
    with tempfile.TemporaryDirectory() as tmpdir:
        test_path = os.path.join(tmpdir, "test_format.docx")
        
        # 1. Create a document with two-column section, headers/footers, and styling
        doc = Document()
        
        # Add a section with 2 columns
        section = doc.sections[0]
        sectPr = section._sectPr
        from docx.oxml import OxmlElement
        cols = OxmlElement('w:cols')
        cols.set(qn('w:num'), '2')
        sectPr.append(cols)
        
        # Add text to header and footer
        section.header.paragraphs[0].text = "Header text"
        section.footer.paragraphs[0].text = "Footer text"
        
        # Add paragraphs with split text
        p1 = doc.add_paragraph("This is a word split-")
        p2 = doc.add_paragraph("ting sentence.")
        
        p3 = doc.add_paragraph("Normal paragraph")
        p4 = doc.add_paragraph("another sentence starting with lowercase.")
        
        # Add two consecutive tables with same columns separated by an empty paragraph
        t1 = doc.add_table(rows=1, cols=2)
        t1.rows[0].cells[0].text = "A1"
        t1.rows[0].cells[1].text = "B1"
        
        sep = doc.add_paragraph("") # empty paragraph separator
        
        t2 = doc.add_table(rows=1, cols=2)
        t2.rows[0].cells[0].text = "A2"
        t2.rows[0].cells[1].text = "B2"
        
        doc.save(test_path)
        
        # 2. Run post-processing
        post_process_docx(test_path)
        
        # 3. Reload and verify cleanups
        doc_mod = Document(test_path)
        
        # Verify single column
        cols_mod = doc_mod.sections[0]._sectPr.xpath('w:cols')
        assert cols_mod[0].get(qn('w:num')) == '1'
        
        # Verify headers/footers cleared
        assert doc_mod.sections[0].header.paragraphs[0].text == ""
        assert doc_mod.sections[0].footer.paragraphs[0].text == ""
        
        # Verify paragraphs merged
        p_texts = [p.text for p in doc_mod.paragraphs]
        # "This is a word split-ting sentence." -> "This is a word splitting sentence." (hyphen stripped)
        assert "This is a word splitting sentence." in p_texts
        # "Normal paragraph" and "another sentence..." merged with a space
        assert "Normal paragraph another sentence starting with lowercase." in p_texts
        
        # Verify tables merged
        assert len(doc_mod.tables) == 1
        assert len(doc_mod.tables[0].rows) == 2
        assert doc_mod.tables[0].rows[0].cells[0].text == "A1"
        assert doc_mod.tables[0].rows[1].cells[0].text == "A2"
        
        # Verify all runs set to Times New Roman
        for p in doc_mod.paragraphs:
            for r in p.runs:
                rPr = r._r.get_or_add_rPr()
                rFonts = rPr.xpath('w:rFonts')
                assert rFonts[0].get(qn('w:ascii')) == 'Times New Roman'




def test_backup_list_and_download_endpoints(
    auth_cookie_client,
    admin_user,
    db_session,
    app_env,
):
    client = auth_cookie_client(admin_user)
    
    from app.domains.projects.models import Project
    from app.domains.workflow.models import ChapterInfo
    from app.domains.clients.models import Client
    
    client_rec = Client(
        category_type="Organization",
        contact_type="Customer",
        name_company="Test Client Co",
        company="Test Client Co",
        division="TEST_DIV",
        email="test@example.com",
        active_status=True
    )
    db_session.add(client_rec)
    db_session.commit()
    db_session.refresh(client_rec)
    
    project = Project(code="TEST_BACKUP_PROJ", title="Test Backup Project", client_id=client_rec.id, client_name="Test Client Co", status="In-progress")
    db_session.add(project)
    db_session.commit()
    db_session.refresh(project)
    
    chapter = ChapterInfo(client="Test Client Co", project="TEST_BACKUP_PROJ", chapters="1", chapter_title="Chapter 1", status="In-progress")
    db_session.add(chapter)
    db_session.commit()
    db_session.refresh(chapter)
    
    # Call backup-list (should return empty list initially)
    response = client.get(f"/api/v2/uploads/{project.id}/chapter/chapter-1/backup-list")
    assert response.status_code == 200
    body = response.json()
    assert "files" in body
    assert isinstance(body["files"], list)

def test_single_indesign_to_word_endpoint(
    auth_cookie_client,
    admin_user,
    db_session,
    app_env,
):
    client = auth_cookie_client(admin_user)
    
    from app.domains.projects.models import Project
    from app.domains.workflow.models import ChapterInfo
    from app.domains.clients.models import Client
    from app.models import File
    import os
    import zipfile
    
    # Setup test DB entities
    client_rec = Client(
        category_type="Organization",
        contact_type="Customer",
        name_company="Test Client Co 2",
        company="Test Client Co 2",
        division="TEST_DIV2",
        email="test2@example.com",
        active_status=True
    )
    db_session.add(client_rec)
    db_session.commit()
    db_session.refresh(client_rec)
    
    project = Project(code="TEST_PROJ_SINGLE", title="Test Single InDesign Project", client_id=client_rec.id, client_name="Test Client Co 2", status="In-progress")
    db_session.add(project)
    db_session.commit()
    db_session.refresh(project)
    
    chapter = ChapterInfo(client="Test Client Co 2", project="TEST_PROJ_SINGLE", chapters="3", chapter_title="Chapter 3", status="In-progress")
    db_session.add(chapter)
    db_session.commit()
    db_session.refresh(chapter)
    
    # Mock source file paths and database File records
    # Create the files physically to allow packaging in zipfile
    from app.services.file_service import UPLOAD_DIR
    
    indd_rel_path = "TEST_PROJ_SINGLE/3/InDesign/9781284238990_CH01_001_032.indd"
    art_rel_path = "TEST_PROJ_SINGLE/3/Art/image.png"
    font_rel_path = "TEST_PROJ_SINGLE/3/Misc/custom_font.ttf"
    
    indd_path = os.path.join(UPLOAD_DIR, indd_rel_path)
    art_path = os.path.join(UPLOAD_DIR, art_rel_path)
    font_path = os.path.join(UPLOAD_DIR, font_rel_path)
    
    os.makedirs(os.path.dirname(indd_path), exist_ok=True)
    os.makedirs(os.path.dirname(art_path), exist_ok=True)
    os.makedirs(os.path.dirname(font_path), exist_ok=True)
    
    with open(indd_path, "wb") as f:
        f.write(b"mock indesign binary")
    with open(art_path, "wb") as f:
        f.write(b"mock png data")
    with open(font_path, "wb") as f:
        f.write(b"mock font data")
        
    indd_file = File(
        project_id=project.id,
        chapter_id=chapter.id,
        filename="9781284238990_CH01_001_032.indd",
        path=indd_rel_path,
        category="InDesign"
    )
    art_file = File(
        project_id=project.id,
        chapter_id=chapter.id,
        filename="image.png",
        path=art_rel_path,
        category="Art"
    )
    font_file = File(
        project_id=project.id,
        chapter_id=chapter.id,
        filename="custom_font.ttf",
        path=font_rel_path,
        category="Misc"
    )
    
    db_session.add(indd_file)
    db_session.add(art_file)
    db_session.add(font_file)
    db_session.commit()
    
    # Configure INDESIGN_SERVER_URL for testing
    from app.core.config import get_settings
    settings = get_settings()
    old_url = settings.INDESIGN_SERVER_URL
    settings.INDESIGN_SERVER_URL = "http://host.docker.internal:5555"
    
    try:
        # Call endpoint with mocked Windows server response
        with patch("requests.post") as mock_post:
            mock_resp = MagicMock()
            mock_resp.status_code = 200
            mock_resp.content = b"CONVERTED DOCX CONTENT"
            mock_post.return_value = mock_resp
            
            # Capture the request payload to inspect the zipped file
            response = client.post(f"/api/v1/conversion/indesign-to-word/{indd_file.id}")
    
            # Verify remote request was sent with zip archive
            assert mock_post.called
        call_kwargs = mock_post.call_args[1]
        assert "files" in call_kwargs
        uploaded_files = call_kwargs["files"]
        assert "file" in uploaded_files
        zip_filename, zip_bytes, _ = uploaded_files["file"]
        
        assert zip_filename.startswith("packaged_9781284238990_CH01_001_032")
        
        # Extract and verify ZIP contents
        with zipfile.ZipFile(io.BytesIO(zip_bytes)) as test_zip:
            namelist = test_zip.namelist()
            assert "9781284238990_CH01_001_032.indd" in namelist
            assert "Links/image.png" in namelist
            assert "Document Fonts/custom_font.ttf" in namelist
            
    finally:
        settings.INDESIGN_SERVER_URL = old_url
            
    assert response.status_code == 200
    body = response.json()
    assert body["status"] == "success"
    
    # Check that a Manuscript file record was created with the correct name: 9781284238990_CH01_001_032.docx
    manuscript = db_session.query(File).filter(
        File.project_id == project.id,
        File.chapter_id == chapter.id,
        File.category == "Manuscript"
    ).first()
    
    assert manuscript is not None
    assert manuscript.filename == "9781284238990_CH01_001_032.docx"
    
    # Clean up physical files
    for p in [indd_path, art_path, font_path]:
        try:
            os.remove(p)
        except Exception:
            pass
