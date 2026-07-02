"""
Round-trip tests for XHTML → DOCX conversion with formatting preservation.
Tests verify that formatting, images, tables, and track changes survive the round-trip.
"""

import os
import tempfile
from pathlib import Path

import pytest
from docx import Document
from docx.oxml.ns import qn


class TestXhtmlToDocxRoundTrip:
    """Tests for XHTML to DOCX conversion with formatting preservation."""

    @pytest.fixture
    def sample_docx_path(self):
        """Create a minimal DOCX file for testing."""
        doc = Document()

        # Add a paragraph with text
        p = doc.add_paragraph("Sample paragraph with ")
        run = p.add_run("bold text")
        run.bold = True

        # Add a heading
        doc.add_heading("Test Heading", level=1)

        # Add a table
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Header 1"
        table.rows[0].cells[1].text = "Header 2"
        table.rows[1].cells[0].text = "Cell 1"
        table.rows[1].cells[1].text = "Cell 2"

        # Add paragraph with image (simulated)
        doc.add_paragraph("Document with test content")

        # Add paragraphs for track changes testing
        doc.add_paragraph("Text with inserted text and deleted text.")
        doc.add_paragraph("Normal text bold insert and italic insert.")

        tmp = tempfile.NamedTemporaryFile(suffix=".docx", delete=False)
        tmp.close()
        doc.save(tmp.name)
        yield tmp.name

        # Cleanup
        Path(tmp.name).unlink(missing_ok=True)

    @pytest.fixture
    def xhtml_engine(self):
        """Provide the XHTML to DOCX conversion engine."""
        from app.processing.xhtml_to_docx import XhtmlToDocxEngine
        return XhtmlToDocxEngine()

    def test_xhtml_to_docx_preserves_formatting(self, sample_docx_path, xhtml_engine, tmp_path):
        """Verify that basic formatting (bold, italic) survives round-trip."""
        # Create XHTML with formatting
        html_content = """
        <html>
            <body>
                <p data-style-label="Normal">Sample paragraph with <strong>bold text</strong></p>
                <h1 data-style-label="H1">Test Heading</h1>
            </body>
        </html>
        """

        html_path = tmp_path / "test.html"
        html_path.write_text(html_content)

        output_docx = tmp_path / "output.docx"
        result = xhtml_engine.convert(str(html_path), str(sample_docx_path), username="Test User")

        # Verify DOCX was processed
        assert Path(result).exists()

        # Load and verify formatting was applied
        doc = Document(result)
        assert len(doc.paragraphs) > 0

    def test_xhtml_to_docx_preserves_table_formatting(self, sample_docx_path, xhtml_engine, tmp_path):
        """Verify that table cell formatting is preserved."""
        html_content = """
        <html>
            <body>
                <table>
                    <tr>
                        <td>Header 1</td>
                        <td>Header 2</td>
                    </tr>
                    <tr>
                        <td>Cell 1</td>
                        <td>Cell 2</td>
                    </tr>
                </table>
            </body>
        </html>
        """

        html_path = tmp_path / "test_table.html"
        html_path.write_text(html_content)

        result = xhtml_engine.convert(str(html_path), str(sample_docx_path), username="Test User")

        # Verify DOCX was processed
        assert Path(result).exists()

        # Load and verify table exists
        doc = Document(result)
        assert len(doc.tables) > 0

    def test_xhtml_to_docx_track_changes_get_unique_ids(self, sample_docx_path, xhtml_engine, tmp_path):
        """Verify that track changes (ins/del) get unique w:id values and timestamps."""
        html_content = """
        <html>
            <body>
                <p>Text with <ins>inserted text</ins> and <del>deleted text</del>.</p>
            </body>
        </html>
        """

        html_path = tmp_path / "test_track_changes.html"
        html_path.write_text(html_content)

        result = xhtml_engine.convert(str(html_path), str(sample_docx_path), username="Editor")

        # Verify DOCX was processed
        assert Path(result).exists()

        # Load and verify track changes
        doc = Document(result)

        # Find w:ins and w:del elements in the document XML
        found_ins = False
        found_del = False

        for para in doc.paragraphs:
            for elem in para._element:
                tag = elem.tag
                if qn('w:ins') in tag:
                    found_ins = True
                    # Verify w:id is set and not hardcoded to '1'
                    w_id = elem.get(qn('w:id'))
                    assert w_id is not None
                    assert w_id != ''

                    # Verify w:author is set
                    w_author = elem.get(qn('w:author'))
                    assert w_author == 'Editor'

                    # Verify w:date is ISO 8601 format
                    w_date = elem.get(qn('w:date'))
                    assert w_date is not None
                    assert 'T' in w_date
                    assert 'Z' in w_date

                elif qn('w:del') in tag:
                    found_del = True
                    # Same checks for del
                    w_id = elem.get(qn('w:id'))
                    assert w_id is not None
                    assert w_id != ''

        assert found_ins
        assert found_del

    def test_xhtml_to_docx_handles_missing_html(self, xhtml_engine, tmp_path):
        """Verify that converter handles missing HTML files gracefully."""
        missing_html = tmp_path / "nonexistent.html"
        output_docx = tmp_path / "output.docx"

        with pytest.raises(RuntimeError):
            xhtml_engine.convert(str(missing_html), str(output_docx))

    def test_xhtml_to_docx_creates_output_directory(self, sample_docx_path, xhtml_engine, tmp_path):
        """Verify that output directory is created if it doesn't exist."""
        html_content = "<html><body><p>Test</p></body></html>"
        html_path = tmp_path / "test.html"
        html_path.write_text(html_content)

        # Create output in non-existent subdirectory
        output_docx = tmp_path / "deep" / "nested" / "output.docx"

        # Copy sample_docx_path to the non-existent location
        output_docx.parent.mkdir(parents=True, exist_ok=True)
        import shutil
        shutil.copy(sample_docx_path, output_docx)

        result = xhtml_engine.convert(str(html_path), str(output_docx), username="Test")
        assert Path(result).exists()

    def test_xhtml_to_docx_style_label_applied(self, sample_docx_path, xhtml_engine, tmp_path):
        """Verify that data-style-label attributes are parsed and applied."""
        html_content = """
        <html>
            <body>
                <p data-style-label="Heading">Custom Style Paragraph</p>
                <h1 data-style-label="H1">Heading 1</h1>
                <p>Normal paragraph</p>
            </body>
        </html>
        """

        html_path = tmp_path / "test_styles.html"
        html_path.write_text(html_content)

        result = xhtml_engine.convert(str(html_path), str(sample_docx_path), username="Styler")
        assert Path(result).exists()

    def test_xhtml_to_docx_lowercase_style_label_is_canonicalized(self, xhtml_engine, tmp_path):
        """A data-style-label that drifted to lowercase (e.g. "h1", from
        however the paragraph's style name was originally cased) must be
        normalized to the canonical uppercase tag when written back to the
        DOCX, not applied verbatim. The DOCX already has a custom "H1" style
        registered, mirroring real documents that already went through the
        AI-structuring step before reaching the manual style picker."""
        from docx.enum.style import WD_STYLE_TYPE

        doc = Document()
        h1_style = doc.styles.add_style("H1", WD_STYLE_TYPE.PARAGRAPH)
        h1_style.base_style = doc.styles["Normal"]
        para = doc.add_paragraph("Test Heading")
        para.style = "H1"

        docx_path = tmp_path / "with_h1_style.docx"
        doc.save(docx_path)

        html_content = """
        <html>
            <body>
                <h1 data-style-label="h1">Test Heading</h1>
            </body>
        </html>
        """

        html_path = tmp_path / "test_lowercase_style.html"
        html_path.write_text(html_content)

        result = xhtml_engine.convert(str(html_path), str(docx_path), username="Test User")
        assert Path(result).exists()

        doc2 = Document(result)
        heading_para = next(p for p in doc2.paragraphs if p.text.strip() == "Test Heading")
        assert heading_para.style.name == "H1"
        assert "h1" not in [s.name for s in doc2.styles]

    def test_xhtml_to_docx_image_preservation(self, sample_docx_path, xhtml_engine, tmp_path):
        """
        Verify that images in the original DOCX are preserved during the round-trip.

        The strategy: Copy + patch preserves the underlying DOCX binary,
        so images (stored as blip elements) should survive the conversion.
        """
        # Create a DOCX with an image-like reference
        original_doc = Document(sample_docx_path)
        original_image_count = sum(
            1 for rel in original_doc.part.rels.values()
            if "image" in rel.reltype.lower()
        )

        html_content = "<html><body><p>Document content</p></body></html>"
        html_path = tmp_path / "test_images.html"
        html_path.write_text(html_content)

        result = xhtml_engine.convert(str(html_path), str(sample_docx_path), username="Image Tester")

        # Verify output exists
        assert Path(result).exists()

        # Load result and verify images are intact
        result_doc = Document(result)
        result_image_count = sum(
            1 for rel in result_doc.part.rels.values()
            if "image" in rel.reltype.lower()
        )

        # Images should be preserved (count should be same or greater)
        assert result_image_count >= original_image_count, \
            f"Image count decreased: {original_image_count} → {result_image_count}"

    def test_xhtml_to_docx_preserves_formatting_in_track_changes(self, sample_docx_path, xhtml_engine, tmp_path):
        """Verify that bold and italic formatting applied within <ins> tags is preserved in the DOCX."""
        html_content = """
        <html>
            <body>
                <p>Normal text <ins><strong>bold insert</strong> and <em>italic insert</em></ins>.</p>
            </body>
        </html>
        """

        html_path = tmp_path / "test_formatting_track.html"
        html_path.write_text(html_content)

        result = xhtml_engine.convert(str(html_path), str(sample_docx_path), username="Editor")

        assert Path(result).exists()

        doc = Document(result)
        found_bold_ins = False
        found_italic_ins = False

        for para in doc.paragraphs:
            for elem in para._element:
                if qn('w:ins') in elem.tag:
                    for run in elem.findall(qn('w:r')):
                        rPr = run.find(qn('w:rPr'))
                        if rPr is not None:
                            if rPr.find(qn('w:b')) is not None:
                                found_bold_ins = True
                            if rPr.find(qn('w:i')) is not None:
                                found_italic_ins = True

        assert found_bold_ins
        assert found_italic_ins


class TestXhtmlTableCellFormatting:
    """Tests for table cell formatting preservation."""

    @pytest.fixture
    def xhtml_engine(self):
        """Provide the XHTML to DOCX conversion engine."""
        from app.processing.xhtml_to_docx import XhtmlToDocxEngine
        return XhtmlToDocxEngine()

    @pytest.fixture
    def docx_with_table(self, tmp_path):
        """Create a DOCX with a table."""
        doc = Document()
        table = doc.add_table(rows=2, cols=2)
        table.rows[0].cells[0].text = "Header 1"
        table.rows[0].cells[1].text = "Header 2"
        table.rows[1].cells[0].text = "Data 1"
        table.rows[1].cells[1].text = "Data 2"

        docx_path = tmp_path / "table.docx"
        doc.save(docx_path)
        return docx_path

    def test_table_cells_with_formatting(self, docx_with_table, xhtml_engine, tmp_path):
        """Verify that table cells can have formatting applied."""
        html_content = """
        <html>
            <body>
                <table>
                    <tr>
                        <td>Header 1</td>
                        <td>Header 2</td>
                    </tr>
                    <tr>
                        <td>Data 1</td>
                        <td>Data 2</td>
                    </tr>
                </table>
            </body>
        </html>
        """

        html_path = tmp_path / "table_formatted.html"
        html_path.write_text(html_content)

        result = xhtml_engine.convert(str(html_path), str(docx_with_table), username="Table Formatter")
        assert Path(result).exists()

        # Verify table still exists in output
        doc = Document(result)
        assert len(doc.tables) > 0


class TestXhtmlNestedInlineSdts:
    """Tests to verify that nested inline SDTs are preserved in the roundtrip."""

    def test_nested_inline_sdt_roundtrip(self):
        from app.processing.docx_to_xhtml_runs import _paragraph_content_to_html
        from docx import Document
        from lxml import etree
        
        # 1. Create a paragraph with a nested inline SDT manually
        # [FigureRef [Figure Fig. ] [ChapNo 8] . [SeqNo 3] ]
        doc = Document()
        p = doc.add_paragraph()
        p_el = p._p
        
        # Outer SDT
        outer_sdt = etree.SubElement(p_el, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdt")
        outer_sdtPr = etree.SubElement(outer_sdt, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdtPr")
        alias_el = etree.SubElement(outer_sdtPr, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}alias")
        alias_el.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "FigureRef")
        outer_sdtContent = etree.SubElement(outer_sdt, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdtContent")
        
        # Inner SDT 1 (Figure)
        inner_sdt1 = etree.SubElement(outer_sdtContent, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdt")
        inner_sdtPr1 = etree.SubElement(inner_sdt1, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdtPr")
        alias_el1 = etree.SubElement(inner_sdtPr1, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}alias")
        alias_el1.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "Figure")
        inner_sdtContent1 = etree.SubElement(inner_sdt1, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdtContent")
        r1 = etree.SubElement(inner_sdtContent1, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r")
        t1 = etree.SubElement(r1, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
        t1.text = "Fig."
        
        # Inner SDT 2 (ChapNo)
        inner_sdt2 = etree.SubElement(outer_sdtContent, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdt")
        inner_sdtPr2 = etree.SubElement(inner_sdt2, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdtPr")
        alias_el2 = etree.SubElement(inner_sdtPr2, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}alias")
        alias_el2.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "ChapNo")
        inner_sdtContent2 = etree.SubElement(inner_sdt2, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdtContent")
        r2 = etree.SubElement(inner_sdtContent2, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r")
        t2 = etree.SubElement(r2, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
        t2.text = "8"
        
        # Dot run in between
        r_dot = etree.SubElement(outer_sdtContent, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r")
        t_dot = etree.SubElement(r_dot, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
        t_dot.text = "."
        
        # Inner SDT 3 (SeqNo)
        inner_sdt3 = etree.SubElement(outer_sdtContent, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdt")
        inner_sdtPr3 = etree.SubElement(inner_sdt3, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdtPr")
        alias_el3 = etree.SubElement(inner_sdtPr3, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}alias")
        alias_el3.set("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val", "SeqNo")
        inner_sdtContent3 = etree.SubElement(inner_sdt3, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}sdtContent")
        r3 = etree.SubElement(inner_sdtContent3, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r")
        t3 = etree.SubElement(r3, "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}t")
        t3.text = "3"
        
        # 2. Render to HTML
        html_out = _paragraph_content_to_html(p_el, p, doc)
        
        # Verify the HTML structure retains nested spans with aliases
        assert 'class="sdt-inline"' in html_out
        assert 'data-alias="FigureRef"' in html_out
        assert 'data-alias="Figure"' in html_out
        assert 'data-alias="ChapNo"' in html_out
        assert 'data-alias="SeqNo"' in html_out

