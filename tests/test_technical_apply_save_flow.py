import os
import json
import tempfile
from pathlib import Path
from lxml import etree
from docx import Document
from docx.oxml.ns import qn

from app.processing.docx_to_xhtml_runs import DocxToXhtmlRunsEngine
from app.processing.xhtml_to_docx_delta import XhtmlToDocxDeltaEngine
from app.domains.processing.technical_editor_service import RESULTS_DIR

def test_technical_apply_save_flow():
    # 1. Create a mock DOCX file containing a yellow-highlighted text run
    doc = Document()
    p = doc.add_paragraph()
    r1 = p.add_run("The ")
    
    # Yellow highlighted run representing a bias term finding
    r2 = p.add_run("elderly")
    r2.font.highlight_color = 7 # Yellow
    
    r3 = p.add_run(" patient was seen.")
    
    fd, docx_path = tempfile.mkstemp(suffix=".docx")
    os.close(fd)
    
    file_id = 99999
    
    # Ensure RESULTS_DIR exists for mock findings scan cache
    RESULTS_DIR.mkdir(parents=True, exist_ok=True)
    cache_path = RESULTS_DIR / f"{file_id}_scan.json"
    
    try:
        doc.save(docx_path)
        
        # 2. Mock the scan results containing the replacement suggestion
        scan_data = {
            "findings": [
                {
                    "para_index": 0,
                    "match_start": 4,
                    "surface": "elderly",
                    "replacement": "older adults",
                    "category": "bias",
                    "rule_id": "bias_elderly"
                }
            ]
        }
        cache_path.write_text(json.dumps(scan_data, ensure_ascii=False), encoding="utf-8")
        
        # 3. Export to XHTML using DocxToXhtmlRunsEngine, passing the file_id
        exporter = DocxToXhtmlRunsEngine()
        html_content = exporter.convert(docx_path, file_id=file_id)
        
        # Verify that HTML has custom attributes and styling for the highlight replacement
        assert "data-replacement" in html_content
        assert "older adults" in html_content
        assert "data-rule-category" in html_content
        assert "bias" in html_content
        assert "occurrence-highlight" in html_content
        assert "occurrence-bias" in html_content
        assert "background-color:yellow" in html_content
        
        # 4. Save the HTML to a file and delta-patch it back to the DOCX
        fd_h, html_path = tempfile.mkstemp(suffix=".html")
        os.close(fd_h)
        try:
            with open(html_path, "w", encoding="utf-8") as f:
                f.write(html_content)
                
            importer = XhtmlToDocxDeltaEngine()
            importer.convert(html_path, docx_path, username="Test Author")
            
            # 5. Read the patched DOCX and verify it has track changes w:ins and w:del
            doc_patched = Document(docx_path)
            p_elem = doc_patched.paragraphs[0]._element
            
            # Inspect children of the paragraph XML element to check for del and ins
            tags = [etree.QName(child.tag).localname for child in p_elem]
            
            assert 'del' in tags
            assert 'ins' in tags
            
            # Verify deleted text is "elderly"
            del_node = p_elem.find(qn('w:del'))
            del_text_node = del_node.find(f".//{qn('w:delText')}")
            assert del_text_node.text == "elderly"
            
            # Verify inserted text is "older adults"
            ins_node = p_elem.find(qn('w:ins'))
            ins_text_node = ins_node.find(f".//{qn('w:t')}")
            assert ins_text_node.text == "older adults"
            
        finally:
            if os.path.exists(html_path):
                os.remove(html_path)
                
    finally:
        if os.path.exists(docx_path):
            os.remove(docx_path)
        if cache_path.exists():
            cache_path.unlink()
