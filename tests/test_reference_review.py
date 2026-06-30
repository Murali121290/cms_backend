import pytest
from pathlib import Path
from docx import Document
from app import models
from app.domains.review import service as structuring_review_service

def _build_docx_with_styles(path: Path, paragraphs_with_styles: list[tuple[str, str]]) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    for text, style_name in paragraphs_with_styles:
        p = doc.add_paragraph(text)
        try:
            p.style = style_name
        except KeyError:
            doc.styles.add_style(style_name, 1) # WD_STYLE_TYPE.PARAGRAPH
            p.style = style_name
    doc.save(path)
    return path

def test_reference_review_detected_style_ama(
    db_session,
    temp_upload_root,
    project_record,
    chapter_record,
    auth_cookie_client,
    admin_user,
):
    # Create an AMA/Numerical document
    file_path = temp_upload_root / project_record.code / chapter_record.number / "Manuscript"
    file_path.mkdir(parents=True, exist_ok=True)
    
    # original file
    original_filename = "chapter01.docx"
    original_docx = _build_docx_with_styles(file_path / original_filename, [("Text with citation [1].", "Normal")])
    
    original_record = models.File(
        project_id=project_record.id,
        chapter_id=chapter_record.id,
        filename=original_filename,
        file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        category="Manuscript",
        path=str(original_docx),
        version=1,
    )
    db_session.add(original_record)
    db_session.commit()
    db_session.refresh(original_record)

    # processed file
    processed_filename = "chapter01_Processed.docx"
    paragraphs = [
        ("Some text with a citation [1].", "Normal"),
        ("1. Smith J. Reference 1.", "REF-N"),
    ]
    processed_docx = _build_docx_with_styles(file_path / processed_filename, paragraphs)
    
    processed_record = models.File(
        project_id=project_record.id,
        chapter_id=chapter_record.id,
        filename=processed_filename,
        file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        category="Manuscript",
        path=str(processed_docx),
        version=1,
    )
    db_session.add(processed_record)
    db_session.commit()
    db_session.refresh(processed_record)

    # Let's call the API client
    client = auth_cookie_client(admin_user)
    response = client.get(f"/api/v2/files/{processed_record.id}/reference-review")
    
    assert response.status_code == 200
    data = response.json()
    assert data["status"] == "ok"
    assert data["validation_logs"]["detected_style"] == "AMA"


def test_reference_review_detected_style_apa(
    db_session,
    temp_upload_root,
    project_record,
    chapter_record,
    auth_cookie_client,
    admin_user,
):
    # Create an APA/Author-Year document with REF-U paragraphs
    file_path = temp_upload_root / project_record.code / chapter_record.number / "Manuscript"
    file_path.mkdir(parents=True, exist_ok=True)
    
    # original file
    original_filename = "chapter02.docx"
    original_docx = _build_docx_with_styles(file_path / original_filename, [("Text with citation (Smith, 2020).", "Normal")])
    
    original_record = models.File(
        project_id=project_record.id,
        chapter_id=chapter_record.id,
        filename=original_filename,
        file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        category="Manuscript",
        path=str(original_docx),
        version=1,
    )
    db_session.add(original_record)
    db_session.commit()
    db_session.refresh(original_record)

    # processed file
    processed_filename = "chapter02_Processed.docx"
    paragraphs = [
        ("Some text with a citation (Smith, 2020).", "Normal"),
        ("Smith, J. (2020). Reference 2.", "REF-U"),
    ]
    processed_docx = _build_docx_with_styles(file_path / processed_filename, paragraphs)
    
    processed_record = models.File(
        project_id=project_record.id,
        chapter_id=chapter_record.id,
        filename=processed_filename,
        file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        category="Manuscript",
        path=str(processed_docx),
        version=1,
    )
    db_session.add(processed_record)
    db_session.commit()
    db_session.refresh(processed_record)

    # Let's call the API client
    client = auth_cookie_client(admin_user)
    response = client.get(f"/api/v2/files/{processed_record.id}/reference-review")
    
    assert response.status_code == 200
    data = response.json()
    assert data["status"] == "ok"
    assert data["validation_logs"]["detected_style"] == "APA"


def test_reference_review_style_override(
    db_session,
    temp_upload_root,
    project_record,
    chapter_record,
    auth_cookie_client,
    admin_user,
):
    file_path = temp_upload_root / project_record.code / chapter_record.number / "Manuscript"
    file_path.mkdir(parents=True, exist_ok=True)
    
    # Create an AMA/Numerical document with REF-N style
    filename = "chapter03_Processed.docx"
    paragraphs = [
        ("Some text with a citation [1].", "Normal"),
        ("1. Smith J. Reference 1.", "REF-N"),
    ]
    docx_file = _build_docx_with_styles(file_path / filename, paragraphs)
    
    file_record = models.File(
        project_id=project_record.id,
        chapter_id=chapter_record.id,
        filename=filename,
        file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        category="Manuscript",
        path=str(docx_file),
        version=1,
    )
    db_session.add(file_record)
    db_session.commit()
    db_session.refresh(file_record)

    client = auth_cookie_client(admin_user)
    
    # 1. Check default auto-detection (should be AMA)
    response = client.get(f"/api/v2/files/{file_record.id}/reference-review")
    assert response.status_code == 200
    data = response.json()
    assert data["validation_logs"]["detected_style"] == "AMA"
    
    # 2. Check override with APA
    response = client.get(f"/api/v2/files/{file_record.id}/reference-review?style=APA")
    assert response.status_code == 200
    data = response.json()
    assert data["validation_logs"]["detected_style"] == "APA"
    
    # 3. Check validate-only endpoint override with APA
    response = client.get(f"/api/v2/files/{file_record.id}/reference-review/validate-only?style=APA")
    assert response.status_code == 200
    data = response.json()
    assert data["detected_style"] == "APA"

    # 4. Check validate-only endpoint override with AMA
    response = client.get(f"/api/v2/files/{file_record.id}/reference-review/validate-only?style=AMA")
    assert response.status_code == 200
    data = response.json()
    assert data["detected_style"] == "AMA"


def test_reference_review_cache_hit(
    db_session,
    temp_upload_root,
    project_record,
    chapter_record,
    auth_cookie_client,
    admin_user,
    caplog,
):
    import logging
    file_path = temp_upload_root / project_record.code / chapter_record.number / "Manuscript"
    file_path.mkdir(parents=True, exist_ok=True)
    
    # Processed file
    filename = "chapter04_Processed.docx"
    paragraphs = [
        ("Some text with a citation (Smith, 2020).", "Normal"),
        ("Smith, J. (2020). Reference 2.", "REF-U"),
    ]
    docx_file = _build_docx_with_styles(file_path / filename, paragraphs)
    
    file_record = models.File(
        project_id=project_record.id,
        chapter_id=chapter_record.id,
        filename=filename,
        file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        category="Manuscript",
        path=str(docx_file),
        version=1,
    )
    db_session.add(file_record)
    db_session.commit()
    db_session.refresh(file_record)

    client = auth_cookie_client(admin_user)
    
    # First request: should be a MISS
    with caplog.at_level(logging.INFO):
        response1 = client.get(f"/api/v2/files/{file_record.id}/reference-review")
        assert response1.status_code == 200
        assert any("Reference review cache MISS" in record.message for record in caplog.records)
    
    caplog.clear()
    
    # Second request: should be a HIT
    with caplog.at_level(logging.INFO):
        response2 = client.get(f"/api/v2/files/{file_record.id}/reference-review")
        assert response2.status_code == 200
        assert any("Reference review cache HIT" in record.message for record in caplog.records)
        assert not any("Reference review cache MISS" in record.message for record in caplog.records)


def test_reference_review_validate_only_renumbering_and_swap(
    db_session,
    temp_upload_root,
    project_record,
    chapter_record,
    auth_cookie_client,
    admin_user,
):
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    
    file_path = temp_upload_root / project_record.code / chapter_record.number / "Manuscript"
    file_path.mkdir(parents=True, exist_ok=True)
    
    filename = "chapter05_Processed.docx"
    doc_path = file_path / filename
    
    # Build a document with a citation and matching references
    doc = Document()
    p = doc.add_paragraph("This is some text with a citation ")
    
    # Ensure 'cite_bib' character style exists
    try:
        cite_style = doc.styles['cite_bib']
    except KeyError:
        cite_style = doc.styles.add_style('cite_bib', WD_STYLE_TYPE.CHARACTER)
        cite_style.font.superscript = True
    
    # Add two citation runs (both must be cited to avoid unused refs abort)
    run_cite = p.add_run("(1).")
    run_cite.style = 'cite_bib'
    
    p.add_run(" and also ")
    run_cite2 = p.add_run("(2).")
    run_cite2.style = 'cite_bib'
    
    # Also add standard REF-N bibliography paragraphs
    try:
        doc.styles['bib_number']
    except KeyError:
        doc.styles.add_style('bib_number', WD_STYLE_TYPE.CHARACTER)
        
    # Duplicate references to test 80% similarity threshold:
    # A: "Smith J. Reference validation and styling systems. J Style. 2026;1(1):10-20."
    # B: "Smith J. Reference validation and styling system. J Style. 2026;1(1):10-20." (fuzzy duplicate)
    bib_texts = [
        "Smith J. Reference validation and styling systems. J Style. 2026;1(1):10-20.",
        "Smith J. Reference validation and styling system. J Style. 2026;1(1):10-20.",
    ]
    
    for i, bib_text in enumerate(bib_texts):
        p_bib = doc.add_paragraph()
        try:
            p_bib.style = 'REF-N'
        except KeyError:
            doc.styles.add_style('REF-N', WD_STYLE_TYPE.PARAGRAPH)
            p_bib.style = 'REF-N'
            
        r_num = p_bib.add_run(f"{i+1}.")
        r_num.style = 'bib_number'
        p_bib.add_run(f" {bib_text}")
        
    doc.save(doc_path)
    
    file_record = models.File(
        project_id=project_record.id,
        chapter_id=chapter_record.id,
        filename=filename,
        file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        category="Manuscript",
        path=str(doc_path),
        version=1,
    )
    db_session.add(file_record)
    db_session.commit()
    db_session.refresh(file_record)
    
    client = auth_cookie_client(admin_user)
    
    # Validate-only endpoint call
    response = client.get(f"/api/v2/files/{file_record.id}/reference-review/validate-only?style=AMA")
    assert response.status_code == 200
    data = response.json()
    
    validation_logs = data["validation_logs"]
    assert validation_logs["detected_style"] == "AMA"
    assert len(validation_logs["duplicates"]) == 0
    
    # Since duplicates are merged, and renumbering is performed, the file version should have incremented
    db_session.refresh(file_record)
    assert file_record.version == 2
    
    # Reload and check that the document was renumbered / formatted
    saved_doc = Document(file_record.path)
    # The citation should have punctuation swapped: (1). -> .(1)
    first_para = saved_doc.paragraphs[0]
    # Extract text including track changes insertions (which standard paragraph.text ignores)
    first_para_text = "".join(t.text for t in first_para._element.xpath('.//w:t[not(ancestor::w:del)]') if t.text)
    assert ".(1)" in first_para_text
    
    # The duplicate reference entry should have been merged/deleted
    bib_paras = [p for p in saved_doc.paragraphs if p.style and p.style.name == 'REF-N']
    assert len(bib_paras) == 1


def test_reference_review_superscript_unicode_validate(
    db_session,
    temp_upload_root,
    project_record,
    chapter_record,
    auth_cookie_client,
    admin_user,
):
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    
    file_path = temp_upload_root / project_record.code / chapter_record.number / "Manuscript"
    file_path.mkdir(parents=True, exist_ok=True)
    
    filename = "chapter06_Processed.docx"
    doc_path = file_path / filename
    
    # Build a document with superscript citation runs
    doc = Document()
    p = doc.add_paragraph("This text has citations: ")
    
    # Let's add a run with superscript text directly
    # We can use unicode superscript characters like ³⁻⁵ and ¹
    p.add_run("some citations")
    run_cite = p.add_run("³⁻⁵") # cites 3, 4, 5
    
    p.add_run(" and another one ")
    run_cite2 = p.add_run("¹") # cites 1
    
    # Also add standard REF-N bibliography paragraphs for 1, 3, 4, 5
    # (Since unused refs are not allowed, we need all of them to be cited)
    # The order in text is 3, 4, 5, 1. So appearance order: 3, 4, 5, 1.
    # Renumbering will map:
    # 3 -> 1
    # 4 -> 2
    # 5 -> 3
    # 1 -> 4
    try:
        doc.styles['bib_number']
    except KeyError:
        doc.styles.add_style('bib_number', WD_STYLE_TYPE.CHARACTER)
        
    bib_entries = {
        1: "First bibliography entry with unique content to avoid similarity threshold.",
        3: "Third bibliography entry with completely different set of words and topics.",
        4: "Fourth bibliography entry which mentions unrelated things like computers and sky.",
        5: "Fifth bibliography entry containing random text about geography and history.",
    }
    
    for ref_id in [1, 3, 4, 5]:
        p_bib = doc.add_paragraph()
        try:
            p_bib.style = 'REF-N'
        except KeyError:
            doc.styles.add_style('REF-N', WD_STYLE_TYPE.PARAGRAPH)
            p_bib.style = 'REF-N'
            
        r_num = p_bib.add_run(f"{ref_id}.")
        r_num.style = 'bib_number'
        p_bib.add_run(f" {bib_entries[ref_id]}")
        
    doc.save(doc_path)
    
    file_record = models.File(
        project_id=project_record.id,
        chapter_id=chapter_record.id,
        filename=filename,
        file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        category="Manuscript",
        path=str(doc_path),
        version=1,
    )
    db_session.add(file_record)
    db_session.commit()
    db_session.refresh(file_record)
    
    client = auth_cookie_client(admin_user)
    
    # Validate-only endpoint call with style=AMA and citation_format=superscript
    response = client.get(f"/api/v2/files/{file_record.id}/reference-review/validate-only?style=AMA&citation_format=superscript")
    assert response.status_code == 200
    data = response.json()
    
    validation_logs = data["validation_logs"]
    assert validation_logs["detected_style"] == "AMA"
    assert validation_logs["total_cites"] == 4 # 3, 4, 5, 1
    
    # Check that renumbering was executed
    db_session.refresh(file_record)
    assert file_record.version == 2
    
    # Reload and check that the document was renumbered/formatted back to superscripts
    saved_doc = Document(file_record.path)
    
    # The first citation run "³⁻⁵" (3-5) should be renumbered to "1-3" -> "¹⁻³"
    # The second citation run "¹" (1) should be renumbered to "4" -> "⁴"
    first_para = saved_doc.paragraphs[0]
    first_para_text = "".join(t.text for t in first_para._element.xpath('.//w:t[not(ancestor::w:del)]') if t.text)
    
    assert "¹⁻³" in first_para_text
    assert "⁴" in first_para_text


def test_reference_review_two_pass_validation(
    db_session,
    temp_upload_root,
    project_record,
    chapter_record,
    auth_cookie_client,
    admin_user,
):
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    
    file_path = temp_upload_root / project_record.code / chapter_record.number / "Manuscript"
    file_path.mkdir(parents=True, exist_ok=True)
    
    filename = "chapter07_Processed.docx"
    doc_path = file_path / filename
    
    # Build a document with sequence issues and duplicates
    # Citations in text: 1, 4, 3, 5
    doc = Document()
    
    # Ensure 'cite_bib' character style exists
    try:
        cite_style = doc.styles['cite_bib']
    except KeyError:
        cite_style = doc.styles.add_style('cite_bib', WD_STYLE_TYPE.CHARACTER)
        cite_style.font.superscript = True

    p = doc.add_paragraph("This is ")
    r1 = p.add_run("1")
    r1.style = 'cite_bib'
    p.add_run(" then ")
    r4 = p.add_run("4")
    r4.style = 'cite_bib'
    p.add_run(" then ")
    r3 = p.add_run("3")
    r3.style = 'cite_bib'
    p.add_run(" then ")
    r5 = p.add_run("5")
    r5.style = 'cite_bib'
    
    # Bibliography REF-N entries: 1, 3, 4, 5
    # Duplicate entry: 4 is duplicate of 3
    try:
        doc.styles['bib_number']
    except KeyError:
        doc.styles.add_style('bib_number', WD_STYLE_TYPE.CHARACTER)
        
    bib_entries = {
        1: "Reference 1 text about medicine and biology.",
        3: "Duplicate reference text about neural network architectures and deep learning.",
        4: "Duplicate reference text about neural network architectures and deep learning.",
        5: "Reference 5 text about database scaling and sharding.",
    }
    
    for ref_id in [1, 3, 4, 5]:
        p_bib = doc.add_paragraph()
        try:
            p_bib.style = 'REF-N'
        except KeyError:
            doc.styles.add_style('REF-N', WD_STYLE_TYPE.PARAGRAPH)
            p_bib.style = 'REF-N'
            
        r_num = p_bib.add_run(f"{ref_id}.")
        r_num.style = 'bib_number'
        p_bib.add_run(f" {bib_entries[ref_id]}")
        
    doc.save(doc_path)
    
    file_record = models.File(
        project_id=project_record.id,
        chapter_id=chapter_record.id,
        filename=filename,
        file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        category="Manuscript",
        path=str(doc_path),
        version=1,
    )
    db_session.add(file_record)
    db_session.commit()
    db_session.refresh(file_record)
    
    client = auth_cookie_client(admin_user)
    
    response = client.get(f"/api/v2/files/{file_record.id}/reference-review/validate-only?style=AMA&citation_format=auto")
    assert response.status_code == 200
    data = response.json()
    
    validation_logs = data["validation_logs"]
    assert any("Pass 1" in log for log in validation_logs.get("pipeline_log", []))
    
    # Check that version bumped
    db_session.refresh(file_record)
    assert file_record.version == 2
    
    # Check document content renumbered
    saved_doc = Document(file_record.path)
    first_para = saved_doc.paragraphs[0]
    first_para_text = "".join(t.text for t in first_para._element.xpath('.//w:t[not(ancestor::w:del)]') if t.text)
    
    # Composition:
    # 1 -> 1
    # 4 -> 2
    # 3 -> 3 (Pass 1) -> 2 (Pass 2)
    # 5 -> 4 (Pass 1) -> 3 (Pass 2)
    # Expected output text: "This is 1 then 2 then 2 then 3"
    assert "This is 1 then 2 then 2 then 3" in first_para_text
    
    # The bibliography should have 3 items left: 1, 2, 3
    bib_paras = [p for p in saved_doc.paragraphs if p.style and p.style.name == 'REF-N']
    assert len(bib_paras) == 3


def test_reference_review_apa_reorder_and_block_sorting(
    db_session,
    temp_upload_root,
    project_record,
    chapter_record,
    auth_cookie_client,
    admin_user,
):
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    
    file_path = temp_upload_root / project_record.code / chapter_record.number / "Manuscript"
    file_path.mkdir(parents=True, exist_ok=True)
    
    filename = "chapter08_Processed.docx"
    doc_path = file_path / filename
    
    doc = Document()
    p = doc.add_paragraph("This is some text with citation ")
    
    # Ensure 'cite_bib' character style exists
    try:
        cite_style = doc.styles['cite_bib']
    except KeyError:
        cite_style = doc.styles.add_style('cite_bib', WD_STYLE_TYPE.CHARACTER)
        cite_style.font.superscript = False
        
    r1 = p.add_run("(Smith, 2020)")
    r1.style = 'cite_bib'
    p.add_run(" and also ")
    r2 = p.add_run("(Smith, 2020; Murali, 2025; Nive, 2020)")
    r2.style = 'cite_bib'
    
    # Sibling entries live between <ref-open> and <ref-close> or have REF-U style
    p_open = doc.add_paragraph("<ref-open>")
    
    bib_entries = [
        ("Smith, J. (2020). Ref A.", "REF-U"),
        ("Murali, A. (2025). Ref B.", "REF-U"),
        ("Nive, K. (2020). Ref C.", "REF-U"),
    ]
    
    for text, style in bib_entries:
        p_bib = doc.add_paragraph(text)
        try:
            p_bib.style = style
        except KeyError:
            doc.styles.add_style(style, 1)
            p_bib.style = style
            
    p_close = doc.add_paragraph("<ref-close>")
    doc.save(doc_path)
    
    file_record = models.File(
        project_id=project_record.id,
        chapter_id=chapter_record.id,
        filename=filename,
        file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        category="Manuscript",
        path=str(doc_path),
        version=1,
    )
    db_session.add(file_record)
    db_session.commit()
    db_session.refresh(file_record)
    
    client = auth_cookie_client(admin_user)
    
    response = client.get(f"/api/v2/files/{file_record.id}/reference-review/validate-only?style=APA&citation_format=auto")
    assert response.status_code == 200
    data = response.json()
    
    # Check that sorting occurred and stats look correct
    validation_logs = data["validation_logs"]
    assert validation_logs["detected_style"] == "APA"
    assert validation_logs["total_refs"] == 3
    
    # Check document content updated
    saved_doc = Document(file_record.path)
    
    # 1. The multi-citation block should be sorted: (Smith, 2020; Murali, 2025; Nive, 2020) -> (Murali, 2025; Nive, 2020; Smith, 2020)
    first_para = saved_doc.paragraphs[0]
    first_para_text = first_para.text
    assert "(Murali, 2025; Nive, 2020; Smith, 2020)" in first_para_text
    
    # 2. Bibliography entries should be sorted alphabetically: Murali, Nive, Smith
    bib_paras = []
    in_bib = False
    for p in saved_doc.paragraphs:
        if "<ref-open>" in p.text:
            in_bib = True
            continue
        if "<ref-close>" in p.text:
            in_bib = False
            continue
        if in_bib and p.text.strip():
            bib_paras.append(p.text)
            
    assert len(bib_paras) == 3
    assert bib_paras[0].startswith("Murali")
    assert bib_paras[1].startswith("Nive")
    assert bib_paras[2].startswith("Smith")


def test_reference_review_apa_duplicate_merge(
    db_session,
    temp_upload_root,
    project_record,
    chapter_record,
    auth_cookie_client,
    admin_user,
):
    from docx import Document
    from docx.enum.style import WD_STYLE_TYPE
    
    file_path = temp_upload_root / project_record.code / chapter_record.number / "Manuscript"
    file_path.mkdir(parents=True, exist_ok=True)
    
    filename = "chapter09_Processed.docx"
    doc_path = file_path / filename
    
    doc = Document()
    p = doc.add_paragraph("This has citations ")
    
    # Ensure 'cite_bib' character style exists
    try:
        cite_style = doc.styles['cite_bib']
    except KeyError:
        cite_style = doc.styles.add_style('cite_bib', WD_STYLE_TYPE.CHARACTER)
        cite_style.font.superscript = False
        
    r1 = p.add_run("(Smith, 2020)")
    r1.style = 'cite_bib'
    p.add_run(" and duplicated ")
    r2 = p.add_run("(Smith, 2020)")
    r2.style = 'cite_bib'
    
    p_open = doc.add_paragraph("<ref-open>")
    
    # Add duplicate bibliography entries (95%+ match)
    bib_entries = [
        ("Smith, J. (2020). Ref A text sample to test duplicate matching algorithms.", "REF-U"),
        ("Smith, J. (2020). Ref A text sample to test duplicate matching algorithm.", "REF-U"),
    ]
    
    for text, style in bib_entries:
        p_bib = doc.add_paragraph(text)
        try:
            p_bib.style = style
        except KeyError:
            doc.styles.add_style(style, 1)
            p_bib.style = style
            
    p_close = doc.add_paragraph("<ref-close>")
    doc.save(doc_path)
    
    file_record = models.File(
        project_id=project_record.id,
        chapter_id=chapter_record.id,
        filename=filename,
        file_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        category="Manuscript",
        path=str(doc_path),
        version=1,
    )
    db_session.add(file_record)
    db_session.commit()
    db_session.refresh(file_record)
    
    client = auth_cookie_client(admin_user)
    
    response = client.get(f"/api/v2/files/{file_record.id}/reference-review/validate-only?style=APA&citation_format=auto")
    assert response.status_code == 200
    data = response.json()
    
    # Verify that the two duplicates were merged into one
    saved_doc = Document(file_record.path)
    
    bib_paras = []
    in_bib = False
    for p in saved_doc.paragraphs:
        if "<ref-open>" in p.text:
            in_bib = True
            continue
        if "<ref-close>" in p.text:
            in_bib = False
            continue
        if in_bib and p.text.strip():
            bib_paras.append(p.text)
            
    # There should only be 1 reference left in the bibliography
    assert len(bib_paras) == 1




