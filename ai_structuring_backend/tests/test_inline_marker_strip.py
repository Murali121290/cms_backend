"""Tests for the inline-marker stripping helper used by reconstruction."""

from __future__ import annotations

from docx import Document

from processor.reconstruction import _strip_inline_marker_prefix


def _para_with_runs(*texts):
    """Build a paragraph in a fresh doc with one run per text fragment."""
    doc = Document()
    p = doc.add_paragraph()
    for t in texts:
        p.add_run(t)
    return p


def test_strip_marker_in_single_run():
    p = _para_with_runs("<CJC-TTL>Clinical Judgment Case")
    stripped = _strip_inline_marker_prefix(p, "<CJC-TTL>")
    assert stripped is True
    assert p.text == "Clinical Judgment Case"


def test_strip_marker_with_trailing_space():
    p = _para_with_runs("<T2> System")
    stripped = _strip_inline_marker_prefix(p, "<T2>")
    assert stripped is True
    assert p.text == "System"


def test_strip_marker_preserves_remaining_runs():
    p = _para_with_runs("<CJC-TTL>", "Clinical ", "Judgment Case")
    stripped = _strip_inline_marker_prefix(p, "<CJC-TTL>")
    assert stripped is True
    assert p.text == "Clinical Judgment Case"
    # Three runs still present (first is now empty).
    assert len(p.runs) == 3
    assert p.runs[0].text == ""
    assert p.runs[1].text == "Clinical "
    assert p.runs[2].text == "Judgment Case"


def test_strip_marker_spans_multiple_runs():
    """Author wrote <CJC-TTL>foo but Word split the marker across runs."""
    p = _para_with_runs("<CJC", "-TTL>", "Foo")
    stripped = _strip_inline_marker_prefix(p, "<CJC-TTL>")
    assert stripped is True
    assert p.text == "Foo"


def test_strip_no_op_when_marker_not_at_start():
    """Marker text appears mid-paragraph; do not strip."""
    p = _para_with_runs("Before <CJC-TTL>after")
    stripped = _strip_inline_marker_prefix(p, "<CJC-TTL>")
    assert stripped is False
    assert p.text == "Before <CJC-TTL>after"


def test_strip_no_op_when_no_marker():
    p = _para_with_runs("Plain text")
    stripped = _strip_inline_marker_prefix(p, "<CJC-TTL>")
    assert stripped is False
    assert p.text == "Plain text"


def test_strip_with_leading_whitespace_run():
    p = _para_with_runs("  ", "<H2>", "Pain Theories")
    stripped = _strip_inline_marker_prefix(p, "<H2>")
    assert stripped is True
    assert p.text == "  Pain Theories"


def test_strip_empty_marker_no_op():
    p = _para_with_runs("<CJC-TTL>Foo")
    stripped = _strip_inline_marker_prefix(p, "")
    assert stripped is False
    assert p.text == "<CJC-TTL>Foo"


def test_strip_empty_paragraph_no_op():
    p = _para_with_runs()
    stripped = _strip_inline_marker_prefix(p, "<CJC-TTL>")
    assert stripped is False
