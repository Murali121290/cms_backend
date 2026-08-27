# -*- coding: utf-8 -*-
import argparse
import json
import sys
import os
import html
from docx import Document

def generate_html_report(docx_path, json_config_path, output_html_path):
    # 1. Load JSON configuration (Springer Styles)
    try:
        with open(json_config_path, 'r', encoding='utf-8') as f:
            rules = json.load(f)
    except Exception as e:
        print(f"[ERROR] Failed to load JSON configuration file: {e}")
        sys.exit(1)

    if isinstance(rules, list):
        allowed_styles = set(rules)
    else:
        allowed_styles = set(rules.get("allowed_styles", []))

    # 2. Open Edited Word Document
    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"[ERROR] Failed to open Word document: {e}")
        sys.exit(1)

    used_styles = set()
    style_locations = {}

    # Build style_id to name mapping from doc.styles
    style_id_to_name = {}
    for style in doc.styles:
        sid = getattr(style, "style_id", None)
        sname = getattr(style, "name", None)
        if sid and sname:
            style_id_to_name[sid] = sname

    def track_style_location(style_name, location_str):
        if style_name not in style_locations:
            style_locations[style_name] = []
        style_locations[style_name].append(location_str)

    def get_style_name(style_obj):
        if style_obj and hasattr(style_obj, 'name') and style_obj.name:
            return style_obj.name
        return "[Unset Style]"

    # 3. Styles only from Active Paragraphs (Content in use word document)
    for i, p in enumerate(doc.paragraphs, start=1):
        # Paragraph style actually applied in text
        style_name = get_style_name(p.style)
        used_styles.add(style_name)
        track_style_location(style_name, f"Paragraph {i}")

        # Character styles actively applied to text runs
        for r_idx, run in enumerate(p.runs, start=1):
            if run.style and run.style.name and run.style.name != "Default Paragraph Font":
                run_style_name = run.style.name
                used_styles.add(run_style_name)
                track_style_location(run_style_name, f"Paragraph {i}, Run {r_idx}")

    # 4. Scan ONLY Active Tables (Content in use word document, including cells, runs, and nested tables)
    def scan_table(table, location_prefix):
        table_style = get_style_name(table.style)
        used_styles.add(table_style)
        track_style_location(table_style, location_prefix)

        scanned_cells = set()
        for r_idx, row in enumerate(table.rows, start=1):
            for c_idx, cell in enumerate(row.cells, start=1):
                if cell._tc in scanned_cells:
                    continue
                scanned_cells.add(cell._tc)

                cell_loc = f"{location_prefix}, Row {r_idx}, Cell {c_idx}"
                for p_idx, p in enumerate(cell.paragraphs, start=1):
                    # Paragraph style
                    p_style_name = get_style_name(p.style)
                    used_styles.add(p_style_name)
                    track_style_location(p_style_name, f"{cell_loc}, Paragraph {p_idx}")
                    
                    # Character styles inside runs
                    for r_run_idx, run in enumerate(p.runs, start=1):
                        if run.style and run.style.name and run.style.name != "Default Paragraph Font":
                            run_style_name = run.style.name
                            used_styles.add(run_style_name)
                            track_style_location(run_style_name, f"{cell_loc}, Paragraph {p_idx}, Run {r_run_idx}")
                
                # Check for nested tables in this cell
                for nt_idx, nested_table in enumerate(cell.tables, start=1):
                    scan_table(nested_table, f"{cell_loc}, Nested Table {nt_idx}")

    for t_idx, table in enumerate(doc.tables, start=1):
        scan_table(table, f"Table {t_idx}")

    # 5. Scan Footnotes and Endnotes XML from document package
    import zipfile
    import xml.etree.ElementTree as ET

    def scan_xml_notes(docx_path, xml_name, note_type):
        namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        try:
            with zipfile.ZipFile(docx_path) as z:
                if xml_name in z.namelist():
                    xml_content = z.read(xml_name)
                    root = ET.fromstring(xml_content)
                    tag_name = f".//w:{note_type}"
                    for note in root.findall(tag_name, namespaces):
                        note_id = note.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}id') or note.attrib.get('id')
                        note_type_attr = note.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}type') or note.attrib.get('type')
                        if note_id in ('-1', '0') or note_type_attr in ('separator', 'continuationSeparator'):
                            continue
                        
                        note_label = f"{note_type.capitalize()} {note_id}"
                        # Scan paragraphs inside the note
                        for p_idx, p in enumerate(note.findall('.//w:p', namespaces), start=1):
                            p_style_elem = p.find('.//w:pPr/w:pStyle', namespaces)
                            if p_style_elem is not None:
                                p_style_id = p_style_elem.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                                if p_style_id:
                                    style_name = style_id_to_name.get(p_style_id, p_style_id)
                                    used_styles.add(style_name)
                                    track_style_location(style_name, f"{note_label}, Paragraph {p_idx}")
                            
                            # Scan runs inside the paragraph
                            for r_idx, r in enumerate(p.findall('.//w:r', namespaces), start=1):
                                r_style_elem = r.find('.//w:rPr/w:rStyle', namespaces)
                                if r_style_elem is not None:
                                    r_style_id = r_style_elem.attrib.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}val')
                                    if r_style_id:
                                        style_name = style_id_to_name.get(r_style_id, r_style_id)
                                        if style_name and style_name != "Default Paragraph Font":
                                            used_styles.add(style_name)
                                            track_style_location(style_name, f"{note_label}, Paragraph {p_idx}, Run {r_idx}")
        except Exception as note_err:
            print(f"[WARNING] Failed to parse {xml_name}: {note_err}")

    scan_xml_notes(docx_path, 'word/footnotes.xml', 'footnote')
    scan_xml_notes(docx_path, 'word/endnotes.xml', 'endnote')

    # Deduplicate and separate unique style lists
    all_unique_styles = sorted(list(used_styles), key=lambda x: str(x))
    allowed_styles_lower = {s.lower(): s for s in allowed_styles}

    unauthorized_styles = []
    allowed_used_styles = []

    for s in used_styles:
        if s.lower() in allowed_styles_lower:
            allowed_used_styles.append(s)
        else:
            unauthorized_styles.append(s)

    unauthorized_styles = sorted(unauthorized_styles, key=lambda x: str(x))
    allowed_used_styles = sorted(allowed_used_styles, key=lambda x: str(x))

    status_pass = len(unauthorized_styles) == 0
    total_styles = len(all_unique_styles)
    approved_count = len(allowed_used_styles)
    compliance_rate = int((approved_count / total_styles) * 100) if total_styles > 0 else 100

    if compliance_rate == 100:
        gauge_color = "#10b981"  # Emerald
        gauge_status = "PASS"
        gauge_status_class = "tag-allowed"
    elif compliance_rate >= 80:
        gauge_color = "#f59e0b"  # Amber
        gauge_status = "WARNING"
        gauge_status_class = "tag-warning"
    else:
        gauge_color = "#f43f5e"  # Rose
        gauge_status = "FAIL"
        gauge_status_class = "tag-unauthorized"

    # Build rows HTML
    rows_html = ""
    for style_name in all_unique_styles:
        is_allowed = style_name in allowed_used_styles
        status_text = "Approved" if is_allowed else "Unauthorized"
        badge_class = "badge-allowed" if is_allowed else "badge-unauthorized"
        tag_class = "tag-allowed" if is_allowed else "tag-unauthorized"
        
        locs = style_locations.get(style_name, [])
        loc_count = len(locs)
        
        max_pills = 6
        pills_html = ""
        for idx, loc in enumerate(locs):
            hidden_class = " pill-hidden" if idx >= max_pills else ""
            hidden_style = " style='display: none;'" if idx >= max_pills else ""
            pills_html += f'<span class="location-pill{hidden_class}"{hidden_style}>{html.escape(loc)}</span>'
        
        if loc_count > max_pills:
            pills_html += f'<button class="expand-btn" data-expanded="false" onclick="togglePills(this)">+{loc_count - max_pills} more...</button>'
        
        rows_html += f"""
                    <tr class="style-row" data-style="{html.escape(str(style_name))}" data-allowed="{'true' if is_allowed else 'false'}">
                        <td style="vertical-align: middle;">
                            <span class="status-tag {tag_class}">{status_text}</span>
                        </td>
                        <td style="vertical-align: middle;">
                            <span class="style-name-badge {badge_class}">{html.escape(str(style_name))}</span>
                        </td>
                        <td style="vertical-align: middle; font-weight: 600;">
                            {loc_count} {'time' if loc_count == 1 else 'times'}
                        </td>
                        <td>
                            <div class="locations-wrapper">
                                {pills_html}
                            </div>
                        </td>
                    </tr>
"""

    # 5. Build Responsive HTML Document
    html_doc = f"""<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Word Document Style Validation Report</title>
    <link href="https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap" rel="stylesheet">
    <style>
        :root {{
            --bg-color: #0f172a;
            --surface-color: #1e293b;
            --surface-hover: #334155;
            --border-color: #475569;
            --text-main: #f8fafc;
            --text-muted: #94a3b8;
            --accent-color: #6366f1;
            --accent-hover: #4f46e5;
            --pass-color: #10b981;
            --pass-bg: rgba(16, 185, 129, 0.1);
            --fail-color: #f43f5e;
            --fail-bg: rgba(244, 63, 94, 0.1);
            --warn-color: #f59e0b;
            --warn-bg: rgba(245, 158, 11, 0.1);
        }}

        body {{
            font-family: 'Inter', -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif;
            background-color: var(--bg-color);
            color: var(--text-main);
            margin: 0;
            padding: 30px 15px;
            line-height: 1.5;
            -webkit-font-smoothing: antialiased;
        }}

        .container {{
            max-width: 1080px;
            margin: 0 auto;
        }}

        /* Header Layout */
        .dashboard-header {{
            background: linear-gradient(135deg, #1e293b 0%, #0f172a 100%);
            border: 1px solid var(--border-color);
            border-radius: 16px;
            padding: 24px 32px;
            margin-bottom: 24px;
            box-shadow: 0 10px 25px -5px rgba(0, 0, 0, 0.3);
            display: flex;
            justify-content: space-between;
            align-items: center;
            flex-wrap: wrap;
            gap: 20px;
        }}

        .header-title-section h1 {{
            margin: 0 0 8px 0;
            font-size: 24px;
            font-weight: 700;
            letter-spacing: -0.025em;
            background: linear-gradient(to right, #f8fafc, #94a3b8);
            -webkit-background-clip: text;
            -webkit-text-fill-color: transparent;
        }}

        .document-name {{
            font-size: 14px;
            color: var(--text-muted);
            display: flex;
            align-items: center;
            gap: 6px;
        }}

        /* Gauge section */
        .gauge-container {{
            display: flex;
            align-items: center;
            gap: 16px;
            background: rgba(15, 23, 42, 0.5);
            padding: 12px 20px;
            border-radius: 12px;
            border: 1px solid rgba(255, 255, 255, 0.05);
        }}

        .radial-gauge {{
            position: relative;
            width: 70px;
            height: 70px;
        }}

        .circular-chart {{
            display: block;
            width: 100%;
            height: 100%;
            transform: rotate(-90deg);
        }}

        .circle-bg {{
            fill: none;
            stroke: #334155;
            stroke-width: 3.8;
        }}

        .circle {{
            fill: none;
            stroke-width: 3.8;
            stroke-linecap: round;
        }}

        .gauge-percentage {{
            position: absolute;
            top: 50%;
            left: 50%;
            transform: translate(-50%, -50%);
            font-size: 15px;
            font-weight: 700;
            color: var(--text-main);
        }}

        .gauge-info {{
            display: flex;
            flex-direction: column;
        }}

        .gauge-label {{
            font-size: 11px;
            color: var(--text-muted);
            text-transform: uppercase;
            font-weight: 600;
            letter-spacing: 0.05em;
        }}

        .gauge-status {{
            font-size: 15px;
            font-weight: 700;
            margin-top: 2px;
        }}

        /* Stats Grid */
        .stats-grid {{
            display: grid;
            grid-template-columns: repeat(auto-fit, minmax(200px, 1fr));
            gap: 16px;
            margin-bottom: 24px;
        }}

        .stat-card {{
            background-color: var(--surface-color);
            border: 1px solid var(--border-color);
            border-radius: 14px;
            padding: 18px 24px;
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
            transition: transform 0.2s, border-color 0.2s;
        }}

        .stat-card:hover {{
            transform: translateY(-2px);
            border-color: var(--surface-hover);
        }}

        .stat-label {{
            font-size: 11px;
            color: var(--text-muted);
            text-transform: uppercase;
            letter-spacing: 0.05em;
            font-weight: 600;
            margin-bottom: 6px;
        }}

        .stat-value {{
            font-size: 26px;
            font-weight: 700;
            line-height: 1.1;
        }}

        /* Controls Section */
        .controls-card {{
            background-color: var(--surface-color);
            border: 1px solid var(--border-color);
            border-radius: 14px;
            padding: 16px 24px;
            margin-bottom: 24px;
            display: flex;
            justify-content: space-between;
            align-items: center;
            flex-wrap: wrap;
            gap: 16px;
        }}

        .tab-bar {{
            display: flex;
            gap: 8px;
        }}

        .tab-btn {{
            background: none;
            border: none;
            color: var(--text-muted);
            padding: 8px 16px;
            font-size: 13px;
            font-weight: 600;
            border-radius: 8px;
            cursor: pointer;
            transition: all 0.2s;
        }}

        .tab-btn:hover {{
            color: var(--text-main);
            background-color: var(--surface-hover);
        }}

        .tab-btn.active {{
            color: var(--text-main);
            background-color: var(--accent-color);
        }}

        .search-box {{
            position: relative;
            max-width: 320px;
            width: 100%;
        }}

        .search-input {{
            width: 100%;
            background-color: var(--bg-color);
            border: 1px solid var(--border-color);
            color: var(--text-main);
            padding: 8px 16px 8px 36px;
            border-radius: 8px;
            font-size: 13px;
            outline: none;
            box-sizing: border-box;
            transition: border-color 0.2s;
        }}

        .search-input:focus {{
            border-color: var(--accent-color);
        }}

        .search-icon {{
            position: absolute;
            left: 12px;
            top: 50%;
            transform: translateY(-50%);
            width: 16px;
            height: 16px;
            color: var(--text-muted);
            pointer-events: none;
        }}

        /* Styles List Card */
        .list-card {{
            background-color: var(--surface-color);
            border: 1px solid var(--border-color);
            border-radius: 16px;
            padding: 24px 32px;
            box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.1);
        }}

        .list-title {{
            font-size: 18px;
            font-weight: 700;
            margin: 0 0 20px 0;
        }}

        .styles-table {{
            width: 100%;
            border-collapse: collapse;
            font-size: 14px;
        }}

        .styles-table th {{
            text-align: left;
            padding: 12px 16px;
            border-bottom: 1px solid var(--border-color);
            color: var(--text-muted);
            font-weight: 600;
            font-size: 11px;
            text-transform: uppercase;
            letter-spacing: 0.05em;
        }}

        .styles-table td {{
            padding: 16px 16px;
            border-bottom: 1px solid rgba(255, 255, 255, 0.05);
            vertical-align: top;
        }}

        .styles-table tr:hover td {{
            background-color: rgba(255, 255, 255, 0.01);
        }}

        .style-name-badge {{
            font-family: 'Courier New', Courier, monospace;
            font-weight: 700;
            font-size: 13px;
            padding: 4px 8px;
            border-radius: 6px;
            display: inline-block;
        }}

        .badge-allowed {{
            background-color: var(--pass-bg);
            color: var(--pass-color);
            border: 1px solid rgba(16, 185, 129, 0.2);
        }}

        .badge-unauthorized {{
            background-color: var(--fail-bg);
            color: var(--fail-color);
            border: 1px solid rgba(244, 63, 94, 0.2);
        }}

        .status-tag {{
            font-size: 11px;
            font-weight: 700;
            text-transform: uppercase;
            padding: 4px 10px;
            border-radius: 6px;
            display: inline-block;
        }}

        .tag-allowed {{
            background-color: var(--pass-bg);
            color: var(--pass-color);
        }}

        .tag-unauthorized {{
            background-color: var(--fail-bg);
            color: var(--fail-color);
        }}

        .tag-warning {{
            background-color: var(--warn-bg);
            color: var(--warn-color);
        }}

        .locations-wrapper {{
            display: flex;
            flex-wrap: wrap;
            gap: 6px;
            align-items: center;
        }}

        .location-pill {{
            background-color: rgba(255, 255, 255, 0.05);
            border: 1px solid rgba(255, 255, 255, 0.08);
            color: var(--text-muted);
            font-size: 11px;
            padding: 2px 8px;
            border-radius: 6px;
            cursor: default;
            transition: all 0.2s;
        }}

        .location-pill:hover {{
            background-color: rgba(255, 255, 255, 0.1);
            color: var(--text-main);
        }}

        .expand-btn {{
            background: none;
            border: none;
            color: var(--accent-color);
            font-size: 11px;
            font-weight: 600;
            cursor: pointer;
            padding: 2px 6px;
            border-radius: 4px;
            transition: background-color 0.2s;
        }}

        .expand-btn:hover {{
            background-color: rgba(99, 102, 241, 0.15);
        }}

        .no-records {{
            text-align: center;
            color: var(--text-muted);
            padding: 40px 0;
            font-size: 14px;
        }}
    </style>
</head>
<body>
    <div class="container">
        <!-- Dashboard Header -->
        <div class="dashboard-header">
            <div class="header-title-section">
                <h1>S4carlisle Word Style Validation Report</h1>
                <div class="document-name">
                    <svg style="width:16px;height:16px;" fill="none" stroke="currentColor" viewBox="0 0 24 24" xmlns="http://www.w3.org/2000/svg"><path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M7 21h10a2 2 0 002-2V9.414a1 1 0 00-.293-.707l-5.414-5.414A1 1 0 0012.586 3H7a2 2 0 00-2 2v14a2 2 0 002 2z"></path></svg>
                    <span>{html.escape(os.path.basename(docx_path))}</span>
                </div>
            </div>
            
            <!-- Radial Compliance Gauge -->
            <div class="gauge-container">
                <div class="radial-gauge">
                    <svg viewBox="0 0 36 36" class="circular-chart">
                        <path class="circle-bg" d="M18 2.0845 a 15.9155 15.9155 0 0 1 0 31.831 a 15.9155 15.9155 0 0 1 0 -31.831" />
                        <path class="circle" stroke="{gauge_color}" stroke-dasharray="{compliance_rate}, 100" d="M18 2.0845 a 15.9155 15.9155 0 0 1 0 31.831 a 15.9155 15.9155 0 0 1 0 -31.831" />
                    </svg>
                    <div class="gauge-percentage">{compliance_rate}%</div>
                </div>
                <div class="gauge-info">
                    <span class="gauge-label">Compliance</span>
                    <span class="gauge-status {gauge_status_class}">{gauge_status}</span>
                </div>
            </div>
        </div>

        <!-- Stats Cards Grid -->
        <div class="stats-grid">
            <div class="stat-card">
                <div class="stat-label">Unique Styles Used</div>
                <div class="stat-value">{total_styles}</div>
            </div>
            <div class="stat-card">
                <div class="stat-label">Approved Styles</div>
                <div class="stat-value" style="color: var(--pass-color);">{approved_count}</div>
            </div>
            <div class="stat-card">
                <div class="stat-label">Unauthorized Styles</div>
                <div class="stat-value" style="color: {"var(--fail-color)" if len(unauthorized_styles) > 0 else "var(--pass-color)"};">
                    {len(unauthorized_styles)}
                </div>
            </div>
            <div class="stat-card">
                <div class="stat-label">Configuration Rules</div>
                <div class="stat-value" style="font-size: 14px; word-break: break-all; margin-top: 6px;">
                    {html.escape(os.path.basename(json_config_path))}
                </div>
            </div>
        </div>

        <!-- Search and Filters Section -->
        <div class="controls-card">
            <div class="tab-bar">
                <button class="tab-btn active" data-tab="all" onclick="switchTab(this)">All Styles</button>
                <button class="tab-btn" data-tab="unauthorized" onclick="switchTab(this)" style="color: {"var(--fail-color)" if len(unauthorized_styles) > 0 else "var(--text-muted)"};">
                    Unauthorized ({len(unauthorized_styles)})
                </button>
                <button class="tab-btn" data-tab="approved" onclick="switchTab(this)">Approved ({approved_count})</button>
            </div>
            
            <div class="search-box">
                <svg class="search-icon" fill="none" stroke="currentColor" viewBox="0 0 24 24" xmlns="http://www.w3.org/2000/svg"><path stroke-linecap="round" stroke-linejoin="round" stroke-width="2" d="M21 21l-6-6m2-5a7 7 0 11-14 0 7 7 0 0114 0z"></path></svg>
                <input type="text" id="search-input" class="search-input" placeholder="Search style name..." oninput="filterStyles()">
            </div>
        </div>

        <!-- Styles Table Card -->
        <div class="list-card">
            <div class="list-title">Style Analysis Details</div>
            
            <table class="styles-table">
                <thead>
                    <tr>
                        <th style="width: 15%;">Status</th>
                        <th style="width: 30%;">Style Name</th>
                        <th style="width: 15%;">Usage</th>
                        <th style="width: 40%;">Locations Found</th>
                    </tr>
                </thead>
                <tbody id="styles-tbody">
                    {rows_html}
                </tbody>
            </table>
            
            <div id="no-records" class="no-records" style="display: none;">
                No styles match the current filters.
            </div>
        </div>
    </div>

    <!-- Script for Dynamic Search & Filter -->
    <script>
        function filterStyles() {{
            const query = document.getElementById('search-input').value.toLowerCase();
            const activeTab = document.querySelector('.tab-btn.active').dataset.tab;
            const rows = document.querySelectorAll('.style-row');
            let visibleCount = 0;

            rows.forEach(row => {{
                const styleName = row.dataset.style.toLowerCase();
                const isAllowed = row.dataset.allowed === 'true';
                
                let matchesSearch = styleName.includes(query);
                let matchesTab = activeTab === 'all' || 
                                 (activeTab === 'unauthorized' && !isAllowed) || 
                                 (activeTab === 'approved' && isAllowed);

                if (matchesSearch && matchesTab) {{
                    row.style.display = '';
                    visibleCount++;
                }} else {{
                    row.style.display = 'none';
                }}
            }});

            const noRecords = document.getElementById('no-records');
            if (visibleCount === 0) {{
                noRecords.style.display = '';
            }} else {{
                noRecords.style.display = 'none';
            }}
        }}

        function switchTab(btn) {{
            document.querySelectorAll('.tab-btn').forEach(b => b.classList.remove('active'));
            btn.classList.add('active');
            filterStyles();
        }}

        function togglePills(btn) {{
            const parent = btn.parentElement;
            const hiddenPills = parent.querySelectorAll('.pill-hidden');
            const isExpanded = btn.dataset.expanded === 'true';

            hiddenPills.forEach(pill => {{
                pill.style.display = isExpanded ? 'none' : 'inline-block';
            }});

            if (isExpanded) {{
                btn.textContent = `+${{hiddenPills.length}} more...`;
                btn.dataset.expanded = 'false';
            }} else {{
                btn.textContent = 'Show Less';
                btn.dataset.expanded = 'true';
            }}
        }}
    </script>
</body>
</html>
"""

    # Save HTML output
    try:
        with open(output_html_path, 'w', encoding='utf-8') as f:
            f.write(html_doc)
        print(f"[SUCCESS] In-use HTML report saved to: {output_html_path}")
        return 0 if status_pass else 1
    except Exception as e:
        print(f"[ERROR] Failed to save HTML report file: {e}")
        sys.exit(1)


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Validate in-use Word document (.docx) styles against allowed JSON list."
    )
    parser.add_argument("-d", "--doc", required=True, help="Path to the Word document (.docx)")
    parser.add_argument("-c", "--config", required=True, help="Path to allowed styles JSON file")
    parser.add_argument("-o", "--output", default="style_report.html", help="Path for output HTML report")

    args = parser.parse_args()
    exit_code = generate_html_report(args.doc, args.config, args.output)
    sys.exit(exit_code)