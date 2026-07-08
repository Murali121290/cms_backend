import os
import re
import logging
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

logger = logging.getLogger("app.conversion.postprocessor")

def set_section_columns_to_one(section):
    """
    Clears section columns and sets w:cols to num="1".
    This forces a single-column layout.
    """
    sectPr = section._sectPr
    for cols in sectPr.xpath('w:cols'):
        sectPr.remove(cols)
    cols = OxmlElement('w:cols')
    cols.set(qn('w:num'), '1')
    sectPr.append(cols)

def set_font_times_new_roman(doc):
    """
    Sets document default and run-level fonts to Times New Roman.
    """
    if 'Normal' in doc.styles:
        style = doc.styles['Normal']
        style.font.name = 'Times New Roman'
        
    def force_run_font(run):
        rPr = run._r.get_or_add_rPr()
        for rFonts in rPr.xpath('w:rFonts'):
            rPr.remove(rFonts)
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        rFonts.set(qn('w:cs'), 'Times New Roman')
        rPr.append(rFonts)

    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            force_run_font(run)
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        force_run_font(run)

def remove_headers_footers(doc):
    """
    Clears text from headers and footers across all sections.
    """
    for section in doc.sections:
        for p in section.header.paragraphs:
            p.text = ""
        for p in section.footer.paragraphs:
            p.text = ""
            
        try:
            if section.first_page_header:
                for p in section.first_page_header.paragraphs:
                    p.text = ""
            if section.first_page_footer:
                for p in section.first_page_footer.paragraphs:
                    p.text = ""
        except (AttributeError, TypeError):
            pass

def should_merge(p1, p2):
    """
    Rules for merging paragraph p2 into p1.
    """
    t1 = p1.text.strip()
    t2 = p2.text.strip()
    
    if not t1 or not t2:
        return False
        
    if re.search(r'[\.\!\?]["”]?$', t1):
        return False
        
    if t2.startswith(('•', '-', '*')):
        return False
        
    if t2[0].islower():
        return True
        
    if t1.endswith('-'):
        return True
        
    return False

def merge_paragraphs(p1, p2):
    """
    Merges paragraph p2 content into p1 and removes p2.
    """
    t1 = p1.text.rstrip()
    t2 = p2.text.strip()
    
    if t1.endswith('-') and t2 and t2[0].islower():
        for run in reversed(p1.runs):
            if '-' in run.text:
                run.text = run.text.rsplit('-', 1)[0]
                break
    elif t1 and not t1.endswith(' ') and t2:
        if p1.runs:
            p1.runs[-1].text += ' '
            
    for run in p2.runs:
        p1._p.append(run._r)
        
    p2._p.getparent().remove(p2._p)

def clean_document_flow(doc):
    """
    Reconstructs continuous text flow by scanning and merging split paragraphs.
    """
    body_paragraphs = list(doc.paragraphs)
    i = 0
    while i < len(body_paragraphs) - 1:
        p1 = body_paragraphs[i]
        p2 = body_paragraphs[i+1]
        
        if p1._p.getparent() == p2._p.getparent() and should_merge(p1, p2):
            merge_paragraphs(p1, p2)
            body_paragraphs.pop(i+1)
        else:
            i += 1
            
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_paragraphs = list(cell.paragraphs)
                j = 0
                while j < len(cell_paragraphs) - 1:
                    p1 = cell_paragraphs[j]
                    p2 = cell_paragraphs[j+1]
                    if should_merge(p1, p2):
                        merge_paragraphs(p1, p2)
                        cell_paragraphs.pop(j+1)
                    else:
                        j += 1

def merge_consecutive_tables(doc):
    """
    Merges consecutive tables separated only by empty paragraphs.
    """
    body = doc.element.body
    children = list(body)
    
    i = 0
    while i < len(children) - 2:
        el1 = children[i]
        el2 = children[i+1]
        el3 = children[i+2]
        
        if el1.tag.endswith('tbl') and el3.tag.endswith('tbl'):
            if el2.tag.endswith('p'):
                p_text = el2.text or "".join(el2.itertext()).strip()
                if not p_text:
                    from docx.table import Table
                    t1 = Table(el1, doc)
                    t2 = Table(el3, doc)
                    
                    if len(t1.columns) == len(t2.columns):
                        for row in t2.rows:
                            t1._tbl.append(row._tr)
                        body.remove(el2)
                        body.remove(el3)
                        children = list(body)
                        continue
        i += 1

def post_process_docx(docx_path: str):
    """
    Main post-processor entry point.
    Loads a DOCX file, applies layout/format cleanups, and saves it.
    """
    if not os.path.exists(docx_path):
        logger.error(f"DOCX post-processing failed: file not found at {docx_path}")
        return
        
    logger.info(f"Post-processing DOCX file: {docx_path}")
    try:
        doc = Document(docx_path)
        
        for section in doc.sections:
            set_section_columns_to_one(section)
            
        remove_headers_footers(doc)
        merge_consecutive_tables(doc)
        clean_document_flow(doc)
        set_font_times_new_roman(doc)
        
        doc.save(docx_path)
        logger.info(f"Post-processing completed successfully: {docx_path}")
    except Exception as ex:
        logger.error(f"Error during DOCX post-processing: {str(ex)}")
