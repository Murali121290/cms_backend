"""
Step 0 â€” Pre-Conversion Analysis
Validates file format, audits styles, counts document elements, detects markup tags.
Report-only step; makes no modifications.
"""

import re
import zipfile
from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph
from collections import defaultdict

from docx_pipeline.utils.report import ReportLogger
from docx_pipeline.config import (
    SEMANTIC_BOLD_STYLES,
    HEADING_STYLES,
    TXT_FLUSH_STYLE,
    FIGURE_CAPTION_STYLE,
    TABLE_CAPTION_STYLE,
)


def pre_flight_check(path: str, logger: ReportLogger) -> bool:
    """
    File-level validation before Document() is opened.
    Returns True if safe to proceed, False if critical issue found.
    """
    logger.set_step("0-preconversion")
    try:
        with zipfile.ZipFile(path, "r") as zf:
            # Check if word/settings.xml exists (basic DOCX structure)
            if "word/settings.xml" not in zf.namelist():
                logger.error("Not a valid DOCX: missing word/settings.xml")
                return False

            # Check for password protection
            settings_xml = zf.read("word/settings.xml").decode("utf-8")
            if 'w:documentProtection' in settings_xml:
                if 'w:edit="readOnly"' in settings_xml or 'w:cryptProviderType' in settings_xml:
                    logger.error("Document is password-protected; cannot process")
                    return False

    except zipfile.BadZipFile:
        logger.error("Not a valid DOCX: file is corrupted or not a ZIP")
        return False
    except Exception as e:
        logger.error(f"File validation error: {str(e)}")
        return False

    logger.info("File validation passed: valid DOCX, not password-protected")
    return True


def run(doc: Document, logger: ReportLogger) -> Document:
    """
    Audit document: styles, stats, markup tags. Never modifies the document.
    """
    logger.set_step("0-preconversion")

    _run_style_audit(doc, logger)
    _run_doc_stats(doc, logger)
    _run_markup_detection(doc, logger)

    return doc  # unchanged


def _run_style_audit(doc: Document, logger: ReportLogger) -> None:
    """Inventory paragraph and character styles; flag missing required styles."""

    # Collect paragraph styles used
    para_styles_used = defaultdict(int)
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "Normal"
        para_styles_used[style_name] += 1

    # Collect character styles used
    char_styles_used = set()
    for para in doc.paragraphs:
        for run in para.runs:
            if run.style:
                char_styles_used.add(run.style.name)

    # Required publisher styles from config
    required_styles = [
        "H1", "H2", "H3", "H4", "H5", "H6",
        "TXT", "TXT-FLUSH", "TXL",
        "CN", "CT", "CAU", "T1", "T2",
        "REFH1", "REFH2"
    ]

    # Flag missing required styles
    for style_name in required_styles:
        if style_name not in para_styles_used or para_styles_used[style_name] == 0:
            logger.flag(f"Required style '{style_name}' not used in document")

    # Log all found paragraph styles
    style_summary = "\n".join(
        f"  {name}: {count} paragraphs"
        for name, count in sorted(para_styles_used.items(), key=lambda x: -x[1])
    )
    logger.info(f"Paragraph styles found:\n{style_summary}")

    # Log character styles
    if char_styles_used:
        char_summary = ", ".join(sorted(char_styles_used))
        logger.info(f"Character styles found: {char_summary}")
    else:
        logger.info("No character styles found (direct formatting only)")


def _run_doc_stats(doc: Document, logger: ReportLogger) -> None:
    """Count paragraphs, words, tables, images, footnotes, endnotes."""

    para_count = len(doc.paragraphs)

    # Word count
    word_count = 0
    for para in doc.paragraphs:
        words = para.text.split()
        word_count += len(words)

    # Table count
    table_count = len(doc.tables)

    # Image count (inline shapes)
    image_count = len(doc.inline_shapes)

    # Footnote count
    footnote_count = 0
    try:
        footnotes_part = doc.part.footnotes
        if footnotes_part:
            footnote_count = len(footnotes_part._element.findall(qn('w:fn')))
    except:
        pass

    # Endnote count
    endnote_count = 0
    try:
        endnotes_part = doc.part.endnotes
        if endnotes_part:
            endnote_count = len(endnotes_part._element.findall(qn('w:en')))
    except:
        pass

    stats = (
        f"Paragraphs: {para_count} | Words: {word_count} | "
        f"Tables: {table_count} | Images: {image_count} | "
        f"Footnotes: {footnote_count} | Endnotes: {endnote_count}"
    )
    logger.info(stats)


def _run_markup_detection(doc: Document, logger: ReportLogger) -> None:
    """Scan for <TAG> markup patterns; categorize and flag issues."""

    markup_pattern = re.compile(r"<([A-Z][A-Z0-9\-]*)>")

    found_tags = defaultdict(int)  # tag_name -> count
    paragraphs_with_tags = 0
    unrecognized_tags = set()

    # Known tag categories
    heading_tags = {"H1", "H2", "H3", "H4", "H5", "H6"}
    body_tags = {"TXT", "TXT-FLUSH", "TXL"}
    special_tags = {
        "CAU", "CN", "CT", "BX1", "BX2", "BX3", "BX4", "BX5",
        "BX6", "BX7", "BX8", "BX9", "BX10", "BX11", "BX12",
        "BX13", "BX14", "BX15", "BX16", "BX17", "BX18", "BX19",
        "BX20", "BX21", "BX22", "BX23", "BX24", "BX25",
        "FigureCaption", "TableCaption", "ImageCaption",
    }

    for para in doc.paragraphs:
        matches = markup_pattern.findall(para.text)
        if matches:
            paragraphs_with_tags += 1
            for tag in matches:
                found_tags[tag] += 1
                # Check if tag is recognized
                if tag not in heading_tags and tag not in body_tags and tag not in special_tags:
                    unrecognized_tags.add(tag)

    # Flag unrecognized tags
    for tag in sorted(unrecognized_tags):
        count = found_tags[tag]
        logger.flag(f"Unrecognized markup tag: <{tag}> (found {count} times)")

    # Overall tagging status
    total_paragraphs = len(doc.paragraphs)
    non_empty_paragraphs = len([p for p in doc.paragraphs if p.text.strip()])

    if paragraphs_with_tags == 0:
        tagging_status = "UNTAGGED"
    elif paragraphs_with_tags == non_empty_paragraphs:
        tagging_status = "FULLY TAGGED"
    else:
        tagging_status = "PARTIALLY TAGGED"

    logger.info(
        f"Markup tagging status: {tagging_status} "
        f"({paragraphs_with_tags} of {non_empty_paragraphs} non-empty paragraphs tagged)"
    )

    # Log tag summary if any
    if found_tags:
        tag_summary = ", ".join(f"{tag}({found_tags[tag]})" for tag in sorted(found_tags.keys()))
        logger.info(f"Tags found: {tag_summary}")

