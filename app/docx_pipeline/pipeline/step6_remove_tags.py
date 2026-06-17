"""
pipeline/step6_remove_tags.py â€” Remove pipeline validation markers and markup tags.

Stripping actions:
  - XML elements: w:bookmarkStart, w:bookmarkEnd, w:proofErr
  - String markup matching `<Tag>` patterns across all paragraphs (except "Image" style).
"""

import re
from docx import Document
from docx.oxml.ns import qn
from lxml import etree
from docx_pipeline.utils.report import ReportLogger

W_ = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"

# Tags to unconditionally strip from the document body
STRIP_XML_TAGS = [
    W_ + "proofErr",
    W_ + "bookmarkStart",
    W_ + "bookmarkEnd",
]

# Regex for markup tags (e.g., <H1>, <CAU>). Matches anything in angle brackets without spaces/newlines.
MARKUP_REGEX = re.compile(r"(\<[^\>\<\s\n]+\>)")


def _strip_elements(body: etree._Element, tag: str) -> int:
    count = 0
    for el in body.iter(tag):
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)
            count += 1
    return count


def run(doc: Document, logger: ReportLogger) -> Document:
    logger.set_step("6-remove-tags")

    # 1. Strip XML Elements (Bookmarks, Proofing)
    body = doc.element.body
    xml_total = 0

    for tag in STRIP_XML_TAGS:
        n = _strip_elements(body, tag)
        xml_total += n
        if n:
            short = tag.split("}")[-1]
            logger.info(f"Removed {n} <{short}> XML elements.")

    # 2. Strip string markup tags matching REGEX
    text_tags_removed = 0
    unique_tags_found = set()

    all_paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)

    for para in all_paras:
        # User Macro Rule: Ignore paragraphs with style "Image"
        if para.style and para.style.name == "Image":
            continue

        text = para.text
        matches = MARKUP_REGEX.findall(text)
        if not matches:
            continue

        for m in matches:
            unique_tags_found.add(m)

        # Replace tags within runs directly to preserve run formatting natively
        for run_obj in para.runs:
            for m in set(matches):
                if m in run_obj.text:
                    text_tags_removed += run_obj.text.count(m)
                    run_obj.text = run_obj.text.replace(m, "")

        # Fallback for cross-run split tags
        remaining = MARKUP_REGEX.findall(para.text)
        if remaining:
            # If the tag is split across runs, reconstructing the paragraph text 
            # into the first run guarantees its removal
            cleaned_text = para.text
            for m in set(matches):
                cleaned_text = cleaned_text.replace(m, "")
                
            if para.runs:
                para.runs[0].text = cleaned_text
                for r in para.runs[1:]:
                    r.text = ""
                text_tags_removed += len(remaining)

    if xml_total == 0 and text_tags_removed == 0:
        logger.info("No validation markers or text tags to remove.")
    else:
        logger.info(f"Stripped {text_tags_removed} text tags matching regex markup.")
        if unique_tags_found:
            logger.info(f"Unique tags discovered and removed: {', '.join(unique_tags_found)}")

    return doc

