"""
pipeline/step3_remove_bold.py â€” Remove bold from semantic styles.

For paragraphs whose style is in SEMANTIC_BOLD_STYLES, explicitly
overrides and disables bold formatting on all runs (Range.Font.Bold = False).
"""

from docx import Document
from docx_pipeline.config import SEMANTIC_BOLD_STYLES
from docx_pipeline.utils.report import ReportLogger


def run(doc: Document, logger: ReportLogger) -> Document:
    logger.set_step("3-remove-bold")

    target_styles = {s.lower() for s in SEMANTIC_BOLD_STYLES}
    count = 0

    # Collect all paragraphs from the main document and tables
    all_paras = list(doc.paragraphs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                all_paras.extend(cell.paragraphs)

    for para in all_paras:
        if not para.style:
            continue
            
        sname = para.style.name.lower()
        if sname not in target_styles:
            continue

        # In Word, para.Range.Font.Bold = False explicitly disables bold 
        # by inserting `<w:b w:val="0"/>` over the entire range.
        # This overrides the intrinsic boldness of heading styles like H1.
        for r in para.runs:
            # Setting to False explicitly injects the val="0" structure needed.
            # Do NOT remove the w:b element, otherwise it reverts to the style's default (bold)!
            r.bold = False
            count += 1

    logger.info(
        f"Bold explicitly disabled on {count} runs across target heading styles."
    )
    return doc

