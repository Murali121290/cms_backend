"""
pipeline/runner.py — Orchestrates all steps for a single DOCX file.

Steps 1–7 pass a python-docx Document object through a chain.
Step 8 operates on the saved file via lxml (no python-docx SDT API).
Step 9 re-opens the saved file and converts Symbol/Math chars to HTML entities.
Step 10 re-opens the saved file, applies Math/Symbol character styles, and
       restores entity strings back to their original Unicode characters.
"""

import logging
import shutil
import traceback
from pathlib import Path
from datetime import datetime

from docx import Document

from docx_pipeline.config import PIPELINE_HALT_ON_ERROR
from docx_pipeline.utils.report import ReportLogger
from docx_pipeline.pipeline import (
    step0_preconversion,
    step1_cleanup,
    step2_unused_styles,
    step3_remove_bold,
    step4_heading_validation,
    step5_txt_flush,
    step6_remove_tags,
    step7_char_styles,
    step8_content_controls,
    step9_symbol_math,
    step10_math_symbol_styles,
)

# Steps 1–7: each receives and returns a python-docx Document
DOC_STEPS = [
    ("1-cleanup",           step1_cleanup),
    ("2-unused-styles",     step2_unused_styles),
    ("3-remove-bold",       step3_remove_bold),
    ("4-heading-validation",step4_heading_validation),
    ("5-txt-flush",         step5_txt_flush),
    ("6-remove-tags",       step6_remove_tags),
    ("7-char-styles",       step7_char_styles),
]

# Friendly labels emitted by the per-step logger.
STEP_TITLES = {
    0:  "Preconversion",
    1:  "Cleanup",
    2:  "Remove Unused Styles",
    3:  "Remove Bold",
    4:  "Heading Validation",
    5:  "Text Flush",
    6:  "Remove Tags",
    7:  "Character Styles",
    8:  "Content Controls",
    9:  "Symbol & Math",
    10: "Math Symbol Styles",
}

step_log = logging.getLogger("docx_pipeline.runner")


def process_file(input_path: Path, output_dir: Path, on_event=None) -> dict:
    """
    Run the full pipeline on one DOCX file.

    Returns a dict with keys:
      file     — filename
      status   — "ok" | "error"
      issues   — list of {"level", "step", "message"}
      output   — path to processed file (or None on error)
      report   — path to HTML sidecar report
    """
    logger = ReportLogger(input_path.name)
    result = {
        "file":   input_path.name,
        "status": "ok",
        "issues": [],
        "output": None,
        "report": None,
    }

    def _emit(event: dict) -> None:
        if on_event is None:
            return
        try:
            on_event(event)
        except Exception:
            step_log.debug("on_event callback raised", exc_info=True)

    def _log_start(n: int) -> None:
        step_log.info("→ Step %d - %s", n, STEP_TITLES[n])
        _emit({"type": "step", "step": n, "title": STEP_TITLES[n], "status": "start"})

    def _log_ok(n: int) -> None:
        step_log.info("✓ Step %d - %s", n, STEP_TITLES[n])
        _emit({"type": "step", "step": n, "title": STEP_TITLES[n], "status": "ok"})

    def _log_skip(n: int, reason: str) -> None:
        step_log.info("⏭ Step %d - %s (skipped: %s)", n, STEP_TITLES[n], reason)
        _emit({"type": "step", "step": n, "title": STEP_TITLES[n], "status": "skip", "reason": reason})

    def _log_fail(n: int, exc: BaseException) -> None:
        step_log.warning("✗ Step %d - %s (%s)", n, STEP_TITLES[n], exc)
        _emit({"type": "step", "step": n, "title": STEP_TITLES[n], "status": "fail", "error": str(exc)})

    def _make_backup(out_path: Path, step_label: str) -> None:
        """Create a timestamped backup copy before modifying the file."""
        if not out_path.exists():
            return
        backup_dir = output_dir / ".backups"
        backup_dir.mkdir(exist_ok=True)
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        stem = out_path.stem
        backup_path = backup_dir / f"{stem}_{step_label}_{timestamp}.docx"
        shutil.copy2(str(out_path), str(backup_path))

    try:
        # Copy to output dir so we never touch the original
        out_path = output_dir / input_path.name
        shutil.copy2(str(input_path), str(out_path))

        # ── Step 0: Pre-conversion file validation ─────────────
        logger.set_step("0-preconversion")
        _log_start(0)
        if not step0_preconversion.pre_flight_check(str(out_path), logger):
            _log_fail(0, RuntimeError("pre-flight check failed"))
            step_log.warning("Pipeline halted at Step 0 — pre-flight check failed.")
            result["status"] = "error"
            logger.to_html(str(output_dir))
            result["report"] = str(output_dir / f"{input_path.stem}_report.html")
            result["issues"] = logger.get_issues()
            return result

        # ── Steps 1–7: python-docx chain ──────────────────────
        doc = Document(str(out_path))
        doc = step0_preconversion.run(doc, logger)
        _log_ok(0)
        failed = False

        for label, step_mod in DOC_STEPS:
            step_num = int(label.split("-", 1)[0])
            logger.set_step(label)
            if failed and PIPELINE_HALT_ON_ERROR:
                _log_skip(step_num, "earlier step failed")
                logger.warning(f"Skipping {label} — earlier step failed")
                continue

            _log_start(step_num)
            try:
                _make_backup(out_path, label)
                doc = step_mod.run(doc, logger)
                _log_ok(step_num)
            except Exception as exc:
                _log_fail(step_num, exc)
                logger.error(f"{label} failed: {exc}")
                logger.error(traceback.format_exc())
                failed = True

        doc.save(str(out_path))

        # ── Step 8: lxml-based SDT grouping ───────────────────
        if failed and PIPELINE_HALT_ON_ERROR:
            _log_skip(8, "earlier step failed")
        else:
            _log_start(8)
            try:
                _make_backup(out_path, "8-content-controls")
                step8_content_controls.run(str(out_path), logger)
                _log_ok(8)
            except Exception as exc:
                _log_fail(8, exc)
                logger.error(f"8-content-controls failed: {exc}")
                logger.error(traceback.format_exc())
                failed = True

        # ── Step 9: Symbol / Math → Unicode entity conversion ─
        if failed and PIPELINE_HALT_ON_ERROR:
            _log_skip(9, "earlier step failed")
        else:
            _log_start(9)
            try:
                _make_backup(out_path, "9-symbol-math")
                doc2 = Document(str(out_path))
                doc2 = step9_symbol_math.run(doc2, logger)
                doc2.save(str(out_path))
                _log_ok(9)
            except Exception as exc:
                _log_fail(9, exc)
                logger.error(f"9-symbol-math failed: {exc}")
                logger.error(traceback.format_exc())
                failed = True

        # ── Step 10: Math/Symbol entity → character style + restore char ──
        if failed and PIPELINE_HALT_ON_ERROR:
            _log_skip(10, "earlier step failed")
        else:
            _log_start(10)
            try:
                _make_backup(out_path, "10-math-symbol-styles")
                doc3 = Document(str(out_path))
                doc3 = step10_math_symbol_styles.run(doc3, logger)
                doc3.save(str(out_path))
                _log_ok(10)
            except Exception as exc:
                _log_fail(10, exc)
                logger.error(f"10-math-symbol-styles failed: {exc}")
                logger.error(traceback.format_exc())
                failed = True

        if failed:
            step_log.warning("Pipeline completed with errors.")
            _emit({"type": "summary", "status": "error"})
        else:
            step_log.info("Pipeline completed successfully.")
            _emit({"type": "summary", "status": "ok"})

        result["output"] = str(out_path)

    except Exception as exc:
        result["status"] = "error"
        logger.error(f"Fatal error: {exc}")
        logger.error(traceback.format_exc())

    # Write sidecar HTML report
    try:
        report_path     = logger.to_html(str(output_dir))
        result["report"] = report_path
    except Exception:
        pass

    result["issues"] = logger.get_issues()
    if logger.has_errors():
        result["status"] = "error"

    return result
