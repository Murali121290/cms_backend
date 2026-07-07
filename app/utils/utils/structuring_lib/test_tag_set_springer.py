"""Client tag-set overlay: canonical <-> Springer tag translation.

Verifies tag_set_loader against the real tag_sets/springer.yaml, and an
end-to-end styler.process_docx run confirming that (a) tag_set=None
behaves exactly as before (canonical LL-/OL- tags), and (b)
tag_set="springer" writes Springer's case-aware Lc-/Uc- style names based
on each list item's actual marker case.
"""

import os
import tempfile

from docx import Document

from app.utils.utils.structuring_lib.tag_set_loader import (
    get_tag_map,
    get_reverse_tag_map,
    translate_tag,
)
from app.utils.utils.structuring_lib.styler import process_docx


def test_get_tag_map_missing_tag_set_returns_empty():
    assert get_tag_map(None) == {}
    assert get_tag_map("does-not-exist") == {}


def test_get_tag_map_springer_has_expected_entries():
    tag_map = get_tag_map("springer")
    assert tag_map["H1"] == "Head1"
    assert tag_map["BL-MID"] == "BulletList1"
    assert tag_map["BL-FIRST"] == "BulletList1_first"
    assert tag_map["LL-MID"] == {"lower": "Lc-AlphaList1", "upper": "Uc-AlphaList1"}
    assert tag_map["OL-FIRST"] == {
        "lower": "Lc-RomanList1_first",
        "upper": "Uc-RomanList1_first",
    }
    # NBX-* family intentionally has no Springer equivalent ("No" in the
    # source sheet) and must be absent so it falls through to canonical.
    assert "NBX-TTL" not in tag_map
    # Corrected copy/paste typo from the source spreadsheet.
    assert tag_map["EXT-LAST"] == "eXtractTxt1_last"


def test_translate_tag_exact_match():
    tag_map = get_tag_map("springer")
    assert translate_tag("H1", tag_map, []) == "Head1"


def test_translate_tag_unmapped_tag_passes_through():
    tag_map = get_tag_map("springer")
    assert translate_tag("NBX-TTL", tag_map, []) == "NBX-TTL"


def test_translate_tag_box_heading_levels_3_to_5():
    tag_map = get_tag_map("springer")
    assert translate_tag("BX1-H3", tag_map, []) == "Box-01-Head3"
    assert translate_tag("BX1-H4", tag_map, []) == "Box-01-Head4"
    assert translate_tag("BX1-H5", tag_map, []) == "Box-01-Head5"


def test_translate_tag_numbered_box_falls_back_to_box1_style():
    # box_prefixer.py numbers boxes by document order (BX2-, BX3-, ...),
    # but Springer's template has one generic box style family - any
    # numbered box without its own tag-set entry renders as Box 1's style.
    tag_map = get_tag_map("springer")
    assert translate_tag("BX2-H1", tag_map, []) == "Box-01-Head1"
    assert translate_tag("BX3-TXT", tag_map, []) == "Box-01-ParaFirstLine-Ind"
    assert translate_tag("BX12-H3", tag_map, []) == "Box-01-Head3"


def test_translate_tag_case_dependent_selects_variant():
    tag_map = get_tag_map("springer")
    assert translate_tag("LL-MID", tag_map, [], case="upper") == "Uc-AlphaList1"
    assert translate_tag("LL-MID", tag_map, [], case="lower") == "Lc-AlphaList1"


def test_translate_tag_case_dependent_without_case_falls_back_to_canonical():
    tag_map = get_tag_map("springer")
    assert translate_tag("LL-MID", tag_map, [], case=None) == "LL-MID"


def test_get_reverse_tag_map_inverts_springer():
    reverse = get_reverse_tag_map("springer")
    assert reverse["Head1"] == "H1"
    assert reverse["Uc-AlphaList1"] == "LL-MID"
    assert reverse["Lc-RomanList1_last"] == "OL-LAST"


def test_get_reverse_tag_map_duplicate_value_keeps_first_declared():
    # NL-TXT and BL-TXT both map forward to "ListItemPara-FL1"; NL-TXT is
    # declared first in springer.yaml, so it wins the (unavoidably
    # ambiguous) reverse mapping.
    reverse = get_reverse_tag_map("springer")
    assert reverse["ListItemPara-FL1"] == "NL-TXT"


def _build_doc(lines):
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    return doc


def _process(doc, tag_set=None):
    with tempfile.TemporaryDirectory() as tmp:
        in_path = os.path.join(tmp, "in.docx")
        out_path = os.path.join(tmp, "out.docx")
        doc.save(in_path)
        result = process_docx(in_path, out_path, mode="style", tag_set=tag_set)
        out_doc = Document(out_path)
        return result, [p.style.name for p in out_doc.paragraphs]


def test_process_docx_canonical_by_default():
    doc = _build_doc(["A. Apple is red", "B. Bravery under pressure"])
    result, styles = _process(doc, tag_set=None)
    assert result["success"] or result["paragraphs_processed"] == 2
    assert styles == ["LL-FIRST", "LL-LAST"]


def test_process_docx_translates_to_springer_by_marker_case():
    doc = _build_doc(["A. Apple is red", "B. Bravery under pressure"])
    _, styles = _process(doc, tag_set="springer")
    assert styles == ["Uc-AlphaList1_first", "Uc-AlphaList1_last"]


def test_process_docx_translates_lowercase_alpha_to_springer():
    doc = _build_doc(["a. apple is red", "b. bravery under pressure"])
    _, styles = _process(doc, tag_set="springer")
    assert styles == ["Lc-AlphaList1_first", "Lc-AlphaList1_last"]


def test_process_docx_translates_roman_to_springer():
    doc = _build_doc(["i. First point", "ii. Second point"])
    _, styles = _process(doc, tag_set="springer")
    assert styles == ["Lc-RomanList1_first", "Lc-RomanList1_last"]


def test_match_springer_box_marker_pairs_open_and_close():
    # Recognized in place by annotate_document (recognize_springer_box_markers) -
    # no paragraph text is ever rewritten to achieve this pairing.
    from app.utils.utils.structuring_lib.annotator import _match_springer_box_marker

    assert _match_springer_box_marker("Box1-open") == ("BX1", False, "BX1-open", "Box1")
    assert _match_springer_box_marker("Box1-close") == ("BX1", True, "BX1-open", "Box1")
    assert _match_springer_box_marker("box12-OPEN") == ("BX12", False, "BX12-open", "Box12")
    assert _match_springer_box_marker("BX1-Header") is None
    assert _match_springer_box_marker("TXT") is None


def test_springer_box_markers_not_recognized_without_tag_set():
    # Recognition is gated on tag_set being active (styler.process_docx
    # only opts annotate_document in when tag_set is truthy) - with no
    # tag_set, "<Box1-open>"/"<Box1-close>" fall through unrecognized,
    # exactly as before this feature existed.
    doc = _build_doc(["<Box1-open>", "Some box content.", "<Box1-close>"])
    _, styles = _process(doc, tag_set=None)
    assert styles[1] != "Box-01-ParaFirstLine-Ind"


def test_process_docx_pairs_springer_box_open_close_markers():
    doc = _build_doc([
        "<Box1-open>",
        "Some box content.",
        "<Box1-close>",
    ])
    _, styles = _process(doc, tag_set="springer")
    # Content between the markers is correctly paired (in memory, without
    # rewriting either marker's text) and prefixed by box_prefixer.py, then
    # translated to Springer's Box-01-* style.
    assert styles[1] == "Box-01-ParaFirstLine-Ind"
    # The markers' own applied style keeps Springer's "Box"-spelled wording
    # (matching the text the author actually typed) rather than switching
    # to the canonical "BX1-open"/"BX1-close" spelling.
    assert styles[0] == "Box1-open"
    assert styles[2] == "Box1-close"
