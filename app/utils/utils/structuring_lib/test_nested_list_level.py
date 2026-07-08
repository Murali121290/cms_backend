"""Nested/sub-list items (indented one level under a parent list item) get
a distinct "level 2" family - BL2-MID/NL2-MID - that is always MID, never
FIRST/LAST, regardless of the sub-list's own length or position. Mirrors
the existing table-nested convention (TBL2-MID/TNL2-MID in springer.yaml)
but for lists nested in ordinary body text.

Top-level list items keep their existing FIRST/MID/LAST behavior, and a
nested sub-list sitting in the middle of a parent list doesn't break the
parent's FIRST/MID/LAST run continuity (it's a transparent bridge, like
EMPTY/PMI)."""

from docx import Document

from app.utils.utils.structuring_lib.annotator import annotate_document, get_list_indent_level
from app.utils.utils.structuring_lib.list_normalizer import normalize_list_positions
from app.utils.utils.structuring_lib.tag_set_loader import get_tag_map, translate_tag


def _set_indent_level(paragraph, ilvl: int, num_id: int = 1) -> None:
    numPr = paragraph._p.get_or_add_pPr().get_or_add_numPr()
    numId = numPr.get_or_add_numId()
    numId.val = num_id
    ilvl_el = numPr.get_or_add_ilvl()
    ilvl_el.val = ilvl


def test_get_list_indent_level_reads_ilvl():
    doc = Document()
    p = doc.add_paragraph("Some item")
    assert get_list_indent_level(p) == 0
    _set_indent_level(p, 1)
    assert get_list_indent_level(p) == 1


def test_single_nested_bullet_gets_bl2_mid():
    doc = Document()
    doc.add_paragraph("1. First item")
    sub = doc.add_paragraph("• Sub point")
    _set_indent_level(sub, 1)
    doc.add_paragraph("2. Second item")

    annotations = normalize_list_positions(annotate_document(doc))
    tags = [a["tag"] for a in annotations]
    assert tags == ["NL-FIRST", "BL2-MID", "NL-LAST"]


def test_nested_bullet_sublist_of_three_all_mid_never_first_or_last():
    doc = Document()
    doc.add_paragraph("1. First item")
    doc.add_paragraph("2. Second item")
    subs = [doc.add_paragraph(f"• Sub point {n}") for n in (1, 2, 3)]
    for sub in subs:
        _set_indent_level(sub, 1)
    doc.add_paragraph("3. Third item")

    annotations = normalize_list_positions(annotate_document(doc))
    tags = [a["tag"] for a in annotations]

    # Parent numbered list stays continuous FIRST -> MID -> LAST even
    # though a nested sub-list interrupts it in the middle.
    assert tags == [
        "NL-FIRST",
        "NL-MID",
        "BL2-MID", "BL2-MID", "BL2-MID",
        "NL-LAST",
    ]


def test_bl2_mid_translates_to_springer_bulletlist2():
    tag_map = get_tag_map("springer")
    assert translate_tag("BL2-MID", tag_map, []) == "BulletList2"
    assert translate_tag("NL2-MID", tag_map, []) == "NumberList2"
    assert translate_tag("UL2-MID", tag_map, []) == "UnNumberedList2"


def test_nested_lettered_sublist_gets_ll2_mid():
    doc = Document()
    doc.add_paragraph("1. First item")
    sub = doc.add_paragraph("a. Sub point")
    _set_indent_level(sub, 1)
    doc.add_paragraph("2. Second item")

    annotations = normalize_list_positions(annotate_document(doc))
    tags = [a["tag"] for a in annotations]
    assert tags == ["NL-FIRST", "LL2-MID", "NL-LAST"]


def test_nested_roman_sublist_gets_ol2_mid():
    doc = Document()
    doc.add_paragraph("1. First item")
    sub = doc.add_paragraph("i. Sub point")
    _set_indent_level(sub, 1)
    doc.add_paragraph("2. Second item")

    annotations = normalize_list_positions(annotate_document(doc))
    tags = [a["tag"] for a in annotations]
    assert tags == ["NL-FIRST", "OL2-MID", "NL-LAST"]


def test_ll2_and_ol2_mid_translate_to_springer_case_aware():
    tag_map = get_tag_map("springer")
    assert translate_tag("LL2-MID", tag_map, [], case="lower") == "Lc-AlphaList2"
    assert translate_tag("LL2-MID", tag_map, [], case="upper") == "Uc-AlphaList2"
    assert translate_tag("OL2-MID", tag_map, [], case="lower") == "Lc-RomanList2"
    assert translate_tag("OL2-MID", tag_map, [], case="upper") == "Uc-RomanList2"
