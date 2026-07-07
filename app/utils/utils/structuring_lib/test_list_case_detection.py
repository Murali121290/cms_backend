"""Case-aware list detection for alphabetical (LL-) and roman-numeral (OL-)
lists: detect_list_kind() must report not just the list kind but also
whether the marker was uppercase or lowercase, so a client tag set can pick
between e.g. Springer's Lc-AlphaList1 / Uc-AlphaList1 style variants.

Also covers the roman-vs-lettered ordering fix: single-letter roman
numerals (i, v, x, l, c, d, m) must resolve to "roman", not "lettered",
since roman's character set is a strict subset of "any single letter"."""

from docx import Document

from app.utils.utils.structuring_lib.annotator import annotate_document, detect_list_kind


def test_lowercase_alpha_marker_detected_with_lower_case():
    assert detect_list_kind("a. apple is red") == ("lettered", "lower")


def test_uppercase_alpha_marker_detected_with_upper_case():
    assert detect_list_kind("A. Apple is red") == ("lettered", "upper")


def test_uppercase_alpha_marker_with_paren_detected():
    assert detect_list_kind("B) Bravery under pressure") == ("lettered", "upper")


def test_uppercase_roman_marker_detected_with_upper_case():
    assert detect_list_kind("I. Introduction") == ("roman", "upper")


def test_lowercase_multiletter_roman_marker_detected_with_lower_case():
    assert detect_list_kind("ii. Body section") == ("roman", "lower")


def test_uppercase_multiletter_roman_marker_detected_with_upper_case():
    assert detect_list_kind("III. Conclusion") == ("roman", "upper")


def test_single_letter_lowercase_roman_wins_over_lettered():
    # "i", "v", "x", "l", "c", "d", "m" are valid roman numerals as well as
    # ordinary single letters - roman must win the collision.
    assert detect_list_kind("i. Introduction") == ("roman", "lower")
    assert detect_list_kind("v. Fifth item") == ("roman", "lower")


def test_single_letter_lowercase_non_roman_letter_stays_lettered():
    assert detect_list_kind("b. Bravery") == ("lettered", "lower")


def test_alpha_list_items_tagged_ll_mid_with_list_case_recorded():
    doc = Document()
    doc.add_paragraph("A. Apple is red")
    doc.add_paragraph("B. Bravery under pressure")
    annotations = annotate_document(doc)
    assert [a["tag"] for a in annotations] == ["LL-MID", "LL-MID"]
    assert [a["list_case"] for a in annotations] == ["upper", "upper"]


def test_roman_list_items_tagged_ol_mid_with_list_case_recorded():
    doc = Document()
    doc.add_paragraph("i. First point")
    doc.add_paragraph("ii. Second point")
    annotations = annotate_document(doc)
    assert [a["tag"] for a in annotations] == ["OL-MID", "OL-MID"]
    assert [a["list_case"] for a in annotations] == ["lower", "lower"]


def test_non_list_paragraph_has_no_list_case():
    doc = Document()
    doc.add_paragraph("Just an ordinary sentence.")
    annotations = annotate_document(doc)
    assert annotations[0]["list_case"] is None
