"""Regression test for a realistic, multi-level document classified
end-to-end through the standalone pipeline (reader -> features -> signature
-> classifier). Catches whole-document regressions that the narrower unit
tests (test_signature.py, test_classifier.py) miss when testing isolated
features - e.g. a comparator change that looks fine on hand-built tuples
but produces an unexpected document-wide ranking in practice."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.utils.utils.structuring_lib.heading_classifier.pipeline import classify_paragraphs
from app.utils.utils.structuring_lib.heading_classifier.reader import read_paragraphs


def _add(doc, text, size, bold=False, center=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p


def test_mixed_level_document_classifies_correctly():
    doc = Document()
    _add(doc, "DOCUMENT MANAGEMENT SYSTEM", 24, bold=True, center=True)
    _add(doc, "Introduction", 18, bold=True)
    _add(doc, "Purpose", 16, bold=True)
    _add(doc, "1. Requirements", 18, bold=True)
    _add(doc, "Functional Requirements", 16, bold=True)

    body = doc.add_paragraph()
    body.add_run("This paragraph contains ")
    bold_run = body.add_run("bold text")
    bold_run.bold = True
    body.add_run(" but is not a heading.")
    for run in body.runs:
        run.font.size = Pt(12)

    _add(doc, "APPENDIX", 18, bold=True)
    _add(doc, "Abbreviations", 16, bold=True)

    raw_paragraphs = read_paragraphs(doc)
    results = classify_paragraphs(doc, raw_paragraphs)

    # Heading level is determined by font-size rank only. All 18pt paragraphs
    # (Introduction, "1. Requirements", APPENDIX) are H2. All 16pt paragraphs
    # (Purpose, Functional Requirements, Abbreviations) are H3, regardless of
    # bold/ALL CAPS differences between them. The 24pt title is H1. The 12pt
    # body paragraph ties the document's baseline (smallest font, no other
    # senior signal - its inline bold run isn't the dominant run by length)
    # and stays Body. Secondary signals (ALL CAPS, Title Case, etc.) still
    # affect candidacy and signature ordering within a font-size group but
    # do not split paragraphs of the same font-size rank into different levels.
    expected = ["H1", "H2", "H3", "H2", "H3", "Body", "H2", "H3"]
    assert [r.classification for r in results] == expected
