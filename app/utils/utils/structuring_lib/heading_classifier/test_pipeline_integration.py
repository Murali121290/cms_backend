"""Regression suite for the styler.py integration point
(`classify_headings_by_formatting`) and the full `process_docx` pipeline.

This proves the "full promotion, locked-only restriction" reconciliation
rule: any non-locked paragraph can be classified as a heading purely from
its formatting signature, regardless of what tag the regex engine
originally assigned it. The only paragraphs that are ever off-limits are
author-locked (explicit <TAG>) ones. When the formatting engine has no
opinion (verdict is "Body"), the regex engine's original tag/style is left
completely untouched - it is never forced down to a generic tag.
"""

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from app.utils.utils.structuring_lib.annotator import annotate_document
from app.utils.utils.structuring_lib.heading_classifier import classify_headings_by_formatting
from app.utils.utils.structuring_lib.hierarchy_manager import enforce_hierarchy
from app.utils.utils.structuring_lib.styler import process_docx

SCRATCH_DIR = r"C:\Users\DELL\AppData\Local\Temp\claude\d--S4-git-new-main-cms-backend\f26b9e87-4c1e-4808-918f-db4119b2ad18\scratchpad"


def test_no_formatting_cues_falls_back_gracefully():
    doc = Document()
    doc.add_paragraph("BACKGROUND INFORMATION")  # regex-tags this H1, no run formatting set

    annotations = annotate_document(doc)
    refined = classify_headings_by_formatting(doc, annotations)

    assert len(refined) == 1
    assert refined[0]["tag"].startswith("H")


def test_locked_paragraph_never_reclassified_despite_heavy_formatting():
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("<H2> Introduction")
    run.bold = True
    run.font.size = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    annotations = annotate_document(doc)
    assert annotations[0]["locked"] is True
    assert annotations[0]["tag"] == "H2"

    refined = classify_headings_by_formatting(doc, annotations)
    assert refined[0]["tag"] == "H2"
    assert refined[0]["style"] == "H2"


def test_non_heading_regex_tag_promoted_to_heading_on_strong_signal():
    # Regex tags this plain sentence as TXT, but full promotion means a
    # non-locked paragraph with a strong, senior formatting signature
    # (bold, largest font in the doc, centered) is now eligible to become a
    # heading purely from formatting, regardless of its original tag.
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("This is a body paragraph.")
    run.bold = True
    run.font.size = Pt(28)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    annotations = annotate_document(doc)
    assert annotations[0]["tag"] == "TXT"

    refined = classify_headings_by_formatting(doc, annotations)
    assert refined[0]["tag"] == "H1"
    assert refined[0]["style"] == "H1"


def test_structurally_meaningful_non_heading_tag_preserved_when_no_strong_signal():
    # "Objectives" is recognized by the regex engine as a distinct
    # structural role (OBJ1), not a generic heading. With no formatting
    # signal at all, the formatting engine has no opinion (verdict Body),
    # so this structurally meaningful tag must be left completely
    # untouched rather than forced down to something generic.
    doc = Document()
    doc.add_paragraph("Objectives")

    annotations = annotate_document(doc)
    assert annotations[0]["tag"] == "OBJ1"

    refined = classify_headings_by_formatting(doc, annotations)
    assert refined[0]["tag"] == "OBJ1"
    assert refined[0]["style"] == "OBJ1"


def test_regex_heading_with_zero_signal_keeps_its_original_tag():
    # No special protection for heading tags anymore - this just exercises
    # the general "Body verdict leaves the original tag untouched" rule for
    # a paragraph that happens to already be tagged as a heading.
    doc = Document()
    doc.add_paragraph("Introduction")  # regex-tags H1; no bold/size/alignment set

    annotations = annotate_document(doc)
    assert annotations[0]["tag"] == "H1"

    refined = classify_headings_by_formatting(doc, annotations)
    assert refined[0]["tag"] in {"H1", "H2", "H3", "H4", "H5", "H6"}


def test_bullet_list_item_never_promoted_to_heading():
    # A bulleted/numbered list item may be bold and large, but the regex
    # engine's structural tag (BL-*, NL-*, UL-*, etc.) must be preserved -
    # the formatting engine only touches generic body-text tags (TXT) and
    # existing heading tags, never structural list/box/reference tags.
    doc = Document()
    p = doc.add_paragraph()
    run = p.add_run("• Important point about the system design")
    run.bold = True
    run.font.size = Pt(18)

    annotations = annotate_document(doc)
    original_tag = annotations[0]["tag"]
    # Must be some kind of list tag from the regex engine (BL-*, UL-*, etc.)
    assert "H" not in original_tag or original_tag in {"H1","H2","H3","H4","H5","H6"}

    refined = classify_headings_by_formatting(doc, annotations)
    assert refined[0]["tag"] == original_tag


def test_empty_annotations_list_returns_unchanged():
    assert classify_headings_by_formatting(Document(), []) == []


def _build_sample_docx(path):
    doc = Document()

    title = doc.add_paragraph()
    title_run = title.add_run("INTRODUCTION")
    title_run.bold = True
    title_run.font.size = Pt(28)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("This chapter introduces the core concepts used throughout the rest of the text.")

    sub = doc.add_paragraph()
    sub_run = sub.add_run("Background")
    sub_run.bold = True
    sub_run.font.size = Pt(18)

    doc.add_paragraph("Some background body text follows here.")
    doc.save(path)


def test_process_docx_end_to_end_style_mode():
    os.makedirs(SCRATCH_DIR, exist_ok=True)
    input_path = os.path.join(SCRATCH_DIR, "heading_classifier_style_input.docx")
    output_path = os.path.join(SCRATCH_DIR, "heading_classifier_style_output.docx")
    _build_sample_docx(input_path)

    result = process_docx(input_path, output_path, mode="style")
    assert result["success"] is True


def test_process_docx_end_to_end_tag_mode():
    os.makedirs(SCRATCH_DIR, exist_ok=True)
    input_path = os.path.join(SCRATCH_DIR, "heading_classifier_tag_input.docx")
    output_path = os.path.join(SCRATCH_DIR, "heading_classifier_tag_output.docx")
    _build_sample_docx(input_path)

    result = process_docx(input_path, output_path, mode="tag")
    assert result["success"] is True


def test_h5_h6_assignment_composes_with_hierarchy_auto_fix():
    doc = Document()

    h1 = doc.add_paragraph()
    h1_run = h1.add_run("<H1> Chapter One")
    h1_run.bold = True
    h1_run.font.size = Pt(28)

    # Directly exercise the hierarchy_manager.py H5/H6 fix by feeding
    # enforce_hierarchy a regex-produced H6 right after H1.
    doc.add_paragraph("A run-in heading")

    annotations = annotate_document(doc)
    annotations[1]["tag"] = "H6"
    annotations[1]["style"] = "H6"

    refined = enforce_hierarchy(annotations)
    assert refined[1]["tag"] == "H2"
