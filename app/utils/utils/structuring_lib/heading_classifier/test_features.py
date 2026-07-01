"""Tests for Phase 2 (feature extraction) of the formatting-based heading
classifier."""

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt

from app.utils.utils.structuring_lib.heading_classifier.features import (
    build_feature_context,
    extract_features,
)
from app.utils.utils.structuring_lib.heading_classifier.reader import read_paragraphs


def _features_for(doc, index):
    raw_paragraphs = read_paragraphs(doc)
    context = build_feature_context(raw_paragraphs)
    raw = raw_paragraphs[index]
    prev = raw_paragraphs[index - 1] if index > 0 else None
    next_ = raw_paragraphs[index + 1] if index + 1 < len(raw_paragraphs) else None
    return extract_features(raw, prev, next_, context, doc)


def test_text_shape_features():
    doc = Document()
    doc.add_paragraph("HELLO WORLD")
    doc.add_paragraph("hello world")
    doc.add_paragraph("Hello World Title")
    doc.add_paragraph("A sentence that ends with a period.")
    doc.add_paragraph("A label that ends with a colon:")
    doc.add_paragraph("")

    f0 = _features_for(doc, 0)
    f1 = _features_for(doc, 1)
    f2 = _features_for(doc, 2)
    f3 = _features_for(doc, 3)
    f4 = _features_for(doc, 4)
    f5 = _features_for(doc, 5)

    assert f0["is_all_uppercase"] is True
    assert f1["is_lowercase"] is True
    assert f2["is_title_case"] is True
    assert f3["ends_with_period"] is True
    assert f4["ends_with_colon"] is True
    assert f5["is_empty"] is True
    assert f0["word_count"] == 2
    assert f0["char_count"] == len("HELLO WORLD")


def test_font_size_rank_largest_is_zero_and_ties_share_rank():
    doc = Document()
    p1 = doc.add_paragraph()
    p1.add_run("Big").font.size = Pt(24)
    p2 = doc.add_paragraph()
    p2.add_run("Also Big").font.size = Pt(24)
    p3 = doc.add_paragraph()
    p3.add_run("Smaller").font.size = Pt(12)

    f0 = _features_for(doc, 0)
    f1 = _features_for(doc, 1)
    f2 = _features_for(doc, 2)

    assert f0["font_size_rank"] == 0
    assert f1["font_size_rank"] == 0
    assert f2["font_size_rank"] == 1


def test_font_features_use_dominant_run():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("1.2  ")  # short leading run, not bold
    main_run = p.add_run("Heading Text")
    main_run.bold = True
    main_run.font.size = Pt(18)

    f = _features_for(doc, 0)
    assert f["bold"] is True
    assert f["font_size_pt"] == 18.0


def test_paragraph_alignment_and_spacing_features():
    doc = Document()
    p = doc.add_paragraph("Centered Heading")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    body = doc.add_paragraph("Body with default spacing.")
    body.paragraph_format.space_before = Pt(6)
    doc.add_paragraph("Another body paragraph.").paragraph_format.space_before = Pt(6)

    f = _features_for(doc, 0)
    assert f["alignment"] == "CENTER"
    assert f["space_before_pt"] == 24.0
    assert f["is_large_space_before"] is True


def test_word_list_and_numbering_detection_delegates_to_annotator_helpers():
    doc = Document()
    p = doc.add_paragraph("1. First numbered item")
    numPr = p._p.get_or_add_pPr().get_or_add_numPr()
    numId = numPr.get_or_add_numId()
    numId.val = 1
    ilvl = numPr.get_or_add_ilvl()
    ilvl.val = 0

    f = _features_for(doc, 0)
    assert f["is_list"] is True


def test_outline_level_read_from_raw_xml():
    doc = Document()
    p = doc.add_paragraph("Outlined Paragraph")
    pPr = p._p.get_or_add_pPr()
    outline_lvl = pPr.makeelement(qn("w:outlineLvl"), {qn("w:val"): "2"})
    pPr.append(outline_lvl)

    f = _features_for(doc, 0)
    assert f["outline_level"] == 2

    doc2 = Document()
    doc2.add_paragraph("No outline level here")
    f2 = _features_for(doc2, 0)
    assert f2["outline_level"] is None


def test_context_features_track_neighbors_and_position():
    doc = Document()
    doc.add_paragraph("First")
    doc.add_paragraph("")
    doc.add_paragraph("Last")

    f0 = _features_for(doc, 0)
    f1 = _features_for(doc, 1)

    assert f0["prev_is_empty"] is None
    assert f0["next_is_empty"] is True
    assert f1["position_index"] == 1
    assert f0["position_ratio"] == 0.0
