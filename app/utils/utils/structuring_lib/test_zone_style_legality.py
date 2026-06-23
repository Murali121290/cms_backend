"""Gap #6: zone-style legality check. Out-of-zone styles are warned about
(counted in result['zone_warnings']) but never rewritten - warn-only."""

from docx import Document

from app.utils.utils.structuring_lib.styler import process_docx


def test_out_of_zone_style_is_warned_but_not_rewritten(tmp_path):
    doc = Document()
    doc.add_paragraph("REFERENCES")
    # KT is a legal style elsewhere, but not inside REFERENCES_BLOCK.
    doc.add_paragraph("<KT> Forced Term")

    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    doc.save(input_path)

    result = process_docx(str(input_path), str(output_path))

    assert result["success"] is True
    assert result["zone_warnings"] == 1

    out_doc = Document(str(output_path))
    # warn-only: the style is still applied as-is, not rewritten.
    assert out_doc.paragraphs[1].style.name == "KT"


def test_clean_references_block_has_no_zone_warnings(tmp_path):
    doc = Document()
    doc.add_paragraph("REFERENCES")
    doc.add_paragraph("1. Smith J. Title of paper. Journal. 2020.")

    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    doc.save(input_path)

    result = process_docx(str(input_path), str(output_path))

    assert result["success"] is True
    assert result["zone_warnings"] == 0
