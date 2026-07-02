"""P8 (optional enhancement): mnemonic/lettered lists - a single capital
letter followed by a tab/space run (e.g. "A   Apple is for resilience")
should be detected as a list item and tagged UL-MID, not generic body text.

Also covers the em-dash/en-dash variant (e.g. "S—Is your sleep disturbed?"),
as used in mnemonic checklists like SIGECAPS."""

from docx import Document

from app.utils.utils.structuring_lib.annotator import annotate_document, detect_list_kind
from app.utils.utils.structuring_lib.list_normalizer import normalize_list_positions


def test_detect_list_kind_recognizes_mnemonic_lettered_pattern():
    assert detect_list_kind("A   Apple is for resilience") == "lettered"


def test_lettered_list_items_tagged_ul_mid():
    doc = Document()
    doc.add_paragraph("A   Apple is for resilience")
    doc.add_paragraph("B   Bravery under pressure")
    annotations = annotate_document(doc)
    tags = [a["tag"] for a in annotations]
    assert tags == ["UL-MID", "UL-MID"]


def test_ordinary_sentence_is_not_treated_as_lettered_list():
    assert detect_list_kind("A patient presented with fever.") is None


def test_detect_list_kind_recognizes_em_dash_mnemonic_pattern():
    assert detect_list_kind("S—Is your sleep disturbed?") == "lettered"


def test_detect_list_kind_recognizes_en_dash_mnemonic_pattern():
    assert detect_list_kind("I–Have you noted a loss of libido?") == "lettered"


def test_detect_list_kind_recognizes_em_dash_with_trailing_space():
    assert detect_list_kind("A— Have you experienced changes in your appetite?") == "lettered"


def test_sigecaps_em_dash_list_normalized_to_first_mid_last():
    lines = [
        "S—Is your sleep disturbed?",
        "I—Have you noted a loss of libido or interest in your usual activities?",
        "G—Are you feeling guilty or having self-deprecatory thoughts?",
        "E—Have you noticed a decrease in your energy level?",
        "C—Have you been having trouble concentrating?",
        "A—Have you experienced changes in your appetite and weight?",
        "P—Have you been physically slowed down or sped up?",
        "S—Have you had thoughts of suicide, feelings of hopelessness?",
    ]
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    annotations = normalize_list_positions(annotate_document(doc))
    tags = [a["tag"] for a in annotations]
    assert tags == [
        "UL-FIRST",
        "UL-MID",
        "UL-MID",
        "UL-MID",
        "UL-MID",
        "UL-MID",
        "UL-MID",
        "UL-LAST",
    ]


def test_interrupting_text_paragraph_splits_em_dash_list_into_two_runs():
    lines = [
        "S—Is your sleep disturbed?",
        "I—Have you noted a loss of libido or interest in your usual activities?",
        "G—Are you feeling guilty or having self-deprecatory thoughts?",
        "asdfgh",
        "E—Have you noticed a decrease in your energy level?",
        "C—Have you been having trouble concentrating?",
        "A— Have you experienced changes in your appetite and weight?",
        "P—Have you been physically slowed down or sped up?",
        "S—Have you had thoughts of suicide, feelings of hopelessness?",
    ]
    doc = Document()
    for line in lines:
        doc.add_paragraph(line)
    annotations = normalize_list_positions(annotate_document(doc))
    tags = [a["tag"] for a in annotations]
    # "asdfgh" is just a stand-in non-list paragraph to verify the list
    # splits around it - its own tag isn't the point of this test. It now
    # becomes H2 via the short-standalone-phrase fallback (single word, no
    # terminal punctuation), an accepted tradeoff of that fallback.
    assert tags == [
        "UL-FIRST",
        "UL-MID",
        "UL-LAST",
        "H2",
        "UL-FIRST",
        "UL-MID",
        "UL-MID",
        "UL-MID",
        "UL-LAST",
    ]
