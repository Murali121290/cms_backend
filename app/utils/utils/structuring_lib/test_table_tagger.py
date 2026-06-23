"""Gap #4 (source/footnote detection) and gap #5 (T4 stub-heading heuristic)
in styler.tag_tables()."""

from docx import Document

from app.utils.utils.structuring_lib.styler import tag_tables


def _build_table_doc(rows):
    doc = Document()
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    for r_idx, row_texts in enumerate(rows):
        for c_idx, text in enumerate(row_texts):
            table.cell(r_idx, c_idx).text = text
    return doc, table


def test_header_row_still_detected():
    doc, table = _build_table_doc([["NAME", "VALUE"], ["Item 1", "10"]])
    tag_tables(doc, mode="style")
    assert table.cell(0, 0).paragraphs[0].style.name == "TBCH"


def test_source_line_tagged_tsn():
    doc, table = _build_table_doc([["NAME", "VALUE"], ["Item 1", "Source: WHO, 2020"]])
    tag_tables(doc, mode="style")
    assert table.cell(1, 1).paragraphs[0].style.name == "TSN"


def test_footnote_line_tagged_tfn():
    doc, table = _build_table_doc([["NAME", "VALUE"], ["Item 1", "*Excludes outliers"]])
    tag_tables(doc, mode="style")
    assert table.cell(1, 1).paragraphs[0].style.name == "TFN"


def test_stub_heading_first_column_tagged_t4():
    doc, table = _build_table_doc([["NAME", "VALUE"], ["CAR T-Cells", "95%"]])
    tag_tables(doc, mode="style")
    assert table.cell(1, 0).paragraphs[0].style.name == "T4"


def test_sentence_like_first_column_stays_body():
    doc, table = _build_table_doc([["NAME", "VALUE"], ["Administer 500 mg daily", "Yes"]])
    tag_tables(doc, mode="style")
    assert table.cell(1, 0).paragraphs[0].style.name == "TB"


def test_numeric_first_column_stays_body():
    doc, table = _build_table_doc([["NAME", "VALUE"], ["95%", "Yes"]])
    tag_tables(doc, mode="style")
    assert table.cell(1, 0).paragraphs[0].style.name == "TB"
