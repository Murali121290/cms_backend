"""
Core document styling and tagging module.
Handles applying styles and tags to DOCX documents.
"""

import logging
import re
from typing import Literal, Dict, Any, Optional, Callable
from docx import Document
from docx.oxml.ns import qn
from .annotator import annotate_document, detect_list_kind, is_list_paragraph, parse_leading_style_hint
from .logger_config import get_logger
from .box_prefixer import apply_box_tag_prefixes
from .heading_classifier import classify_headings_by_formatting
from .hierarchy_manager import enforce_hierarchy, demote_long_headings
from .list_normalizer import normalize_list_positions
from .reference_normalizer import normalize_reference_numbers
from .zone_styles import check_zone_style_legality
from .enhanced_processor import DocumentProcessor
from .tag_set_loader import get_tag_map, get_reverse_tag_map, translate_tag

logger = get_logger(__name__)


from .rules_loader import get_rules_loader

_SOURCE_LINE_RE = re.compile(
    r"^(?:Source|Adapted from|Data from|Reproduced from|Courtesy of)\b", re.IGNORECASE
)
_FOOTNOTE_LINE_RE = re.compile(r"^(?:Note:|\*|†|[a-z]\))")
_STUB_HEADING_RE = re.compile(r"^[A-Z0-9][A-Z0-9\s/&\-]{1,59}$")
_MEASUREMENT_RE = re.compile(
    r"^\d+(?:\.\d+)?\s*(?:mg|g|kg|ml|l|mmHg|bpm|kcal|IU|mol|cm|mm|m|km)\b", re.IGNORECASE
)


def _looks_like_stub_heading(text: str) -> bool:
    """T4 heuristic: short, punctuation-free, mostly-capitalized cell text
    that reads like a row label rather than a full sentence or a number."""
    text = text.strip()
    if not text or len(text) > 60:
        return False
    if re.search(r"[.!?;:]\s*$", text):
        return False
    if re.match(r"^\d", text) or "%" in text:
        return False
    if _MEASUREMENT_RE.match(text):
        return False
    if _STUB_HEADING_RE.match(text):
        return True

    words = text.split()
    if not words:
        return False
    title_cased = sum(1 for w in words if w[:1].isupper())
    return (title_cased / len(words)) >= 0.7


def tag_tables(doc: Document, mode: Literal["style", "tag"] = "style", tag_map: Dict[str, Any] = None) -> None:
    """
    Apply styles or tags to table cells.

    Args:
        doc: python-docx Document object
        mode: "style" to apply Word styles, "tag" to prefix text with [TAG]
        tag_map: optional canonical -> client tag map (see tag_set_loader);
            when provided, every tag/style written to a table cell is
            translated through it before use. Table cells never carry
            case-dependent (Lc-/Uc-) tags, so translation always passes
            case=None.

    Raises:
        ValueError: If mode is invalid
    """
    if mode not in ("style", "tag"):
        raise ValueError(f"Invalid mode: {mode}. Must be 'style' or 'tag'")

    rules_loader = get_rules_loader()
    prefixes = rules_loader.get_structural_tags().get("prefixes", []) if tag_map else []
    table_config = rules_loader.get_table_config()
    header_style_name = table_config.get("header_style", "T2")
    body_style_name = table_config.get("body_style", "T")
    bullet_style_name = table_config.get("bullet_style", "TBL-MID")
    number_style_name = table_config.get("number_style", "TNL-MID")
    roman_style_name = table_config.get("roman_style", "TOL-MID")
    source_note_style_name = table_config.get("source_note_style", "TSN")
    footnote_style_name = table_config.get("footnote_style", "TFN")
    stub_heading_style_name = table_config.get("stub_heading_style", "T4")
    header_threshold = table_config.get("header_threshold", 0.7)
    doc_processor = DocumentProcessor()

    for table_idx, table in enumerate(doc.tables):
        try:
            for row_idx, row in enumerate(table.rows):
                for cell_idx, cell in enumerate(row.cells):
                    for para in cell.paragraphs:
                        text = para.text.strip()
                        if not text:
                            continue

                        style_name = para.style.name if para.style else ""

                        if _SOURCE_LINE_RE.match(text):
                            tag = "TSN"
                            style = source_note_style_name
                        elif _FOOTNOTE_LINE_RE.match(text):
                            tag = "TFN"
                            style = footnote_style_name
                        else:
                            list_kind_result = detect_list_kind(text, style_name, is_list_paragraph(para), para, doc)
                            list_kind = list_kind_result[0] if list_kind_result is not None else None

                            if list_kind == "bullet":
                                tag = "TBL-MID"
                                style = bullet_style_name
                            elif list_kind == "number":
                                tag = "TNL-MID"
                                style = number_style_name
                            elif list_kind == "roman":
                                tag = "TOL-MID"
                                style = roman_style_name
                            else:
                                score = doc_processor.detect_table_header_smart(
                                    text, row_idx, cell_idx, len(table.rows)
                                )
                                if score >= header_threshold:
                                    tag = "T2"
                                    style = header_style_name
                                elif cell_idx == 0 and row_idx > 0 and _looks_like_stub_heading(text):
                                    tag = "T4"
                                    style = stub_heading_style_name
                                else:
                                    tag = "T"
                                    style = body_style_name

                        if tag_map:
                            tag = translate_tag(tag, tag_map, prefixes)
                            style = translate_tag(style, tag_map, prefixes)

                        if mode == "style":
                            try:
                                para.style = style
                                logger.debug(f"Applied style '{style}' to table[{table_idx}].row[{row_idx}].cell[{cell_idx}]")
                            except KeyError:
                                logger.warning(f"Style '{style}' not found in template. Creating it as default Paragraph Style.")
                                try:
                                    # Create the style if missing
                                    styles = doc.styles
                                    new_style = styles.add_style(style, 1) # 1 is WD_STYLE_TYPE.PARAGRAPH
                                    # Optional: Set base style
                                    new_style.base_style = styles['Normal']
                                    para.style = style
                                except Exception as e2:
                                    logger.error(f"Failed to create style '{style}': {e2}")
                        else:  # tag mode
                            if para.runs:
                                try:
                                    para.runs[0].text = f"[{tag}] " + para.runs[0].text
                                    logger.debug(f"Added tag '{tag}' to table[{table_idx}].row[{row_idx}].cell[{cell_idx}]")
                                except Exception as e:
                                    logger.error(f"Failed to tag table cell: {e}")
        except Exception as e:
            logger.error(f"Error processing table {table_idx}: {e}")


from docx.enum.style import WD_STYLE_TYPE

def process_cross_references(doc: Document) -> None:
    """
    Process cross-references: wrap in brackets and apply character style.
    Warning: This recreates paragraph runs, identifying plain text matches. 
    Existing run-level formatting (bold/italic) inside the specific paragraph *might* be reset 
    if strictly text-based splitting is used. 
    """
    rules_loader = get_rules_loader()
    xref_config = rules_loader.get_cross_references()
    
    if not xref_config:
        return

    # Collect all paragraphs
    all_paras = list(doc.paragraphs)
    for t in doc.tables:
        for r in t.rows:
            for c in r.cells:
                all_paras.extend(c.paragraphs)

    for rule_name, rule in xref_config.items():
        pattern = rule.get("pattern")
        style_name = rule.get("style", "Hyperlink")
        
        # Validate style exists or create it (Character style)
        try:
            doc.styles[style_name]
        except KeyError:
            # Create character style
            logger.info(f"Creating missing character style: {style_name}")
            try:
                styles = doc.styles
                char_style = styles.add_style(style_name, WD_STYLE_TYPE.CHARACTER)
                # Dictionary check for base style?
                # char_style.base_style = styles['Default Paragraph Font'] 
            except Exception as e:
                logger.error(f"Failed to create character style {style_name}: {e}")

        # Compile regex once
        regex = re.compile(pattern)
        
        count = 0
        total_paras = len(all_paras)
        logger.info(f"Scanning {total_paras} paragraphs for pattern: {pattern[:20]}...")
        
        for para in all_paras:
            count += 1
            if count % 500 == 0:
                logger.info(f"Scanned {count}/{total_paras} paragraphs...")
                
            text = para.text
            if not text:
                continue
            
            if not regex.search(text):
                continue
                
            # Iterate over runs to preserve formatting
            # We must iterate a copy because we might modify the list of runs
            
            # --- Better Implementation below ---
            # 1. Capture existing state
            original_runs = []
            for r in para.runs:
                original_runs.append({
                    "text": r.text,
                    "style": r.style,
                    "bold": r.bold,
                    "italic": r.italic,
                    "underline": r.underline,
                    "superscript": r.font.superscript if r.font else None,
                    "subscript": r.font.subscript if r.font else None,
                    "strike": r.font.strike if r.font else None,
                    "font": r.font.name if r.font else None,
                    "size": r.font.size if r.font else None,
                    "color": r.font.color.rgb if r.font and r.font.color else None,
                    "highlight": r.font.highlight_color if r.font else None
                })
            
            # 2. Remove only run elements, preserving bookmarks/footnotes/field codes
            p_elem = para._p
            for r_elem in list(p_elem.findall(qn('w:r'))):
                p_elem.remove(r_elem)

            # 3. Rebuild with Splitting
            for run_data in original_runs:
                text_chunk = run_data["text"]
                matches = list(regex.finditer(text_chunk))
                
                if not matches:
                    # No cross-ref in this run, restore as is
                    new_run = para.add_run(text_chunk)
                    # Restore props
                    _restore_run_props(new_run, run_data)
                else:
                    # Split this run's text
                    cursor = 0
                    for m in matches:
                        m_start, m_end = m.span()
                        
                        # Pre-match text (inherit origin props)
                        if m_start > cursor:
                            r = para.add_run(text_chunk[cursor:m_start])
                            _restore_run_props(r, run_data)
                        
                        # Match text (Apply XREF style, but keep boldness if it was bold?)
                        # Usually Xref style overrides, but let's keep bold/italic valid if implicit.
                        r_match = para.add_run(text_chunk[m_start:m_end])
                        try:
                            r_match.style = style_name
                        except: pass
                        # Optional: Force keep bold if original was bold? 
                        # Usually character style handles format. User asked to "retain".
                        # If original was bold, we should probably keep it bold even if Hyperlinked.
                        if run_data["bold"]: r_match.bold = True
                        if run_data["italic"]: r_match.italic = True
                        
                        cursor = m_end
                    
                    # Post-match text
                    if cursor < len(text_chunk):
                        r = para.add_run(text_chunk[cursor:])
                        _restore_run_props(r, run_data)

def _restore_run_props(run, data):
    # Helper to restore ALL formatting
    if data["style"]: 
        try: run.style = data["style"]
        except: pass
    
    # Boolean properties
    if data["bold"] is not None: run.bold = data["bold"]
    if data["italic"] is not None: run.italic = data["italic"]
    if data["underline"] is not None: run.underline = data["underline"]
    if data.get("superscript") is not None: run.font.superscript = data["superscript"]
    if data.get("subscript") is not None: run.font.subscript = data["subscript"]
    if data.get("strike") is not None: run.font.strike = data["strike"]
    # Double strike is rare but possible
    # if data.get("double_strike") is not None: run.font.double_strike = data["double_strike"]
    
    # Advanced Font properties
    if data["font"]: run.font.name = data["font"]
    if data["size"]: run.font.size = data["size"]
    if data["color"]: run.font.color.rgb = data["color"]
    if data.get("highlight"): run.font.highlight_color = data["highlight"]


def _snapshot_effective_run_formatting(para, doc):
    """Resolve each run's effective font size and name, walking the style chain so
    direct formatting overrides survive a paragraph-style swap."""
    snapshots = []
    for run in para.runs:
        size_pt = None
        if run.font.size is not None:
            size_pt = run.font.size.pt
        else:
            style = para.style
            while style is not None:
                if style.font.size is not None:
                    size_pt = style.font.size.pt
                    break
                style = style.base_style
            if size_pt is None:
                try:
                    normal = doc.styles['Normal']
                    if normal.font.size:
                        size_pt = normal.font.size.pt
                except Exception:
                    pass

        font_name = run.font.name
        if not font_name:
            style = para.style
            while style is not None:
                if style.font.name:
                    font_name = style.font.name
                    break
                style = style.base_style

        snapshots.append({'size_pt': size_pt, 'font_name': font_name})
    return snapshots


def _apply_formatting_snapshot(para, snapshots):
    """Re-stamp direct run-level font/size after a style change to lock in original appearance."""
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement
    for run, snap in zip(para.runs, snapshots):
        rPr = run._element.get_or_add_rPr()
        if snap['size_pt']:
            val = str(int(snap['size_pt'] * 2))
            for tag in ('w:sz', 'w:szCs'):
                el = rPr.find(_qn(tag))
                if el is None:
                    el = OxmlElement(tag)
                    rPr.append(el)
                el.set(_qn('w:val'), val)
        if snap['font_name']:
            rFonts = rPr.find(_qn('w:rFonts'))
            if rFonts is None:
                rFonts = OxmlElement('w:rFonts')
                rPr.insert(0, rFonts)
            rFonts.set(_qn('w:ascii'), snap['font_name'])
            rFonts.set(_qn('w:hAnsi'), snap['font_name'])


def _ensure_style_exists(doc, style_name: str) -> bool:
    """Create *style_name* as a Normal-based paragraph style if it is absent.
    Returns True when the style is ready to use, False on failure."""
    try:
        doc.styles[style_name]
        return True
    except KeyError:
        pass
    try:
        new_style = doc.styles.add_style(style_name, 1)  # 1 = WD_STYLE_TYPE.PARAGRAPH
        new_style.base_style = doc.styles['Normal']
        return True
    except Exception as e:
        logger.error(f"Could not create style '{style_name}': {e}")
        return False


def _normalize_client_tags_to_canonical(doc: Document, reverse_map: Dict[str, str]) -> None:
    """Rewrite a document's explicit tag markers and pre-existing paragraph
    styles from client-facing tag names back to canonical, in place, before
    annotation runs.

    Covers reprocessing a document this pipeline already styled with a
    client's tag names (para.style.name carries the client name), and
    hand-typed explicit markers (e.g. "<HEAD-1>") written using client tag
    names. No-op if reverse_map is empty. Best-effort: any single
    paragraph that can't be normalized is left as-is rather than aborting
    the whole pass.
    """
    if not reverse_map:
        return

    for para in doc.paragraphs:
        try:
            style_name = para.style.name if para.style else None
            if style_name and style_name in reverse_map:
                canonical = reverse_map[style_name]
                if _ensure_style_exists(doc, canonical):
                    para.style = canonical
        except Exception as e:
            logger.warning(f"Could not normalize paragraph style back to canonical: {e}")

        try:
            # Springer's "Box<N>-open"/"Box<N>-close" box markers are
            # recognized directly by annotate_document (see
            # recognize_springer_box_markers) without ever rewriting the
            # paragraph's text, so they're intentionally not handled here -
            # only other hand-typed explicit markers using client tag
            # names (e.g. "<HEAD-1>") still need rewriting back to
            # canonical before annotation.
            full_match, token, _ = parse_leading_style_hint(para.text)
            if not token:
                continue
            is_close = token.startswith("/")
            bare_token = token[1:] if is_close else token
            canonical_token = reverse_map.get(bare_token)
            if canonical_token is None:
                continue
            replacement_token = f"/{canonical_token}" if is_close else canonical_token
            new_marker = full_match.replace(token, replacement_token, 1)
            first_run = para.runs[0] if para.runs else None
            if first_run is not None and full_match in first_run.text:
                first_run.text = first_run.text.replace(full_match, new_marker, 1)
        except Exception as e:
            logger.warning(f"Could not normalize explicit tag marker back to canonical: {e}")


def process_docx(
    input_path: str,
    output_path: str,
    mode: Literal["style", "tag"] = "style",
    tag_set: Optional[str] = None,
    on_progress: Optional[Callable[[str, int], None]] = None,
) -> Dict[str, Any]:
    """
    Process a DOCX document by annotating and styling paragraphs and tables.

    Args:
        input_path: Path to input DOCX file
        output_path: Path to save output DOCX file
        mode: "style" to apply Word styles, "tag" to prefix text with [TAG]
        tag_set: optional client tag-set key (see tag_sets/*.yaml).
        on_progress: optional progress tracking callback: (step_name, percentage) -> None

    Returns:
        Dictionary with processing results:
        - success: bool
        - paragraphs_processed: int
        - tables_processed: int
        - errors: list of error messages
    """
    if mode not in ("style", "tag"):
        raise ValueError(f"Invalid mode: {mode}. Must be 'style' or 'tag'")
    
    result = {
        "success": False,
        "paragraphs_processed": 0,
        "tables_processed": 0,
        "errors": [],
        "zone_warnings": 0
    }

    try:
        if on_progress:
            on_progress("Opening document", 10)
        logger.info(f"Opening document: {input_path}")
        doc = Document(input_path)

        if tag_set:
            if on_progress:
                on_progress("Normalizing client tags to canonical", 20)
            logger.info(f"Normalizing client tag set '{tag_set}' back to canonical before annotation")
            _normalize_client_tags_to_canonical(doc, get_reverse_tag_map(tag_set))

        if on_progress:
            on_progress("Annotating document styles and structure", 30)
        logger.info(f"Annotating document with mode: {mode}")
        annotations = annotate_document(doc, recognize_springer_box_markers=bool(tag_set))
        annotations = apply_box_tag_prefixes(annotations)

        try:
            if on_progress:
                on_progress("Classifying heading levels via formatting", 45)
            logger.info("Classifying heading levels via formatting engine")
            annotations = classify_headings_by_formatting(doc, annotations)
        except Exception as e:
            logger.error(
                f"Heading classification by formatting failed; keeping regex-engine heading levels unchanged: {e}",
                exc_info=True,
            )

        if mode == "style":
            if on_progress:
                on_progress("Enforcing heading hierarchy and list validation", 60)
            logger.info("Enforcing heading hierarchy and validation")
            annotations = enforce_hierarchy(annotations)
            annotations = demote_long_headings(annotations)
            annotations = normalize_list_positions(annotations)
            annotations = normalize_reference_numbers(annotations)
            result["zone_warnings"] = check_zone_style_legality(annotations)

        tag_map = get_tag_map(tag_set)
        prefixes = get_rules_loader().get_structural_tags().get("prefixes", []) if tag_map else []

        if on_progress:
            on_progress("Applying styles and tags to paragraphs", 75)
        # Process annotations
        for idx, item in enumerate(annotations):
            try:
                para = item["para"]
                tag = item["tag"]
                style = item["style"]

                if tag_map:
                    list_case = item.get("list_case")
                    tag = translate_tag(tag, tag_map, prefixes, case=list_case)
                    style = translate_tag(style, tag_map, prefixes, case=list_case)

                if mode == "style":
                    try:
                        _ensure_style_exists(doc, style)
                        snapshots = _snapshot_effective_run_formatting(para, doc)
                        para.style = style
                        _apply_formatting_snapshot(para, snapshots)
                        result["paragraphs_processed"] += 1
                        logger.debug(f"Paragraph {idx}: Applied style '{style}' (tag: {tag})")
                    except Exception as style_err:
                        logger.warning(
                            "Paragraph %d: could not apply style '%s': %s",
                            idx, style, style_err,
                        )
                        result["errors"].append(f"Paragraph {idx} style error: {style_err}")
                else:  # tag mode
                    if para.runs:
                        try:
                            para.runs[0].text = f"[{tag}] " + para.runs[0].text
                            result["paragraphs_processed"] += 1
                            logger.debug(f"Paragraph {idx}: Added tag '{tag}'")
                        except Exception as e:
                            logger.error(f"Paragraph {idx}: Failed to add tag: {e}")
                            result["errors"].append(f"Paragraph {idx} tag error: {str(e)}")
                    else:
                        logger.warning(f"Paragraph {idx}: No runs found (empty paragraph)")
            
            except KeyError as e:
                logger.error(f"Annotation {idx} missing key: {e}")
                result["errors"].append(f"Invalid annotation format: {str(e)}")
            except Exception as e:
                logger.error(f"Error processing paragraph {idx}: {e}")
                result["errors"].append(f"Paragraph {idx} error: {str(e)}")
        
        # Process tables
        try:
            if on_progress:
                on_progress("Processing tables", 85)
            logger.info(f"Processing tables (mode: {mode})")
            tag_tables(doc, mode, tag_map=tag_map)
            result["tables_processed"] = len(doc.tables)
        except Exception as e:
            logger.error(f"Error processing tables: {e}")
            result["errors"].append(f"Table processing error: {str(e)}")

        # Process cross-references
        if mode == "style":
            try:
                if on_progress:
                    on_progress("Processing cross-references", 90)
                logger.info("Processing cross-references")
                process_cross_references(doc)
            except Exception as e:
                logger.error(f"Error processing cross-references: {e}")
                result["errors"].append(f"Xref error: {str(e)}")
        
        # Save document
        if on_progress:
            on_progress("Saving document", 95)
        logger.info(f"Saving document to: {output_path}")
        doc.save(output_path)
        
        result["success"] = True
        logger.info(f"Document processed successfully. Paragraphs: {result['paragraphs_processed']}, Tables: {result['tables_processed']}")
        
    except FileNotFoundError:
        msg = f"Input file not found: {input_path}"
        logger.error(msg)
        result["errors"].append(msg)
    except Exception as e:
        msg = f"Unexpected error processing document: {e}"
        logger.error(msg, exc_info=True)
        result["errors"].append(msg)
    
    return result


if __name__ == "__main__":
    # Example usage
    result = process_docx(
        input_path="input.docx",
        output_path="processed_input.docx",
        mode="style"
    )
    print(f"Processing result: {result}")

