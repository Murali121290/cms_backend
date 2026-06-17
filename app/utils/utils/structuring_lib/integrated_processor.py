"""

Integration module connecting annotation, styling, and enhanced processing.

Combines all components into a cohesive pipeline for document processing.

"""



import logging
import re

from typing import List, Dict, Any, Optional, Tuple

from dataclasses import dataclass

from docx.document import Document

from docx.oxml.ns import qn



from .logger_config import get_logger

from .rules_loader import get_rules_loader

from .annotator import (
    classify_explicit_context,
    detect_list_kind,
    is_explicit_context_closer,
    is_list_paragraph,
    normalize_style_token,
    parse_leading_style_hint,
)

from .enhanced_processor import DocumentProcessor, BlockState



logger = get_logger(__name__)





@dataclass

class ProcessingResult:

    """Result from document processing"""

    success: bool

    paragraphs_processed: int

    styles_applied: int

    warnings: List[str]

    errors: List[str]

    state_transitions: List[Tuple[BlockState, BlockState]]

    cross_references_found: int

    validation_issues: List[str]


@dataclass
class ExplicitStyleContext:

    """Bounded explicit-tag inheritance context."""

    kind: str

    first_body_pending: bool = False


class IntegratedDocumentProcessor:

    """

    Unified processor combining annotation, styling, and enhancement.

    Orchestrates the document processing pipeline.

    """

    

    def __init__(self, rules_path: Optional[str] = None):

        """Initialize integrated processor"""

        self.logger = logger

        self.rules_loader = get_rules_loader(rules_path)

        self.doc_processor = DocumentProcessor()

        self.result: Optional[ProcessingResult] = None

    

    def process_document_integrated(self, document: Document) -> ProcessingResult:

        """

        Process document through integrated pipeline.

        

        Pipeline steps:

        1. Validate document structure

        2. Validate template styles

        3. Process paragraphs with state machine

        4. Detect tables and headers

        5. Preserve formatting when applying styles

        6. Validate heading hierarchy

        7. Find cross-references

        8. Handle empty paragraphs

        

        Args:

            document: python-docx Document object

            

        Returns:

            ProcessingResult with detailed processing information

        """

        self.doc = document
        self._obj_item_count = 0
        self.logger.info(f"Starting integrated processing for document with {len(document.paragraphs)} paragraphs")

        

        errors: List[str] = []

        warnings: List[str] = []

        styles_applied = 0

        state_transitions: List[Tuple[BlockState, BlockState]] = []

        

        try:

            # Step 1: Validate template styles

            self.logger.debug("Validating template styles")

            style_warnings = self.doc_processor.validate_styles_in_template(document)

            warnings.extend(style_warnings)

            self.logger.debug(f"Template validation found {len(style_warnings)} warnings")

            

            # Step 2: Process each paragraph

            self.logger.debug(f"Processing {len(document.paragraphs)} paragraphs")
            previous_style: Optional[str] = None
            explicit_context: Optional[ExplicitStyleContext] = None

            for para_idx, para in enumerate(document.paragraphs):

                try:

                    para_text = para.text or ""
                    full_hint, explicit_token, normalized_text = parse_leading_style_hint(para_text)
                    existing_style_raw = self._get_existing_style_name(para)
                    existing_style = self._canonicalize_source_style(
                        existing_style_raw,
                        normalized_text,
                        explicit_context.kind if explicit_context else None,
                    )

                    # Update block state

                    old_state = self.doc_processor.state

                    self.doc_processor.update_block_state(normalized_text)

                    new_state = self.doc_processor.state

                    

                    if old_state != new_state:

                        state_transitions.append((old_state, new_state))

                        self.logger.debug(f"State transition: {old_state} -> {new_state} at paragraph {para_idx}")

                    

                    # Handle empty paragraphs

                    if not normalized_text.strip() and not explicit_token:

                        handling = self.doc_processor.handle_empty_paragraph(para)

                        if handling:

                            self.logger.debug(f"Handled empty paragraph at index {para_idx}: {handling}")

                            continue

                    

                    # Normalize explicit tags/source styles first, then preserve
                    # existing structured styles before falling back to rules.yaml matching.
                    style = None
                    inferred_style = self._infer_style_for_paragraph(
                        para,
                        normalized_text,
                        previous_style,
                        explicit_context,
                        existing_style_raw,
                        para_idx,
                        document,
                    )
                    explicit_style = self._canonicalize_explicit_style(
                        explicit_token,
                        normalized_text,
                        explicit_context,
                    )
                    # Override check: if text is "OBJECTIVES", always use OBJ1
                    objectives_synonyms = ["objectives", "learning objectives", "learningobjectives", "lesson objectives"]
                    if normalized_text.strip().lower() in objectives_synonyms:
                        style = "OBJ1"
                        self.logger.debug(f"Para {para_idx}: Detected learning objectives heading '{normalized_text}', setting to OBJ1")
                    elif explicit_token:
                        if (
                            inferred_style
                            and explicit_style
                            and self._is_overrideable_base_style(explicit_style)
                            and not self._is_overrideable_base_style(inferred_style)
                        ):
                            style = inferred_style
                        else:
                            style = (
                                existing_style
                                if self._should_preserve_existing_style(existing_style, explicit_style)
                                else explicit_style
                            )
                        # Keep tags in the document - don't strip them
                        # if full_hint:
                        #     self._strip_leading_hint_from_runs(para, full_hint)
                    elif self._should_preserve_existing_style(existing_style, inferred_style):
                        style = existing_style
                    elif explicit_context is not None:
                        style = self._get_explicit_context_style(
                            explicit_context,
                            normalized_text,
                            inferred_style,
                            existing_style_raw,
                        )
                    else:
                        style = inferred_style

                    if style:
                        style = self._canonicalize_source_style(
                            style,
                            normalized_text,
                            explicit_context.kind if explicit_context else None,
                        )
                        # Rule 4: First Paragraph Rule
                        if self.doc_processor.is_first_para_after_heading and style == "TXT":
                            style = "TXT-FLUSH"
                        
                        # Set heading flag for next para
                        # Also track CT, CN, OBJ1, etc. as headings that cause the next para to be TXT-FLUSH
                        is_heading = False
                        if style.startswith("H") and len(style) <= 2:
                            is_heading = True
                        elif style in ("CN", "CT", "OBJ1", "REFH1", "ABS"):
                            is_heading = True
                        elif "TTL" in style or "-H" in style:
                            is_heading = True
                            
                        if is_heading:
                            self.doc_processor.is_first_para_after_heading = True
                        elif style not in ("PMI", "EMPTY"):
                            self.doc_processor.is_first_para_after_heading = False

                        # Ensure style exists
                        try:
                            _ = document.styles[style]
                        except KeyError:
                            from docx.enum.style import WD_STYLE_TYPE
                            self.logger.warning(f"Style '{style}' not found. Creating it dynamically.")
                            document.styles.add_style(style, WD_STYLE_TYPE.PARAGRAPH)

                        # Preserve formatting while applying
                        # Preserve formatting while applying. 
                        # We use clear_formatting and set the paragraph style.
                        # Setting para.text wiped out runs, we fixed it via regex, but para.text wipe still happened.
                        # Wait, we changed para.text earlier: para.text = local_re.sub(...). That destroys runs!
                        # We should strip the tag from the runs instead of para.text.

                        # Preserve formatting while applying. 
                        # We use clear_formatting and set the paragraph style.
                        # Setting para.text wiped out runs, we fixed it via regex, but para.text wipe still happened.
                        # Wait, we changed para.text earlier: para.text = local_re.sub(...). That destroys runs!
                        # We should strip the tag from the runs instead of para.text.

                        if para.runs:
                            para.style = style
                            styles_applied += 1
                        else:
                            para.style = style
                            styles_applied += 1
                        if style != "EMPTY":
                            previous_style = style
                        explicit_context = self._advance_explicit_context(
                            explicit_context,
                            explicit_token,
                            style,
                            normalized_text,
                        )
                        self.logger.debug(f"Applied style '{style}' to paragraph {para_idx}")
                    else:
                        previous_style = existing_style or previous_style
                        explicit_context = self._advance_explicit_context(
                            explicit_context,
                            explicit_token,
                            style,
                            normalized_text,
                        )

                    

                    # Find cross-references

                    refs = {}

                    if refs:

                        for ref_type, numbers in refs.items():

                            self.logger.debug(f"Found {ref_type} references: {numbers}")

                

                except Exception as e:

                    error_msg = f"Error processing paragraph {para_idx}: {str(e)}"

                    self.logger.error(error_msg)

                    errors.append(error_msg)

            

            # Step 3: Process tables

            self.logger.debug(f"Processing {len(document.tables)} tables")

            for table_idx, table in enumerate(document.tables):

                try:

                    for row_idx, row in enumerate(table.rows):

                        for cell_idx, cell in enumerate(row.cells):

                            for para in cell.paragraphs:

                                cell_text = para.text.strip()

                                if not cell_text:
                                    continue

                                t_style = self._get_table_style_for_paragraph(
                                    para,
                                    cell_text,
                                    row_idx,
                                    cell_idx,
                                    len(table.rows),
                                )

                                try:
                                    try:
                                        _ = document.styles[t_style]
                                    except KeyError:
                                        from docx.enum.style import WD_STYLE_TYPE
                                        document.styles.add_style(t_style, WD_STYLE_TYPE.PARAGRAPH)

                                    para.style = t_style
                                    styles_applied += 1
                                except Exception as e:
                                    self.logger.debug(f"Could not apply table style: {e}")

                

                except Exception as e:

                    error_msg = f"Error processing table {table_idx}: {str(e)}"

                    self.logger.warning(error_msg)

                    warnings.append(error_msg)

            

            # Step 4: Validate heading hierarchy

            self.logger.debug("Validating heading hierarchy")

            hierarchy_issues = self.doc_processor.validation_warnings

            if hierarchy_issues:

                warnings.extend(hierarchy_issues)

                self.logger.debug(f"Heading hierarchy validation found {len(hierarchy_issues)} issues")

            

            # Compile result

            self.result = ProcessingResult(

                success=len(errors) == 0,

                paragraphs_processed=len(document.paragraphs),

                styles_applied=styles_applied,

                warnings=warnings,

                errors=errors,

                state_transitions=state_transitions,

                cross_references_found=sum(len(v) for v in self.doc_processor.cross_references.values()),

                validation_issues=self.doc_processor.validation_warnings

            )

            

            self.logger.info(

                f"Processing complete: {styles_applied} styles applied, "

                f"{len(errors)} errors, {len(warnings)} warnings"

            )

            

            return self.result

        

        except Exception as e:

            self.logger.error(f"Critical error during processing: {str(e)}")

            return ProcessingResult(

                success=False,

                paragraphs_processed=0,

                styles_applied=styles_applied,

                warnings=warnings,

                errors=[f"Critical error: {str(e)}"] + errors,

                state_transitions=state_transitions,

                cross_references_found=0,

                validation_issues=[]

            )

    

    def _get_style_for_paragraph(self, text: str) -> Optional[str]:

        """

        Determine appropriate style for paragraph text.

        

        Checks rules in priority order and returns first matching style.

        

        Args:

            text: Paragraph text to analyze

            

        Returns:

            Style name if match found, None otherwise

        """

        try:

            paragraphs = self.rules_loader.get_paragraphs()

            

            # Lower numeric priority runs first, as declared in rules.yaml.
            for para_rule in sorted(paragraphs, key=lambda x: x.get("priority", 999)):

                pattern = para_rule.get("pattern")

                style = para_rule.get("style")

                

                if pattern and style:

                    try:

                        if __import__('re').match(pattern, text):

                            return style

                    except Exception as e:

                        self.logger.debug(f"Pattern matching error: {e}")

            defaults = self.rules_loader.rules.get("defaults", {})
            return defaults.get("unmatched_text", {}).get("style")

        except Exception as e:

            self.logger.error(f"Error determining style: {e}")

            return None

    def _infer_style_for_paragraph(
        self,
        para,
        text: str,
        previous_style: Optional[str],
        explicit_context: Optional[ExplicitStyleContext] = None,
        existing_style_raw: Optional[str] = None,
        para_idx: Optional[int] = None,
        document: Optional[Any] = None,
    ) -> Optional[str]:

        """Infer the style from rules.yaml and paragraph context."""

        if previous_style == "CN":
            return "CT"

        style = self._get_list_style_for_paragraph(para, text, previous_style, explicit_context, para_idx, document)
        if style is not None:
            if style == "OBJ1":
                self._obj_item_count = 0
            return style

        style = self._get_style_for_paragraph(text)
        if style == "OBJ1":
            self._obj_item_count = 0
            
        if (
            explicit_context is not None
            and style in {"TXT", "TXT-FLUSH", "H2", "H3", "H4"}
        ):
            contextual_style = self._get_explicit_context_style(
                explicit_context,
                text,
                style,
                existing_style_raw,
            )
            if contextual_style:
                return contextual_style
        if self.doc_processor.state == BlockState.REFERENCES and style != "REFH1":
            return self._get_reference_style_for_text(text, None, para)
            
        if self.doc_processor.state == BlockState.LEARNING_OBJECTIVES and style in {"TXT", "TXT-FLUSH"}:
            # Keep count for OBJ-TXT-FIRST logic
            if not hasattr(self, '_obj_item_count'):
                self._obj_item_count = 0
            self._obj_item_count += 1
            if self._obj_item_count == 1:
                return "OBJ-TXT-FIRST"
            return "OBJ-TXT"
            
        return style

    def _get_existing_style_name(self, para) -> Optional[str]:

        """Return the paragraph's current style name, if available."""

        try:
            return para.style.name if para.style is not None else None
        except Exception:
            return None

    def _canonicalize_source_style(
        self,
        style_name: Optional[str],
        text: str = "",
        context_kind: Optional[str] = None,
    ) -> Optional[str]:

        """Normalize incoming source styles to canonical output styles."""

        if not style_name:
            return None

        style_name = style_name.strip()
        mapped = normalize_style_token(style_name, context_kind)
        if mapped and mapped != style_name:
            return mapped

        if context_kind == "box" and style_name == "TITLE":
            return self.rules_loader.get_box_config().get("title_style", "NBX1-TTL")

        return style_name

    def _canonicalize_explicit_style(
        self,
        explicit_token: Optional[str],
        text: str,
        explicit_context: Optional[ExplicitStyleContext],
    ) -> Optional[str]:

        """Normalize explicit tag/style tokens to canonical output styles."""

        if not explicit_token:
            return None

        context_kind = explicit_context.kind if explicit_context else None
        explicit_style = normalize_style_token(explicit_token, context_kind)

        if explicit_style == "TITLE" and context_kind == "box":
            return self.rules_loader.get_box_config().get("title_style", "NBX1-TTL")

        return explicit_style or explicit_token

    def _normalize_paragraph_text_for_matching(self, text: str) -> str:

        """Strip leading explicit style markers before state/rule matching."""

        _, _, stripped_text = parse_leading_style_hint(text or "")
        return stripped_text

    def _get_reference_style_for_text(self, text: str, list_kind: Optional[str] = None, para: Optional[Any] = None) -> str:

        """Classify a reference paragraph as numbered or unnumbered author-year."""

        text = (text or "").strip()
        
        # 1. Use the explicit numbered list detection logic
        is_numbered = bool(re.match(r'^\[?\d+\]?[\.\)\t\s]', text))
        if not is_numbered and para is not None:
            try:
                from docx.oxml.ns import qn
                pPr = para._element.find(qn('w:pPr'))
                if pPr is not None and pPr.find(qn('w:numPr')) is not None:
                    is_numbered = True
            except Exception:
                pass
                
        if is_numbered or list_kind == "number":
            return "REF-N"

        author_year_patterns = [
            r"^[A-Z][A-Za-z'`\-]+(?:,\s*(?:[A-Z]\.\s*)+).*\(\d{4}[a-z]?\)",
            r"^[A-Z][A-Za-z'`\-]+(?:\s+et\s+al\.)?.*\(\d{4}[a-z]?\)",
            r"^[A-Z][A-Za-z'`\-]+(?:,\s*[A-Z][A-Za-z'`\-]+)*(?:\s*&\s*[A-Z][A-Za-z'`\-]+)?.*\b\d{4}[a-z]?\b",
        ]
        if any(re.match(pattern, text) for pattern in author_year_patterns):
            return "REF-U"

        return "REF-U"

    def _strip_leading_hint_from_runs(self, para, hint_text: str) -> None:

        """Remove a leading explicit tag without flattening the paragraph runs."""

        chars_removed = 0
        hint_len = len(hint_text)
        for run in para.runs:
            if chars_removed >= hint_len:
                break
            if not run.text:
                continue
            remaining = hint_len - chars_removed
            if len(run.text) <= remaining:
                chars_removed += len(run.text)
                run.text = ""
            else:
                run.text = run.text[remaining:]
                chars_removed = hint_len

    def _get_explicit_context_style(
        self,
        explicit_context: Optional[ExplicitStyleContext],
        text: str,
        inferred_style: Optional[str],
        existing_style_raw: Optional[str],
    ) -> Optional[str]:

        """Resolve bounded explicit-style inheritance for boxes, objectives, and key terms."""

        if explicit_context is None or not text:
            return inferred_style

        box_cfg = self.rules_loader.get_box_config()
        existing_style = self._canonicalize_source_style(
            existing_style_raw,
            text,
            explicit_context.kind,
        )

        if explicit_context.kind == "objective":
            if inferred_style in {None, "TXT", "TXT-FLUSH"}:
                return self.rules_loader.get_list_patterns().get(
                    "learning_objectives_body",
                    {},
                ).get("style", "OBJ-TXT")
            return inferred_style

        if explicit_context.kind == "keyterm":
            if inferred_style in {None, "TXT", "TXT-FLUSH"}:
                return "KT"
            return inferred_style

        if explicit_context.kind == "box":
            if existing_style == box_cfg.get("title_style", "NBX1-TTL") or re.match(r"^Box\s+\d+", text):
                return box_cfg.get("title_style", "NBX1-TTL")

            if inferred_style in {None, "TXT", "TXT-FLUSH", "H2", "H3", "H4"}:
                if explicit_context.first_body_pending:
                    return box_cfg.get("first_body_style", "NBX-TXT-FIRST")
                return box_cfg.get("body_style", "NBX-TXT")
            return inferred_style

        return inferred_style

    def _advance_explicit_context(
        self,
        explicit_context: Optional[ExplicitStyleContext],
        explicit_token: Optional[str],
        resolved_style: Optional[str],
        text: str,
    ) -> Optional[ExplicitStyleContext]:

        """Advance, start, or clear bounded explicit inheritance contexts."""

        if explicit_token and is_explicit_context_closer(explicit_token, explicit_context.kind if explicit_context else None):
            return None

        if explicit_token:
            new_kind = classify_explicit_context(explicit_token)
            if new_kind == "box":
                return ExplicitStyleContext(kind="box", first_body_pending=True)
            if new_kind == "objective":
                return ExplicitStyleContext(kind="objective", first_body_pending=False)
            if new_kind == "keyterm":
                return ExplicitStyleContext(kind="keyterm", first_body_pending=False)

        if explicit_context is None:
            return None

        if explicit_context.kind == "box":
            box_cfg = self.rules_loader.get_box_config()
            if resolved_style == box_cfg.get("title_style", "NBX1-TTL"):
                return explicit_context
            if resolved_style in {
                box_cfg.get("body_style", "NBX-TXT"),
                box_cfg.get("first_body_style", "NBX-TXT-FIRST"),
            }:
                explicit_context.first_body_pending = False

        if resolved_style in {"H1", "H2", "H3", "H4", "REFH1", "CN", "CT"}:
            if not (explicit_context.kind == "objective" and resolved_style == "OBJ1"):
                return None

        if not text:
            return explicit_context

        return explicit_context

    def _should_preserve_existing_style(
        self,
        style_name: Optional[str],
        inferred_style: Optional[str] = None,
    ) -> bool:

        """Preserve existing semantic styles so the processor is safe to rerun."""

        if not style_name:
            return False

        if self._is_generic_style(style_name):
            return False

        if (
            inferred_style
            and inferred_style != style_name
            and self._is_overrideable_base_style(style_name)
            and not self._is_overrideable_base_style(inferred_style)
        ):
            return False

        return self._is_semantic_style(style_name)

    def _is_generic_style(self, style_name: str) -> bool:

        """Identify source/template styles that should not block reclassification."""

        generic_styles = {
            "Normal",
            "No Spacing",
            "Body Text",
            "List Paragraph",
            "Quote",
            "Intense Quote",
            "Subtitle",
            "Title",
            "Text",
            "TextHelvetica",
        }
        if style_name in generic_styles:
            return True
        if style_name in {f"Heading {idx}" for idx in range(1, 10)}:
            return True
        lowered = style_name.lower()
        return (
            lowered.startswith("text")
            or lowered.startswith("body")
            or lowered.startswith("heading ")
        )

    def _is_overrideable_base_style(self, style_name: str) -> bool:

        """Base heading/body styles may be replaced by more specific matches."""

        return style_name in {
            "CN",
            "CT",
            "H1",
            "H2",
            "H3",
            "H4",
            "TXT",
            "TXT-FLUSH",
            "Intro",
        }

    def _is_semantic_style(self, style_name: str) -> bool:

        """Recognize semantic processor styles that should survive reruns."""

        configured_styles = {
            rule.get("style")
            for rule in self.rules_loader.get_paragraphs()
            if rule.get("style")
        }
        configured_styles.update(
            cfg.get("style")
            for cfg in self.rules_loader.get_list_patterns().values()
            if isinstance(cfg, dict) and cfg.get("style")
        )

        if style_name in configured_styles:
            return True

        if re.match(r"^[A-Z][A-Z0-9]*(?:-[A-Z0-9]+)+$", style_name):
            return True

        structured_prefixes = (
            "OBJ",
            "REF",
            "BX",
            "NBX",
            "FIG",
            "TB",
            "TBL",
            "TNL",
            "BL",
            "NL",
            "PMI",
            "TXT",
            "ABS",
            "EPI",
            "BIB",
            "KT",
            "INTRO",
            "ref-",
        )
        return style_name.startswith(structured_prefixes)

    def _get_list_style_for_paragraph(
        self,
        para,
        text: str,
        previous_style: Optional[str],
        explicit_context: Optional[ExplicitStyleContext] = None,
        para_idx: Optional[int] = None,
        document: Optional[Any] = None,
    ) -> Optional[str]:

        """Map list-like paragraphs using the list section of rules.yaml."""

        list_config = self.rules_loader.get_list_patterns()
        style_name = self._get_existing_style_name(para) or ""

        list_kind = detect_list_kind(text, style_name, is_list_paragraph(para), para, self.doc)

        if list_kind is None:
            self.doc_processor.list_tracker.evaluate(False, "")
            return None

        # Determine if this is the first or last item in a continuous list
        is_first_in_list = previous_style not in ("NL-MID", "NL-FIRST", "NL-LAST", "BL-MID", "BL-FIRST", "BL-LAST", "RL-MID", "RL-FIRST", "RL-LAST",
                                                  "OBJ-NL-MID", "OBJ-NL-FIRST", "OBJ-NL-LAST", "OBJ-BL-MID")
        is_last_in_list = False

        # Look ahead to see if next paragraph is also a list
        if document is not None and para_idx is not None:
            if para_idx + 1 < len(document.paragraphs):
                next_para = document.paragraphs[para_idx + 1]
                next_text = next_para.text.strip()
                next_style_name = self._get_existing_style_name(next_para) or ""
                next_list_kind = detect_list_kind(next_text, next_style_name, is_list_paragraph(next_para), next_para, self.doc)
                # If next paragraph is NOT a list, this is the last item
                if next_list_kind is None:
                    is_last_in_list = True
            else:
                # Last paragraph in document
                is_last_in_list = True

        objective_context_styles = {
            "OBJ1",
            "OBJ-BL-MID",
            "OBJ-NL-FIRST",
            "OBJ-NL-MID",
            "OBJ-NL-LAST",
            "OBJ-TXT",
        }
        if (
            self.doc_processor.state == BlockState.LEARNING_OBJECTIVES
            or previous_style in objective_context_styles
            or (explicit_context is not None and explicit_context.kind == "objective")
        ):
            self.doc_processor.list_tracker.evaluate(True, "objective")
            if list_kind in {"number", "roman"}:
                if is_first_in_list:
                    return "OBJ-NL-FIRST"
                elif is_last_in_list:
                    return "OBJ-NL-LAST"
                return "OBJ-NL-MID"
            if is_first_in_list:
                return list_config.get("learning_objectives", {}).get("first_style", "OBJ-BL-FIRST")
            elif is_last_in_list:
                return list_config.get("learning_objectives", {}).get("last_style", "OBJ-BL-LAST")
            return list_config.get("learning_objectives", {}).get("style", "OBJ-BL-MID")

        if self.doc_processor.state == BlockState.REFERENCES:
            self.doc_processor.list_tracker.evaluate(True, "references")
            reference_style = self._get_reference_style_for_text(text, list_kind, para)
            if reference_style == "REF-N":
                return list_config.get("references_numbered", {}).get("style", "REF-N")
            return list_config.get("references_unnumbered", {}).get("style", "REF-U")

        self.doc_processor.list_tracker.evaluate(True, list_kind)
        if list_kind == "number":
            if is_first_in_list:
                return list_config.get("general_numbered", {}).get("first_style", "NL-FIRST")
            elif is_last_in_list:
                return list_config.get("general_numbered", {}).get("last_style", "NL-LAST")
            return list_config.get("general_numbered", {}).get("style", "NL-MID")
        if list_kind == "roman":
            if is_first_in_list:
                return list_config.get("general_roman", {}).get("first_style", "RL-FIRST")
            elif is_last_in_list:
                return list_config.get("general_roman", {}).get("last_style", "RL-LAST")
            return list_config.get("general_roman", {}).get("style", "RL-MID")
        # Bullet list
        if is_first_in_list:
            return list_config.get("general_bulleted", {}).get("first_style", "BL-FIRST")
        elif is_last_in_list:
            return list_config.get("general_bulleted", {}).get("last_style", "BL-LAST")
        return list_config.get("general_bulleted", {}).get("style", "BL-MID")

    def _get_table_style_for_paragraph(
        self,
        para,
        text: str,
        row_idx: int,
        cell_idx: int,
        total_rows: int,
    ) -> str:

        """Resolve table paragraph style from YAML-driven list and header rules."""

        table_config = self.rules_loader.get_table_config()
        style_name = self._get_existing_style_name(para) or ""
        list_kind = detect_list_kind(text, style_name, is_list_paragraph(para), para, self.doc)

        if list_kind == "bullet":
            return table_config.get("bullet_style", "TBL-MID")
        if list_kind == "number":
            return table_config.get("number_style", "TNL-MID")
        if list_kind == "roman":
            return table_config.get("roman_style", "TRL-MID")

        header_threshold = table_config.get("header_threshold", 1.0)
        score = self.doc_processor.detect_table_header_smart(text, row_idx, cell_idx, total_rows)
        if score >= header_threshold:
            return table_config.get("header_style", "TBCH")
        return table_config.get("body_style", "TB")

    def process_document(self, document: Document) -> ProcessingResult:

        """Backward-compatible wrapper used by older app_server call sites."""

        return self.process_document_integrated(document)

    

    def get_processing_summary(self) -> Dict[str, Any]:

        """

        Get summary of last processing result.

        

        Returns:

            Dictionary with processing statistics and details

        """

        if not self.result:

            return {"status": "No processing completed yet"}

        

        return {

            "status": "SUCCESS" if self.result.success else "FAILED",

            "paragraphs_processed": self.result.paragraphs_processed,

            "styles_applied": self.result.styles_applied,

            "cross_references_found": self.result.cross_references_found,

            "state_transitions_count": len(self.result.state_transitions),

            "warnings_count": len(self.result.warnings),

            "errors_count": len(self.result.errors),

            "warnings": self.result.warnings[:5],  # First 5

            "errors": self.result.errors[:5],  # First 5

        }





# Global instance

_integrated_processor: Optional[IntegratedDocumentProcessor] = None





def get_integrated_processor() -> IntegratedDocumentProcessor:

    """Get or create global integrated processor instance."""

    global _integrated_processor

    if _integrated_processor is None:

        _integrated_processor = IntegratedDocumentProcessor()

    return _integrated_processor





def reset_processor():

    """Reset global processor instance (useful for testing)."""

    global _integrated_processor

    _integrated_processor = None












