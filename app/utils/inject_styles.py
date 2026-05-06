"""
inject_styles.py
================
Injects publisher-defined paragraph styles into a DOCX file so they appear
in Collabora Online's Style dropdown.

Usage (standalone):
    python inject_styles.py input.docx output.docx

Usage (in Celery task):
    from inject_styles import inject_publisher_styles
    inject_publisher_styles("/path/to/uploaded.docx")   # modifies in-place
"""

import sys
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

# ---------------------------------------------------------------------------
# Default formatting applied to every publisher style
# ---------------------------------------------------------------------------
STYLE_FONT_NAME = "Times New Roman"
STYLE_FONT_SIZE = Pt(12)


def _apply_formatting(style):
    """Apply Times New Roman 12pt 2.0 line spacing to a style."""
    style.font.name             = STYLE_FONT_NAME
    style.font.size             = STYLE_FONT_SIZE
    pf = style.paragraph_format
    pf.line_spacing             = 2.0
    pf.line_spacing_rule        = WD_LINE_SPACING.MULTIPLE


# ---------------------------------------------------------------------------
# Complete list of publisher style tags
# ---------------------------------------------------------------------------
PUBLISHER_STYLES = [
    "ACK1", "ACKTXT", "ANS-NL-FIRST", "ANS-NL-MID", "APX", "APX-TXT",
    "APX-TXT-FLUSH", "APXAU", "APXH1", "APXH2", "APXH3", "APXN", "APXST",
    "APXT", "BIB", "BIBH1", "BIBH2", "BL-FIRST", "BL-LAST", "BL-MID",
    "BL2-MID", "BL3-MID", "BL4-MID", "BL5-MID", "BL6-MID",
    "BX1-BL-FIRST", "BX1-BL-LAST", "BX1-BL-MID", "BX1-BL2-MID",
    "BX1-EXT-ONLY", "BX1-EQ-FIRST", "BX1-EQ-MID", "BX1-EQ-LAST",
    "BX1-EQ-ONLY", "BX1-FN", "BX1-H1", "BX1-H2", "BX1-H3", "BX1-L1",
    "BX1-MCUL-FIRST", "BX1-MCUL-LAST", "BX1-MCUL-MID", "BX1-NL-FIRST",
    "BX1-NL-LAST", "BX1-NL-MID", "BX1-OUT1-FIRST", "BX1-OUT1-MID",
    "BX1-OUT2", "BX1-OUT2-LAST", "BX1-OUT3", "BX1-QUO", "BX1-QUO-AU",
    "BX1-TTL", "BX1-TXT", "BX1-TXT-DC", "BX1-TXT-FIRST", "BX1-TYPE",
    "BX1-UL-FIRST", "BX1-UL-LAST", "BX1-UL-MID",
    "CAU", "CHAP", "CN", "COQ", "COQA", "COUT-1", "COUT-2", "COUTH1",
    "COUT-NL-FIRST", "COUT-NL-MID", "CPAU", "CPT", "CST", "CT",
    "DF", "DIA-FIRST", "DIA-LAST", "DIA-MID",
    "EQ-FIRST", "EQ-MID", "EQ-LAST", "EQ-ONLY",
    "EQN-FIRST", "EQN-MID", "EQN-LAST", "EQN-ONLY",
    "EXER-AB-NL-FIRST", "EXER-AB-NL-MID", "EXER-AT-NL-FIRST",
    "EXER-AT-NL-MID", "EXER-AT-T2", "EXER-CS-AU", "EXER-CS-NL-FIRST",
    "EXER-CS-NL-MID", "EXER-CS-T", "EXER-CS-T2", "EXER-DIR",
    "EXER-DT-NL-FIRST", "EXER-DT-NL-MID", "EXER-FB-NL-FIRST",
    "EXER-FB-NL-MID", "EXER-H1", "EXER-L-UL", "EXER-M-NL-FIRST",
    "EXER-M-NL-MID", "EXER-MC-NL-FIRST", "EXER-MC-NL-MID",
    "EXER-MC-NL2-FIRST", "EXER-MC-NL2-MID", "EXER-SA-NL-FIRST",
    "EXER-SA-NL-MID", "EXER-SP-NL-FIRST", "EXER-SP-NL-MID",
    "EXER-SP-NL2-FIRST", "EXER-SP-NL2-MID", "EXER-TF-NL-FIRST",
    "EXER-TF-NL-MID", "EXER-WB-NL-FIRST", "EXER-WB-NL-MID",
    "EXER-WP-NL-FIRST", "EXER-WP-NL-MID", "EXER-TTL", "EXER-WP-L",
    "EX-NL-FIRST", "EX-NL-LAST", "EX-NL-MID", "EX-H1",
    "EXT-FIRST", "EXT-LAST", "EXT-MID", "EXT-ONLY",
    "FIG-CRED", "FIG-LEG", "FN", "FN-BL-FIRST", "FN-BL-MID",
    "FN-BL-LAST", "FN-LAST", "FOLIO-RECTO", "FOLIO-VERSO",
    "H1", "H1-BL", "H2", "H2 after H1", "H3", "H3 after H2",
    "H4", "H5", "H6", "INTRO",
    "KP1", "KP-BL-FIRST", "KP-BL-LAST", "KP-BL-MID",
    "KP-NL-FIRST", "KP-NL-LAST", "KP-NL-MID",
    "KT-BL-FIRST", "KT-BL-LAST", "KT-BL-MID",
    "KT-NL-FIRST", "KT-NL-LAST", "KT-NL-MID",
    "KT-TXT", "KT-UL-FIRST", "KT-UL-LAST", "KT-UL-MID", "KT1",
    "L1", "L2",
    "MCUL-FIRST", "MCUL-LAST", "MCUL-MID",
    "NBX-BL-FIRST", "NBX-BL-LAST", "NBX-BL-MID", "NBX-BL2-MID",
    "NBX-TXT-DC", "NBX-EQ-FIRST", "NBX-EQ-MID", "NBX-EQ-LAST",
    "NBX-EQ-ONLY", "NBX-EXT-ONLY", "NBX-FN", "NBX-H1", "NBX-H2",
    "NBX-H3", "NBX-L1", "NBX-MCUL-FIRST", "NBX-MCUL-LAST",
    "NBX-MCUL-MID", "NBX-NL-FIRST", "NBX-NL-LAST", "NBX-NL-MID",
    "NBX-OUT1-FIRST", "NBX-OUT1-MID", "NBX-OUT2", "NBX-OUT2-LAST",
    "NBX-OUT3", "NBX-QUO", "NBX-QUO-AU", "NBX-TTL", "NBX-TXT",
    "NBX-TXT-FIRST", "NBX-TYPE", "NBX-UL-FIRST", "NBX-UL-LAST",
    "NBX-UL-MID",
    "NL-FIRST", "NL-LAST", "NL-MID", "NL-MID following L1",
    "OBJ-BL-FIRST", "OBJ-BL-LAST", "OBJ-BL-MID",
    "OBJ-NL-FIRST", "OBJ-NL-LAST", "OBJ-NL-MID",
    "OBJ-TXT", "OBJ-UL-FIRST", "OBJ-UL-LAST", "OBJ-UL-MID", "OBJ1",
    "OUT1-FIRST", "OUT1-LAST", "OUT1-MID",
    "OUT2-FIRST", "OUT2-LAST", "OUT2-MID",
    "OUT3-FIRST", "OUT3-MID",
    "PART", "PAU", "PN", "POC", "POC-FIRST", "POS", "POINT-BLURB",
    "POUT-1", "POUT-2", "POUTH1", "PQUOTE", "PST", "PT", "PTXT", "PTXT-DC",
    "QUES-NL-FIRST", "QUES-NL-MID", "QUES-SUB-FIRST", "QUES-SUB-MID",
    "QUES-SUB-LAST", "QUO", "QUOA",
    "REF-N", "REF-N-FIRST", "REF-U", "REF-H1", "REF-H2",
    "RHR", "RHV",
    "RQ-H1", "RQ-H2", "RQ-NL-FIRST", "RQ-NL-MID", "RQ-NL-LAST",
    "SAU", "SBBL-FIRST", "SBBL-LAST", "SBBL-MID",
    "SBNL-FIRST", "SBNL-MID", "SBT", "SBTXT", "SBUL", "SBUL-FIRST",
    "SECTION", "SN", "SOC", "SOC-FIRST", "SOS",
    "SOUT-1", "SOUT-2", "SOUTH1", "SOUT-NL-FIRST", "SOUT-NL-MID",
    "SP1", "SP2", "SQUOTE", "SR", "SRH1", "SRH2", "SST", "ST",
    "STXT", "STXT-DC", "STYLE LABEL",
    "T", "T-DIR", "T1", "T2", "T2-C", "T3", "T4", "T5",
    "TBL", "TFN", "TFN-FIRST", "TFN-MID", "TFN-LAST",
    "TFN-BL-FIRST", "TFN-BL-MID", "TFN-BL-LAST",
    "TMATH", "TN", "TNL-FIRST", "TNL-MID", "TTXT",
    "TUL-FIRST", "TUL-MID", "TXT", "TXT-DC", "TXT-FLUSH",
    "TXT-SPACE ABOVE", "UAU",
    "UL-FIRST", "UL-LAST", "UL-MID", "UN",
    "UNBX-BL", "UNBX-NL", "UNBX-T", "UNBX-T2", "UNBX-TT", "UNBX-UL",
    "UNFIG", "UNIT", "UNT", "UNT-BL", "UNT-FN",
    "UNT-NL-FIRST", "UNT-NL-MID", "UNT-T2", "UNT-T3", "UNT-TTL", "UNT-UL",
    "UOC", "UOC-FIRST", "UOS",
    "UOUT-1", "UOUT-2", "UOUTH1", "UQUOTE", "UST", "UT", "UTXT",
    "WEBTXT", "WL1",
    "HTTLPG-TTL", "HTTLPG-SUBTTL", "HTTLPG-ED",
    "TTLPG-TTL", "TTLPG-SUBTTL", "TTLPG-ED", "TTLPG-VOL",
    "TTLPG-AU", "TTLPG-AU-AFFIL",
    "CPY", "DED", "DED-AU",
    "FM-TTL", "FM-AU", "FM-AU-AFFIL",
    "CONTRIB-AU", "CONTRIB-AU-AFFIL",
    "REV-AU", "REV-AU-AFFIL",
    "TOC-FM", "TOC-UN", "TOC-UT", "TOC-SN", "TOC-ST",
    "TOC-CN", "TOC-CT", "TOC-CAU", "TOC-H1", "TOC-H2",
    "TOC-BM-FIRST", "TOC-BM",
    "BM-TTL",
    "GLOS-UL-FIRST", "GLOS-UL-MID", "GLOS-NL-FIRST", "GLOS-NL-MID",
    "GLOS-BL-FIRST", "GLOS-BL-MID",
    "IDX-TXT", "IDX-ALPHA", "IDX-1", "IDX-2", "IDX-3",
]


def _style_name_to_id(name: str) -> str:
    return name.replace(" ", "_").replace("-", "_")


def inject_publisher_styles(docx_path: str, save_path: str = None) -> None:
    """
    Inject all publisher styles into the DOCX as paragraph styles.
    Each style is set to Times New Roman, 12pt, 2.0 line spacing.
    Skips styles that already exist (safe to call multiple times).
    """
    doc = Document(docx_path)
    existing = {s.name for s in doc.styles}

    added = 0
    skipped = 0

    for tag in PUBLISHER_STYLES:
        if tag in existing:
            skipped += 1
            continue

        try:
            style = doc.styles.add_style(tag, WD_STYLE_TYPE.PARAGRAPH)
            style.hidden = False
            style.quick_style = True
            style.priority = 100

            try:
                style.base_style = doc.styles["Normal"]
            except KeyError:
                pass

            _apply_formatting(style)
            added += 1
        except Exception as e:
            print(f"  [WARN] Could not add style '{tag}': {e}")

    out_path = save_path or docx_path
    doc.save(out_path)
    print(f"Done — added {added} styles, skipped {skipped} existing. Saved to: {out_path}")


def create_template_dotx(output_path: str = "publisher_template.docx") -> None:
    """
    Create a blank DOCX containing all publisher styles.
    """
    doc = Document()

    for para in doc.paragraphs:
        p = para._element
        p.getparent().remove(p)

    existing = {s.name for s in doc.styles}
    added = 0

    for tag in PUBLISHER_STYLES:
        if tag in existing:
            continue
        try:
            style = doc.styles.add_style(tag, WD_STYLE_TYPE.PARAGRAPH)
            style.hidden = False
            style.quick_style = True
            style.priority = 100
            try:
                style.base_style = doc.styles["Normal"]
            except KeyError:
                pass
            _apply_formatting(style)
            added += 1
        except Exception as e:
            print(f"  [WARN] Could not add style '{tag}': {e}")

    doc.save(output_path)
    print(f"Template created with {added} styles at: {output_path}")


if __name__ == "__main__":
    if len(sys.argv) == 3:
        inject_publisher_styles(sys.argv[1], sys.argv[2])
    elif len(sys.argv) == 2:
        if sys.argv[1] == "--template":
            create_template_dotx("publisher_template.docx")
        else:
            inject_publisher_styles(sys.argv[1])
    else:
        print("Usage:")
        print("  python inject_styles.py input.docx output.docx   # inject into existing doc")
        print("  python inject_styles.py input.docx               # overwrite in-place")
        print("  python inject_styles.py --template               # create blank template")
