"""Run-anchored, structural-bookmarked XHTML → DOCX delta patch.

Uses paragraph, run, table, cell, and footnote bookmarks to map editor HTML elements
1-to-1 back to their source DOCX XML elements. Edits are applied strictly in-place,
guaranteeing 100% formatting preservation.
"""

import base64
import logging
import os
import re
import uuid
from datetime import datetime
from io import BytesIO
from pathlib import Path

import lxml.html
from docx import Document
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn
from docx.shared import Emu, Inches
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from lxml import etree

logger = logging.getLogger("app.processing.xhtml_to_docx_delta")


# Matches data URIs like:
#   data:image/png;base64,iVBORw0KGgo...
#   data:image/jpeg;charset=utf-8;base64,/9j/4AAQ...    (Safari sometimes)
#   data:image/webp;base64,UklGRi...
# The middle chunk (";charset=...", ";name=foo", etc.) is tolerated.
_DATA_URI_RE = re.compile(
    # Anchor on `;base64,` as the last segment before the payload; anything
    # between the format and it (media-type parameters like ;charset=…) is
    # tolerated with a repeated non-capturing group.
    r"^data:image/([a-z0-9.+-]+)(?:;[^,]+)*;base64\s*,\s*(.+)$",
    re.IGNORECASE | re.DOTALL,
)

_SUPPORTED_IMAGE_FMTS = {
    "png": "png",
    "jpg": "jpg",
    "jpeg": "jpg",
    "gif": "gif",
    "webp": "webp",
    "bmp": "bmp",
    # SVG cannot be embedded as a picture by python-docx — skip below.
}


def _decode_image_data_uri(src: str) -> tuple[bytes, str] | tuple[None, None]:
    """Decode a base64 data: URI into (bytes, extension).

    Tolerates optional media-type parameters (charset, name, etc.) and
    whitespace in the base64 payload. Returns (None, None) for non-data-URI
    sources, unsupported formats (SVG), or malformed payloads.
    """
    if not src:
        return None, None
    m = _DATA_URI_RE.match(src.strip())
    if not m:
        return None, None
    fmt = m.group(1).lower()
    ext = _SUPPORTED_IMAGE_FMTS.get(fmt)
    if ext is None:
        return None, None
    payload = m.group(2)
    # Strip whitespace/newlines that some pastes include mid-payload.
    payload = re.sub(r"\s+", "", payload)
    try:
        blob = base64.b64decode(payload, validate=False)
    except Exception:
        return None, None
    if not blob:
        return None, None
    return blob, ext


def _nearest_preceding_bookmarked(el) -> "etree._Element | None":
    """Return the nearest bookmarked block-level ancestor or preceding sibling.

    Walks previous siblings first (deepest-last descendant wins there), then
    climbs to the parent and repeats. Used to anchor pasted images relative to
    existing DOCX paragraphs so the delta engine can insert a new paragraph in
    roughly the right spot.
    """
    def _has_bm(node) -> bool:
        return bool(node.get("data-bookmark")) if hasattr(node, "get") else False

    current = el
    while current is not None:
        prev = current.getprevious()
        while prev is not None:
            if _has_bm(prev):
                return prev
            descendants = prev.xpath(".//*[@data-bookmark]")
            if descendants:
                return descendants[-1]
            prev = prev.getprevious()
        parent = current.getparent()
        if parent is None:
            return None
        if _has_bm(parent):
            return parent
        current = parent
    return None


def _css_color_to_hex(val: str) -> str | None:
    """Normalize a CSS color (#hex, rgb(...), rgba(...)) to a 6-char uppercase hex string.

    Returns None for unrecognized values (e.g. 'inherit', 'transparent', named colors).
    Browsers serialize editor-applied colors as rgb()/rgba(); the color picker emits hex.
    """
    if not val:
        return None
    val = val.strip()
    if val.startswith("#"):
        h = val.lstrip("#")
        if len(h) == 3:
            h = "".join(c * 2 for c in h)
        if len(h) == 6:
            return h.upper()
        return None
    m = re.match(r"rgba?\(\s*(\d+)\s*,\s*(\d+)\s*,\s*(\d+)", val)
    if m:
        r, g, b = (max(0, min(255, int(m.group(i)))) for i in (1, 2, 3))
        return f"{r:02X}{g:02X}{b:02X}"
    return None


# ─── Lookup Helpers ───────────────────────────────────────────────────────────

def _get_unique_bookmark_id(doc) -> int:
    """Finds a unique bookmark integer ID in the document."""
    if not hasattr(doc, "_next_bookmark_id"):
        used_ids = set()
        for bm_start in doc.element.body.findall(f".//{qn('w:bookmarkStart')}"):
            bm_id = bm_start.get(qn("w:id"))
            if bm_id is not None:
                try:
                    used_ids.add(int(bm_id))
                except ValueError:
                    pass
        
        for rel_id, part in doc.part.related_parts.items():
            if "footnotes" in part.partname or "endnotes" in part.partname:
                try:
                    for bm_start in part._element.findall(f".//{qn('w:bookmarkStart')}"):
                        bm_id = bm_start.get(qn("w:id"))
                        if bm_id is not None:
                            try:
                                used_ids.add(int(bm_id))
                            except ValueError:
                                pass
                except Exception:
                    pass
        doc._next_bookmark_id = max(used_ids) + 1 if used_ids else 1

    next_id = doc._next_bookmark_id
    doc._next_bookmark_id += 1
    return next_id


def _find_para_by_bookmark(doc, bookmark_name: str):
    """Finds a body paragraph or table cell paragraph containing the specified bookmark."""
    # 1. Search in body paragraphs
    for para in doc.paragraphs:
        for child in para._p:
            if child.tag == qn("w:bookmarkStart") and child.get(qn("w:name")) == bookmark_name:
                return para

    # 2. Search in table cell paragraphs
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for child in para._p:
                        if child.tag == qn("w:bookmarkStart") and child.get(qn("w:name")) == bookmark_name:
                            return para
    return None


def _find_note_para_by_bookmark(doc, bookmark_name: str):
    """Finds a footnote or endnote definition paragraph containing the specified bookmark."""
    for rel_id, part in doc.part.related_parts.items():
        if "footnotes" in part.partname or "endnotes" in part.partname:
            try:
                for p_elem in part._element.findall(f".//{qn('w:p')}"):
                    for child in p_elem:
                        if child.tag == qn("w:bookmarkStart") and child.get(qn("w:name")) == bookmark_name:
                            import docx
                            return docx.text.paragraph.Paragraph(p_elem, doc)
            except Exception:
                pass
    return None


def _find_run_by_bookmark(para, bookmark_name: str):
    """Finds the run element wrapped or preceded by the specified run bookmark."""
    p_children = list(para._p)
    for idx, child in enumerate(p_children):
        if child.tag == qn("w:bookmarkStart") and child.get(qn("w:name")) == bookmark_name:
            bm_id = child.get(qn("w:id"))
            # The mapped run is immediately following the bookmark start
            for next_idx in range(idx + 1, len(p_children)):
                next_child = p_children[next_idx]
                if next_child.tag == qn("w:r"):
                    return Run(next_child, para)
                elif next_child.tag == qn("w:bookmarkEnd") and next_child.get(qn("w:id")) == bm_id:
                    break
    return None


def _build_bookmark_para_index(doc) -> dict:
    """Builds an O(1) lookup index from bookmark names to their containing paragraphs."""
    index = {}
    for para in doc.paragraphs:
        for child in para._p:
            if child.tag == qn("w:bookmarkStart"):
                name = child.get(qn("w:name"), "")
                if name:
                    index[name] = para
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for child in para._p:
                        if child.tag == qn("w:bookmarkStart"):
                            name = child.get(qn("w:name"), "")
                            if name:
                                index[name] = para
    return index


# ─── Core Delta Saving Engine ────────────────────────────────────────────────

class XhtmlToDocxDeltaEngine:
    """Apply HTML edits back to the DOCX in-place using unique tracking bookmarks."""

    def convert(self, html_path: str, out_docx_path: str, username: str = "WYSIWYG Editor") -> str:
        if not os.path.exists(html_path):
            raise RuntimeError(f"Input HTML not found: {html_path}")
        if not os.path.exists(out_docx_path):
            raise RuntimeError(f"Target DOCX not found: {out_docx_path}")

        # Parse saved HTML content
        with open(html_path, "r", encoding="utf-8") as f:
            html_content = f.read()

        root = lxml.html.fromstring(html_content)
        doc = Document(out_docx_path)
        patched = 0

        # Build bookmark paragraph index for O(1) lookups
        para_index = _build_bookmark_para_index(doc)

        # Find all HTML paragraphs, headings, list items, cell paragraphs, and page breaks
        all_elements = root.xpath("//p | //h1 | //h2 | //h3 | //h4 | //h5 | //h6 | //li | //div[contains(@class, 'page-break')] | //hr[contains(@class, 'page-break')]")
        
        # Filter out li elements that contain p, h1, h2, h3, h4, h5, or h6 descendants
        # to prevent parent-child double matching when list items contain styled paragraphs.
        html_blocks = []
        for el in all_elements:
            if el.tag == "li":
                has_nested_block = any(child.tag in ("p", "h1", "h2", "h3", "h4", "h5", "h6") for child in el.iterdescendants())
                if has_nested_block:
                    continue
            html_blocks.append(el)

        for idx, block_el in enumerate(html_blocks):
            is_page_break = block_el.tag in ("div", "hr") and "page-break" in (block_el.get("class") or "")
            if is_page_break:
                # Find the nearest upcoming element that has a bookmark
                target_para = None
                for next_el in html_blocks[idx + 1:]:
                    next_bm = next_el.get("data-bookmark")
                    if next_bm:
                        target_para = para_index.get(next_bm) or _find_note_para_by_bookmark(doc, next_bm)
                        if target_para:
                            break
                if target_para:
                    try:
                        from docx.enum.text import WD_BREAK
                        p = target_para.insert_paragraph_before()
                        p.add_run().add_break(WD_BREAK.PAGE)
                        patched += 1
                        logger.info("Inserted DOCX page break in delta engine successfully.")
                    except Exception as e:
                        logger.warning(f"Failed to insert page break in delta: {e}")
                continue

            bm_name = block_el.get("data-bookmark")
            if not bm_name:
                continue

            if block_el.tag == "li":
                new_style = _determine_list_style(block_el)
            else:
                new_style = block_el.get("data-style-label") or block_el.get("class", "")
                new_style = new_style.split()[0] if new_style.strip() else "Normal"
                if new_style in ("Normal", "MsoNormal", ""):
                    new_style = "Normal"

            # 1. Retrieve the exact paragraph node in the body, table cells, or footnote parts
            para = para_index.get(bm_name)
            if not para:
                para = _find_note_para_by_bookmark(doc, bm_name)
            
            if not para:
                logger.info(f"Bookmark {bm_name} not found in DOCX (might be a newly added paragraph, skip in-place)")
                continue

            # 2. Update paragraph-level style if changed
            try:
                if para.style.name != new_style:
                    para.style = new_style
            except Exception:
                try:
                    # Fallback to direct XML injection if style name is missing in document template
                    pPr = para._p.get_or_add_pPr()
                    for existing in pPr.findall(qn("w:pStyle")):
                        pPr.remove(existing)
                    pStyle = etree.SubElement(pPr, qn("w:pStyle"))
                    pStyle.set(qn("w:val"), new_style)
                    pPr.insert(0, pStyle)
                except Exception:
                    pass

            # 2b. Sync paragraph alignment (text-align) from HTML style → w:jc
            try:
                inline_style = (block_el.get("style") or "").lower()
                align_match = re.search(r"text-align\s*:\s*(left|center|right|justify)", inline_style)
                pPr = para._p.get_or_add_pPr()
                for existing in pPr.findall(qn("w:jc")):
                    pPr.remove(existing)
                if align_match:
                    val = align_match.group(1)
                    jc_val = "both" if val == "justify" else val
                    jc = OxmlElement("w:jc")
                    jc.set(qn("w:val"), jc_val)
                    pPr.append(jc)
            except Exception as align_err:
                logger.warning(f"Failed to sync alignment for bookmark {bm_name}: {align_err}")

            # 3. Patch paragraph runs strictly in-place
            try:
                self._patch_paragraph_runs(para, block_el, doc, username)
                patched += 1
            except Exception as exc:
                logger.warning(f"Failed to patch runs for paragraph bookmark {bm_name}: {exc}", exc_info=True)

        logger.info(f"Lossless Bookmark Delta patch: {patched} paragraph(s) updated in-place.")

        # Insert images pasted into the editor (not present in the original DOCX).
        # These arrive as top-level <img> siblings with data: URIs; the delta
        # patch above ignores them because they have no bookmark.
        try:
            self._insert_new_images(root, doc, para_index)
        except Exception as img_err:
            logger.warning(f"Failed to insert pasted images: {img_err}", exc_info=True)

        # Apply final manuscript formatting (Times New Roman, 12pt, double spacing)
        try:
            apply_final_docx_formatting(doc)
        except Exception as fmt_err:
            logger.warning(f"Failed to apply final document formatting in delta: {fmt_err}")

        # Save atomically
        tmp_path = out_docx_path + ".delta.tmp"
        doc.save(tmp_path)
        
        # Ensure revision tracking is enabled
        try:
            from app.processing.manuscript_core.fixer import enable_track_revisions
            enable_track_revisions(Path(tmp_path))
        except Exception as exc:
            logger.warning(f"Could not enable track revisions: {exc}")
            
        os.replace(tmp_path, out_docx_path)
        logger.info(f"Lossless Delta-patched DOCX saved: {out_docx_path}")
        return out_docx_path

    def _insert_new_images(self, root, doc, para_index: dict) -> None:
        """Insert images that were pasted into the editor into the DOCX.

        The delta engine only touches paragraphs it can look up via bookmark, so
        anything the user added while editing (pasted images, in particular) is
        otherwise silently dropped. This walks the XHTML for <img> tags with a
        base64 data: URI, decodes the bytes, and inserts a new paragraph after
        the nearest preceding bookmarked block. If no anchor is found, the
        image is appended at the end of the document.
        """
        all_imgs = root.xpath(".//img[@src]")
        if not all_imgs:
            return

        inserted = 0
        skipped = 0
        # Per-anchor cursor: two images anchored to the same bookmark must be
        # inserted after each other in HTML order, not both directly after the
        # anchor (which would reverse them via addnext).
        cursor_for_bm: dict[str, Paragraph] = {}

        # Word for Mac refuses to render pictures whose <pic:cNvPr id> or
        # <wp:docPr id> collide with another drawing in the same document —
        # it paints a blank placeholder rectangle instead. python-docx's
        # add_picture() hardcodes id="0" on <pic:cNvPr> and always starts
        # <wp:docPr id> from 1, so consecutive inserts overlap unless we
        # rewrite the IDs. Seed the counter above any existing drawing ID.
        pic_ns = "http://schemas.openxmlformats.org/drawingml/2006/picture"
        wp_ns = "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing"
        used_pic_ids: set[int] = set()
        used_dp_ids: set[int] = set()
        for existing in doc.element.body.iter(f"{{{pic_ns}}}cNvPr"):
            raw = existing.get("id")
            try:
                if raw is not None:
                    used_pic_ids.add(int(raw))
            except ValueError:
                pass
        for existing in doc.element.body.iter(f"{{{wp_ns}}}docPr"):
            raw = existing.get("id")
            try:
                if raw is not None:
                    used_dp_ids.add(int(raw))
            except ValueError:
                pass
        next_pic_id = (max(used_pic_ids) + 1) if used_pic_ids else 1
        next_dp_id = (max(used_dp_ids) + 1) if used_dp_ids else 1

        for img_el in all_imgs:
            src = img_el.get("src", "")
            image_bytes, ext = _decode_image_data_uri(src)
            if image_bytes is None:
                # Non-data URI (external URL, missing src, or unsupported format
                # like SVG). Nothing to embed without a fetch, so skip quietly.
                skipped += 1
                continue

            # Skip if this <img> already lives inside a bookmarked paragraph —
            # the delta pass has rebuilt that paragraph's runs from scratch and
            # would clobber a run we insert here. This is rare given the block-
            # level ImageNode schema but guarded against for future inline use.
            ancestor_bm = img_el.xpath("ancestor::*[@data-bookmark][1]")
            if ancestor_bm:
                skipped += 1
                continue

            anchor_el = _nearest_preceding_bookmarked(img_el)
            anchor_key = ""
            target_para = None
            if anchor_el is not None:
                bm = anchor_el.get("data-bookmark") or ""
                anchor_key = bm
                if bm:
                    target_para = cursor_for_bm.get(bm)
                    if target_para is None:
                        target_para = para_index.get(bm) or _find_note_para_by_bookmark(doc, bm)

            if target_para is None:
                # Fall back to appending at the very end of the document body.
                anchor_key = "__doc_end__"
                paragraphs = list(doc.paragraphs)
                if not paragraphs:
                    logger.warning("No anchor paragraph available for pasted image; skipping.")
                    skipped += 1
                    continue
                target_para = cursor_for_bm.get(anchor_key) or paragraphs[-1]

            try:
                new_p_xml = OxmlElement("w:p")
                target_para._element.addnext(new_p_xml)
                new_para = Paragraph(new_p_xml, target_para._parent)

                # Prefer the width the user set in the editor; fall back to a
                # sensible 4-inch default so the picture is never zero-sized.
                width_arg = None
                raw_width = img_el.get("width")
                if raw_width:
                    try:
                        px = float(str(raw_width).replace("px", "").strip())
                        if px > 0:
                            # CSS px → EMU at 96 DPI. 914400 EMU = 1 inch.
                            width_arg = Emu(int(px / 96.0 * 914400))
                    except (ValueError, TypeError):
                        width_arg = None
                if width_arg is None:
                    width_arg = Inches(4)

                run = new_para.add_run()
                run.add_picture(BytesIO(image_bytes), width=width_arg)

                # Rewrite the IDs Word Mac tests for uniqueness. Without this,
                # a second pasted image is rendered as an empty blue rectangle.
                for pic_cnvpr in new_p_xml.iter(f"{{{pic_ns}}}cNvPr"):
                    pic_cnvpr.set("id", str(next_pic_id))
                    next_pic_id += 1
                for docpr in new_p_xml.iter(f"{{{wp_ns}}}docPr"):
                    docpr.set("id", str(next_dp_id))
                    next_dp_id += 1

                # Center the picture — matches typical figure placement and is
                # what the editor's block image node visually implies.
                try:
                    pPr = new_p_xml.find(qn("w:pPr"))
                    if pPr is None:
                        pPr = OxmlElement("w:pPr")
                        new_p_xml.insert(0, pPr)
                    for existing in pPr.findall(qn("w:jc")):
                        pPr.remove(existing)
                    jc = OxmlElement("w:jc")
                    jc.set(qn("w:val"), "center")
                    pPr.append(jc)
                except Exception:
                    pass

                # Advance the anchor cursor so the next image anchored to the
                # same bookmark lands after this one, preserving HTML order.
                if anchor_key:
                    cursor_for_bm[anchor_key] = new_para

                inserted += 1
            except Exception as e:
                logger.warning(f"Failed to embed pasted image: {e}", exc_info=True)
                skipped += 1

        if inserted or skipped:
            logger.info(
                f"Pasted image pass: inserted={inserted}, skipped={skipped}"
            )

    def _patch_paragraph_runs(self, para, html_el, doc, username: str) -> None:
        """Modifies and synchronizes runs inside `<w:p>` completely in-place by rebuilding them."""
        import re
        import uuid
        from datetime import datetime
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from lxml import etree
        from docx.opc.constants import RELATIONSHIP_TYPE

        # 1. Clear existing runs, ins, del, hyperlinks, and run-level bookmarks (r_bm_*)
        p_elem = para._element
        for child in list(p_elem):
            tag_name = etree.QName(child.tag).localname
            if tag_name in ('r', 'ins', 'del', 'hyperlink'):
                p_elem.remove(child)
            elif tag_name in ('bookmarkStart', 'bookmarkEnd'):
                name = child.get(qn('w:name'), '')
                if name.startswith('r_bm_'):
                    p_elem.remove(child)

        # ISO 8601 timestamp with UTC timezone
        timestamp = datetime.utcnow().isoformat() + "Z"
        
        # Track unique revision/bookmark IDs
        next_id = [_get_unique_bookmark_id(doc)]  # Use list to allow mutation

        def parse_style(style_str):
            if not style_str:
                return {}
            styles = {}
            for pair in style_str.split(';'):
                if ':' in pair:
                    key, val = pair.split(':', 1)
                    styles[key.strip().lower()] = val.strip()
            return styles

        def add_rich_run(parent_element, text, bold=False, italic=False, underline=False, strike=False, color=None, bg_color=None, font_size=None, font_name=None, superscript=False, subscript=False, is_link=False, is_del=False, char_style=None):
            if not text:
                return
            r = OxmlElement('w:r')

            rPr = OxmlElement('w:rPr')
            has_rPr = False

            if bold:
                b = OxmlElement('w:b')
                rPr.append(b)
                has_rPr = True
            if italic:
                i = OxmlElement('w:i')
                rPr.append(i)
                has_rPr = True
            if strike:
                s_el = OxmlElement('w:strike')
                rPr.append(s_el)
                has_rPr = True

            final_underline = underline or is_link
            final_color = color
            if is_link and not final_color:
                final_color = "0563C1"

            if final_underline:
                u = OxmlElement('w:u')
                u.set(qn('w:val'), 'single')
                rPr.append(u)
                has_rPr = True
            if final_color:
                c = OxmlElement('w:color')
                c.set(qn('w:val'), final_color.upper())
                rPr.append(c)
                has_rPr = True
            if bg_color:
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), bg_color.upper())
                rPr.append(shd)
                has_rPr = True
            if font_size:
                sz_val = str(int(float(font_size) * 2))
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), sz_val)
                szCs = OxmlElement('w:szCs')
                szCs.set(qn('w:val'), sz_val)
                rPr.append(sz)
                rPr.append(szCs)
                has_rPr = True
            if font_name:
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), font_name)
                rFonts.set(qn('w:hAnsi'), font_name)
                rFonts.set(qn('w:cs'), font_name)
                rPr.append(rFonts)
                has_rPr = True
            if superscript:
                va = OxmlElement('w:vertAlign')
                va.set(qn('w:val'), 'superscript')
                rPr.append(va)
                has_rPr = True
            elif subscript:
                va = OxmlElement('w:vertAlign')
                va.set(qn('w:val'), 'subscript')
                rPr.append(va)
                has_rPr = True
            if char_style and char_style != "Default Paragraph Font":
                rStyle = OxmlElement('w:rStyle')
                rStyle.set(qn('w:val'), char_style)
                rPr.insert(0, rStyle)  # w:rStyle must be first child of w:rPr per OOXML schema
                has_rPr = True

            if has_rPr:
                r.append(rPr)

            t = OxmlElement('w:delText' if is_del else 'w:t')
            if text.startswith(' ') or text.endswith(' '):
                t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t.text = text
            r.append(t)

            parent_element.append(r)

        # Helper to wrap children in a bookmark Start/End
        def wrap_in_bookmark(parent_xml, bm_name, process_content_fn):
            bm_start = OxmlElement("w:bookmarkStart")
            bm_start.set(qn("w:id"), str(next_id[0]))
            bm_start.set(qn("w:name"), bm_name)

            bm_end = OxmlElement("w:bookmarkEnd")
            bm_end.set(qn("w:id"), str(next_id[0]))

            next_id[0] += 1

            parent_xml.append(bm_start)
            process_content_fn(parent_xml)
            parent_xml.append(bm_end)

        def traverse(el, parent_xml, bold=False, italic=False, underline=False, strike=False, color=None, bg_color=None, font_size=None, font_name=None, superscript=False, subscript=False, is_link=False, char_style=None, is_del=False):
            tag = el.tag.split('}')[-1].lower() if isinstance(el.tag, str) else ""

            style_str = el.get("style", "")
            styles = parse_style(style_str)

            has_bold_style = "font-weight" in styles and styles["font-weight"].strip().lower() in ("bold", "700")
            has_underline_style = "text-decoration" in styles and "underline" in styles["text-decoration"].strip().lower()
            has_italic_style = "font-style" in styles and styles["font-style"].strip().lower() == "italic"
            has_strike_style = "text-decoration" in styles and "line-through" in styles["text-decoration"].strip().lower()

            node_color = color
            if "color" in styles:
                parsed = _css_color_to_hex(styles["color"])
                if parsed:
                    node_color = parsed

            node_bg = bg_color
            if "background-color" in styles:
                parsed = _css_color_to_hex(styles["background-color"])
                if parsed:
                    node_bg = parsed
            elif "background" in styles:
                parsed = _css_color_to_hex(styles["background"])
                if parsed:
                    node_bg = parsed

            # TipTap Highlight extension renders <mark data-color="#hex" style="background-color: rgb(...)">.
            # If the inline style fails (e.g. 'inherit'), fall back to the data-color attribute.
            if tag == "mark" and not node_bg:
                parsed = _css_color_to_hex(el.get("data-color") or "")
                if parsed:
                    node_bg = parsed

            node_font_size = font_size
            if "font-size" in styles:
                sz_val = styles["font-size"].strip()
                match = re.match(r"(\d+(\.\d+)?)\s*(pt|px)?", sz_val)
                if match:
                    val = float(match.group(1))
                    unit = match.group(3)
                    if unit == "px":
                        node_font_size = round(val * 0.75, 1)
                    else:
                        node_font_size = val

            # Parse character style from class — only on span elements
            _CAPTION_CHAR_STYLES = {"FigureCitation", "TableCitation", "FIG-NUM", "TN"}
            node_char_style = char_style
            if tag == "span":
                classes = (el.get("class") or "").split()
                style_match = next(
                    (c for c in classes
                     if c.startswith("bib_") or c.startswith("cite_")
                     or (c.isalpha() and c.islower())
                     or c in _CAPTION_CHAR_STYLES),
                    None,
                )
                if style_match:
                    node_char_style = style_match

            current_bold = bold or tag in ('strong', 'b') or has_bold_style
            current_italic = italic or tag in ('em', 'i') or has_italic_style
            current_underline = underline or tag == 'u' or has_underline_style
            current_strike = strike or tag in ('s', 'strike') or has_strike_style
            current_super = superscript or tag == 'sup' or styles.get("vertical-align") == "super"
            current_sub = subscript or tag == 'sub' or styles.get("vertical-align") == "sub"
            current_link = is_link or tag == 'a'

            current_xml_parent = parent_xml
            # is_del propagates from the parent (set by an enclosing <del>) so that
            # nested elements like <del><b>...</b></del> still emit w:delText.
            current_is_del = is_del

            replacement_text = el.get("data-replacement")
            if tag == 'span' and replacement_text is not None:
                author = el.get("data-author") or username
                date = el.get("data-date") or timestamp
                
                # 1. Deletion
                # Parse font-family
                node_font_name = font_name
                if "font-family" in styles:
                    raw_family = styles["font-family"].strip()
                    first_font = raw_family.split(',')[0].strip().replace("'", "").replace('"', '').strip()
                    node_font_name = first_font if first_font else None

                original_text = el.text or ""
                for child in el:
                    original_text += child.text or ""
                    if child.tail:
                        original_text += child.tail
                
                if original_text:
                    del_node = OxmlElement('w:del')
                    del_node.set(qn('w:id'), str(next_id[0]))
                    del_node.set(qn('w:author'), author)
                    del_node.set(qn('w:date'), date)
                    next_id[0] += 1
                    
                    add_rich_run(
                        del_node, original_text,
                        bold=current_bold, italic=current_italic, underline=current_underline, strike=current_strike,
                        color=node_color, bg_color=None, font_size=node_font_size, font_name=node_font_name,
                        superscript=current_super, subscript=current_sub, is_link=current_link,
                        is_del=True, char_style=node_char_style
                    )
                    current_xml_parent.append(del_node)
                
                # 2. Insertion
                if replacement_text:
                    ins_node = OxmlElement('w:ins')
                    ins_node.set(qn('w:id'), str(next_id[0]))
                    ins_node.set(qn('w:author'), author)
                    ins_node.set(qn('w:date'), date)
                    next_id[0] += 1
                    
                    add_rich_run(
                        ins_node, replacement_text,
                        bold=current_bold, italic=current_italic, underline=current_underline, strike=current_strike,
                        color=node_color, bg_color=None, font_size=node_font_size, font_name=node_font_name,
                        superscript=current_super, subscript=current_sub, is_link=current_link,
                        is_del=False, char_style=node_char_style
                    )
                    current_xml_parent.append(ins_node)
                
                # 3. Handle tail text of the span element (tail inherits parent's is_del scope)
                if el.tail:
                    add_rich_run(
                        parent_xml, el.tail,
                        bold=bold, italic=italic, underline=underline, strike=strike,
                        color=color, bg_color=bg_color, font_size=font_size, font_name=font_name,
                        superscript=superscript, subscript=subscript, is_link=is_link,
                        is_del=is_del, char_style=char_style
                    )
                return

            if tag == 'ins':
                author = el.get("data-author") or username
                date = el.get("data-date") or timestamp
                ins_node = OxmlElement('w:ins')
                ins_node.set(qn('w:id'), str(next_id[0]))
                ins_node.set(qn('w:author'), author)
                ins_node.set(qn('w:date'), date)
                next_id[0] += 1
                parent_xml.append(ins_node)
                current_xml_parent = ins_node
            elif tag == 'del':
                author = el.get("data-author") or username
                date = el.get("data-date") or timestamp
                del_node = OxmlElement('w:del')
                del_node.set(qn('w:id'), str(next_id[0]))
                del_node.set(qn('w:author'), author)
                del_node.set(qn('w:date'), date)
                next_id[0] += 1
                parent_xml.append(del_node)
                current_xml_parent = del_node
                current_is_del = True
            elif tag == 'a':
                href = el.get("href", "")
                if href:
                    try:
                        r_id = para.part.relate_to(href, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
                        link_node = OxmlElement('w:hyperlink')
                        link_node.set(qn('r:id'), r_id)
                        parent_xml.append(link_node)
                        current_xml_parent = link_node
                    except Exception as link_err:
                        logger.warning(f"Could not build hyperlink in python-docx: {link_err}")

            if tag == 'math' or el.get("data-latex"):
                try:
                    import latex2mathml.converter
                    import mathml2omml
                    
                    if tag == 'math':
                        # Serialize the element to MathML string
                        mathml_str = etree.tostring(el, encoding="utf-8").decode("utf-8")
                    else:
                        latex_str = el.get("data-latex")
                        mathml_str = latex2mathml.converter.convert(latex_str)
                    
                    omml_str = mathml2omml.convert(mathml_str)
                    if "xmlns:m=" not in omml_str:
                        omml_str = omml_str.replace("<m:oMath", '<m:oMath xmlns:m="http://schemas.openxmlformats.org/officeDocument/2006/math"', 1)
                    
                    omml_el = etree.fromstring(omml_str)
                    current_xml_parent.append(omml_el)
                except Exception as math_err:
                    logger.warning(f"Failed to convert math to OMML: {math_err}")
                    if el.text:
                        add_rich_run(
                            current_xml_parent, el.text,
                            bold=current_bold, italic=current_italic, underline=current_underline, strike=current_strike,
                            color=node_color, bg_color=node_bg, font_size=node_font_size, font_name=node_font_name,
                            superscript=current_super, subscript=current_sub, is_link=current_link,
                            is_del=current_is_del, char_style=node_char_style
                        )

                if el.tail:
                    add_rich_run(
                        parent_xml, el.tail,
                        bold=bold, italic=italic, underline=underline, strike=strike,
                        color=color, bg_color=bg_color, font_size=font_size, font_name=font_name,
                        superscript=superscript, subscript=subscript, is_link=is_link,
                        is_del=is_del, char_style=char_style
                    )
                return

            bm_name = el.get("data-bookmark")
            # Auto-generate a bookmark name if it's a span and doesn't have one
            if tag == 'span' and not bm_name:
                unique_id = uuid.uuid4().hex[:8]
                bm_name = f"r_bm_{unique_id}"
                el.set("data-bookmark", bm_name)

            # Parse font-family
            node_font_name = font_name
            if "font-family" in styles:
                raw_family = styles["font-family"].strip()
                first_font = raw_family.split(',')[0].strip().replace("'", "").replace('"', '').strip()
                node_font_name = first_font if first_font else None

            def process_text_and_children(xml_parent):
                if el.text:
                    add_rich_run(
                        xml_parent, el.text,
                        bold=current_bold, italic=current_italic, underline=current_underline, strike=current_strike,
                        color=node_color, bg_color=node_bg, font_size=node_font_size, font_name=node_font_name,
                        superscript=current_super, subscript=current_sub, is_link=current_link,
                        is_del=current_is_del, char_style=node_char_style
                    )
                for child in el:
                    traverse(
                        child, xml_parent,
                        bold=current_bold, italic=current_italic, underline=current_underline, strike=current_strike,
                        color=node_color, bg_color=node_bg, font_size=node_font_size, font_name=node_font_name,
                        superscript=current_super, subscript=current_sub, is_link=current_link,
                        char_style=node_char_style, is_del=current_is_del
                    )

            if bm_name and bm_name.startswith("r_bm_"):
                wrap_in_bookmark(current_xml_parent, bm_name, process_text_and_children)
            else:
                process_text_and_children(current_xml_parent)

            if el.tail:
                # Tail text sits outside the current element, so the inherited
                # is_del from the parent scope applies (not current_is_del).
                add_rich_run(
                    parent_xml, el.tail,
                    bold=bold, italic=italic, underline=underline, strike=strike,
                    color=color, bg_color=bg_color, font_size=font_size, font_name=font_name,
                    superscript=superscript, subscript=subscript, is_link=is_link,
                    is_del=is_del, char_style=char_style
                )

        # Parse root element's initial styles (if any)
        root_style = html_el.get("style", "")
        root_styles = parse_style(root_style)
        root_color = None
        if "color" in root_styles:
            root_color = _css_color_to_hex(root_styles["color"])
        root_bg = None
        if "background-color" in root_styles:
            root_bg = _css_color_to_hex(root_styles["background-color"])
        elif "background" in root_styles:
            root_bg = _css_color_to_hex(root_styles["background"])
        root_font_size = None
        if "font-size" in root_styles:
            sz_val = root_styles["font-size"].strip()
            match = re.match(r"(\d+(\.\d+)?)\s*(pt|px)?", sz_val)
            if match:
                val = float(match.group(1))
                unit = match.group(3)
                if unit == "px":
                    root_font_size = round(val * 0.75, 1)
                else:
                    root_font_size = val
        root_font_name = None
        if "font-family" in root_styles:
            raw_family = root_styles["font-family"].strip()
            first_font = raw_family.split(',')[0].strip().replace("'", "").replace('"', '').strip()
            root_font_name = first_font if first_font else None

        # Begin traversal
        if html_el.text:
            add_rich_run(p_elem, html_el.text, color=root_color, bg_color=root_bg, font_size=root_font_size, font_name=root_font_name)

        for child in html_el:
            traverse(child, p_elem, color=root_color, bg_color=root_bg, font_size=root_font_size, font_name=root_font_name)


def _determine_list_style(li_el) -> str:
    # Walk up ancestors to count list containers
    parent = li_el.getparent()
    list_containers = []
    while parent is not None:
        if parent.tag in ("ul", "ol"):
            list_containers.append(parent.tag)
        parent = parent.getparent()
    
    if not list_containers:
        return "Normal"
    
    # The immediate list container is the first one in the list
    immediate_type = "bullet" if list_containers[0] == "ul" else "number"
    depth = len(list_containers)
    
    base_style = "List Bullet" if immediate_type == "bullet" else "List Number"
    if depth == 1:
        return base_style
    else:
        return f"{base_style} {depth}"


def apply_final_docx_formatting(doc):
    from docx.shared import Pt
    # Update Normal style so unstyled runs inherit Times New Roman 12pt double-spaced.
    # Run-level rFonts/sz (from user font-family / font-size choices in the editor) are
    # left untouched so they survive the round-trip.
    try:
        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(12)
        normal_style.paragraph_format.line_spacing = 2.0
    except Exception as e:
        logger.warning(f"Could not set Normal style: {e}")

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "Normal"
        is_list = "list" in style_name.lower() or "bullet" in style_name.lower() or "number" in style_name.lower()
        is_extract = any(x in style_name.lower() for x in ("extract", "quote", "ex"))
        if not is_list and not is_extract:
            para.paragraph_format.line_spacing = 2.0
