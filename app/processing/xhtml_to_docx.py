"""
Convert XHTML back to DOCX.

Strategy (preferred):
    Read paragraph style labels from saved HTML (data-style-label / class attributes)
    and apply them directly to the paragraphs of the **existing** processed DOCX file.
    This is done with python-docx and is lossless — formatting, images, tables, and
    track changes are completely preserved.  Only paragraph style names are updated.

Fallback (legacy):
    If the primary strategy fails (e.g. missing python-docx / lxml), fall back to
    converting the HTML via pandoc then LibreOffice.  This fallback *does* lose styles
    but is better than a hard error.
"""

import os
import shutil
import subprocess
import logging
import re
from pathlib import Path

logger = logging.getLogger("app.processing.xhtml_to_docx")


class XhtmlToDocxEngine:
    """Apply HTML-editor style changes back to the processed DOCX file."""

    # -------------------------------------------------------------------------
    # Public API
    # -------------------------------------------------------------------------

    def convert(self, html_path: str, out_docx_path: str, username: str = "WYSIWYG Editor") -> str:
        """
        Apply paragraph style labels from *html_path* to the DOCX at
        *out_docx_path*.

        If *out_docx_path* already exists (the processed DOCX), the style patch
        is applied in-place on a tmp copy and then atomically moved.

        If *out_docx_path* does not exist the legacy pandoc/LibreOffice fallback
        is invoked to create a brand-new DOCX from the HTML.

        Args:
            html_path:      Absolute path to the saved HTML editor content.
            out_docx_path:  Absolute path to the target DOCX file.
            username:       Logged-in username to attribute the track changes to.

        Returns:
            Path to the written DOCX file.
        """
        if not os.path.exists(html_path):
            raise RuntimeError(f"Input HTML not found: {html_path}")

        os.makedirs(os.path.dirname(out_docx_path), exist_ok=True)

        # ── Primary: patch styles directly on the existing DOCX ──────────────
        if os.path.exists(out_docx_path):
            try:
                return self._apply_styles_to_docx(html_path, out_docx_path, username=username)
            except Exception as e:
                logger.warning(
                    f"Style-patch failed ({e}); falling back to full conversion."
                )

        # ── Fallback: full HTML → DOCX conversion ────────────────────────────
        logger.warning(
            "Falling back to pandoc/LibreOffice conversion (style names will be lost)."
        )
        if shutil.which("pandoc"):
            try:
                return self._pandoc(html_path, out_docx_path)
            except Exception as e:
                logger.warning(f"pandoc failed, trying LibreOffice: {e}")

        try:
            return self._libreoffice(html_path, out_docx_path)
        except Exception as e:
            raise RuntimeError(f"Both pandoc and LibreOffice failed: {e}")

    # -------------------------------------------------------------------------
    # Primary strategy — python-docx style patching
    # -------------------------------------------------------------------------

    def _apply_styles_to_docx(self, html_path: str, docx_path: str, username: str = "WYSIWYG Editor") -> str:
        """
        Parse the HTML, extract (text, style-label) pairs for every paragraph /
        heading, then apply those style labels to the matching paragraphs of the
        DOCX using a greedy sliding-window sequence alignment.

        The DOCX is written atomically via a .tmp file.
        """
        import lxml.html
        from docx import Document
        from docx.oxml.ns import qn

        # 1. ── Parse HTML, collect (clean_text, style_name) pairs ───────────
        with open(html_path, "r", encoding="utf-8") as f:
            html_content = f.read()

        root = lxml.html.fromstring(html_content)
        # Select p, h1-h6, li, and page-break elements
        all_elements = root.xpath("//p | //h1 | //h2 | //h3 | //h4 | //h5 | //h6 | //li | //div[contains(@class, 'page-break')] | //hr[contains(@class, 'page-break')]")
        
        # Filter out li elements that contain p, h1, h2, h3, h4, h5, or h6 descendants
        # to prevent parent-child double matching when list items contain styled paragraphs.
        html_paras = []
        for el in all_elements:
            if el.tag == "li":
                has_nested_block = any(child.tag in ("p", "h1", "h2", "h3", "h4", "h5", "h6") for child in el.iterdescendants())
                if has_nested_block:
                    continue
            html_paras.append(el)

        html_entries = []
        for el in html_paras:
            is_page_break = el.tag in ("div", "hr") and "page-break" in (el.get("class") or "")
            raw_text = el.text_content().strip()
            if not raw_text and not is_page_break:
                continue

            if is_page_break:
                style_label = "PageBreak"
            elif el.tag == "li":
                style_label = _determine_list_style(el)
            else:
                # Prefer explicit data-style-label, fall back to class, then "Normal"
                style_label = (
                    el.get("data-style-label")
                    or el.get("class", "")
                    or "Normal"
                )
                # class may contain multiple tokens — take the first non-empty one
                style_label = style_label.split()[0] if style_label.strip() else "Normal"
                if style_label in ("Normal", "MsoNormal", ""):
                    style_label = "Normal"

            html_entries.append({
                "clean": self._clean(raw_text) if not is_page_break else "__PAGE_BREAK__",
                "style": style_label,
                "element": el,
            })

        # Also collect table cells
        html_table_cells = []
        for td_el in root.xpath("//table//td | //table//th"):
            cell_text = td_el.text_content().strip()
            if cell_text:
                html_table_cells.append({
                    "clean": self._clean(cell_text),
                    "element": td_el,
                })

        if not html_entries:
            logger.warning("No paragraphs found in HTML — nothing to patch.")
            return docx_path

        # 2. ── Load DOCX and build paragraph index ───────────────────────────
        doc = Document(docx_path)

        docx_entries = []
        para_map: list = []  # maps docx_entries index → doc.paragraphs index

        for i, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            docx_entries.append({"clean": self._clean(text), "para": para})
            para_map.append(i)

        if not docx_entries:
            logger.warning("DOCX has no non-empty paragraphs — nothing to patch.")
            return docx_path

        # 3. ── Greedy sliding-window alignment ────────────────────────────────
        n_docx = len(docx_entries)
        docx_cursor = 0
        changes_applied = 0

        for html_entry in html_entries:
            h_clean = html_entry["clean"]
            new_style = html_entry["style"]

            # If it is a page break, insert it before the next DOCX paragraph!
            if new_style == "PageBreak":
                if docx_cursor < n_docx:
                    para = docx_entries[docx_cursor]["para"]
                    try:
                        from docx.enum.text import WD_BREAK
                        p = para.insert_paragraph_before()
                        p.add_run().add_break(WD_BREAK.PAGE)
                        changes_applied += 1
                        logger.info("Inserted DOCX page break successfully.")
                    except Exception as e:
                        logger.warning(f"Failed to insert page break: {e}")
                continue

            # Scan a window of 20 DOCX paragraphs ahead
            best_idx = -1
            for offset in range(20):
                i = docx_cursor + offset
                if i >= n_docx:
                    break
                d_clean = docx_entries[i]["clean"]
                if h_clean == d_clean or (d_clean and d_clean in h_clean) or (h_clean and h_clean in d_clean):
                    best_idx = i
                    break

            if best_idx == -1:
                # No match found — advance cursor conservatively
                continue

            para = docx_entries[best_idx]["para"]
            docx_cursor = best_idx + 1

            # Apply the style
            try:
                if not para.style or para.style.name != new_style:
                    para.style = new_style
                    changes_applied += 1
            except Exception as style_err:
                # Style name does not exist in the document's style table.
                # Inject a linked style reference directly via XML so Word
                # will recognise it from its own built-in / attached template.
                try:
                    self._force_style_xml(para, new_style)
                    changes_applied += 1
                except Exception as xml_err:
                    logger.warning(
                        f"Could not apply style '{new_style}' to paragraph "
                        f"'{para.text[:40]}': {xml_err}"
                    )
            
            # Rebuild runs & text (bold, italic, track changes) with username
            # Skip run rebuild if the paragraph contains drawings/images to preserve them losslessly
            has_drawing = False
            for draw_tag in ('w:drawing', 'w:pict'):
                if para._element.find(f'.//{qn(draw_tag)}') is not None:
                    has_drawing = True
                    break

            if not has_drawing:
                try:
                    self._rebuild_paragraph_runs(para, html_entry["element"], username=username)
                except Exception as run_err:
                    logger.warning(f"Could not rebuild runs for paragraph: {run_err}")
            else:
                logger.info(f"Preserved drawing/image runs in paragraph: '{para.text[:30]}'")

        logger.info(
            f"Style patch: {changes_applied} paragraph(s) updated in {docx_path}"
        )

        # 3.5 ── Process table cells ──────────────────────────────────────────
        table_cells_processed = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        text = para.text.strip()
                        if not text:
                            continue
                        para_clean = self._clean(text)

                        # Find matching HTML cell
                        for html_cell in html_table_cells:
                            if para_clean == html_cell["clean"]:
                                # Rebuild cell paragraph with HTML formatting
                                try:
                                    self._rebuild_paragraph_runs(para, html_cell["element"], username=username)
                                    table_cells_processed += 1
                                except Exception as cell_err:
                                    logger.warning(f"Could not rebuild table cell paragraph: {cell_err}")
                                break

        if table_cells_processed > 0:
            logger.info(f"Table cell formatting: {table_cells_processed} cell(s) updated")

        # 4. ── Apply final manuscript formatting (Times New Roman, 12pt, double spacing) ──
        try:
            apply_final_docx_formatting(doc)
        except Exception as fmt_err:
            logger.warning(f"Failed to apply final document formatting: {fmt_err}")

        # 4. ── Save atomically ────────────────────────────────────────────────
        tmp_path = docx_path + ".stylepatch.tmp"
        doc.save(tmp_path)
        
        # Inject w:trackRevisions to Word settings XML so Microsoft Word activates Track Changes by default
        from app.processing.manuscript_core.fixer import enable_track_revisions
        enable_track_revisions(Path(tmp_path))
        
        os.replace(tmp_path, docx_path)
        logger.info(f"Style-patched DOCX saved: {docx_path}")
        return docx_path

    # -------------------------------------------------------------------------
    # Helpers
    # -------------------------------------------------------------------------

    @staticmethod
    def _clean(text: str) -> str:
        """Normalise text for alignment, stripping visual prefix tags like <CN> or <TXT-FLUSH>."""
        import re
        # Remove any leading or visual prefix tag like <TXT-FLUSH> or </BXOBJ>
        cleaned = re.sub(r'<[/A-Za-z0-9_.-]+>', '', text.strip())
        return "".join(c for c in cleaned if c.isalnum()).lower()

    @staticmethod
    def _force_style_xml(para, style_name: str) -> None:
        """
        Forcibly write a <w:pStyle> element into the paragraph's pPr block.
        This allows custom publisher styles (e.g. 'CN', 'ACKTXT') that exist in
        the attached template but are not yet loaded in the document object model.
        """
        from docx.oxml.ns import qn
        from lxml import etree

        pPr = para._p.get_or_add_pPr()
        # Remove any existing pStyle
        for existing in pPr.findall(qn("w:pStyle")):
            pPr.remove(existing)
        # Insert new pStyle as the very first child
        pStyle = etree.SubElement(pPr, qn("w:pStyle"))
        pStyle.set(qn("w:val"), style_name)
        pPr.insert(0, pStyle)

    def _rebuild_paragraph_runs(self, para, html_el, username: str = "WYSIWYG Editor") -> None:
        """
        Clear existing runs in para and rebuild them from html_el,
        supporting bold, italic, underline, color, highlight, font-size,
        superscript, subscript, links, ins (track changes insert), and del (track changes delete).
        Each ins/del element gets a unique w:id and ISO 8601 timestamp.
        """
        import re
        from datetime import datetime
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from lxml import etree
        from docx.opc.constants import RELATIONSHIP_TYPE

        # Track unique revision IDs for this paragraph
        revision_id = [0]  # Use list to allow mutation in nested function

        # 1. Clear all existing runs, ins, del, and hyperlink elements
        p_elem = para._element
        for child in list(p_elem):
            tag_name = etree.QName(child.tag).localname
            if tag_name in ('r', 'ins', 'del', 'hyperlink'):
                p_elem.remove(child)

        # ISO 8601 timestamp with UTC timezone
        timestamp = datetime.utcnow().isoformat() + "Z"

        def parse_style(style_str):
            if not style_str:
                return {}
            styles = {}
            for pair in style_str.split(';'):
                if ':' in pair:
                    key, val = pair.split(':', 1)
                    styles[key.strip().lower()] = val.strip()
            return styles

        # Helper to add a run with formatting
        def add_rich_run(parent_element, text, bold=False, italic=False, underline=False, color=None, bg_color=None, font_size=None, font_name=None, superscript=False, subscript=False, is_link=False, is_del=False):
            if not text:
                return
            r = OxmlElement('w:r')

            # Add run properties if formatted or inside link
            rPr = OxmlElement('w:rPr')
            has_rPr = False

            if bold:
                b = OxmlElement('w:b')
                rPr.append(b)
                has_rPr = True
            if italic:
                i = OxmlElement('w:i')
                rPr.append(i)
                has_rPr = True
            
            # Link style default styling: blue + underline
            final_underline = underline or is_link
            final_color = color
            if is_link and not final_color:
                final_color = "0563C1" # Microsoft Word default hyperlink blue

            if final_underline:
                u = OxmlElement('w:u')
                u.set(qn('w:val'), 'single')
                rPr.append(u)
                has_rPr = True
            if final_color:
                c = OxmlElement('w:color')
                c.set(qn('w:val'), final_color.upper())
                rPr.append(c)
                has_rPr = True
            if bg_color:
                # w:shd for premium customizable highlights
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), bg_color.upper())
                rPr.append(shd)
                has_rPr = True
            if font_size:
                sz_val = str(int(float(font_size) * 2))
                sz = OxmlElement('w:sz')
                sz.set(qn('w:val'), sz_val)
                szCs = OxmlElement('w:szCs')
                szCs.set(qn('w:val'), sz_val)
                rPr.append(sz)
                rPr.append(szCs)
                has_rPr = True
            if font_name:
                rFonts = OxmlElement('w:rFonts')
                rFonts.set(qn('w:ascii'), font_name)
                rFonts.set(qn('w:hAnsi'), font_name)
                rFonts.set(qn('w:cs'), font_name)
                rPr.append(rFonts)
                has_rPr = True
            if superscript:
                va = OxmlElement('w:vertAlign')
                va.set(qn('w:val'), 'superscript')
                rPr.append(va)
                has_rPr = True
            elif subscript:
                va = OxmlElement('w:vertAlign')
                va.set(qn('w:val'), 'subscript')
                rPr.append(va)
                has_rPr = True

            if has_rPr:
                r.append(rPr)

            # Add text node
            t = OxmlElement('w:delText' if is_del else 'w:t')
            if text.startswith(' ') or text.endswith(' '):
                t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
            t.text = text
            r.append(t)

            parent_element.append(r)

        # 2. Traverse HTML element children and text recursively
        def traverse(el, parent_xml, bold=False, italic=False, underline=False, color=None, bg_color=None, font_size=None, font_name=None, superscript=False, subscript=False, is_link=False):
            tag = el.tag.split('}')[-1].lower() if isinstance(el.tag, str) else ""

            style_str = el.get("style", "")
            styles = parse_style(style_str)

            has_bold_style = "font-weight" in styles and styles["font-weight"].strip().lower() in ("bold", "700")
            has_underline_style = "text-decoration" in styles and "underline" in styles["text-decoration"].strip().lower()
            has_italic_style = "font-style" in styles and styles["font-style"].strip().lower() == "italic"

            # Parse text color
            node_color = color
            if "color" in styles:
                c_val = styles["color"].strip()
                if c_val.startswith("#"):
                    node_color = c_val.replace("#", "")
            
            # Parse background color (highlight)
            node_bg = bg_color
            if "background-color" in styles:
                bg_val = styles["background-color"].strip()
                if bg_val.startswith("#"):
                    node_bg = bg_val.replace("#", "")
            elif "background" in styles:
                bg_val = styles["background"].strip()
                if bg_val.startswith("#"):
                    node_bg = bg_val.replace("#", "")

            # Parse font-size
            node_font_size = font_size
            if "font-size" in styles:
                sz_val = styles["font-size"].strip()
                match = re.match(r"(\d+(\.\d+)?)\s*(pt|px)?", sz_val)
                if match:
                    val = float(match.group(1))
                    unit = match.group(3)
                    if unit == "px":
                        node_font_size = round(val * 0.75, 1)
                    else:
                        node_font_size = val

            current_bold = bold or tag in ('strong', 'b') or has_bold_style
            current_italic = italic or tag in ('em', 'i') or has_italic_style
            current_underline = underline or tag == 'u' or has_underline_style
            current_super = superscript or tag == 'sup' or styles.get("vertical-align") == "super"
            current_sub = subscript or tag == 'sub' or styles.get("vertical-align") == "sub"
            current_link = is_link or tag == 'a'

            current_xml_parent = parent_xml
            is_del = False

            if tag == 'ins':
                revision_id[0] += 1
                author = el.get("data-author") or username
                date = el.get("data-date") or timestamp
                ins_node = OxmlElement('w:ins')
                ins_node.set(qn('w:id'), str(revision_id[0]))
                ins_node.set(qn('w:author'), author)
                ins_node.set(qn('w:date'), date)
                parent_xml.append(ins_node)
                current_xml_parent = ins_node
            elif tag == 'del':
                revision_id[0] += 1
                author = el.get("data-author") or username
                date = el.get("data-date") or timestamp
                del_node = OxmlElement('w:del')
                del_node.set(qn('w:id'), str(revision_id[0]))
                del_node.set(qn('w:author'), author)
                del_node.set(qn('w:date'), date)
                parent_xml.append(del_node)
                current_xml_parent = del_node
                is_del = True
            elif tag == 'a':
                href = el.get("href", "")
                if href:
                    try:
                        r_id = para.part.relate_to(href, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
                        link_node = OxmlElement('w:hyperlink')
                        link_node.set(qn('r:id'), r_id)
                        parent_xml.append(link_node)
                        current_xml_parent = link_node
                    except Exception as link_err:
                        logger.warning(f"Could not build hyperlink in python-docx: {link_err}")

            # Parse font-family
            node_font_name = font_name
            if "font-family" in styles:
                node_font_name = styles["font-family"].strip().replace("'", "").replace('"', '')

            # Append direct text
            if el.text:
                add_rich_run(
                    current_xml_parent, el.text,
                    bold=current_bold, italic=current_italic, underline=current_underline,
                    color=node_color, bg_color=node_bg, font_size=node_font_size, font_name=node_font_name,
                    superscript=current_super, subscript=current_sub, is_link=current_link,
                    is_del=is_del
                )

            # Recurse children
            for child in el:
                traverse(
                    child, current_xml_parent,
                    bold=current_bold, italic=current_italic, underline=current_underline,
                    color=node_color, bg_color=node_bg, font_size=node_font_size, font_name=node_font_name,
                    superscript=current_super, subscript=current_sub, is_link=current_link
                )

            # Append tail text
            if el.tail:
                add_rich_run(
                    parent_xml, el.tail,
                    bold=bold, italic=italic, underline=underline,
                    color=color, bg_color=bg_color, font_size=font_size, font_name=font_name,
                    superscript=superscript, subscript=subscript, is_link=is_link,
                    is_del=False
                )

        # Parse root element's initial styles (if any)
        root_style = html_el.get("style", "")
        root_styles = parse_style(root_style)
        root_color = None
        if "color" in root_styles:
            c_val = root_styles["color"].strip()
            if c_val.startswith("#"):
                root_color = c_val.replace("#", "")
        root_bg = None
        if "background-color" in root_styles:
            bg_val = root_styles["background-color"].strip()
            if bg_val.startswith("#"):
                root_bg = bg_val.replace("#", "")
        elif "background" in root_styles:
            bg_val = root_styles["background"].strip()
            if bg_val.startswith("#"):
                root_bg = bg_val.replace("#", "")
        root_font_size = None
        if "font-size" in root_styles:
            sz_val = root_styles["font-size"].strip()
            match = re.match(r"(\d+(\.\d+)?)\s*(pt|px)?", sz_val)
            if match:
                val = float(match.group(1))
                unit = match.group(3)
                if unit == "px":
                    root_font_size = round(val * 0.75, 1)
                else:
                    root_font_size = val
        root_font_name = None
        if "font-family" in root_styles:
            root_font_name = root_styles["font-family"].strip().replace("'", "").replace('"', '')

        # Begin traversal
        if html_el.text:
            add_rich_run(p_elem, html_el.text, color=root_color, bg_color=root_bg, font_size=root_font_size, font_name=root_font_name)

        for child in html_el:
            traverse(child, p_elem, color=root_color, bg_color=root_bg, font_size=root_font_size, font_name=root_font_name)

    # -------------------------------------------------------------------------
    # Legacy fallback — full HTML → DOCX conversion via external tools
    # -------------------------------------------------------------------------

    def _pandoc(self, html_path: str, out_docx_path: str) -> str:
        """Convert XHTML to DOCX using pandoc."""
        tmp_path = out_docx_path + ".tmp"

        try:
            result = subprocess.run(
                [
                    "pandoc",
                    html_path,
                    "-f", "html",
                    "-t", "docx",
                    "--output", tmp_path,
                ],
                capture_output=True,
                text=True,
                timeout=120,
                check=True,
            )

            if not os.path.exists(tmp_path):
                raise RuntimeError(f"pandoc produced no output: {result.stderr}")

            os.replace(tmp_path, out_docx_path)
            logger.info(f"pandoc conversion succeeded: {out_docx_path}")
            return out_docx_path

        except subprocess.TimeoutExpired:
            raise RuntimeError("pandoc conversion timed out (>120s)")
        except subprocess.CalledProcessError as e:
            raise RuntimeError(f"pandoc failed: {e.stderr}")
        finally:
            if os.path.exists(tmp_path):
                os.remove(tmp_path)

    def _libreoffice(self, html_path: str, out_docx_path: str) -> str:
        """Convert XHTML to DOCX using LibreOffice headless."""
        out_dir = os.path.dirname(out_docx_path)
        base_name = Path(html_path).stem

        try:
            result = subprocess.run(
                [
                    "libreoffice",
                    "--headless",
                    "--convert-to", "docx",
                    "--outdir", out_dir,
                    html_path,
                ],
                capture_output=True,
                text=True,
                timeout=120,
                check=True,
            )

            produced = os.path.join(out_dir, f"{base_name}.docx")

            if not os.path.exists(produced):
                raise RuntimeError(
                    f"LibreOffice produced no output at {produced}: {result.stderr}"
                )

            if produced != out_docx_path:
                os.replace(produced, out_docx_path)

            logger.info(f"LibreOffice conversion succeeded: {out_docx_path}")
            return out_docx_path

        except subprocess.TimeoutExpired:
            raise RuntimeError("LibreOffice conversion timed out (>120s)")
        except subprocess.CalledProcessError as e:
            raise RuntimeError(f"LibreOffice failed: {e.stderr}")


def _determine_list_style(li_el) -> str:
    # Walk up ancestors to count list containers
    parent = li_el.getparent()
    list_containers = []
    while parent is not None:
        if parent.tag in ("ul", "ol"):
            list_containers.append(parent.tag)
        parent = parent.getparent()
    
    if not list_containers:
        return "Normal"
    
    # The immediate list container is the first one in the list
    immediate_type = "bullet" if list_containers[0] == "ul" else "number"
    depth = len(list_containers)
    
    base_style = "List Bullet" if immediate_type == "bullet" else "List Number"
    if depth == 1:
        return base_style
    else:
        return f"{base_style} {depth}"


def apply_final_docx_formatting(doc):
    from docx.shared import Pt
    # 1. Update Normal style
    try:
        normal_style = doc.styles['Normal']
        normal_style.font.name = 'Times New Roman'
        normal_style.font.size = Pt(12)
        normal_style.paragraph_format.line_spacing = 2.0
    except Exception as e:
        logger.warning(f"Could not set Normal style formatting: {e}")

    # 2. Iterate over all paragraphs (body + tables)
    # Ensure they are Times New Roman, 12pt, double spacing, unless they have explicit overrides
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else "Normal"
        
        is_list = "list" in style_name.lower() or "bullet" in style_name.lower() or "number" in style_name.lower()
        is_extract = any(x in style_name.lower() for x in ("extract", "quote", "ex"))
        
        if not is_list and not is_extract:
            para.paragraph_format.line_spacing = 2.0
            
        for run in para.runs:
            if not run.font.name:
                run.font.name = 'Times New Roman'
            if not run.font.size:
                run.font.size = Pt(12)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        if not run.font.name:
                            run.font.name = 'Times New Roman'
                        if not run.font.size:
                            run.font.size = Pt(12)
