"""Run-anchored, structural-bookmarked DOCX → XHTML export for the WYSIWYG editor.

Ensures every paragraph, run, table, table cell, footnote, and endnote is mapped
using unique standard Word bookmarks to support precise, lossless, and in-place updates.
"""

import base64
import html
import logging
import uuid
import os
import re

import docx
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

from app.processing.omml_to_mathml import convert_omml_element

logger = logging.getLogger("app.processing.docx_to_xhtml_runs")

# Namespace URI for Office Math (OMML)
_M_NS_URI = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def _omml_element_to_math_span(omml_el, display: str = "inline") -> str:
    """Render an <m:oMath> or <m:oMathPara> as a TipTap math-node span.

    Emits: <span class="math-node" data-mathml="..." data-omml="<b64 xml>"
    data-display="inline|block">MathML</span>

    - data-mathml: inline MathML string, used by the editor to display the eq
      and by the delta engine as a fallback conversion source.
    - data-omml: base64 of the raw OMML XML — used by the delta engine for
      byte-perfect round-trip when the equation was not edited.
    """
    try:
        raw_xml = etree.tostring(omml_el, encoding="utf-8")
        omml_b64 = base64.b64encode(raw_xml).decode("ascii")
    except Exception as e:
        logger.warning(f"Failed to serialize OMML for preservation: {e}")
        omml_b64 = ""

    try:
        mathml_str = convert_omml_element(omml_el)
    except Exception as e:
        logger.warning(f"OMML→MathML conversion failed: {e}")
        # Fallback: render text content so the equation isn't invisible.
        text_fallback = html.escape("".join(omml_el.itertext()).strip() or "[equation]")
        mathml_str = (
            '<math xmlns="http://www.w3.org/1998/Math/MathML">'
            f'<mtext>{text_fallback}</mtext></math>'
        )

    mathml_attr = html.escape(mathml_str, quote=True)
    omml_attr = html.escape(omml_b64, quote=True)

    return (
        f'<span class="math-node" '
        f'data-mathml="{mathml_attr}" '
        f'data-omml="{omml_attr}" '
        f'data-display="{display}">'
        f"{mathml_str}"
        f"</span>"
    )

# Import canonical character styles to build an ID-to-Name map
try:
    from app.processing.reference_char_style_applicator import REFERENCE_CHAR_STYLES
except ImportError:
    REFERENCE_CHAR_STYLES = []

_STYLE_ID_MAP = {}
for _name in REFERENCE_CHAR_STYLES:
    _clean_id = _name.replace("_", "").replace("-", "").lower()
    _STYLE_ID_MAP[_clean_id] = _name


# ─── Track Changes Helpers ────────────────────────────────────────────────────

def _get_run_track_change_info(run):
    """
    Check if a run is inside a <w:ins> or <w:del> element.
    Returns: dict with 'type', 'author', 'date' or None if no track change.
    """
    if run._element is None:
        return None
    parent = run._element.getparent()
    if parent is not None:
        parent_tag = parent.tag
        if parent_tag == qn('w:ins'):
            return {
                'type': 'insertion',
                'author': parent.get(qn('w:author'), ''),
                'date': parent.get(qn('w:date'), ''),
            }
        elif parent_tag == qn('w:del'):
            return {
                'type': 'deletion',
                'author': parent.get(qn('w:author'), ''),
                'date': parent.get(qn('w:date'), ''),
            }
        # Also check grandparent (ins/del might wrap paragraph, not individual runs)
        grandparent = parent.getparent() if parent is not None else None
        if grandparent is not None:
            grandparent_tag = grandparent.tag
            if grandparent_tag == qn('w:ins'):
                return {
                    'type': 'insertion',
                    'author': grandparent.get(qn('w:author'), ''),
                    'date': grandparent.get(qn('w:date'), ''),
                }
            elif grandparent_tag == qn('w:del'):
                return {
                    'type': 'deletion',
                    'author': grandparent.get(qn('w:author'), ''),
                    'date': grandparent.get(qn('w:date'), ''),
                }
    return None


def _get_all_paragraph_runs(para):
    """
    Get all runs from a paragraph, including those inside w:ins, w:del, and w:sdt elements.
    Returns a list of (run, track_change_parent, sdt_alias, sdt_tag) tuples.
    sdt_alias/sdt_tag are set when the run comes from inside an inline w:sdt, else None.
    """
    from docx.text.run import Run

    all_runs = []
    p_elem = para._p

    def _collect_sdt_runs(sdt_elem, outer_alias, outer_tag):
        """Collect all leaf runs from an inline w:sdt, tagged with the outermost alias/tag."""
        sdt_content = sdt_elem.find(qn('w:sdtContent'))
        if sdt_content is None:
            return
        for sdt_child in sdt_content:
            if sdt_child.tag == qn('w:r'):
                all_runs.append((Run(sdt_child, para), None, outer_alias, outer_tag))
            elif sdt_child.tag == qn('w:ins'):
                for r_elem in sdt_child.findall(qn('w:r')):
                    all_runs.append((Run(r_elem, para), sdt_child, outer_alias, outer_tag))
            elif sdt_child.tag == qn('w:del'):
                for r_elem in sdt_child.findall(qn('w:r')):
                    all_runs.append((Run(r_elem, para), sdt_child, outer_alias, outer_tag))
            elif sdt_child.tag == qn('w:sdt'):
                # Nested inline SDT — flatten runs under the outermost alias/tag
                _collect_sdt_runs(sdt_child, outer_alias, outer_tag)

    for child in p_elem:
        child_tag = child.tag

        if child_tag == qn('w:r'):
            all_runs.append((Run(child, para), None, None, None))

        elif child_tag == qn('w:ins'):
            for r_elem in child.findall(qn('w:r')):
                all_runs.append((Run(r_elem, para), child, None, None))

        elif child_tag == qn('w:del'):
            for r_elem in child.findall(qn('w:r')):
                all_runs.append((Run(r_elem, para), child, None, None))

        elif child_tag == qn('w:sdt'):
            alias, tag = _sdt_props(child)
            _collect_sdt_runs(child, alias, tag)

    return all_runs


def _sdt_props(sdt_elem) -> tuple:
    """Extract (alias, tag) strings from a w:sdt element's w:sdtPr. Both default to ''."""
    W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    W_ = "{%s}" % W_NS
    sdtPr = sdt_elem.find(W_ + "sdtPr")
    if sdtPr is None:
        return "", ""
    alias_el = sdtPr.find(W_ + "alias")
    tag_el = sdtPr.find(W_ + "tag")
    alias = alias_el.get(W_ + "val", "") if alias_el is not None else ""
    tag = tag_el.get(W_ + "val", "") if tag_el is not None else ""
    return alias, tag


def _render_runs_with_sdt(items, para, doc, findings_by_para=None, para_idx=0) -> str:
    """
    Render a list of (run, tc_elem, sdt_alias, sdt_tag) items to HTML.
    Consecutive runs sharing the same (sdt_alias, sdt_tag) are wrapped in one
    <span class="sdt-inline"> element. Runs with sdt_alias=None render directly.
    """
    para_findings = findings_by_para.get(para_idx, []) if findings_by_para else []
    parts = []
    i = 0
    while i < len(items):
        run, tc_elem, sdt_alias, sdt_tag = items[i]
        if sdt_alias is None:
            parts.append(_run_to_html(run, para, doc, track_change_element=tc_elem, para_findings=para_findings))
            i += 1
        else:
            # Collect consecutive runs with the same (alias, tag)
            group = []
            while i < len(items) and (items[i][2], items[i][3]) == (sdt_alias, sdt_tag):
                r, tc, _, _ = items[i]
                group.append(_run_to_html(r, para, doc, track_change_element=tc, para_findings=para_findings))
                i += 1
            esc_alias = html.escape(sdt_alias, quote=True)
            esc_tag = html.escape(sdt_tag, quote=True)
            parts.append(
                f'<span class="sdt-inline" data-alias="{esc_alias}" data-tag="{esc_tag}">'
                f'{"".join(group)}</span>'
            )
    return "".join(parts)


# ─── Bookmark Generation Helpers ──────────────────────────────────────────────

def _get_unique_bookmark_id(doc) -> int:
    """Finds a unique bookmark integer ID in the document."""
    if not hasattr(doc, "_next_bookmark_id"):
        used_ids = set()
        for bm_start in doc.element.body.findall(f".//{qn('w:bookmarkStart')}"):
            bm_id = bm_start.get(qn("w:id"))
            if bm_id is not None:
                try:
                    used_ids.add(int(bm_id))
                except ValueError:
                    pass
        
        # Also scan footnote/endnote parts for bookmark IDs
        for rel_id, part in doc.part.related_parts.items():
            if "footnotes" in part.partname or "endnotes" in part.partname:
                try:
                    for bm_start in part._element.findall(f".//{qn('w:bookmarkStart')}"):
                        bm_id = bm_start.get(qn("w:id"))
                        if bm_id is not None:
                            try:
                                used_ids.add(int(bm_id))
                            except ValueError:
                                pass
                except Exception:
                    pass
        doc._next_bookmark_id = max(used_ids) + 1 if used_ids else 1

    next_id = doc._next_bookmark_id
    doc._next_bookmark_id += 1
    return next_id


def _get_or_create_para_bookmark(para, doc, prefix="p_bm_") -> str:
    """Finds or creates a unique bookmark for a paragraph."""
    # Search existing bookmarks in this paragraph
    for elem in para._p:
        if elem.tag == qn("w:bookmarkStart"):
            name = elem.get(qn("w:name"), "")
            if name.startswith(prefix):
                return name

    # Generate new bookmark
    doc._dirty = True
    unique_id = uuid.uuid4().hex[:8]
    bm_name = f"{prefix}{unique_id}"
    next_id = _get_unique_bookmark_id(doc)

    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), str(next_id))
    bm_start.set(qn("w:name"), bm_name)

    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), str(next_id))

    # Fix 4: Insert bookmark after w:pPr to prevent OOXML validation issues
    pPr = para._p.find(qn("w:pPr"))
    if pPr is not None:
        children = list(para._p)
        ppr_idx = children.index(pPr)
        para._p.insert(ppr_idx + 1, bm_start)
        para._p.insert(ppr_idx + 2, bm_end)
    else:
        para._p.insert(0, bm_start)
        para._p.insert(1, bm_end)
    return bm_name


def _get_or_create_run_bookmark(run, para, doc) -> str:
    """Finds or creates a unique bookmark wrapping/preceding a run."""
    r_elem = run._element
    p_children = list(para._p)
    
    # Track-changed runs are wrapped in w:ins/w:del. We search/place bookmarks around the parent element.
    target_elem = r_elem
    parent = r_elem.getparent() if r_elem is not None else None
    if parent is not None and parent.tag in (qn("w:ins"), qn("w:del")):
        target_elem = parent

    try:
        r_idx = p_children.index(target_elem)
        for idx in range(r_idx - 1, -1, -1):
            elem = p_children[idx]
            if elem.tag == qn("w:bookmarkStart"):
                name = elem.get(qn("w:name"), "")
                if name.startswith("r_bm_"):
                    return name
            elif elem.tag in (qn("w:r"), qn("w:ins"), qn("w:del")):
                break
    except ValueError:
        pass

    doc._dirty = True
    unique_id = uuid.uuid4().hex[:8]
    bm_name = f"r_bm_{unique_id}"
    next_id = _get_unique_bookmark_id(doc)

    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), str(next_id))
    bm_start.set(qn("w:name"), bm_name)

    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), str(next_id))

    try:
        r_idx = p_children.index(target_elem)
        para._p.insert(r_idx, bm_start)
        para._p.insert(r_idx + 2, bm_end)
    except ValueError:
        para._p.append(bm_start)
        para._p.append(bm_end)

    return bm_name


def _get_or_create_table_bookmark(table, doc) -> str:
    """Finds or creates a unique bookmark wrapping a table, scoped to its parent container."""
    tbl_elem = table._element
    parent = tbl_elem.getparent()
    parent_children = list(parent) if parent is not None else []
    try:
        t_idx = parent_children.index(tbl_elem)
        for idx in range(t_idx - 1, -1, -1):
            elem = parent_children[idx]
            if elem.tag == qn("w:bookmarkStart"):
                name = elem.get(qn("w:name"), "")
                if name.startswith("tbl_bm_"):
                    return name
            elif elem.tag in (qn("w:tbl"), qn("w:p")):
                break
    except ValueError:
        pass

    doc._dirty = True
    unique_id = uuid.uuid4().hex[:8]
    bm_name = f"tbl_bm_{unique_id}"
    next_id = _get_unique_bookmark_id(doc)

    bm_start = OxmlElement("w:bookmarkStart")
    bm_start.set(qn("w:id"), str(next_id))
    bm_start.set(qn("w:name"), bm_name)

    bm_end = OxmlElement("w:bookmarkEnd")
    bm_end.set(qn("w:id"), str(next_id))

    if parent is not None:
        try:
            t_idx = parent_children.index(tbl_elem)
            parent.insert(t_idx, bm_start)
            parent.insert(t_idx + 2, bm_end)
        except ValueError:
            parent.append(bm_start)
            parent.append(tbl_elem)
            parent.append(bm_end)
    else:
        doc.element.body.append(bm_start)
        doc.element.body.append(tbl_elem)
        doc.element.body.append(bm_end)

    return bm_name


# ─── Formatting Helpers ───────────────────────────────────────────────────────

def _rpr_b64(run) -> str:
    """Base64-encode a run's <w:rPr> element XML."""
    rpr = run._element.find(qn("w:rPr"))
    if rpr is None:
        return ""
    try:
        raw = etree.tostring(rpr)
        return base64.b64encode(raw).decode("ascii")
    except Exception:
        return ""


def _run_inline_style(fmt: dict) -> str:
    """Build a CSS style string for visible run formatting."""
    parts = []
    if fmt["size_pt"] is not None:
        parts.append(f"font-size:{fmt['size_pt']:g}pt")
    if fmt["font_name"]:
        parts.append(f"font-family:'{fmt['font_name']}'")
    if fmt["color_rgb"]:
        parts.append(f"color:#{fmt['color_rgb']}")
    if fmt["highlight"]:
        parts.append("background-color:yellow")
    return ";".join(parts)


def _get_run_formatting(run) -> dict:
    """
    Directly parse formatting from <w:rPr> XML element to bypass slow python-docx property descriptors.
    Returns a dict with: superscript, subscript, underline, italic, bold, size_pt, font_name, color_rgb, highlight
    """
    rpr = run._element.find(qn("w:rPr")) if run._element is not None else None
    res = {
        "superscript": False,
        "subscript": False,
        "underline": False,
        "italic": False,
        "bold": False,
        "size_pt": None,
        "font_name": None,
        "color_rgb": None,
        "highlight": None
    }
    if rpr is None:
        return res

    for child in rpr:
        tag = child.tag
        if tag == qn("w:vertAlign"):
            val = child.get(qn("w:val"))
            if val == "superscript":
                res["superscript"] = True
            elif val == "subscript":
                res["subscript"] = True
        elif tag == qn("w:u"):
            val = child.get(qn("w:val"))
            if val != "none":
                res["underline"] = True
        elif tag == qn("w:i"):
            val = child.get(qn("w:val"))
            if val not in ("false", "0"):
                res["italic"] = True
        elif tag == qn("w:b"):
            val = child.get(qn("w:val"))
            if val not in ("false", "0"):
                res["bold"] = True
        elif tag == qn("w:sz"):
            val = child.get(qn("w:val"))
            if val:
                try:
                    res["size_pt"] = float(val) / 2.0
                except ValueError:
                    pass
        elif tag == qn("w:rFonts"):
            ascii_font = child.get(qn("w:ascii"))
            if ascii_font:
                res["font_name"] = ascii_font
        elif tag == qn("w:color"):
            val = child.get(qn("w:val"))
            if val and val != "auto":
                res["color_rgb"] = val
        elif tag == qn("w:highlight"):
            val = child.get(qn("w:val"))
            if val and val != "none":
                res["highlight"] = val

    return res


def _run_to_html(run, para, doc, track_change_element=None, para_findings=None) -> str:
    """Render a run as a run-anchored span with direct formatting and its run bookmark.

    Args:
        run: The Run object to convert
        para: The parent Paragraph
        doc: The Document
        track_change_element: The w:ins or w:del XML element if this run is inside one, else None
        para_findings: Optional list of scan findings in this paragraph
    """
    text = run.text or ""

    # For deleted runs, python-docx returns empty text. Check for w:delText element directly
    if not text and run._element is not None:
        delText = run._element.find(qn("w:delText"))
        if delText is not None and delText.text:
            text = delText.text
    
    # Render footnote or endnote reference markers
    ftn_ref = run._element.find(qn("w:footnoteReference")) if run._element is not None else None
    etn_ref = run._element.find(qn("w:endnoteReference")) if run._element is not None else None
    if ftn_ref is not None:
        ftn_id = ftn_ref.get(qn("w:id"))
        bm_name = _get_or_create_run_bookmark(run, para, doc)
        return f'<span class="FootnoteRef" data-bookmark="{bm_name}" data-id="{ftn_id}"><sup>{ftn_id}</sup></span>'
    if etn_ref is not None:
        etn_id = etn_ref.get(qn("w:id"))
        bm_name = _get_or_create_run_bookmark(run, para, doc)
        return f'<span class="EndnoteRef" data-bookmark="{bm_name}" data-id="{etn_id}"><sup>{etn_id}</sup></span>'

    if text == "":
        return ""

    inner = html.escape(text).replace("\n", "<br>")

    # Parse formatting directly from XML elements
    fmt = _get_run_formatting(run)

    if fmt["superscript"]:
        inner = f"<sup>{inner}</sup>"
    if fmt["subscript"]:
        inner = f"<sub>{inner}</sub>"
    if fmt["underline"]:
        inner = f"<u>{inner}</u>"
    if fmt["italic"]:
        inner = f"<em>{inner}</em>"
    if fmt["bold"]:
        inner = f"<strong>{inner}</strong>"

    style = _run_inline_style(fmt)
    rpr = _rpr_b64(run)
    bm_name = _get_or_create_run_bookmark(run, para, doc)
    
    # Highlight metadata mapping
    is_yellow_highlighted = (fmt["highlight"] == "yellow")

    finding_replacement = None
    finding_rule_id = None
    finding_category = None
    if is_yellow_highlighted and para_findings:
        for f in para_findings:
            if f.get("surface", "").lower() == text.lower() and f.get("category") in ("bias", "article", "compound", "spelling"):
                finding_replacement = f.get("replacement")
                finding_rule_id = f.get("rule_id")
                finding_category = f.get("category")
                break

    attrs = f' data-run="1" data-bookmark="{bm_name}"'
    if rpr:
        attrs += f' data-rpr="{rpr}"'
    if style:
        attrs += f' style="{html.escape(style, quote=True)}"'
    if finding_replacement is not None:
        attrs += f' data-replacement="{html.escape(finding_replacement, quote=True)}"'
    if finding_rule_id:
        attrs += f' data-rule-id="{html.escape(finding_rule_id, quote=True)}"'
    if finding_category:
        attrs += f' data-rule-category="{html.escape(finding_category, quote=True)}"'
    
    # Export character style name as class if present, plus highlight classes if applicable
    char_style = ""
    rpr_elem = run._element.find(qn("w:rPr")) if run._element is not None else None
    if rpr_elem is not None:
        rstyle = rpr_elem.find(qn("w:rStyle"))
        if rstyle is not None:
            val = rstyle.get(qn("w:val"), "")
            if val:
                val_lower = val.lower()
                if val_lower in _STYLE_ID_MAP:
                    char_style = _STYLE_ID_MAP[val_lower]
                else:
                    try:
                        char_style = doc.styles[val].name
                    except Exception:
                        char_style = val
    if not char_style and run.style and run.style.name != "Default Paragraph Font":
        char_style = run.style.name

    classes = []
    if char_style:
        classes.append(char_style)
    if finding_category:
        classes.append("occurrence-highlight")
        classes.append(f"occurrence-{finding_category}")
    if classes:
        attrs += f' class="{" ".join(html.escape(c, quote=True) for c in classes)}"'

    span_html = f"<span{attrs}>{inner}</span>"

    # Check for track changes and wrap accordingly
    # First try the passed track_change_element (runs inside w:ins/w:del are passed explicitly)
    if track_change_element is not None:
        tc_tag = track_change_element.tag
        if tc_tag == qn('w:ins'):
            author_attr = f' data-author="{html.escape(track_change_element.get(qn("w:author"), ""))}"'
            date_attr = f' data-date="{html.escape(track_change_element.get(qn("w:date"), ""))}"'
            return f"<ins{author_attr}{date_attr}>{span_html}</ins>"
        elif tc_tag == qn('w:del'):
            author_attr = f' data-author="{html.escape(track_change_element.get(qn("w:author"), ""))}"'
            date_attr = f' data-date="{html.escape(track_change_element.get(qn("w:date"), ""))}"'
            return f"<del{author_attr}{date_attr}>{span_html}</del>"

    # Fall back to checking the run's parent hierarchy
    tc_info = _get_run_track_change_info(run)
    if tc_info:
        if tc_info['type'] == 'insertion':
            author_attr = f' data-author="{html.escape(tc_info["author"])}"' if tc_info["author"] else ""
            date_attr = f' data-date="{html.escape(tc_info["date"])}"' if tc_info["date"] else ""
            return f"<ins{author_attr}{date_attr}>{span_html}</ins>"
        elif tc_info['type'] == 'deletion':
            author_attr = f' data-author="{html.escape(tc_info["author"])}"' if tc_info["author"] else ""
            date_attr = f' data-date="{html.escape(tc_info["date"])}"' if tc_info["date"] else ""
            return f"<del{author_attr}{date_attr}>{span_html}</del>"

    return span_html


def _block_sdt_to_html(sdt_elem, doc, body_p_map=None, findings_by_para=None) -> str:
    """Render a block-level w:sdt as <div class="sdt-block" data-alias="..." data-tag="...">."""
    alias, tag = _sdt_props(sdt_elem)
    esc_alias = html.escape(alias, quote=True)
    esc_tag = html.escape(tag, quote=True)
    open_tag = f'<div class="sdt-block" data-alias="{esc_alias}" data-tag="{esc_tag}">'

    sdt_content = sdt_elem.find(qn("w:sdtContent"))
    if sdt_content is None:
        return f'{open_tag}<p><br></p></div>'

    inner_blocks = []
    current_list = []

    for child in sdt_content:
        child_localname = etree.QName(child.tag).localname

        if child_localname == "p":
            para = docx.text.paragraph.Paragraph(child, doc)
            para_idx = body_p_map.get(child, 0) if body_p_map else 0
            is_list, list_type, ilvl = _get_list_info(para)
            if is_list:
                current_list.append((para, para_idx, list_type, ilvl))
            else:
                if current_list:
                    inner_blocks.append(_nested_list_to_html(current_list, doc, findings_by_para=findings_by_para))
                    current_list = []
                inner_blocks.append(_paragraph_to_html(para, para_idx, doc, findings_by_para=findings_by_para))

        elif child_localname == "tbl":
            if current_list:
                inner_blocks.append(_nested_list_to_html(current_list, doc, findings_by_para=findings_by_para))
                current_list = []
            table = docx.table.Table(child, doc)
            inner_blocks.append(_table_to_html(table, doc, body_p_map=body_p_map, findings_by_para=findings_by_para))

        elif child_localname == "sdt":
            if current_list:
                inner_blocks.append(_nested_list_to_html(current_list, doc, findings_by_para=findings_by_para))
                current_list = []
            inner_blocks.append(_block_sdt_to_html(child, doc, body_p_map=body_p_map, findings_by_para=findings_by_para))

    if current_list:
        inner_blocks.append(_nested_list_to_html(current_list, doc, findings_by_para=findings_by_para))

    if not inner_blocks:
        inner_blocks.append("<p><br></p>")

    return f'{open_tag}{"".join(inner_blocks)}</div>'


def _paragraph_content_to_html(p_elem, para, doc, findings_by_para=None, para_idx=0) -> str:
    """
    Recursively renders the children of a paragraph element to HTML.
    Supports nested w:sdt, w:ins, w:del, w:hyperlink, and w:r elements, preserving hierarchy.
    """
    from docx.text.run import Run

    para_findings = findings_by_para.get(para_idx, []) if findings_by_para else []
    parts = []
    
    for child in p_elem:
        qname = etree.QName(child.tag)
        tag_local = qname.localname
        ns_uri = qname.namespace

        # Office Math (OMML): m:oMath (inline) or m:oMathPara (block).
        # Serialize raw OMML for lossless round-trip; render MathML for display.
        if ns_uri == _M_NS_URI and tag_local in ("oMath", "oMathPara"):
            display = "block" if tag_local == "oMathPara" else "inline"
            parts.append(_omml_element_to_math_span(child, display=display))
            continue

        if tag_local == 'r':
            run = Run(child, para)
            parts.append(_run_to_html(run, para, doc, track_change_element=None, para_findings=para_findings))

        elif tag_local == 'ins':
            author_attr = f' data-author="{html.escape(child.get(qn("w:author"), ""))}"' if child.get(qn("w:author")) else ""
            date_attr = f' data-date="{html.escape(child.get(qn("w:date"), ""))}"' if child.get(qn("w:date")) else ""
            inner_html = _paragraph_content_to_html(child, para, doc, findings_by_para, para_idx)
            parts.append(f"<ins{author_attr}{date_attr}>{inner_html}</ins>")
            
        elif tag_local == 'del':
            author_attr = f' data-author="{html.escape(child.get(qn("w:author"), ""))}"' if child.get(qn("w:author")) else ""
            date_attr = f' data-date="{html.escape(child.get(qn("w:date"), ""))}"' if child.get(qn("w:date")) else ""
            inner_html = _paragraph_content_to_html(child, para, doc, findings_by_para, para_idx)
            parts.append(f"<del{author_attr}{date_attr}>{inner_html}</del>")
            
        elif tag_local == 'sdt':
            alias, tag_val = _sdt_props(child)
            esc_alias = html.escape(alias, quote=True)
            esc_tag = html.escape(tag_val, quote=True)
            sdt_content = child.find(qn('w:sdtContent'))
            if sdt_content is not None:
                inner_html = _paragraph_content_to_html(sdt_content, para, doc, findings_by_para, para_idx)
                parts.append(
                    f'<span class="sdt-inline" data-alias="{esc_alias}" data-tag="{esc_tag}">'
                    f'{inner_html}</span>'
                )
            else:
                parts.append(f'<span class="sdt-inline" data-alias="{esc_alias}" data-tag="{esc_tag}"></span>')
                
        elif tag_local == 'hyperlink':
            inner_html = _paragraph_content_to_html(child, para, doc, findings_by_para, para_idx)
            rId = child.get(qn('r:id'))
            if rId and rId in para.part.rels:
                url = para.part.rels[rId].target_ref
                parts.append(f'<a href="{html.escape(url, quote=True)}">{inner_html}</a>')
            else:
                parts.append(inner_html)
                
    return "".join(parts)


def _paragraph_to_html(para, para_idx: int, doc, findings_by_para=None) -> str:
    """Render a standard body paragraph to HTML."""
    style_name = para.style.name if para.style else "Normal"
    tag = "p"
    if style_name.lower().startswith("heading "):
        level = style_name.lower().replace("heading ", "").strip()
        if level in {"1", "2", "3", "4", "5", "6"}:
            tag = f"h{level}"

    runs_html = _paragraph_content_to_html(para._p, para, doc, findings_by_para, para_idx)
    if not runs_html:
        runs_html = "<br>"

    label = html.escape(style_name, quote=True)
    bm_name = _get_or_create_para_bookmark(para, doc)

    return (
        f'<{tag} class="{label}" data-style-label="{label}" data-para-idx="{para_idx}" data-bookmark="{bm_name}">'
        f"{runs_html}</{tag}>"
    )


def _table_to_html(table, doc, body_p_map=None, findings_by_para=None) -> str:
    """Render a table and its cells to HTML with matching table and cell bookmarks."""
    tbl_bm = _get_or_create_table_bookmark(table, doc)
    rows_html = []
    
    for row in table.rows:
        cells_html = []
        for cell in row.cells:
            # We uniquely identify the cell by bookmarking its first paragraph
            cell_bm = ""
            if cell.paragraphs:
                cell_bm = _get_or_create_para_bookmark(cell.paragraphs[0], doc, prefix="cell_bm_")
            
            cell_inner = ""
            for p in cell.paragraphs:
                # Format each paragraph in the cell
                p_style = p.style.name if p.style else "Normal"
                para_idx = body_p_map.get(p._element, 0) if body_p_map else 0
                runs = _paragraph_content_to_html(p._p, p, doc, findings_by_para, para_idx)
                if not runs:
                    runs = "<br>"
                p_bm = _get_or_create_para_bookmark(p, doc)
                cell_inner += f'<p class="{p_style}" data-style-label="{p_style}" data-bookmark="{p_bm}">{runs}</p>'

            attrs = f' data-bookmark="{cell_bm}"' if cell_bm else ""
            cells_html.append(f"<td{attrs}>{cell_inner}</td>")
        rows_html.append(f"<tr>{''.join(cells_html)}</tr>")
    
    return f'<table class="TableGrid" data-bookmark="{tbl_bm}"><tbody>{"".join(rows_html)}</tbody></table>'


def _footnote_endnote_to_html(doc, findings_by_para=None) -> str:
    """Extract footnote and endnote definition content paragraphs and render them to XHTML."""
    blocks = []
    
    # 1. Parse Footnotes
    ftn_part = None
    for rel_id, part in doc.part.related_parts.items():
        if "footnotes" in part.partname:
            ftn_part = part
            break
            
    if ftn_part is not None:
        try:
            for ftn_elem in ftn_part._element.findall(qn("w:footnote")):
                ftn_id = ftn_elem.get(qn("w:id"))
                # Filter out separator/placeholder footnotes
                if ftn_id in (None, "-1", "0"):
                    continue
                
                ftn_paras = []
                for p_elem in ftn_elem.findall(qn("w:p")):
                    para = docx.text.paragraph.Paragraph(p_elem, doc)
                    p_style = para.style.name if para.style else "FootnoteText"
                    para_idx = 10000 + int(ftn_id)
                    all_items = _get_all_paragraph_runs(para)
                    runs = _render_runs_with_sdt(all_items, para, doc, findings_by_para, para_idx) or html.escape(para.text)
                    if not runs:
                        runs = "<br>"
                    p_bm = _get_or_create_para_bookmark(para, doc, prefix="fnpara_bm_")
                    ftn_paras.append(
                        f'<p class="{p_style}" data-style-label="{p_style}" data-bookmark="{p_bm}">{runs}</p>'
                    )
                
                blocks.append(
                    f'<div class="FootnoteContent" data-id="{ftn_id}">'
                    f'{"".join(ftn_paras)}'
                    f'</div>'
                )
        except Exception as e:
            logger.warning(f"Failed to parse footnotes: {e}")

    # 2. Parse Endnotes
    etn_part = None
    for rel_id, part in doc.part.related_parts.items():
        if "endnotes" in part.partname:
            etn_part = part
            break
            
    if etn_part is not None:
        try:
            for etn_elem in etn_part._element.findall(qn("w:endnote")):
                etn_id = etn_elem.get(qn("w:id"))
                if etn_id in (None, "-1", "0"):
                    continue
                
                etn_paras = []
                for p_elem in etn_elem.findall(qn("w:p")):
                    para = docx.text.paragraph.Paragraph(p_elem, doc)
                    p_style = para.style.name if para.style else "EndnoteText"
                    para_idx = 20000 + int(etn_id)
                    all_items = _get_all_paragraph_runs(para)
                    runs = _render_runs_with_sdt(all_items, para, doc, findings_by_para, para_idx) or html.escape(para.text)
                    if not runs:
                        runs = "<br>"
                    p_bm = _get_or_create_para_bookmark(para, doc, prefix="enpara_bm_")
                    etn_paras.append(
                        f'<p class="{p_style}" data-style-label="{p_style}" data-bookmark="{p_bm}">{runs}</p>'
                    )
                
                blocks.append(
                    f'<div class="EndnoteContent" data-id="{etn_id}">'
                    f'{"".join(etn_paras)}'
                    f'</div>'
                )
        except Exception as e:
            logger.warning(f"Failed to parse endnotes: {e}")

    return "\n".join(blocks)


def _get_list_info(para) -> tuple[bool, str, int]:
    """
    Determine if a paragraph is a list item, its type (bullet/number), and level (0-indexed).
    """
    style_name = para.style.name if para.style else "Normal"
    pPr = para._p.pPr
    ilvl = 0
    numId = None
    
    if pPr is not None:
        numPr = pPr.find(qn("w:numPr"))
        if numPr is not None:
            ilvl_el = numPr.find(qn("w:ilvl"))
            numId_el = numPr.find(qn("w:numId"))
            if ilvl_el is not None:
                try:
                    ilvl = int(ilvl_el.get(qn("w:val")))
                except ValueError:
                    pass
            if numId_el is not None:
                numId = numId_el.get(qn("w:val"))
                
    is_list = False
    list_type = "bullet"
    
    if numId is not None:
        is_list = True
        if any(x in style_name.lower() for x in ("number", "num", "enum", "ordered")):
            list_type = "number"
    elif any(x in style_name.lower() for x in ("bullet", "listbullet")):
        is_list = True
        list_type = "bullet"
        match = re.search(r"List Bullet\s*(\d+)", style_name, re.IGNORECASE)
        if match:
            ilvl = int(match.group(1)) - 1
    elif any(x in style_name.lower() for x in ("number", "num", "enum", "ordered")):
        is_list = True
        list_type = "number"
        match = re.search(r"(List Number|List Num|Number)\s*(\d+)", style_name, re.IGNORECASE)
        if match:
            ilvl = int(match.group(2)) - 1
            
    if ilvl < 0:
        ilvl = 0
        
    return is_list, list_type, ilvl


def _nested_list_to_html(list_items, doc, findings_by_para=None) -> str:
    if not list_items:
        return ""
    
    html_parts = []
    stack = []
    
    for para, para_idx, list_type, ilvl in list_items:
        target_tag = "ul" if list_type == "bullet" else "ol"
        
        if not stack:
            html_parts.append(f"<{target_tag}>")
            stack.append((target_tag, 0))
        
        current_level = len(stack) - 1
        
        if ilvl > current_level:
            while len(stack) - 1 < ilvl:
                html_parts.append(f"<{target_tag}>")
                stack.append((target_tag, len(stack)))
        elif ilvl < current_level:
            while len(stack) - 1 > ilvl:
                closed_tag, _ = stack.pop()
                html_parts.append(f"</li></{closed_tag}>")
            if stack and stack[-1][0] != target_tag:
                closed_tag, _ = stack.pop()
                html_parts.append(f"</li></{closed_tag}>")
                html_parts.append(f"<{target_tag}>")
                stack.append((target_tag, ilvl))
            else:
                html_parts.append("</li>")
        else:
            if stack[-1][0] != target_tag:
                closed_tag, _ = stack.pop()
                html_parts.append(f"</li></{closed_tag}>")
                html_parts.append(f"<{target_tag}>")
                stack.append((target_tag, ilvl))
            else:
                html_parts.append("</li>")
                
        style_name = para.style.name if para.style else "Normal"
        label = html.escape(style_name, quote=True)
        bm_name = _get_or_create_para_bookmark(para, doc)
        
        all_items = _get_all_paragraph_runs(para)
        runs_html = _render_runs_with_sdt(all_items, para, doc, findings_by_para, para_idx)
        if not runs_html:
            runs_html = "<br>"

        html_parts.append(
            f'<li>'
            f'<p class="{label}" data-style-label="{label}" data-para-idx="{para_idx}" data-bookmark="{bm_name}">'
            f'{runs_html}'
            f'</p>'
        )
        
    while stack:
        closed_tag, _ = stack.pop()
        html_parts.append(f"</li></{closed_tag}>")
        
    return "\n".join(html_parts)


# ─── Core Converter Engine ───────────────────────────────────────────────────

class DocxToXhtmlRunsEngine:
    """Export a DOCX as run-anchored, dual-bookmarked XHTML for the WYSIWYG editor."""

    def convert(self, docx_path: str, file_id: int | None = None) -> str:
        doc = Document(docx_path)
        doc._dirty = False
        
        # Load scan findings cache if file_id is provided
        findings_by_para = {}
        if file_id is not None:
            try:
                from app.domains.processing.technical_editor_service import RESULTS_DIR
                import json
                from pathlib import Path
                cache_path = RESULTS_DIR / f"{file_id}_scan.json"
                if cache_path.exists():
                    scan_data = json.loads(cache_path.read_text(encoding="utf-8"))
                    findings = scan_data.get("findings", [])
                    for f in findings:
                        p_idx = f.get("para_index")
                        if p_idx is not None:
                            if p_idx not in findings_by_para:
                                findings_by_para[p_idx] = []
                            findings_by_para[p_idx].append(f)
            except Exception as e:
                logger.warning(f"Could not load scan cache for file_id {file_id}: {e}")

        # Build map of body paragraph element -> sequential index (matching extractor.py)
        body_p_map = {}
        for idx, p_elem in enumerate(doc.element.body.iter(qn("w:p"))):
            body_p_map[p_elem] = idx

        blocks = []
        current_list = []

        # Export standard body body blocks
        for element in doc.element.body:
            tag = etree.QName(element.tag).localname
            if tag == "p":
                para = docx.text.paragraph.Paragraph(element, doc)
                para_idx = body_p_map.get(element, 0)

                is_list, list_type, ilvl = _get_list_info(para)
                if is_list:
                    current_list.append((para, para_idx, list_type, ilvl))
                else:
                    if current_list:
                        blocks.append(_nested_list_to_html(current_list, doc, findings_by_para=findings_by_para))
                        current_list = []
                    blocks.append(_paragraph_to_html(para, para_idx, doc, findings_by_para=findings_by_para))
            elif tag == "tbl":
                if current_list:
                    blocks.append(_nested_list_to_html(current_list, doc, findings_by_para=findings_by_para))
                    current_list = []
                table = docx.table.Table(element, doc)
                blocks.append(_table_to_html(table, doc, body_p_map=body_p_map, findings_by_para=findings_by_para))

            elif tag == "sdt":
                if current_list:
                    blocks.append(_nested_list_to_html(current_list, doc, findings_by_para=findings_by_para))
                    current_list = []
                blocks.append(_block_sdt_to_html(element, doc, body_p_map=body_p_map, findings_by_para=findings_by_para))

        # Flush any remaining lists
        if current_list:
            blocks.append(_nested_list_to_html(current_list, doc, findings_by_para=findings_by_para))

        # Append Footnote and Endnote definition contents
        notes_html = _footnote_endnote_to_html(doc, findings_by_para=findings_by_para)
        if notes_html:
            blocks.append(f'<div class="NotesContainer">{notes_html}</div>')

        # Persist all newly generated paragraph & run bookmarks permanently back into the DOCX source file!
        if getattr(doc, "_dirty", False):
            doc.save(docx_path)
            logger.info(f"Assigned structural bookmarks and saved DOCX: {docx_path}")
        else:
            logger.info(f"No new bookmarks added. Skipped saving DOCX: {docx_path}")

        body = "\n".join(blocks)
        return f"<!DOCTYPE html>\n<html><body>{body}</body></html>"
