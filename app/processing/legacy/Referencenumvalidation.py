import re
import os
import sys
import io
import zipfile
from collections import defaultdict

# Add app directory to path for legacy module imports
_app_dir = os.path.dirname(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))
if _app_dir not in sys.path:
    sys.path.insert(0, _app_dir)

from flask import Flask, request, send_file, render_template, redirect, url_for, session
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.text.paragraph import Paragraph
from docx.table import Table
from docx.shared import RGBColor
from app.utils import track_changes
import logging

TRACK_CHANGES_ENABLED = False

app = Flask(__name__)
app.secret_key = "secret_key_for_session_encryption"
UPLOAD_DIR = "temp_reports"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# =====================================================
# Helpers & Core Logic
# =====================================================

def iter_document_paragraphs(doc):
    """
    Iterate through all paragraphs in the document body in order,
    including those inside tables.
    """
    body = doc._element.body
    for child in body:
        if isinstance(child, CT_P):
            yield Paragraph(child, doc)
        elif isinstance(child, CT_Tbl):
            table = Table(child, doc)
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        yield p


_SUPER_TO_NORMAL = {
    '⁰': '0', '¹': '1', '²': '2', '³': '3', '⁴': '4',
    '⁵': '5', '⁶': '6', '⁷': '7', '⁸': '8', '⁹': '9',
    '⁺': '+', '⁻': '-', '⁼': '=', '⁽': '(', '⁾': ')',
    'ⁿ': 'n', '–': '-', '—': '-'
}

_NORMAL_TO_SUPER = {
    '0': '⁰', '1': '¹', '2': '²', '3': '³', '4': '⁴',
    '5': '⁵', '6': '⁶', '7': '⁷', '8': '⁸', '9': '⁹',
    '+': '⁺', '-': '⁻', '=': '⁼', '(': '⁽', ')': '⁾',
    'n': 'ⁿ'
}

def convert_superscript_to_normal(text):
    if not text:
        return ""
    return "".join(_SUPER_TO_NORMAL.get(c, c) for c in text)

def convert_normal_to_superscript(text):
    if not text:
        return ""
    return "".join(_NORMAL_TO_SUPER.get(c, c) for c in text)


def get_numbers(text):
    """
    Extract numbers from text like '1', '2-5', '1, 3, 5'.
    Handles ranges "1-5" -> [1, 2, 3, 4, 5].
    """
    if not text:
        return []
    text = convert_superscript_to_normal(text)
    nums = []
    # Matches: (start)-(end) OR (single)
    # Allows hyphen, en dash, em dash
    pattern = re.compile(r'(\d+)\s*[-–—]\s*(\d+)|(\d+)')
    
    for start, end, single in pattern.findall(text):
        if start and end:
            try:
                s, e = int(start), int(end)
                if s <= e:
                    nums.extend(range(s, e + 1))
            except ValueError:
                pass
        elif single:
            try:
                nums.append(int(single))
            except ValueError:
                pass
    return nums


def format_numbers(nums):
    """
    Format a list of numbers into a string like '1-3, 5'.
    Collapses ranges of 3 or more (e.g. 1,2,3 -> 1-3).
    """
    nums = sorted(set(nums))
    if not nums:
        return ""

    parts = []
    if not nums:
        return ""

    start = prev = nums[0]

    for n in nums[1:]:
        if n == prev + 1:
            prev = n
        else:
            length = prev - start + 1
            if length >= 3:
                parts.append(f"{start}-{prev}")
            elif length == 2:
                parts.append(f"{start},{prev}")
            else:
                parts.append(str(start))
            start = prev = n

    length = prev - start + 1
    if length >= 3:
        parts.append(f"{start}-{prev}")
    elif length == 2:
        parts.append(f"{start},{prev}")
    else:
        parts.append(str(start))

    return ", ".join(parts)


def is_citation_run(run):
    """
    Determine if a run is part of a citation.
    Strictly checks for 'cite_bib' styles.
    """
    if run.style and run.style.name in ["cite_bib"]:
        return True
    return False


BIB_NUMBER_PATTERNS = [
    re.compile(r'^\s*\[(\d+)\]\.\s+'),   # [1].
    re.compile(r'^\s*\((\d+)\)\.\s+'),   # (1).
    re.compile(r'^\s*\[(\d+)\]\s+'),     # [1]
    re.compile(r'^\s*\((\d+)\)\s+'),     # (1)
    re.compile(r'^\s*(\d+)\.\s+'),       # 1.
    re.compile(r'^\s*(\d+)\s+'),         # 1  (space only)
]

# ── New validation patterns ───────────────────────────────────────────────────
_BROKEN_RANGE_PAT = re.compile(
    r'(\d+)\s*[-–—]\s*(?!\d)'  # number-dash with nothing after
    r'|(?<!\d)\s*[-–—]\s*(\d+)'  # dash-number with nothing before
)
_INLINE_TEXT_CITE_PAT = re.compile(
    r'\b(?:see\s+)?(?:references?\s+|refs?\s+|ref\.\s*)(\d+)(?!\s*[a-zA-Z])\b'
    r'|\bsee\s+(\d+)\b',
    re.IGNORECASE
)
_ROMAN_IN_BRACKET = re.compile(
    r'[\[(](i{1,3}|iv|vi{0,3}|ix|x{1,3})[,\-\]]',
    re.IGNORECASE
)


def extract_bib_number(text):
    """Extract bibliography number from text, handling all formats: 1., [1]., (1)., etc."""
    for pat in BIB_NUMBER_PATTERNS:
        m = pat.match(text)
        if m:
            return int(m.group(1)), pat
    return None, None


class ReferenceProcessor:
    def __init__(self, doc, citation_format=None):
        self.doc = doc
        self.citation_format = citation_format or "auto"
        
    def is_citation_run(self, run):
        """
        Determine if a run is part of a citation based on citation_format.
        """
        fmt = self.citation_format
        
        # 1. Always match styled runs
        if run.style and run.style.name in ["cite_bib"]:
            return True
            
        # 2. Check by format
        if fmt == "styled":
            return False
            
        has_unicode_superscript = any(c in "⁰¹²³⁴⁵⁶⁷⁸⁹" for c in (run.text or ""))
        
        if fmt == "superscript":
            if run.font and run.font.superscript:
                return True
            if has_unicode_superscript:
                return True
            return False
            
        elif fmt == "bracket":
            if run.text and re.search(r'\[\s*\d+\s*(?:[-–—]\s*\d+\s*)?(?:,\s*\d+\s*(?:[-–—]\s*\d+\s*)?)*\]', run.text):
                return True
            return False
            
        elif fmt == "paren":
            if run.text and re.search(r'\(\s*\d+\s*(?:[-–—]\s*\d+\s*)?(?:,\s*\d+\s*(?:[-–—]\s*\d+\s*)?)*\)', run.text):
                return True
            return False
            
        elif fmt == "plain":
            if run.text and re.search(r'\b\d+\b', run.text):
                return True
            return False
            
        elif fmt == "auto":
            if run.font and run.font.superscript:
                return True
            if has_unicode_superscript:
                return True
            if run.text and re.search(r'\[\s*\d+\s*(?:[-–—]\s*\d+\s*)?(?:,\s*\d+\s*(?:[-–—]\s*\d+\s*)?)*\]', run.text):
                return True
            if run.text and re.search(r'\(\s*\d+\s*(?:[-–—]\s*\d+\s*)?(?:,\s*\d+\s*(?:[-–—]\s*\d+\s*)?)*\)', run.text):
                return True
            return False
            
        return False
        
    def get_references_in_bibliography(self):
        """
        Returns a Set of IDs found in the bibliography sections (REF-N style).
        Also returns a list of objects for reordering later.
        """
        refs_found = set()
        ref_objects = [] # list of dicts: {'id': int, 'para': p, 'run': r}

        for para in self.doc.paragraphs:
            if para.style and para.style.name == "REF-N":
                found_id = None
                bib_run = None
                
                # Try finding styled run
                for run in para.runs:
                    if run.style and run.style.name == "bib_number":
                        nums = get_numbers(run.text)
                        if nums:
                            found_id = nums[0]
                            bib_run = run
                            break
                
                # Fallback: extract from plain text using multi-format patterns
                if found_id is None:
                    found_id, pattern = extract_bib_number(para.text)

                if found_id is not None:
                    refs_found.add(found_id)
                    ref_objects.append({
                        'id': found_id,
                        'para': para,
                        'run': bib_run
                    })
                    
        return refs_found, ref_objects

    def get_citations_in_text(self):
        """
        Scans document for citations.
        Returns:
            all_cited_ids: list of all IDs in order of appearance (with duplicates)
            appearance_order: list of unique IDs in order of first appearance
        """
        all_cited_ids = []
        appearance_order = []
        seen = set()
        
        for para in iter_document_paragraphs(self.doc):
            if para.style and para.style.name == "REF-N":
                continue
            # 1. Process runs
            current_group = []
            
            for run in para.runs:
                if self.is_citation_run(run):
                    current_group.append(run)
                else:
                    if current_group:
                        # Flush group
                        text = "".join(r.text for r in current_group)
                        nums = get_numbers(text)
                        all_cited_ids.extend(nums)
                        for n in nums:
                            if n not in seen:
                                seen.add(n)
                                appearance_order.append(n)
                        current_group = []
            
            # Flush trailing group
            if current_group:
                text = "".join(r.text for r in current_group)
                nums = get_numbers(text)
                all_cited_ids.extend(nums)
                for n in nums:
                    if n not in seen:
                        seen.add(n)
                        appearance_order.append(n)
                        
        return all_cited_ids, appearance_order

    def find_duplicates(self, ref_objects):
        """
        Finds duplicate references using fuzzy matching (difflib).
        Returns a list of dicts: {'id': int, 'text': str, 'duplicate_of': int, 'score': float}
        """
        import difflib
        
        duplicates = []
        processed_refs = [] # list of (id, clean_text)
        
        # 1. Pre-process all candidates
        for obj in ref_objects:
            full_text = obj['para'].text.strip()
            # Remove leading numbering like "1. ", "[1] "
            clean_text = re.sub(r'^\[?\d+\]?[\.\s]*', '', full_text)
            processed_refs.append({'id': obj['id'], 'text': clean_text})
            
        # 2. Compare O(N^2)
        # We only check forward to avoid double reporting (A=B, B=A)
        # We assume the *earlier* ID is the "original" and later is "duplicate"
        n = len(processed_refs)
        matcher = difflib.SequenceMatcher(None, "", "")
        
        for i in range(n):
            ref_a = processed_refs[i]
            text_a = ref_a['text']
            len_a = len(text_a)
            
            if len_a == 0:
                continue
                
            matcher.set_seq1(text_a)
            
            for j in range(i + 1, n):
                ref_b = processed_refs[j]
                text_b = ref_b['text']
                len_b = len(text_b)
                
                if len_b == 0: 
                    continue
                    
                # Optimization: Length ratio check
                # If lengths differ significantly, they can't be high matches
                # If ratio > 0.85, then min_len / max_len must be roughly > 0.85
                # We use 0.80 as a safe filter threshold.
                if min(len_a, len_b) / max(len_a, len_b) < 0.80:
                    continue
                
                matcher.set_seq2(text_b)

                # Performance Optimization: Check cheap upper bounds first
                if matcher.real_quick_ratio() < 0.80:
                    continue
                if matcher.quick_ratio() < 0.80:
                    continue

                ratio = matcher.ratio()

                # Threshold: 0.80 (80% similar)
                if ratio >= 0.80:
                    duplicates.append({
                        # Legacy keys
                        'id': ref_b['id'],
                        'text': ref_b['text'][:100] + "...",
                        'duplicate_of': ref_a['id'],
                        'score': round(ratio * 100, 1),
                        # Frontend-expected keys
                        'num1': ref_a['id'],
                        'num2': ref_b['id'],
                        'text1': ref_a['text'][:200],
                        'text2': ref_b['text'][:200],
                        'similarity': ratio,
                    })
                    
        return duplicates

    def _collect_broken_ranges(self):
        """Scan citation-styled runs for broken range patterns like [1-3,5-] or [-4]."""
        results = []
        for para in iter_document_paragraphs(self.doc):
            if para.style and para.style.name == "REF-N":
                continue
            current_group = []
            for run in para.runs:
                if self.is_citation_run(run):
                    current_group.append(run)
                else:
                    if current_group:
                        raw = "".join(r.text for r in current_group)
                        for m in _BROKEN_RANGE_PAT.finditer(raw):
                            results.append({
                                "type": "broken_range",
                                "raw": raw.strip(),
                                "match": m.group(0).strip(),
                            })
                        current_group = []
            if current_group:
                raw = "".join(r.text for r in current_group)
                for m in _BROKEN_RANGE_PAT.finditer(raw):
                    results.append({
                        "type": "broken_range",
                        "raw": raw.strip(),
                        "match": m.group(0).strip(),
                    })
        return results

    def detect_invalid_numbers(self, all_cited_ids):
        """Detect citation numbers less than 1 (e.g. [0], [-1])."""
        invalid = []
        seen = set()
        for n in all_cited_ids:
            if n < 1 and n not in seen:
                seen.add(n)
                invalid.append({
                    "type": "invalid_number",
                    "number": n,
                    "message": f"Citation [{n}] is invalid — reference numbers must be ≥ 1.",
                })
        return invalid

    def detect_mixed_citation_styles(self):
        """Detect when both superscript and bracket/paren citation formats coexist in the document."""
        seen_styles = set()
        for para in iter_document_paragraphs(self.doc):
            if para.style and para.style.name == "REF-N":
                continue
            for run in para.runs:
                txt = run.text or ""
                has_super = (
                    (run.font and run.font.superscript) or
                    any(c in "⁰¹²³⁴⁵⁶⁷⁸⁹" for c in txt)
                )
                has_bracket = bool(re.search(r'\[\s*\d+', txt))
                has_paren = bool(re.search(r'\(\s*\d+', txt))
                if has_super:
                    seen_styles.add("superscript")
                if has_bracket:
                    seen_styles.add("bracket")
                if has_paren:
                    seen_styles.add("paren")
        if len(seen_styles) > 1:
            return {
                "type": "mixed_citation_style",
                "styles_found": sorted(seen_styles),
                "message": (
                    f"Mixed citation styles detected: {', '.join(sorted(seen_styles))}. "
                    "Standardize to one format throughout the document."
                ),
            }
        return None

    def detect_inline_text_citations(self):
        """Detect unformatted references like 'see reference 5' or 'ref. 10' outside citation spans."""
        findings = []
        for para in iter_document_paragraphs(self.doc):
            if para.style and para.style.name == "REF-N":
                continue
            for run in para.runs:
                if self.is_citation_run(run):
                    continue
                txt = run.text or ""
                for m in _INLINE_TEXT_CITE_PAT.finditer(txt):
                    num_str = m.group(1) or m.group(2)
                    if num_str:
                        findings.append({
                            "type": "inline_text_citation",
                            "number": int(num_str),
                            "raw": m.group(0),
                            "message": (
                                f"Inline text reference '{m.group(0)}' found outside a citation span. "
                                "Apply the cite_bib character style so this citation is tracked."
                            ),
                        })
        return findings

    def _detect_roman_numerals(self):
        """Detect Roman numeral citations inside brackets/parens, e.g. [iv]."""
        findings = []
        for para in iter_document_paragraphs(self.doc):
            if para.style and para.style.name == "REF-N":
                continue
            txt = para.text or ""
            for m in _ROMAN_IN_BRACKET.finditer(txt):
                findings.append({
                    "type": "roman_numeral_citation",
                    "raw": m.group(0),
                    "message": (
                        f"Possible Roman numeral citation '{m.group(0)}' detected. "
                        "Verify this is intentional or convert to Arabic numerals."
                    ),
                })
        return findings

    def get_validation_stats(self):
        bib_refs, ref_objects = self.get_references_in_bibliography()
        all_cited, _ = self.get_citations_in_text()
        
        unique_cited = set(all_cited)
        
        # Missing: Cited but not in Bib
        missing = sorted(unique_cited - bib_refs)
        
        # Unused: In Bib but not Cited
        unused = sorted(bib_refs - unique_cited)
        
        # Duplicates
        duplicates = self.find_duplicates(ref_objects)
        
        # Sequence Issues
        sequence_issues = []
        seen_in_seq = []
        previous_max = 0
        
        for n in all_cited:
            if n not in seen_in_seq:
                if n < previous_max:
                     pass
                
                if n != len(seen_in_seq) + 1:
                     sequence_issues.append({
                         "position": len(seen_in_seq) + 1,
                         "current": n,
                         "expected": len(seen_in_seq) + 1
                     })
                
                seen_in_seq.append(n)
                previous_max = max(previous_max, n)
                
        broken_ranges = self._collect_broken_ranges()
        invalid_numbers = self.detect_invalid_numbers(all_cited)
        mixed_style = self.detect_mixed_citation_styles()
        inline_cites = self.detect_inline_text_citations()
        roman_cites = self._detect_roman_numerals()

        return {
            "total_references": len(bib_refs),
            "total_citations": len(all_cited),
            "missing_references": missing,
            "unused_references": unused,
            "duplicate_references": duplicates,
            "sequence_issues": sequence_issues,
            "is_perfect": (
                not missing and not unused and not sequence_issues
                and not duplicates and not broken_ranges and not invalid_numbers
            ),
            # New fields
            "broken_ranges": broken_ranges,
            "invalid_numbers": invalid_numbers,
            "mixed_citation_style": mixed_style,
            "inline_text_citations": inline_cites,
            "roman_numeral_citations": roman_cites,
            "summary": {
                "missing_references": len(missing),
                "unused_references": len(unused),
                "sequence_issues": len(sequence_issues),
                "broken_ranges": len(broken_ranges),
                "invalid_numbers": len(invalid_numbers),
                "format_warnings": (1 if mixed_style else 0) + len(inline_cites),
            },
        }

    def renumber(self, ignore_duplicates=False):
        """
        Renumber citations and reorder bibliography.
        Returns: mapping (Old -> New)
        """
        _, appearance_order = self.get_citations_in_text()
        
        # Ensure 'cite_bib' style exists
        from docx.enum.style import WD_STYLE_TYPE
        styles = self.doc.styles
        try:
            styles['cite_bib']
        except KeyError:
            s = styles.add_style('cite_bib', WD_STYLE_TYPE.CHARACTER)
            s.font.superscript = True

        # Get duplicates to resolve duplicate IDs to original IDs
        _, ref_objects = self.get_references_in_bibliography()
        if ignore_duplicates:
            duplicates = []
        else:
            duplicates = self.find_duplicates(ref_objects)
        dup_to_orig = {dup['id']: dup['duplicate_of'] for dup in duplicates}

        # Create Mapping
        mapping = {} 
        new_id = 1
        for old_id in appearance_order:
            # Resolve actual ID by resolving duplicate chain
            actual_old_id = old_id
            while actual_old_id in dup_to_orig:
                actual_old_id = dup_to_orig[actual_old_id]
                
            if actual_old_id not in mapping:
                mapping[actual_old_id] = new_id
                new_id += 1
                
            mapping[old_id] = mapping[actual_old_id]
            
        for para in iter_document_paragraphs(self.doc):
            if para.style and para.style.name == "REF-N":
                continue
            i = 0
            while i < len(para.runs):
                run = para.runs[i]
                
                if self.is_citation_run(run):
                    txt = run.text
                    nums = get_numbers(txt)
                    if nums:
                        new_nums = [mapping.get(n, n) for n in nums]
                        new_nums_formatted = format_numbers(new_nums)
                        
                        # Find the first and last digit in the original text to preserve prefix/suffix
                        match_digits = list(re.finditer(r'[\d⁰¹²³⁴⁵⁶⁷⁸⁹]+', txt))
                        
                        has_orig_super = any(c in "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾ⁿ" for c in (txt or ""))
                        if has_orig_super:
                            new_nums_formatted = convert_normal_to_superscript(new_nums_formatted)
                            
                        if match_digits:
                            first_start = match_digits[0].start()
                            last_end = match_digits[-1].end()
                            prefix = txt[:first_start]
                            suffix = txt[last_end:]
                            new_text = prefix + new_nums_formatted + suffix
                        else:
                            new_text = new_nums_formatted
                        
                        is_renumbered = (nums != new_nums)
                        highlight_color = "008000" if is_renumbered else None
                        
                        style_name = run.style.name if run.style else "cite_bib"
                        
                        if TRACK_CHANGES_ENABLED:
                            # Must replace the whole run
                            track_changes.delete_tracked_run(para, run)
                            
                            run_del = run._element.getparent()
                            anchor = run_del if run_del.tag == track_changes.qn('w:del') else run._element
                            
                            ins_new = track_changes.add_tracked_text(para, new_text, style=style_name, color=highlight_color)
                            anchor.addnext(ins_new)
                        else:
                            run.text = new_text
                            if is_renumbered:
                                run.font.color.rgb = RGBColor(0, 128, 0)

                i += 1

        # 2. Reorder Bibliography
        _, ref_objects = self.get_references_in_bibliography()
        
        # Sort objects into Cited and Uncited
        cited_refs = []
        uncited_refs = []
        
        seen_new_ids = set()
        for obj in ref_objects:
            if obj['id'] in mapping:
                obj['new_id'] = mapping[obj['id']]
                if obj['new_id'] not in seen_new_ids:
                    seen_new_ids.add(obj['new_id'])
                    cited_refs.append(obj)
            else:
                uncited_refs.append(obj)
        
        if not ref_objects:
            return mapping

        # Find anchor (min index)
        body = self.doc._element.body
        
        indices = []
        for obj in ref_objects:
            try:
                idx = body.index(obj['para']._element)
                indices.append(idx)
            except ValueError:
                pass 
        
        if not indices:
            return mapping
            
        anchor = min(indices)
        
        # Remove all
        for obj in ref_objects:
             p = obj['para']._element
             if p.getparent() == body:
                 body.remove(p)
                 
        # Insert Cited (Sorted)
        cited_refs.sort(key=lambda x: x['new_id'])
        
        insert_idx = anchor
        for obj in cited_refs:
            # Update ID text
            if obj['run']:
                old_text = obj['run'].text
                new_text = str(obj['new_id'])
                
                if old_text != new_text:
                    if TRACK_CHANGES_ENABLED:
                        style_name = obj['run'].style.name if obj['run'].style else None
                        
                        track_changes.delete_tracked_run(obj['para'], obj['run'])
                        run_del = obj['run']._element.getparent()
                        anchor = run_del if run_del.tag == track_changes.qn('w:del') else obj['run']._element
                        
                        ins_new = track_changes.add_tracked_text(obj['para'], new_text, style=style_name)
                        anchor.addnext(ins_new)
                    else:
                        obj['run'].text = new_text
            
            body.insert(insert_idx, obj['para']._element)
            insert_idx += 1
            
        # Insert Uncited (Appended after cited)
        for obj in uncited_refs:
            body.insert(insert_idx, obj['para']._element)
            insert_idx += 1
            
        return mapping


_SUPERSCRIPT_CHARS = "⁰¹²³⁴⁵⁶⁷⁸⁹⁺⁻⁼⁽⁾ⁿ"
_PUNCT_SWAP_BRACKET = re.compile(r'([\[(])(\d+)([\])])([.,;:?])')
_PUNCT_SWAP_SUPER = re.compile(r'([' + _SUPERSCRIPT_CHARS + r']+)([.,;:?])')


def ensure_bib_number_style(para, doc):
    text = para.text
    # If the paragraph already has a bib_number styled run, skip rebuilding
    already_styled = False
    for run in para.runs:
        if run.style and run.style.name == "bib_number":
            nums = get_numbers(run.text)
            if nums:
                already_styled = True
                break
    if already_styled:
        return None

    for pat in BIB_NUMBER_PATTERNS:
        m = pat.match(text)
        if m:
            full_match = m.group(0) # e.g. "2. "
            number_str = m.group(1) # e.g. "2"
            
            # Find the start and end of the number within full_match
            num_start = full_match.find(number_str)
            num_end = num_start + len(number_str)
            
            prefix_part = full_match[:num_start] # before digits
            num_part = number_str
            suffix_part = full_match[num_end:]  # after digits (e.g. ". ")
            
            if not para.runs:
                continue
                
            first_run = para.runs[0]
            first_run_text = first_run.text
            
            # Check if the first run text starts with full_match
            if first_run_text.startswith(full_match):
                remaining_text = first_run_text[len(full_match):]
                
                # Split runs: insert prefix, number (styled), suffix, and set first_run = remaining
                p_element = para._element
                first_run_element = first_run._element
                first_run_idx = p_element.index(first_run_element)
                
                def insert_run_before(target_idx, run_text, style_name=None):
                    r = para.add_run(run_text)
                    if style_name:
                        r.style = style_name
                    p_element.remove(r._element)
                    p_element.insert(target_idx, r._element)
                    return r
                
                curr_idx = first_run_idx
                if prefix_part:
                    insert_run_before(curr_idx, prefix_part)
                    curr_idx += 1
                
                num_run = insert_run_before(curr_idx, num_part, 'bib_number')
                curr_idx += 1
                
                if suffix_part:
                    insert_run_before(curr_idx, suffix_part)
                    curr_idx += 1
                
                first_run.text = remaining_text
                return num_run
            else:
                # Fallback: if first run doesn't start with full_match, clear runs and rebuild
                remaining_text = text[len(full_match):]
                para.text = "" # clears all runs
                
                if prefix_part:
                    para.add_run(prefix_part)
                num_run = para.add_run(num_part)
                num_run.style = 'bib_number'
                if suffix_part:
                    para.add_run(suffix_part)
                para.add_run(remaining_text)
                return num_run
    return None


def apply_styles_prep(doc, citation_format="auto"):
    # Ensure styles exist in document
    from docx.enum.style import WD_STYLE_TYPE
    try:
        doc.styles['cite_bib']
    except KeyError:
        s = doc.styles.add_style('cite_bib', WD_STYLE_TYPE.CHARACTER)
        s.font.superscript = True

    try:
        doc.styles['bib_number']
    except KeyError:
        doc.styles.add_style('bib_number', WD_STYLE_TYPE.CHARACTER)

    # 1. Apply 'bib_number' style to references in bibliography
    for para in doc.paragraphs:
        if para.style and para.style.name == "REF-N":
            ensure_bib_number_style(para, doc)

    # 2. Apply 'cite_bib' style to body citations
    processor = ReferenceProcessor(doc, citation_format=citation_format)
    for para in iter_document_paragraphs(doc):
        if para.style and para.style.name == "REF-N":
            continue
        for run in para.runs:
            if processor.is_citation_run(run):
                run.style = 'cite_bib'


def swap_citation_punctuation_in_runs(doc):
    """
    Pre-processing step: in AMA/Vancouver style, citations must appear AFTER
    punctuation (e.g. `.¹` not `¹.`).  Scans every run in the document body and
    swaps the order in-place so that `(1).` becomes `.(1)` and `¹.` becomes `.¹`.
    """
    for para in iter_document_paragraphs(doc):
        for run in para.runs:
            text = run.text
            if not text:
                continue
            # Bracket/paren format: (1). -> .(1)
            text = _PUNCT_SWAP_BRACKET.sub(r'\4\1\2\3', text)
            # Superscript format: ¹. -> .¹
            text = _PUNCT_SWAP_SUPER.sub(r'\2\1', text)
            if text != run.text:
                run.text = text


def process_document(file, citation_format="auto"):
    doc = Document(file)
    swap_citation_punctuation_in_runs(doc)
    
    # 1. Apply styles prep before validation check
    apply_styles_prep(doc, citation_format=citation_format)
    
    # 1. Check BEFORE on original state
    processor1 = ReferenceProcessor(doc, citation_format=citation_format)
    before_stats = processor1.get_validation_stats()
    
    # DECISION:
    # 1. If Unused References exist -> ABORT renumbering.
    if before_stats["unused_references"]:
        return doc, before_stats, before_stats, {}, "Aborted: Document validation failed due to unused references."

    # 3. If Missing Refs -> Can't safely renumber usually
    if before_stats["missing_references"]:
         return doc, before_stats, before_stats, {}, "Aborted: Missing references detected."

    # 2. If Perfect -> No need.
    if before_stats["is_perfect"]:
        return doc, before_stats, before_stats, {}, "Validation completed."

    # ── PASS 1: Reorder Sequence (ignore duplicates) ───────────────────────
    pass1_mapping = processor1.renumber(ignore_duplicates=True)
    
    # Apply styles prep again on the intermediate document state
    apply_styles_prep(doc, citation_format=citation_format)
    
    # ── PASS 2: Find Duplicates & Merge/Renumber ────────────────────────────
    # We must instantiate a new ReferenceProcessor since the document elements might have changed/been reordered
    processor2 = ReferenceProcessor(doc, citation_format=citation_format)
    pass2_mapping = processor2.renumber(ignore_duplicates=False)
    
    # 3. Check AFTER final state
    after_stats = processor2.get_validation_stats()
    
    # Composed mapping composition
    mapping = {}
    for old_id, intermediate_id in pass1_mapping.items():
        mapping[old_id] = pass2_mapping.get(intermediate_id, intermediate_id)
        
    # Count duplicates resolved
    dup_before = len(before_stats.get("duplicate_references", []))
    dup_after = len(after_stats.get("duplicate_references", []))
    duplicates_resolved = dup_before - dup_after
    
    # Determine status message
    changes_made = any(k != v for k, v in mapping.items())
    
    if duplicates_resolved > 0:
        status_msg = f"Two-pass validation: Pass 1 renumbered, Pass 2 {duplicates_resolved} duplicate reference{'s' if duplicates_resolved > 1 else ''} removed and renumbered."
    elif changes_made:
        status_msg = "Renumbering completed successfully."
    else:
        status_msg = "Validation completed."
        
    return doc, before_stats, after_stats, mapping, status_msg


# =====================================================
# Flask Routes
# =====================================================
@app.route("/")
def upload_file():
    return render_template("upload.html")


@app.route("/process", methods=["GET", "POST"])
def process():
    if request.method == "POST":
        file = request.files.get("file")
        if not file or not file.filename.endswith(".docx"):
            return "Invalid file", 400

        doc, before, after, mapping, status_msg = process_document(file)

        base = os.path.splitext(file.filename)[0]
        doc_path = os.path.join(UPLOAD_DIR, f"{base}_renumbered.docx")
        report_path = os.path.join(UPLOAD_DIR, f"{base}_validation.txt")

        doc.save(doc_path)

        with open(report_path, "w", encoding="utf-8") as f:
            f.write(f"STATUS: {status_msg}\n")
            f.write("VALIDATION BEFORE\n")
            f.write(str(before) + "\n\n")
            f.write("VALIDATION AFTER\n")
            f.write(str(after) + "\n\n")
            if mapping:
                f.write("RENUMBERING MAPPING (Old -> New)\n")
                for old, new in sorted(mapping.items(), key=lambda x: x[1]):
                    f.write(f"{old} -> {new}\n")

        # Create ZIP package
        zip_filename = f"{base}_results.zip"
        zip_path = os.path.join(UPLOAD_DIR, zip_filename)
        
        # Validation HTML Report (Offline)
        html_report_filename = f"{base}_results.html"
        html_report_path = os.path.join(UPLOAD_DIR, html_report_filename)
        
        # Render the template for offline use
        # Note: We pass offline_mode=True to make links relative
        html_content = render_template(
            "validation_results.html",
            filename=file.filename,
            results=after,
            before=before,
            mapping=mapping,
            status_msg=status_msg,
            report_file=os.path.basename(report_path),
            doc_file=os.path.basename(doc_path),
            zip_file=None, # No zip button in offline report
            offline_mode=True 
        )
        
        with open(html_report_path, "w", encoding="utf-8") as f:
            f.write(html_content)

        with zipfile.ZipFile(zip_path, 'w') as zf:
             # Add Doc
             zf.write(doc_path, arcname=os.path.basename(doc_path))
             # Add Text Report
             zf.write(report_path, arcname=os.path.basename(report_path))
             # Add HTML Report
             zf.write(html_report_path, arcname=os.path.basename(html_report_path))

        # Store data in session for GET request
        session['processing_result'] = {
            'filename': file.filename,
            'before': before,
            'after': after,
            'mapping': mapping,
            'status_msg': status_msg,
            'report_file': os.path.basename(report_path),
            'doc_file': os.path.basename(doc_path),
            'zip_file': zip_filename
        }
        
        return redirect(url_for('process'))

    # GET request - retrieve from session
    result = session.get('processing_result')
    if not result:
        return redirect(url_for('upload_file'))
        
    return render_template(
        "validation_results.html",
        filename=result['filename'],
        results=result['after'],
        before=result['before'],
        mapping=result['mapping'],
        status_msg=result['status_msg'],
        report_file=result['report_file'],
        doc_file=result['doc_file'],
        zip_file=result.get('zip_file')
    )


@app.route("/download/<path:filename>")
def download_file(filename):
    # Security: Ensure filename is in UPLOAD_DIR
    return send_file(os.path.join(UPLOAD_DIR, filename), as_attachment=True)


if __name__ == "__main__":
    app.run(debug=True, port=5000)
