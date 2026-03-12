import re
import logging
from typing import Dict, Optional
from copy import deepcopy
from typing import Callable, Dict, List, Optional, Tuple
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment

try:
    import pdfplumber
except ImportError:
    pdfplumber = None

# Configure logging
logger = logging.getLogger(__name__)
if not logger.handlers:
    handler = logging.StreamHandler()
    formatter = logging.Formatter('%(asctime)s - %(name)s - %(levelname)s - %(message)s')
    handler.setFormatter(formatter)
    logger.addHandler(handler)
    logger.setLevel(logging.DEBUG)


# ======================================================
# CONFIGURATION CONSTANTS
# ======================================================

# Document structure markers (headings that signal end-of-document sections)
END_OF_DOCUMENT_MARKERS = re.compile(
    r'^\s*(References?|Works?\s+Cited|Bibliography|Bibliographies|Appendix|Appendices|'
    r'Figure\s+Legends?|Table\s+Legends?|Figures?|Tables?)\s*$',
    re.IGNORECASE
)

# Extraction window configurations
FIGURE_CREDIT_WINDOW = 2  # Figure credit: same line or next 1 paragraph
TABLE_CREDIT_WINDOW = 5   # Table credit: inside table or next 4 paragraphs (table note/caption line)
TABLE_FOOTNOTE_WINDOW = 15  # Distinguish table footnote vs box credit
BOX_TITLE_LOOKBACK = 15  # Look back for box title
BOX_TABLE_FOOTNOTE_PROXIMITY = 15  # Avoid confusing table footnotes with box credits

# Text length thresholds
MAX_TITLE_LENGTH = 200  # Skip very long candidates
MIN_CREDIT_LENGTH = 10  # Minimum credit line length


# ======================================================
# REGEX SETUP
# ======================================================

def _setup_regex_patterns() -> Dict[str, re.Pattern]:
    patterns = {}
    patterns['single'] = re.compile(
        r'(?:\(|\b)(Figures?|Figs?\.?|Tables?|Tabs?\.?|Boxes?|Images?|Imgs?\.?|Photos?|Illustrations?)\.?\s*'
        r'([0-9]+(?:[.\-][0-9]+)*)([A-Za-z]?)(?:\)|\b)',
        re.IGNORECASE
    )
    patterns['range'] = re.compile(
        r'(?:\(|\b)(Figures?|Figs?\.?|Tables?|Tabs?\.?|Boxes?|Images?|Imgs?\.?|Photos?|Illustrations?)\.?\s+'
        r'([0-9]+(?:[\.\-][0-9]+)+)([A-Za-z]?)\s*'
        r'(?:to|through|–|—|-)\s*'
        r'([0-9]+(?:[\.\-][0-9]+)*)([A-Za-z]?)(?:\)|\b)',
        re.IGNORECASE
    )
    patterns['and'] = re.compile(
        r'(?:\(|\b)(Figures?|Figs?\.?|Tables?|Tabs?\.?|Boxes?|Images?|Imgs?\.?|Photos?|Illustrations?)\.?\s+'
        r'([0-9]+(?:[\.\-][0-9]+)+)([A-Za-z]?)\s+'
        r'(?:and|&)\s*'
        r'([0-9]+(?:[\.\-][0-9]+)*)([A-Za-z]?)(?:\)|\b)',
        re.IGNORECASE
    )
    patterns['unnumbered'] = re.compile(
        r'(?:\(|\b)(Figures?|Figs?\.?|Tables?|Tabs?\.?|Boxes?|Images?|Imgs?\.?|Photos?|Illustrations?)\.?(?:\s+|$)',
        re.IGNORECASE
    )
    return patterns


CAPTION_PATTERNS = _setup_regex_patterns()

CHAPTER_REGEX = re.compile(r"(?i)^(chapter\s+\d+|\d+\.\s+[A-Z][A-Za-z ].+)")

BOX_START_TAGS = re.compile(r'\\?<BX(?:-TITLE)?>|^\s*\<BX(?:-TITLE)?\>', re.IGNORECASE)
BOX_END_TAG = re.compile(r'\\?</BX>|^\s*\</BX\>', re.IGNORECASE)
BOX_TITLE_TAG = re.compile(r'\\?<BX-TITLE>', re.IGNORECASE)


# ======================================================
# CREDIT DETECTION
# ======================================================

CREDIT_KEYWORDS_REGEX = re.compile(
    r"""(?ix)(
        sources?\s*[:;\s]|
        information\s+from|
        data\s+from|
        adapted\s+(?:with\s+permission\s+)?(?:from|of)|
        modified\s+(?:with\s+permission\s+)?(?:from|of)|
        based\s+on|
        reprinted\s+(?:with\s+permission\s+)?(?:from|of)|
        reproduced\s+(?:with\s+permission\s+)?(?:from|of)|
        redrawn\s+(?:with\s+permission\s+)?(?:from|of)?|
        used\s+with\s+permission|
        with\s+permission|
        courtesy\s+of|
        images?\s+courtesy|
        photo\s+credit|
        illustration\s+by|
        illustrated\s+by|
        shutterstock|
        getty|
        retrieved\s+from|
        accessed\s+from|
        https?://|
        doi\.org/|
        image\s+from|
        from\s+[A-Z][a-z]+.*\[\d{4}\]|
        from\s+[A-Z][a-z]+.*\(\d{4}\)|
        \(\s*from\s+[A-Z][a-z]+|
        \bet\s+al\b.*?\b(?:18|19|20)\d{2}\b|
        \b(?:eds?|editors?)\b\.?.*?\b(?:18|19|20)\d{2}\b|
        \b\d+(?:st|nd|rd|th)\s+ed\.?|
        \b(?:18|19|20)\d{2}\b.*?\b(?:p\.|pp\.|page|pages|fig|figure|vol\.|volume|press)\b|
        \b(?:press|university|wiley|elsevier|springer|lippincott|nature|science|journal)\b.*?\b(?:18|19|20)\d{2}\b|
        \b[A-Z][A-Za-z\-]+,\s*\b(?:18|19|20)\d{2}\b|
        ^[A-Z][a-z]+,\s+[A-Z][a-z]*(?:\s+[A-Z][a-z]*)*.*\(\d{4}\)|
        ^[A-Z][a-z]+\s+et\s+al\.?.*\(\d{4}\)|
        \b[A-Z][A-Za-z\-]+,\s*[A-Z]\..*\(\d{4}\)
    )"""
)

CAPTION_START_REGEX = re.compile(
    r'^\s*(Figures?|Figs?\.?|Tables?|Tabs?\.?|Boxes?|Images?|Imgs?\.?|Photos?|Illustrations?)',
    re.IGNORECASE
)

PERMISSION_RISK_REGEX = re.compile(
    r"""(?i)(
        adapted\s+(?:with\s+permission\s+)?from|
        modified\s+(?:with\s+permission\s+)?from|
        based\s+on|
        reproduced\s+(?:with\s+permission\s+)?from|
        reprinted\s+(?:with\s+permission\s+)?from|
        courtesy\s+of|
        copyright|©|
        from\s+another\s+book|
        journal|press|university|
        https?://|doi\.org/|
        sources?[.\s:;]*
    )""",
    re.VERBOSE
)

STANDALONE_CREDIT_REGEX = re.compile(
    r'^[\*_]?\s*(sources?|information\s+from|data\s+from|adapted\s+(?:with\s+permission\s+)?from|modified\s+(?:with\s+permission\s+)?from|reproduced\s+(?:with\s+permission\s+)?from|reprinted\s+(?:with\s+permission\s+)?from)\b[\s:;]*',
    re.IGNORECASE
)

# For TABLE captions, a post-table paragraph is only accepted as a credit when
# it starts with an explicit credit prefix. This prevents table footnote notes
# like "Note: RDA definitions..." or "based on the adverse effect(s)..." from
# being misidentified as the table's credit/source line.
TABLE_CREDIT_PREFIX_REGEX = re.compile(
    r'^\W*(Source[s]?[:\s]|From\s|Adapted|Modified|Reproduced|Reprinted|Courtesy|Copyright|©)',
    re.IGNORECASE
)

# Detects a credit in the last merged row of a table (e.g. "Reprinted from OSHA...")
_LASTROW_CREDIT_RE = re.compile(
    r'(?i)(reprinted|adapted|modified|reproduced|source[s]?[:\s]|from\s|courtesy|copyright|©|https?://|doi\.)',
)


# ======================================================
# LEGEND SECTION PARSER  (formerly legend_module.py)
# ======================================================

_LEGEND_ENTRY_RE = re.compile(
    r'^(Figures?|Figs?\.?|Tables?|Tabs?\.?|Boxes?|Images?|Imgs?\.?|Photos?|Illustrations?|FIGURE|TABLE|FIG)\.?\s*'
    r'([0-9]+(?:[.\-][0-9]+)*)([A-Za-z]?)\s*[:\-\u2013\u2014]?\s*',
    re.IGNORECASE,
)

_SECTION_HEADER_RE = re.compile(
    r'^\s*(Figure\s+Legends?|FIGURE\s+LEGENDS?|Legends?|LEGENDS?'
    r'|Figure\s+and\s+Table\s+Legends?|TABLES?)\s*$',
    re.IGNORECASE,
)

_LEGEND_CREDIT_INDICATORS_RE = re.compile(
    r'From\s|Copyright|\u00a9|https?://|doi\.'
    r'|[Ww]ith\s+permission|[Cc]ourtesy'
    r'|[Aa]dapted|[Mm]odified|[Rr]eproduced|[Rr]eprinted'
    r'|[Ii]llustration\s+(?:by|:|\u00a9)'
    r'|[A-Z][a-z]+.{0,40}\b\d{4}\b',
)

# Matches a parenthetical block that is truly at the END of the string.
# Inline parens like "(measured) compiled by Author.2023" are rejected
# because extra text follows the closing ")".
_TRAILING_PAREN_RE = re.compile(
    r'\(([^()]*(?:\([^()]*\)[^()]*)*)\)\s*\.?\s*$'
)


def _legend_normalize_type(raw: str) -> str:
    r = raw.lower()
    if 'fig' in r:
        return 'Figure'
    if 'tab' in r:
        return 'Table'
    if 'box' in r:
        return 'Box'
    return 'Figure'


def _split_legend_caption_credit(text: str) -> Tuple[str, str]:
    """
    Split a legend remainder into (caption, credit).
    Credit must be a TRAILING parenthetical containing a credit indicator.
    Inline parens (extra text after ')') are intentionally rejected.
    """
    text = text.strip()
    m = _TRAILING_PAREN_RE.search(text)
    if m:
        paren_block = m.group(0).strip()
        before = text[: m.start()].strip().rstrip('.')
        if _LEGEND_CREDIT_INDICATORS_RE.search(paren_block):
            return before, paren_block
    return text, ''


def _has_unmatched_open(text: str) -> bool:
    """True when text has more '(' than ')' — legend entry continues on next line."""
    return text.count('(') > text.count(')')


def parse_legend_section(paragraphs: List[str]) -> Dict[Tuple[str, str], dict]:
    """
    Scan paragraphs for a 'Figure Legends' / 'TABLES' section header,
    then parse every structured legend entry that follows.

    Multi-paragraph legends (credit wraps to the next line) are auto-merged
    by tracking unmatched open parentheses.

    Returns dict keyed by (item_type, item_no) -> {'caption': str, 'credit': str}
    """
    # 1. Find the legend section header
    legend_start: Optional[int] = None
    for i, para in enumerate(paragraphs):
        if _SECTION_HEADER_RE.match(para.strip()):
            legend_start = i
            break

    if legend_start is None:
        return {}

    logger.debug("Legend section detected at paragraph %d: %r",
                 legend_start, paragraphs[legend_start][:60])

    # 2. Collect and merge continuation lines
    merged: List[str] = []
    for para in paragraphs[legend_start + 1:]:
        if _LEGEND_ENTRY_RE.match(para):
            merged.append(para)
        elif merged and _has_unmatched_open(merged[-1]):
            # Credit wrapped to next line — append until parens balance
            merged[-1] = merged[-1].rstrip() + ' ' + para.strip()
        # else: table data, body text — skip

    logger.debug("Legend section: %d entries after merging", len(merged))

    # 3. Parse each merged legend entry
    results: Dict[Tuple[str, str], dict] = {}
    for legend_text in merged:
        m = _LEGEND_ENTRY_RE.match(legend_text)
        if not m:
            continue

        raw_type  = m.group(1)
        item_no   = m.group(2) + (m.group(3) or '')
        item_type = _legend_normalize_type(raw_type)
        remainder = legend_text[m.end():].strip()

        caption, credit = _split_legend_caption_credit(remainder)
        caption = caption.strip().strip('.')

        results[(item_type, item_no)] = {'caption': caption, 'credit': credit}
        logger.debug("  Parsed legend: %s %s | credit=%s",
                     item_type, item_no, 'YES' if credit else 'no')

    return results


def merge_with_legend(
    inline_results: List[dict],
    legend_map: Dict[Tuple[str, str], dict],
    source_filename: str,
    needs_permission_fn: Callable[[str, str], str],
) -> List[dict]:
    """
    Merge inline extraction results with the authoritative legend map.

    - If (item_type, item_no) exists inline: override caption and credit
      with legend values (legend is always more reliable than lookahead).
    - If a legend entry has no inline match: add it as a new row
      (only when it has a non-empty credit string).
    - Inline results with no matching legend entry are left unchanged.
    """
    if not legend_map:
        return inline_results

    output = deepcopy(inline_results)

    # Index inline results by (item_type, item_no)
    inline_index: Dict[Tuple[str, str], int] = {}
    for idx, r in enumerate(output):
        key = (r['item_type'], r['item_no'])
        inline_index[key] = idx  # last occurrence wins if duplicates exist

    for (item_type, item_no), leg in legend_map.items():
        key        = (item_type, item_no)
        leg_cap    = leg['caption']
        leg_credit = leg['credit']

        if key in inline_index:
            row = output[inline_index[key]]

            # Legend caption wins when it's longer / more complete
            if leg_cap and len(leg_cap) >= len(row.get('caption', '')):
                row['caption'] = leg_cap

            # Legend credit is always authoritative — inline lookahead often
            # grabs wrong body text; the legend section is the ground truth
            if leg_credit:
                row['credit'] = leg_credit
                row['needs_permission'] = needs_permission_fn(row['caption'], leg_credit)

        else:
            # Entry only in legend (never found inline) — add it
            if leg_credit:
                output.append({
                    'chapter':          source_filename,
                    'item_type':        item_type,
                    'item_no':          item_no,
                    'caption':          leg_cap,
                    'credit':           leg_credit,
                    'needs_permission': needs_permission_fn(leg_cap, leg_credit),
                })
                logger.debug("Added legend-only entry: %s %s", item_type, item_no)

    return output


# ======================================================
# END-OF-DOCUMENT SECTION DETECTION
# ======================================================

def find_figures_tables_section_start(paragraphs: List[str]) -> int:
    """
    Detect where Figures/Tables/Boxes section starts in the document.
    
    Returns the paragraph index where the end-of-document section (References,
    Appendix, Figures, etc.) begins. If not found, returns 0 (process entire doc).
    
    This optimization skips processing the main document body and focuses on
    the end-of-document section where figures, tables, and boxes are typically placed.
    """
    for i, para in enumerate(paragraphs):
        if END_OF_DOCUMENT_MARKERS.match(para.strip()):
            logger.info(f"Found end-of-document section at paragraph {i}: {para.strip()[:50]}")
            return i
    
    logger.debug("No explicit end-of-document marker found; processing entire document")
    return 0


# ======================================================
# CHAPTER LOOKUP HELPER
# ======================================================

def get_chapter_at(index: int, chapter_map: dict) -> str:
    """Return the chapter active at `index`."""
    if index in chapter_map:
        return chapter_map[index]
    chapter = ""
    for k in sorted(chapter_map.keys()):
        if k <= index:
            chapter = chapter_map[k]
        else:
            break
    return chapter


# ======================================================
# CREDIT EXTRACTION
# ======================================================

def extract_credit_from_text(text: str) -> tuple:
    """
    Returns (caption_without_credit, credit_string).
    Credit preserves full parenthetical or standalone block.
    """
    text = re.sub(r'\s+', ' ', text).strip()

    if not CREDIT_KEYWORDS_REGEX.search(text):
        return text, ""

    # Strategy 1: standalone credit paragraph (Source:, Information from:, Data from:)
    if STANDALONE_CREDIT_REGEX.match(text):
        return "", _clean_italic_markers(text)

    # Strategy 2: parenthetical credit at end
    paren = _extract_paren_credit(text)
    if paren:
        credit_str, start_idx = paren
        caption = text[:start_idx].strip().rstrip('.')
        return caption, credit_str

    # Strategy 3: inline credit after sentence boundary
    inline = _extract_inline_credit(text)
    if inline:
        credit_str, start_idx = inline
        caption = text[:start_idx].strip().rstrip('.')
        return caption, credit_str

    return text, ""


def _clean_italic_markers(text: str) -> str:
    return re.sub(r'[\*_]', '', text).strip()


def _extract_paren_credit(text: str) -> Optional[tuple]:
    """
    Find top-level parenthetical blocks in the text and return the last one 
    that contains credit keywords. This allows nested parentheses to be
    captured as a single block.
    """
    total_len = len(text)
    blocks = []
    
    # Identify all top-level blocks (those starting at depth 1)
    # Scanning from left-to-right to find balanced pairs.
    idx = 0
    while idx < total_len:
        if text[idx] == '(':
            start = idx
            depth = 1
            idx += 1
            while idx < total_len and depth > 0:
                if text[idx] == '(':
                    depth += 1
                elif text[idx] == ')':
                    depth -= 1
                idx += 1
            
            if depth == 0:
                # Successfully found balanced pair
                block_text = text[start:idx]
                blocks.append((block_text, start))
            else:
                # Unbalanced - rest of string is the block
                block_text = text[start:]
                blocks.append((block_text, start))
                break
        else:
            idx += 1

    # Search blocks in reverse (right-to-left) for keywords
    for block_text, start_idx in reversed(blocks):
        if CREDIT_KEYWORDS_REGEX.search(block_text):
            return block_text, start_idx
            
    return None


def _extract_inline_credit(text: str) -> Optional[tuple]:
    matches = []
    
    pattern = re.compile(
        r'(?<=[.!?])\s+((?:re(?:printed|drawn|produced|created)|adapted|modified|used|from|courtesy)\b.*)',
        re.IGNORECASE | re.DOTALL
    )
    m = pattern.search(text)
    if m and CREDIT_KEYWORDS_REGEX.search(m.group(1)):
        matches.append((m.group(1).strip(), m.start(1)))
        
    src = re.search(r'(sources?\s*:.*)', text, re.IGNORECASE)
    if src:
        matches.append((src.group(1).strip(), src.start(1)))

    # Pattern 4: Academic citation starting with author name(s) and year
    author_pattern = re.compile(
        r'\b([A-Z][a-z]+(?:\s+(?:et\s+al\.?|&\s+[A-Z][a-z]+|[A-Z]\.))*\s*(?:\(\d{4}\)|\.?\s*\d{4}\b).*)',
        re.IGNORECASE
    )
    m = author_pattern.search(text)
    if m and re.search(r'\(\d{4}\)|\d{4}', m.group(1)):
        matches.append((m.group(1).strip(), m.start(1)))

    # Pattern 5: Direct mention of adapted/modified/reproduced anywhere
    direct_pattern = re.search(r'\b((?:adapted|modified|reproduced|reprinted|based)\s+(?:with\s+permission\s+)?(?:from|on|of)\s+.*)', text, re.IGNORECASE)
    if direct_pattern and CREDIT_KEYWORDS_REGEX.search(direct_pattern.group(1)):
        matches.append((direct_pattern.group(1).strip(), direct_pattern.start(1)))

    if not matches:
        return None
        
    best_match = min(matches, key=lambda x: x[1])
    return best_match


def needs_permission(caption, credit):
    check = credit if credit else caption
    if PERMISSION_RISK_REGEX.search(check):
        return "YES"
    return "NO"


# ======================================================
# TEXT EXTRACTION
# ======================================================

from docx.text.paragraph import Paragraph
from docx.table import Table

def _para_text_with_urls(para) -> str:
    """
    Return paragraph text with hyperlinks resolved:
    - If a hyperlink's display text already equals the URL → keep as-is
    - If display text differs from the URL (e.g. journal name or "click here")
      → replace the display text with the actual URL so credit lines always
        contain the real link rather than an anchor label.
    Falls back to para.text for any paragraph with no hyperlinks.
    """
    from docx.oxml.ns import qn as _qn
    # Fast path: no hyperlinks in this paragraph
    if not para._element.findall(".//" + _qn("w:hyperlink")):
        return para.text

    parts = []
    for child in para._element:
        tag = child.tag.split("}")[-1]
        if tag == "r":
            # Plain run — collect all <w:t> text preserving spaces
            parts.append("".join(
                t.text or "" for t in child.findall(_qn("w:t"))
            ))
        elif tag == "hyperlink":
            r_id = child.get(_qn("r:id"))
            display = "".join(
                t.text or "" for t in child.findall(".//" + _qn("w:t"))
            )
            url = None
            if r_id and r_id in para.part.rels:
                url = para.part.rels[r_id].target_ref
            # Use the actual URL only when display text is a different label
            if url and display.strip() != url.strip():
                parts.append(url)
            else:
                parts.append(display)
        # bookmarkStart / bookmarkEnd / proofErr etc. → skip (no text content)
    return "".join(parts)


def extract_text_from_docx(path):
    """
    Returns (paragraphs, table_follows_indices).

    paragraphs            : flat list of paragraph strings (same as before).
    table_follows_indices : set of paragraph indices where the very next
                            top-level document element is a real <tbl>.
                            Used by extract_figures_tables() to distinguish
                            a genuine table caption paragraph from a body
                            sentence that merely references a table by number
                            (e.g. "Table 4.1 reflects the range of ...").
    """
    doc = Document(path)
    paras = []
    table_follows_indices = set()
    body_elts = list(doc.element.body)
    para_idx = 0

    for elt_i, elt in enumerate(body_elts):
        next_elt = body_elts[elt_i + 1] if elt_i + 1 < len(body_elts) else None
        next_is_table = next_elt is not None and next_elt.tag.endswith('tbl')

        if elt.tag.endswith('p'):
            para = Paragraph(elt, doc)
            text = _para_text_with_urls(para).strip()
            if text:
                if next_is_table:
                    table_follows_indices.add(para_idx)
                paras.append(text)
                para_idx += 1

        elif elt.tag.endswith('tbl'):
            table = Table(elt, doc)
            seen_cells = set()
            for row in table.rows:
                for cell in row.cells:
                    if cell in seen_cells:
                        continue
                    seen_cells.add(cell)
                    for p in cell.paragraphs:
                        text = _para_text_with_urls(p).strip()
                        if text:
                            paras.append(text)
                            para_idx += 1

    # Build a set of all paragraph indices that are INSIDE a table's cell content.
    # Also build a map: caption_para_idx -> last-row credit string, for tables
    # where the credit appears as the final merged row of the table itself.
    table_cell_indices = set()
    table_lastrow_credits = {}   # {caption_para_idx: credit_string}
    running_idx = 0
    prev_caption_idx = None      # para_idx of the most recent table caption

    for elt in body_elts:
        if elt.tag.endswith('p'):
            para = Paragraph(elt, doc)
            text = _para_text_with_urls(para).strip()
            if text:
                running_idx += 1
        elif elt.tag.endswith('tbl'):
            table = Table(elt, doc)
            seen_cells2 = set()
            all_cell_texts = []
            for row in table.rows:
                for cell in row.cells:
                    if cell in seen_cells2:
                        continue
                    seen_cells2.add(cell)
                    for p in cell.paragraphs:
                        cell_text = _para_text_with_urls(p).strip()
                        if cell_text:
                            table_cell_indices.add(running_idx)
                            all_cell_texts.append((running_idx, cell_text))
                            running_idx += 1

            # Check if the last row is a single merged cell containing a credit.
            # A credit last-row: spans the full width (only one unique cell in the row),
            # contains credit keywords, and is clearly different from data rows.
            if table.rows and all_cell_texts:
                last_row = table.rows[-1]
                last_seen = set()
                last_unique = []
                for cell in last_row.cells:
                    if cell not in last_seen:
                        last_seen.add(cell)
                        last_unique.append(cell.text.strip())
                # Single merged cell in last row = footnote/credit row
                if len(last_unique) == 1 and last_unique[0]:
                    candidate = last_unique[0]
                    # Accept as credit if it starts with a credit keyword or
                    # contains "reprinted", "adapted", "source", "from", etc.
                    if _LASTROW_CREDIT_RE.search(candidate):
                        # Find the para_idx just before this table started
                        if all_cell_texts:
                            first_cell_idx = all_cell_texts[0][0]
                            # The caption paragraph is the one immediately before
                            # the first cell index (para_idx = first_cell_idx - 1
                            # only if that para is in table_follows_indices)
                            cap_idx = first_cell_idx - 1
                            if cap_idx >= 0:
                                table_lastrow_credits[cap_idx] = candidate

    # Build set of paragraph indices where the immediately PRECEDING top-level
    # body element was a <tbl>.  A standalone "Source:" paragraph in this set
    # is a table footnote, not a box/sidebar credit.
    table_precedes_indices = set()
    running_idx2 = 0
    prev_was_table = False
    for elt in body_elts:
        if elt.tag.endswith('p'):
            para = Paragraph(elt, doc)
            if para.text.strip():
                if prev_was_table:
                    table_precedes_indices.add(running_idx2)
                running_idx2 += 1
            prev_was_table = False
        elif elt.tag.endswith('tbl'):
            seen_cp = set()
            for row in Table(elt, doc).rows:
                for cell in row.cells:
                    if cell not in seen_cp:
                        seen_cp.add(cell)
                        for p in cell.paragraphs:
                            if p.text.strip():
                                running_idx2 += 1
            prev_was_table = True

    # Build set of paragraph indices that are INSIDE a <BX>...</BX> tag block,
    # or whose text contains a <BX> opening tag.
    # Used to reject body sentences like "Box 21.1 summarizes..." that are
    # merely inline references, not real box captions.
    _BX_OPEN  = re.compile(r'<BX[^/>\s]*>', re.IGNORECASE)
    _BX_CLOSE = re.compile(r'</BX>', re.IGNORECASE)
    box_tag_indices = set()   # para indices that contain or follow a <BX> open tag
    in_bx_block = False
    for idx, text in enumerate(paras):
        if _BX_OPEN.search(text):
            in_bx_block = True
            box_tag_indices.add(idx)
        elif _BX_CLOSE.search(text):
            in_bx_block = False
        elif in_bx_block:
            box_tag_indices.add(idx)

    return paras, table_follows_indices, table_cell_indices, table_lastrow_credits, table_precedes_indices, box_tag_indices


def extract_text_from_pdf(path):
    if not pdfplumber:
        raise RuntimeError("pdfplumber not installed")
    paras = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            paras.extend([l.strip() for l in text.split("\n") if l.strip()])
    return paras


# ======================================================
# CAPTION MATCHING
# ======================================================

def match_caption(paragraph):
    if not CAPTION_START_REGEX.match(paragraph):
        return None, None
    for ptype, regex in CAPTION_PATTERNS.items():
        m = regex.match(paragraph)
        if m:
            return ptype, m
    return None, None


def normalize_item_type(raw):
    raw = raw.lower()
    if "fig" in raw: return "Figure"
    if "tab" in raw: return "Table"
    if "box" in raw: return "Box"
    if "image" in raw or "img" in raw or "photo" in raw or "illustration" in raw: return "Figure"
    if "exhibit" in raw: return "Exhibit"
    if "appendix" in raw: return "Appendix"
    return "Case Study"


# ======================================================
# BOX/SIDEBAR EXTRACTION
# ======================================================

def extract_boxes(paragraphs: list, current_chapter_map: dict, source_filename: str = "",
                  table_precedes_indices: set = None,
                  table_follows_indices: set = None) -> list:
    """
    Extract box/sidebar content from document.
    
    Uses two detection paths:
    - Path A: Tag-based detection (<BX>, <BX-TITLE>, </BX>)
    - Path B: Tag-free detection via standalone credit lines (Source:, etc.)
    
    Avoids false positives by:
    - Rejecting table footnotes (within TABLE_FOOTNOTE_WINDOW)
    - Rejecting body text references (e.g., "Box 21.1 summarizes...")
    - Rejecting figure/table credit lines (preceding 10 lines check)
    
    Args:
        paragraphs: List of paragraph strings from document
        current_chapter_map: Mapping of para index to chapter
        source_filename: Name of source file (for attribution)
        table_precedes_indices: Para indices immediately after a table
        table_follows_indices: Para indices immediately before a table
    
    Returns:
        List of extracted box items with caption, credit, and permission status
        Structure: {chapter, item_type, item_no, caption, credit, needs_permission}
    """
    if table_precedes_indices is None:
        table_precedes_indices = set()
    if table_follows_indices is None:
        table_follows_indices = set()
    results = []
    used_indices = set()
    logger.info(f"Starting box extraction from file: {source_filename}")

    for i, para in enumerate(paragraphs):
        # ── Path A: tag-based detection ──────────────────────────────
        if BOX_START_TAGS.search(para):
            chapter = source_filename
            box_lines = []
            j = i
            while j < len(paragraphs):
                line = paragraphs[j]
                if j > i and (BOX_START_TAGS.search(line) or CHAPTER_REGEX.match(line)):
                    break
                if BOX_END_TAG.search(line):
                    j += 1
                    break
                clean = BOX_START_TAGS.sub('', line)
                clean = BOX_END_TAG.sub('', clean)
                clean = BOX_TITLE_TAG.sub('', clean)
                clean = clean.strip()
                if clean:
                    box_lines.append(clean)
                j += 1

            if not box_lines:
                continue

            title = box_lines[0].strip('*_').strip()
            credit_line = ""
            for line in box_lines[1:]:
                _, credit = extract_credit_from_text(line)
                if credit:
                    credit_line = _clean_italic_markers(credit)
                    break

            if title and credit_line:
                used_indices.add(i)
                results.append({
                    "chapter": chapter,
                    "item_type": "Box",
                    "item_no": "Unnumbered",
                    "caption": title,
                    "credit": credit_line,
                    "needs_permission": needs_permission(title, credit_line)
                })
            continue

        # ── Path B: tag-free detection via standalone credit line ─────
        if not STANDALONE_CREDIT_REGEX.match(para):
            continue
        if i in used_indices:
            continue

        # Skip if this standalone credit paragraph is a table footnote.
        # Criteria: within TABLE_FOOTNOTE_WINDOW paragraphs ahead there is
        # another table caption (table_follows_indices entry), meaning this
        # "Source:" line sits in the footnote block between two tables.
        # Also skip if the paragraph immediately precedes a table.
        is_table_footnote = i in table_precedes_indices  # immediately after a table
        if not is_table_footnote:
            for _fw in range(i + 1, min(i + BOX_TABLE_FOOTNOTE_PROXIMITY, len(paragraphs))):
                if _fw in table_follows_indices:
                    # There's a table caption just ahead — this Source: is a footnote
                    is_table_footnote = True
                    break
                if match_caption(paragraphs[_fw])[1] or CHAPTER_REGEX.match(paragraphs[_fw]):
                    # Hit a figure/chapter heading first — not a table footnote
                    break
        if is_table_footnote:
            continue

        # Skip if a figure/table caption appears within the preceding 10 lines
        is_fig_table_credit = False
        for k in range(max(0, i - 10), i):
            if match_caption(paragraphs[k])[1]:
                is_fig_table_credit = True
                break
        if is_fig_table_credit:
            continue

        chapter = source_filename

        # Look back for the box title (up to BOX_TITLE_LOOKBACK lines)
        title = ""
        credit_line = ""
        for k in range(i - 1, max(i - BOX_TITLE_LOOKBACK, -1), -1):
            candidate = paragraphs[k].strip().strip('*_').strip()
            if not candidate:
                continue
            if CHAPTER_REGEX.match(candidate):
                break
            if match_caption(candidate)[1]:
                break
            if STANDALONE_CREDIT_REGEX.match(candidate):
                break
            if len(candidate) > MAX_TITLE_LENGTH:
                continue
            title = candidate
            break

        # Extract credit from current paragraph (Path B)
        if title:
            _, credit_line = extract_credit_from_text(para)
            if credit_line:
                used_indices.add(i)
                results.append({
                    "chapter": chapter,
                    "item_type": "Box",
                    "item_no": "Unnumbered",
                    "caption": title,
                    "credit": credit_line,
                    "needs_permission": needs_permission(title, credit_line)
                })

    logger.info(f"Box extraction complete: {len(results)} boxes found")
    return results


# ======================================================
# FIGURE/TABLE EXTRACTION
# ======================================================

def extract_figures_tables(paragraphs: list, current_chapter_map: dict, source_filename: str = "",
                           table_follows_indices: set = None,
                           table_cell_indices: set = None,
                           table_lastrow_credits: dict = None,
                           box_tag_indices: set = None) -> list:
    """
    Extract figures and tables from document paragraphs.
    
    Uses smart heuristics to distinguish real captions from body text references:
    
    For FIGURES:
    - Matches caption patterns (Figure 1, Fig. 2.1, etc.)
    - Looks for credit in same line or next 1 paragraph (FIGURE_CREDIT_WINDOW)
    - Allows implicit credits (no prefix required)
    
    For TABLES:
    - Validates with table_follows_indices (immediately before <tbl> element)
    - Or checks for explicit credit prefix (Source:, Adapted, From, etc.)
    - Captures credits embedded in last row of table (table_lastrow_credits)
    - Looks ahead up to TABLE_CREDIT_WINDOW paragraphs
    - Rejects table footnotes (Note:, etc.) masquerading as credits
    
    Avoids false positives:
    - Skips body sentences like "Table 4.1 shows..." without structural markers
    - Skips table cell indices (data rows, not credit lines)
    - Skips box captions outside <BX> tags
    
    Args:
        paragraphs: List of paragraph strings from document
        current_chapter_map: Mapping of para index to chapter
        source_filename: Name of source file (for attribution)
        table_follows_indices: Set of para indices where next element is a <tbl>
        table_cell_indices: Set of para indices inside table cell content
        table_lastrow_credits: Dict {caption_para_idx: credit_string} from table last rows
        box_tag_indices: Set of para indices inside/at <BX> tagged blocks
    
    Returns:
        List of extracted figure/table items with caption, credit, and permission status
        Structure: {chapter, item_type, item_no, caption, credit, needs_permission}
        Item types: Figure, Table, Box, Exhibit, Appendix, Case Study
    """
    if table_follows_indices is None:
        table_follows_indices = set()
    if table_cell_indices is None:
        table_cell_indices = set()
    if table_lastrow_credits is None:
        table_lastrow_credits = {}
    if box_tag_indices is None:
        box_tag_indices = set()

    results = []
    logger.info(f"Starting extraction from file: {source_filename}")
    logger.info(f"Total paragraphs to process: {len(paragraphs)}")

    for i, para in enumerate(paragraphs):
        ptype, match = match_caption(para)
        if not match:
            continue

        raw_type = match.group(1)
        item_type = normalize_item_type(raw_type)

        # ── Box caption guard ────────────────────────────────────────────────
        # A Box-type paragraph is only a real caption when it is inside or
        # immediately at a <BX>...</BX> tagged block. Body sentences like
        # "Box 21.1 summarizes what can be done..." are skipped entirely.
        if item_type == "Box" and box_tag_indices and i not in box_tag_indices:
            logger.debug(f"Skipping body-text box reference at para {i}: {para[:60]}")
            continue

        # ── Table caption guard ───────────────────────────────────────────────
        # Accept a Table paragraph as a real caption when:
        #   (a) it is immediately followed by a <tbl> element (table_follows_indices), OR
        #   (b) it has a last-row credit already captured (table_lastrow_credits), OR
        #   (c) it is a list-style table (no <tbl>) where a credit paragraph
        #       starting with "Adapted from:" / "Source:" etc. appears within
        #       the next TABLE_CREDIT_WINDOW paragraphs.
        # Body sentences like "Table 4.1 reflects the range of..." are skipped.
        if item_type == "Table" and table_follows_indices:
            is_real_caption = (
                i in table_follows_indices
                or i in table_lastrow_credits
            )
            if not is_real_caption:
                # Check for list-style table: scan ahead for an explicit credit prefix
                for _k in range(i + 1, min(i + TABLE_CREDIT_WINDOW, len(paragraphs))):
                    _p = paragraphs[_k]
                    if match_caption(_p)[1] or CHAPTER_REGEX.match(_p):
                        break
                    if TABLE_CREDIT_PREFIX_REGEX.match(_p):
                        is_real_caption = True
                        break
            if not is_real_caption:
                logger.debug(f"Skipping body-text table reference at para {i}: {para[:60]}")
                continue

        if ptype == "single":
            item_no = match.group(2) + (match.group(3) or "")
        elif ptype == "range":
            item_no = f"{match.group(2)}{match.group(3) or ''}–{match.group(4)}{match.group(5) or ''}"
        elif ptype == "and":
            item_no = f"{match.group(2)}{match.group(3) or ''} & {match.group(4)}{match.group(5) or ''}"
        else:
            item_no = "Unnumbered"

        chapter = source_filename

        full_text = para[match.end():].strip(" :.-")
        caption, credit_line = extract_credit_from_text(full_text)

        # Debug logging for Table 27.3
        if "27.3" in item_no or "27.3" in caption:
            logger.debug(f"[DEBUG] Processing {item_type} {item_no}")
            logger.debug(f"  Caption: {caption[:60]}")
            logger.debug(f"  Credit from same line: {bool(credit_line)}")

        # ── Last-row credit (credit embedded as final merged row of table) ────
        if not credit_line and i in table_lastrow_credits:
            credit_line = table_lastrow_credits[i]
            logger.debug(f"Using last-row credit for {item_type} {item_no}: {credit_line[:60]}")

        # Look ahead for credit
        # Figures: look only next 1 paragraph (same line or next para)
        # Tables: look up to TABLE_CREDIT_WINDOW paragraphs ahead (table note or credit line)
        lookahead_window = FIGURE_CREDIT_WINDOW if item_type == "Figure" else TABLE_CREDIT_WINDOW
        if not credit_line and item_type != "Figure":
            for j in range(i + 1, min(i + lookahead_window, len(paragraphs))):
                next_p = paragraphs[j]
                if match_caption(next_p)[1] or CHAPTER_REGEX.match(next_p):
                    break
                # Skip paragraphs that are inside a table's cell content — they are
                # part of the table data, not a credit line for this caption.
                if j in table_cell_indices:
                    continue

                candidate_cap, candidate = extract_credit_from_text(next_p)
                is_potential_credit = False

                if candidate:
                    is_potential_credit = True
                    credit_line = _clean_italic_markers(candidate)
                elif not candidate and CREDIT_KEYWORDS_REGEX.search(next_p):
                    is_potential_credit = True
                    credit_line = next_p.strip()
                elif re.match(r'^([A-Z][a-z]+(?:\s+et\s+al\.?|,?\s+[A-Z]\.|\s*&\s*[A-Z][a-z]+)*)\s*\(\d{4}\)', next_p):
                    is_potential_credit = True
                    credit_line = next_p.strip()

                if is_potential_credit:
                    # For Table captions, require an explicit credit prefix on the line
                    # (Source:, From, Adapted, etc.).  This prevents table footnotes
                    # like "Note: RDA definitions..." or "based on the adverse effect..."
                    # from being grabbed as the table credit.
                    if item_type == "Table" and not TABLE_CREDIT_PREFIX_REGEX.match(next_p):
                        credit_line = ""
                        is_potential_credit = False

                if is_potential_credit:
                    is_standalone = STANDALONE_CREDIT_REGEX.search(next_p)
                    is_explicit_prefix = re.match(r'^\W*(Adapted|Modified|Reproduced|Reprinted|Courtesy|Sources?|Source|Neuman|Chasin|Bankaitis)\b', next_p, re.IGNORECASE)
                    is_short_line = len(candidate_cap.strip()) <= 15 if candidate_cap else False

                    # Debug logging for Table 27.3
                    if "27.3" in item_no:
                        logger.debug(f"  [Look-ahead {j}] potential_credit={is_potential_credit}, text: {next_p[:60]}")
                        logger.debug(f"    standalone={is_standalone}, explicit_prefix={is_explicit_prefix}, short_line={is_short_line}")
                        logger.debug(f"    has_keywords={bool(CREDIT_KEYWORDS_REGEX.search(next_p))}")

                    if is_standalone or is_explicit_prefix or is_short_line or CREDIT_KEYWORDS_REGEX.search(next_p):
                        if re.match(r'^sources?[\s:;]', next_p, re.IGNORECASE):
                            source_lines = [credit_line]
                            for k in range(j + 1, min(j + 5, len(paragraphs))):
                                nr = paragraphs[k]
                                if match_caption(nr)[1] or CHAPTER_REGEX.match(nr):
                                    break
                                if nr and (nr[0].isupper() or re.match(r'^[A-Z][a-z]+,', nr) or 'http' in nr or 'doi' in nr.lower()):
                                    source_lines.append(nr)
                                elif not nr.strip():
                                    continue
                                else:
                                    break
                            if len(source_lines) > 1:
                                credit_line = " ".join(source_lines)
                        break

        # NOTE: This fallback is intentionally skipped for Figure and Table items.
        if not credit_line and item_no != "Unnumbered" and item_type not in ["Figure", "Table"]:
            search_str = f"{item_type} {item_no}".lower()
            for para_text in paragraphs:
                if search_str in para_text.lower():
                    candidate_cap, candidate = extract_credit_from_text(para_text)
                    if candidate:
                        credit_line = _clean_italic_markers(candidate)
                        break

        caption = caption.strip().strip('.')

        if caption and credit_line:
            results.append({
                "chapter": chapter,
                "item_type": item_type,
                "item_no": item_no,
                "caption": caption,
                "credit": _clean_italic_markers(credit_line),
                "needs_permission": needs_permission(caption, credit_line)
            })

    items_with_credit = sum(1 for r in results if r["credit"])
    logger.info(f"Extraction complete: {len(results)} items found WITH credit lines")
    return results


# ======================================================
# EXTRACTION RESULT VALIDATION
# ======================================================

def validate_extraction_results(results: List[dict]) -> List[dict]:
    """
    Validate and filter extraction results to ensure quality standards.
    
    Removes entries that don't meet minimum requirements:
    - Missing caption or credit
    - Credit line too short (less than MIN_CREDIT_LENGTH chars)
    - Invalid item types
    
    Args:
        results: List of extracted item dictionaries
    
    Returns:
        Filtered list of valid results with logging of rejections
    """
    valid_results = []
    rejected_count = 0
    
    valid_types = {"Figure", "Table", "Box", "Exhibit", "Appendix", "Case Study"}
    
    for r in results:
        # Validate required fields
        if not r.get('caption') or not r.get('credit'):
            rejected_count += 1
            continue
        
        # Validate credit length
        if len(r['credit'].strip()) < MIN_CREDIT_LENGTH:
            logger.debug(f"Rejected {r['item_type']} {r['item_no']}: credit too short")
            rejected_count += 1
            continue
        
        # Validate item type
        if r.get('item_type') not in valid_types:
            logger.debug(f"Rejected {r['item_type']}: unknown type")
            rejected_count += 1
            continue
        
        valid_results.append(r)
    
    if rejected_count > 0:
        logger.info(f"Validation: rejected {rejected_count} results, kept {len(valid_results)}")
    
    return valid_results


# ======================================================
# MAIN EXTRACTION
# ======================================================

def extract_from_file(path):
    if path.lower().endswith(".pdf"):
        paragraphs = extract_text_from_pdf(path)
        table_follows_indices = set()  # PDFs have no structural table markers
        table_cell_indices = set()
        table_lastrow_credits = {}
        table_precedes_indices = set()
        box_tag_indices = set()
    else:
        paragraphs, table_follows_indices, table_cell_indices, table_lastrow_credits, table_precedes_indices, box_tag_indices = extract_text_from_docx(path)

    # Build chapter map: every paragraph index maps to the current chapter title
    current_chapter_map = {}
    current_chapter = ""
    for i, para in enumerate(paragraphs):
        if CHAPTER_REGEX.match(para):
            m = re.search(r'\d+', para)
            current_chapter = m.group(0) if m else para
        current_chapter_map[i] = current_chapter

    import os
    source_filename = os.path.splitext(os.path.basename(path))[0]

    # Detect end-of-document section (References, Bibliography, Appendix, Figures, etc.)
    # This optimization focuses extraction on the section where figures/tables/boxes are typically placed
    section_start = find_figures_tables_section_start(paragraphs)
    if section_start > 0:
        logger.info(f"Processing paragraphs {section_start} to {len(paragraphs)-1}")

    fig_table_results = extract_figures_tables(paragraphs, current_chapter_map, source_filename,
                                               table_follows_indices, table_cell_indices,
                                               table_lastrow_credits, box_tag_indices)
    box_results = extract_boxes(paragraphs, current_chapter_map, source_filename,
                              table_precedes_indices, table_follows_indices)

    all_results = fig_table_results + box_results

    # ── Legend section: parse and merge authoritative captions/credits ─────────
    # Overrides inline lookahead results and adds legend-only entries (e.g.
    # figures whose captions only appear in the end-of-chapter legend section).
    legend_map = parse_legend_section(paragraphs)
    all_results = merge_with_legend(all_results, legend_map, source_filename, needs_permission)

    # ── Validate results ─────────────────────────────────────────────────────
    all_results = validate_extraction_results(all_results)

    def sort_key(r):
        try:
            ch = int(r["chapter"]) if r["chapter"] else 999
        except ValueError:
            ch = 999
        type_order = {"Figure": 1, "Table": 2, "Box": 3}
        t = type_order.get(r["item_type"], 9)
        try:
            num = float(re.split(r'[.\-–]', r["item_no"].replace("Unnumbered", "0"))[0] or 0)
        except Exception:
            num = 0
        return (ch, t, num)

    all_results.sort(key=sort_key)
    logger.info(f"Final results: {len(all_results)} items after validation and sorting")
    return all_results


# ======================================================
# EXCEL OUTPUT
# ======================================================

def write_permission_log(results, output_file):
    wb = Workbook()
    wb.remove(wb.active)

    wb.create_sheet("Sheet1")
    wb.create_sheet("Sheet2")
    log = wb.create_sheet("Permission Log")

    log["B1"] = "Enter Chap. (Fig/Table/Box)"
    log["C1"] = "#Choose Item Type"
    log["D1"] = "Enter Item #"
    log["G1"] = "Enter Figure Legend or Table/Box Title"
    log["H1"] = "Enter Credit Line from Chapter"
    log["I1"] = "Likely Needs Permission"

    log.column_dimensions['B'].width = 22
    log.column_dimensions['C'].width = 16
    log.column_dimensions['D'].width = 14
    log.column_dimensions['G'].width = 48
    log.column_dimensions['H'].width = 55
    log.column_dimensions['I'].width = 24

    row = 2
    for r in results:
        log[f"B{row}"] = r["chapter"]
        log[f"C{row}"] = r["item_type"]
        log[f"D{row}"] = r["item_no"]
        log[f"G{row}"] = r["caption"]
        log[f"H{row}"] = r["credit"]
        log[f"I{row}"] = r["needs_permission"]

        log[f"G{row}"].alignment = Alignment(wrap_text=True)
        log[f"H{row}"].alignment = Alignment(wrap_text=True)

        row += 1

    wb.save(output_file)