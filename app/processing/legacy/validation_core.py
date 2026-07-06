"""
validation_core.py  –  APA 7th Edition Citation Validator
==========================================================
"""
from __future__ import annotations

import re
import logging
import difflib
import threading
import unicodedata
from collections import defaultdict, namedtuple
from typing import Any, Dict, List, Optional, Set, Tuple
from xml.etree import ElementTree as ET
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

try:
    from utils import track_changes
    TRACK_CHANGES_AVAILABLE = True
except ImportError:
    track_changes = None
    TRACK_CHANGES_AVAILABLE = False

try:
    import urllib.request as _urllib_req
    import urllib.parse   as _urllib_parse
    from lxml import etree as _lxml_etree
    _GROBID_DEPS = True
except ImportError:
    _GROBID_DEPS = False

# ── Constants ────────────────────────────────────────────────────────────────
CITE_STYLE         = "cite_bib"
GREEN              = WD_COLOR_INDEX.BRIGHT_GREEN
YELLOW             = WD_COLOR_INDEX.YELLOW
COMMENT_AUTHOR     = "S4C"
COMMENT_INITIALS   = "S4C"
FUZZY_THRESHOLD    = 0.88
SHORT_NAME_MAX_LEN = 6
NEAR_DUP_THRESHOLD = 0.97
ET_AL_MIN          = 3
BIB_TAG_OPEN       = "<ref-open>"
BIB_TAG_CLOSE      = "<ref-close>"
_NAME_PREFIXES: Set[str] = {
    "van","von","de","del","della","di","du",
    "le","la","los","las","den","der","ter","al",
}

# Named tuple returned by match_citation
MatchResult = namedtuple("MatchResult", ["key", "match_type", "score"])

# ── Year helpers ─────────────────────────────────────────────────────────────
_Y             = r"(?:19|20)\d{2}[a-z]?"
_YEAR_SPECIAL  = r"(?:n\.d\.|in\s+press)"
_YEAR_ANY      = rf"(?:{_Y}|{_YEAR_SPECIAL})"
_YEAR_RANGE_P  = r"(?:(?:19|20)\d{2})/(?:(?:19|20)\d{2})"

# ── Regex patterns ───────────────────────────────────────────────────────────
_LOCATOR_TAIL  = r'(?:,\s*pp?\.\s*\d+(?:[\u2013\u2014\-]\d+)?)?'
_LOCATOR_TAIL  = _LOCATOR_TAIL.replace(r'\u2013', '\u2013').replace(r'\u2014', '\u2014')
_RE_PAREN      = re.compile(
    r'\(([^\s\d][^()]{1,250}?' + _YEAR_ANY + _LOCATOR_TAIL + r')\)',
    re.UNICODE | re.IGNORECASE | re.DOTALL)
_RE_NARRATIVE  = re.compile(
    r'(?<!\()([^\s\d]+(?:\s+et\s+al\.)?)\s+\((' + _YEAR_ANY + r')\)',
    re.UNICODE | re.IGNORECASE | re.DOTALL)
_RE_MULTI_Y_P  = re.compile(
    r'\(([^\s\d;][^();]{1,150}?),\s*((?:' + _Y + r')(?:\s*,\s*(?:' + _Y + r')){1,4})\)',
    re.UNICODE | re.DOTALL)
_RE_MULTI_Y_N  = re.compile(
    r'(?<!\()([^\s\d;]+(?:\s+et\s+al\.)?)\s+\(((?:' + _Y + r')(?:\s*,\s*(?:' + _Y + r')){1,4})\)',
    re.UNICODE | re.DOTALL)
_RE_SECONDARY  = re.compile(
    r'\(([^()]+?),\s*(' + _YEAR_ANY + r')\s*,\s*as\s+cited\s+in\s+([^()]+?),\s*(' + _YEAR_ANY + r')\)',
    re.IGNORECASE)
_RE_CITE_UNIT  = re.compile(
    r'^(.*?)\s*,?\s*(' + _YEAR_ANY + r'|' + _YEAR_RANGE_P + r')$',
    re.DOTALL | re.IGNORECASE)
_RE_LOCATOR    = re.compile(r',\s*(pp?\.\s*\d+(?:[–—\-]\d+)?)', re.IGNORECASE)
_RE_AMA        = re.compile(r'\b([A-Z][a-z]+(?:\s+et\s+al\.)?)\s+((?:19|20)\d{2})\b(?!\s*[,;])')
_RE_BAD_ETAL   = re.compile(r'\(([A-Z][a-z]+)\s+et\s+al\s+((?:19|20)\d{2}[a-z]?)\)')
_RE_MISS_COMMA = re.compile(r'\(([^\s\d][^(),]{0,80}?)\s+((?:19|20)\d{2}[a-z]?)\)')
_RE_YR_RANGE   = re.compile(r'\b((?:19|20)\d{2})/((?:19|20)\d{2})\b')
_RE_BAD_ND     = re.compile(r'\(\s*([^\s\d][^()]{0,80}?),?\s*(n\.d\.?|N\.D\.?|nd\.)\s*\)')
_RE_BAD_INPRES = re.compile(r'\(\s*([^\s\d][^()]{0,80}?),\s*(In\s+Press|IN\s+PRESS|In\s+press)\s*\)')
_RE_ETAL_NOPER = re.compile(r'\bet\s+al(?!\.)\b')
_RE_ETDOT_AL   = re.compile(r'\bet\.\s*al\.?', re.IGNORECASE)
_RE_PAREN_AND  = re.compile(r'\(\s*([^\s\d][^(),]{1,120}?)\s+and\s+([^\s\d][^(),]{1,120}?),\s*((?:19|20)\d{2}[a-z]?)\s*\)', re.IGNORECASE | re.UNICODE | re.DOTALL)
_RE_BIB_ABBREV = re.compile(r'\[([A-Z]{2,8})\]')
_HEADING_RE    = re.compile(r"heading\s*\d|title", re.IGNORECASE)
_FOOTNOTE_RE   = re.compile(r"footnote|endnote", re.IGNORECASE)
_CAPTION_RE    = re.compile(r"caption|figure|table\s*(?:title|caption)", re.IGNORECASE)
_BQUOTE_RE     = re.compile(r"block\s*quote|blockquote|quote", re.IGNORECASE)

_RE_PERSONAL_COMM = re.compile(
    r'\b([A-Z]\.\s+[A-Z][a-z]+|[A-Z][a-z]+(?:\s+[A-Z]\.)?)\s*,\s*'
    r'personal\s+communication\s*,\s*'
    r'(?:January|February|March|April|May|June|July|August|'
    r'September|October|November|December)'
    r'\s+\d{1,2}\s*,\s*(?:19|20)\d{2}',
    re.IGNORECASE,
)

_INITIAL_ONLY_RE = re.compile(r'^([A-Z]\.?\s*-?\s*)+[A-Z]?\.?$')

_ORG_UPPER_RUN = re.compile(r'[A-Z]{3,}')
_ORG_KW_RE     = re.compile(
    r'\b(?:association|board|committee|department|society|institute|organization|'
    r'organisation|council|federation|center|centre|national|international|'
    r'academy|university|college|school|hospital|ministry|agency|commission|'
    r'foundation|group|corporation|inc|ltd|company|union|administration|'
    r'authority|office|bureau|world|network|consortium|congress|assembly)\b',
    re.IGNORECASE)

# FIX-41: regex to target only the year token inside a parenthetical citation
# for surgical replacement (avoids replacing page-number digits or other years
# that happen to share the same four digits).
# Matches ", YYYY" or ", YYYYa" at the end of the citation body, before ")" or locator.
_RE_CITE_YEAR_REPLACE = re.compile(
    r'(,\s*)(' + _Y + r')(\s*(?:,\s*pp?\.\s*\d[\d\-–]*)?(?:\s*\)|\s*;))',
    re.IGNORECASE,
)

# ── Advanced normalisation ────────────────────────────────────────────────────
def _strip_diacritics(s: str) -> str:
    return "".join(c for c in unicodedata.normalize("NFKD", s)
                   if not unicodedata.combining(c))

def _to_smart_quotes(text: str) -> str:
    """Convert straight quotes to curly (smart) quotes."""
    if not text:
        return text
    text = re.sub(r'(^|[\s(\[{])"', '\\1\u201c', text)
    text = text.replace('"', '\u201d')
    text = re.sub(r"(^|[\s(\[{])'", '\\1\u2018', text)
    text = text.replace("'", '\u2019')
    return text


def _norm_apos(s: str) -> str:
    return re.sub(r"['\u2019\u02bc\u2018\u201b\u0060\u00b4]", "'", s)

def _sort_key_name(s: str) -> str:
    s = _strip_diacritics(_norm_apos(s))
    words = s.split()
    return " ".join(
        w for w in words if w.lower().rstrip(".,") not in _NAME_PREFIXES
    ).lower().strip()

def _norm(s: str) -> str:
    s = _strip_diacritics(_norm_apos(s))
    s = re.sub(r"\s*et\s+al\.?\s*", " etal ", s, flags=re.IGNORECASE)
    s = re.sub(r"[.,;&]", " ", s)
    return re.sub(r"\s+", " ", s).strip().lower()

def _first_surname(a: str) -> str:
    p = _norm(a).split()
    return p[0] if p else ""

def _primary_surname(author: str) -> str:
    """Return the primary (non-particle) surname for matching.
    For 'van der Berg, J.' returns 'berg'; for 'García Márquez' returns 'garcia'."""
    parts = _norm(author).split()
    non_prefix = [p for p in parts if p.rstrip(".,") not in _NAME_PREFIXES]
    return non_prefix[0] if non_prefix else (parts[0] if parts else "")

def _primary_year(year: str) -> str:
    """For reprint years like '1900/1953', return the later (primary) year."""
    m = _RE_YR_RANGE.match(year)
    if m:
        y1, y2 = int(m.group(1)), int(m.group(2))
        return str(max(y1, y2))
    return year

def _surname_set(a: str) -> Set[str]:
    return (set(re.findall(r"[a-z\u00c0-\u024f]{2,}", _norm(a)))
            - {"et", "al", "etal", "and", "the"})

def _strip_suffix(year: str) -> str:
    m = re.match(r"^((?:19|20)\d{2})[a-z]?$", year)
    return m.group(1) if m else year

def _acronym_of(text: str) -> str:
    return "".join(w[0].upper() for w in text.split() if w and w[0].isupper())

def _is_organization(full: str) -> bool:
    org_keywords = {
        'association','board','committee','department','society','institute','institutes',
        'organization','organisation','council','federation','center','centre',
        'national','international','academy','academies','university','college',
        'school','hospital','ministry','agency','commission','foundation','group',
        'corporation','inc','ltd','company','union','administration','authority',
        'office','bureau','world','network','consortium','collaborators','alliance',
        'task force','working group','services','congress','assembly','trust',
        'disease control','food and','u.s.','uk','united states','health',
        'study','collaborative','program','programme','initiative','project',
        'collaboration','registry','survey','cohort','trial group',
    }
    if '/' in full:
        return True
    return any(re.search(rf"\b{re.escape(kw)}\b", full, re.IGNORECASE) for kw in org_keywords)

def _para_label(idx: int, para_offset: int = 0) -> str:
    if idx < 0:
        return "?"
    if para_offset > 0 and idx >= para_offset:
        table_num = (idx - para_offset) // 1000 + 1
        cell_num  = (idx - para_offset) % 1000 + 1
        return f"Table {table_num}·cell {cell_num}"
    return str(idx + 1)

# ── BUG-04: Robust author counter ───────────────────────────────────────────
def _count_authors(full: str, et_al_min: int = ET_AL_MIN) -> int:
    if _is_organization(full):
        return 1

    has_et_al = bool(re.search(r'\bet\s+al\b', full, re.IGNORECASE))
    cleaned = re.sub(r'\bet\s+al\.?\b', '', full, flags=re.IGNORECASE).strip()

    # Normalize missing comma between a single initial and the next surname,
    # e.g. "Forde, D. Fang, M.L." → "Forde, D., Fang, M.L."
    cleaned = re.sub(r'([A-Z]\.)\s+([A-Z][a-z])', r'\1, \2', cleaned)

    normalised = re.sub(r'\s*[&]\s*|\s+and\s+', ', ', cleaned, flags=re.IGNORECASE)

    raw_parts = re.split(r',\s*', normalised)
    parts = []
    i = 0
    while i < len(raw_parts):
        part = raw_parts[i].strip()
        if not part:
            i += 1
            continue
        if (i + 1 < len(raw_parts)
                and re.match(r'^[A-Z]\.', raw_parts[i + 1].strip())):
            i += 1
        parts.append(part)
        i += 1

    count = 0
    for part in parts:
        part = part.strip().rstrip('.,')
        if not part:
            continue
        if _INITIAL_ONLY_RE.match(part):
            continue
        count += 1

    if has_et_al:
        count = max(count + 1, et_al_min)

    return max(count, 1)

def _is_org_match(cite_auth: str, bib_auth: str) -> bool:
    c_parts = [p.strip() for p in re.split(r',|&', cite_auth) if p.strip()]
    if not c_parts:
        return False

    b_norm = _norm(bib_auth)

    def match_part(cp):
        cp_norm = _norm(cp)
        if re.search(r'\b' + re.escape(cp_norm) + r'\b', b_norm):
            return True
        cp_up = cp.upper().replace('.', '').replace(' ', '')

        b_words = re.sub(r'U\.S\.', 'US ', bib_auth)
        b_words = re.sub(r'U\.K\.', 'UK ', b_words)
        b_words = re.split(r'[\s.,;]+', b_words)

        strict = ''
        loose = ''
        for w in b_words:
            if w in ('US', 'UK'):
                strict += w; loose += w
            elif w:
                if w[0].isupper() and w.lower() not in {'of','and','the','for','in','on'}:
                    strict += w[0].upper()
                if w[0].isalpha():
                    loose += w[0].upper()

        if cp_up in strict or cp_up in loose:
            return True

        b_all = [w for w in b_words if w.lower() not in {'and','for','of','the','in','on'}]
        ac2 = ''.join(w[0].upper() for w in b_all if w)
        if cp_up in ac2:
            return True
        if len(cp_up) >= 2 and any(a.startswith(cp_up) for a in [strict, loose, ac2]):
            return True
        return False

    return all(match_part(cp) for cp in c_parts)

# ── Locator validator ─────────────────────────────────────────────────────────
def check_locator(loc: str) -> List[str]:
    probs: List[str] = []
    is_range   = bool(re.search(r"\d+[–—\-]\d+", loc))
    uses_pp    = loc.lower().startswith("pp")
    uses_p     = loc.lower().startswith("p.") and not uses_pp
    has_endash = "–" in loc or "—" in loc
    has_hyphen = bool(re.search(r"\d-\d", loc))
    has_space  = bool(re.match(r"pp?\.\s\d", loc, re.IGNORECASE))
    if is_range  and uses_p:    probs.append("use 'pp.' (not 'p.') for ranges")
    if not is_range and uses_pp: probs.append("use 'p.' (not 'pp.') for single page")
    if is_range  and has_hyphen and not has_endash:
        probs.append("use en-dash (–) not hyphen for ranges")
    if not has_space:
        spaced = re.sub(r"\.(?=\d)", ". ", loc)
        probs.append(f"add space after period: e.g. '{spaced}'")
    return probs

# ── APA Fixer ─────────────────────────────────────────────────────────────────
class ApaFixer:
    @staticmethod
    def needs_fix(text: str) -> bool:
        return bool(
            _RE_BAD_ETAL.search(text)
            or _RE_MISS_COMMA.search(text) or _RE_YR_RANGE.search(text)
            or _RE_BAD_ND.search(text) or _RE_BAD_INPRES.search(text)
            or _RE_ETAL_NOPER.search(text) or _RE_ETDOT_AL.search(text)
            or _RE_PAREN_AND.search(text)
        )

    @staticmethod
    def fix(text: str) -> Tuple[str, List[Dict]]:
        changes: List[Dict] = []
        r = text

        def _chg(orig, fixed, ft):
            if orig != fixed:
                changes.append({"original": orig, "fixed": fixed, "fix_type": ft})
            return fixed

        def _ama(m):
            return _chg(m.group(0), f"{m.group(1)} ({m.group(2)})", "ama_to_apa_narrative")

        def _comma(m):
            orig, inner, year = m.group(0), m.group(1).rstrip(), m.group(2)
            if (inner.endswith(",")
                    or re.search(r"et\s+al", inner, re.IGNORECASE)
                    or _ORG_UPPER_RUN.search(inner)
                    or _ORG_KW_RE.search(inner)):
                return orig
            return _chg(orig, f"({inner}, {year})", "missing_comma")

        def _etal(m):
            return _chg(
                m.group(0),
                f"({m.group(1).strip()} et al., {m.group(2)})",
                "etal_punctuation"
            )

        def _etalp(m):
            changes.append({
                "original": m.group(0), "fixed": "et al.",
                "fix_type": "etal_missing_period"
            })
            return "et al."

        def _etdotal(m):
            changes.append({
                "original": m.group(0), "fixed": "et al.",
                "fix_type": "etal_bad_spelling"
            })
            return "et al."

        def _paren_and(m):
            if _ORG_UPPER_RUN.search(m.group(0)) or _ORG_KW_RE.search(m.group(0)):
                return m.group(0)
            return _chg(
                m.group(0),
                f"({m.group(1).strip()} & {m.group(2).strip()}, {m.group(3)})",
                "parenthetical_ampersand"
            )

        def _nd(m):
            return _chg(
                m.group(0),
                f"({m.group(1).strip().rstrip(',')}, n.d.)",
                "nd_format"
            )

        def _inp(m):
            return _chg(
                m.group(0),
                f"({m.group(1).strip().rstrip(',')}, in press)",
                "inpress_capitalisation"
            )

        r = _RE_MISS_COMMA.sub(_comma, r)
        r = _RE_BAD_ETAL.sub(_etal, r)
        r = _RE_ETAL_NOPER.sub(_etalp, r)
        r = _RE_ETDOT_AL.sub(_etdotal, r)
        r = _RE_PAREN_AND.sub(_paren_and, r)
        r = _RE_BAD_ND.sub(_nd, r)
        r = _RE_BAD_INPRES.sub(_inp, r)
        for m in _RE_YR_RANGE.finditer(r):
            changes.append({
                "original": m.group(0), "fixed": m.group(0),
                "fix_type": "year_range_not_apa"
            })
        return r, changes

    @staticmethod
    def fix_etal_expansion(cite_author: str, bib: Dict) -> Optional[str]:
        n = bib.get("author_count", 1)
        has_etal = re.search(r"\bet\s+al\b", cite_author, re.IGNORECASE)

        if n < ET_AL_MIN:
            if n == 2 and has_etal:
                return None  # retain et al. as written — do not expand to two-author form
            return None

        if has_etal:
            return None

        first = bib.get("display", "").split(" et al.")[0].split(" &")[0].strip()
        res = f"{first} et al." if first else None
        return _to_smart_quotes(res) if res else None

# ── Citation Extractor ────────────────────────────────────────────────────────
class CitationExtractor:
    @staticmethod
    def extract(text: str) -> List[Dict]:
        hits: List[Dict] = []
        occ:  List[Tuple[int, int]] = []

        def _over(s, e):
            return any(s < oe and e > os for os, oe in occ)

        # Secondary
        for m in _RE_SECONDARY.finditer(text):
            occ.append((m.start(), m.end()))
            hits.append({
                "raw": m.group(0), "author": m.group(3).strip(),
                "year": m.group(4).strip(),
                "original_author": m.group(1).strip(),
                "original_year": m.group(2).strip(),
                "cite_type": "secondary", "start": m.start(), "end": m.end(),
            })

        # Multi-year parenthetical
        for m in _RE_MULTI_Y_P.finditer(text):
            if _over(m.start(), m.end()):
                continue
            author = re.sub(r'\s+', ' ', m.group(1).strip().rstrip(","))
            years  = [re.sub(r'\s+', ' ', y.strip()) for y in m.group(2).split(",") if y.strip()]
            hits.append({
                "raw": m.group(0), "author": author, "years": years,
                "cite_type": "multi_year", "start": m.start(), "end": m.end(),
            })
            occ.append((m.start(), m.end()))

        # Multi-year narrative
        for m in _RE_MULTI_Y_N.finditer(text):
            if _over(m.start(), m.end()):
                continue
            years = [re.sub(r'\s+', ' ', y.strip()) for y in m.group(2).split(",") if y.strip()]
            hits.append({
                "raw": m.group(0), "author": re.sub(r'\s+', ' ', m.group(1).strip()), "years": years,
                "cite_type": "multi_year", "start": m.start(), "end": m.end(),
            })
            occ.append((m.start(), m.end()))

        # Standard parenthetical
        for m in _RE_PAREN.finditer(text):
            if _over(m.start(), m.end()):
                continue
            # Split on semicolon or colon (some authors mistakenly use a colon to separate multiple references)
            segs  = [s.strip() for s in re.split(r'[;:]', m.group(1))]
            auths: List[str] = []
            block: List[Dict] = []
            for seg in segs:
                lm       = _RE_LOCATOR.search(seg)
                locator  = lm.group(1) if lm else None
                seg_core = seg[:lm.start()].strip() if lm else seg
                um = _RE_CITE_UNIT.search(seg_core)
                if um:
                    sa = re.sub(r'\s+', ' ', um.group(1).strip().rstrip(","))
                    sy = re.sub(r'\s+', ' ', um.group(2).strip())
                    
                    # Check if author ends with a year (multi-year citation like "Omura et al., 2019a, 2019b")
                    # Pattern: author ends with ", YYYY" or ", YYYYa, YYYYb"
                    year_in_author = re.search(r',\s*((?: ' + _Y + r')(?:\s*,\s*(?:' + _Y + r'))*)$', sa)
                    if year_in_author:
                        # Multi-year: extract all years and clean author
                        all_years_str = year_in_author.group(1) + ', ' + sy
                        years = [re.sub(r'\s+', ' ', y.strip()) for y in re.split(r',', all_years_str) if y.strip()]
                        sa = sa[:year_in_author.start()].strip()
                        block.append({
                            "raw": f"({seg})" if len(segs) > 1 else m.group(0),
                            "author": sa, "years": years,
                            "cite_type": "parenthetical",
                            "start": m.start(), "end": m.end(),
                            "block_raw": m.group(0), "block_size": len(segs),
                            "block_order_ok": True,
                        })
                        auths.append(_first_surname(sa))
                    else:
                        block.append({
                            "raw": f"({seg})" if len(segs) > 1 else m.group(0),
                            "author": sa, "year": sy,
                            "locator": locator,
                            "cite_type": "parenthetical",
                            "start": m.start(), "end": m.end(),
                            "block_raw": m.group(0), "block_size": len(segs),
                            "block_order_ok": True,
                        })
                        auths.append(_first_surname(sa))
            order_ok = all(auths[i] >= auths[i - 1] for i in range(1, len(auths)))
            for bc in block:
                bc["block_order_ok"] = order_ok
            hits.extend(block)
            occ.append((m.start(), m.end()))

        # Narrative
        for m in _RE_NARRATIVE.finditer(text):
            if _over(m.start(), m.end()):
                continue
            lm = _RE_LOCATOR.search(m.group(0))
            hits.append({
                "raw": m.group(0), "author": m.group(1).strip(),
                "year": m.group(2).strip(),
                "locator": lm.group(1) if lm else None,
                "cite_type": "narrative", "start": m.start(), "end": m.end(),
            })
            occ.append((m.start(), m.end()))

        return hits

# ── GROBID NLP Reference Parser ──────────────────────────────────────────────
class GrobidClient:
    """
    Thin client for GROBID's /api/processReferences endpoint.
    (See original docstring for architecture notes.)
    """

    _TEI_NS   = "http://www.tei-c.org/ns/1.0"
    _XML_NS   = "http://www.w3.org/XML/1998/namespace"
    _ENDPOINT = "/api/processReferences"

    def __init__(self, base_url: str = "http://localhost:8070",
                 timeout: int = 10):
        self.base_url = base_url.rstrip("/")
        self.timeout  = timeout
        self._log     = logging.getLogger(f"{__name__}.GrobidClient")

    def is_alive(self) -> bool:
        if not _GROBID_DEPS:
            return False
        try:
            url = f"{self.base_url}/api/isalive"
            req = _urllib_req.Request(url, method="GET")
            with _urllib_req.urlopen(req, timeout=self.timeout) as r:
                return r.status == 200
        except Exception:
            return False

    def parse_references(self, raw_strings: List[str]
                         ) -> Optional[List[Optional[Dict]]]:
        if not _GROBID_DEPS:
            return None
        if not raw_strings:
            return []
        try:
            combined = "\n".join(raw_strings)
            payload  = _urllib_parse.urlencode({
                "citations":            combined,
                "consolidateCitations": "0",
            }).encode("utf-8")
            url = f"{self.base_url}{self._ENDPOINT}"
            req = _urllib_req.Request(
                url, data=payload, method="POST",
                headers={"Content-Type":
                         "application/x-www-form-urlencoded; charset=utf-8"},
            )
            with _urllib_req.urlopen(req, timeout=self.timeout) as resp:
                if resp.status != 200:
                    self._log.warning("GROBID returned HTTP %s", resp.status)
                    return None
                xml_bytes = resp.read()
        except Exception as exc:
            self._log.warning("GROBID request failed: %s", exc)
            return None

        return self._parse_tei(xml_bytes, len(raw_strings))

    def _parse_tei(self, xml_bytes: bytes,
                   expected: int) -> List[Optional[Dict]]:
        NS  = {"tei": self._TEI_NS}
        try:
            root = _lxml_etree.fromstring(xml_bytes)
        except Exception as exc:
            self._log.warning("GROBID TEI parse error: %s", exc)
            return [None] * expected

        results: List[Optional[Dict]] = []
        for bib in root.findall(".//tei:biblStruct", NS):
            try:
                results.append(self._extract_bib(bib, NS))
            except Exception as exc:
                self._log.debug("Could not extract biblStruct: %s", exc)
                results.append(None)

        while len(results) < expected:
            results.append(None)
        return results[:expected]

    def _extract_bib(self, bib, NS: dict) -> Optional[Dict]:
        author_nodes = (bib.findall(".//tei:analytic/tei:author", NS)
                        or bib.findall(".//tei:monogr/tei:author", NS)
                        or bib.findall(".//tei:author", NS))

        org_nodes = bib.findall(".//tei:orgName", NS)
        is_org    = False

        author_parts: List[str] = []
        for an in author_nodes:
            pn = an.find("tei:persName", NS)
            if pn is not None:
                surname  = (pn.findtext("tei:surname",  "", NS) or "").strip()
                forename = (pn.findtext("tei:forename", "", NS) or "").strip()
                if surname:
                    initials = self._to_initials(forename)
                    author_parts.append(
                        f"{surname}, {initials}" if initials else surname)
            org = an.find("tei:orgName", NS)
            if org is not None and org.text:
                author_parts.append(org.text.strip())
                is_org = True

        if not author_parts and org_nodes:
            for on in org_nodes:
                if on.text:
                    author_parts.append(on.text.strip())
                    is_org = True

        if not author_parts:
            return None

        if is_org:
            full_author = author_parts[0]
        elif len(author_parts) == 1:
            full_author = author_parts[0]
        elif len(author_parts) == 2:
            full_author = f"{author_parts[0]}, & {author_parts[1]}"
        else:
            inner = ", ".join(author_parts[:-1])
            full_author = f"{inner}, & {author_parts[-1]}"

        year = ""
        for date_el in bib.findall(".//tei:date", NS):
            when = date_el.get("when", "")
            if re.match(r"^(19|20)\d{2}", when):
                year = when[:4]
                break
        if not year:
            raw_txt = "".join(bib.itertext())
            ym = re.search(r"\b((?:19|20)\d{2})\b", raw_txt)
            year = ym.group(1) if ym else "n.d."

        if year != "n.d." and len(year) == 4:
            raw_txt = "".join(bib.itertext())
            sm = re.search(rf"\b{year}([a-z])\b", raw_txt)
            if sm:
                year = f"{year}{sm.group(1)}"

        title = (bib.findtext(".//tei:title[@level='a']", None, NS)
              or bib.findtext(".//tei:title[@level='m']", None, NS)
              or bib.findtext(".//tei:title[@level='j']", None, NS)
              or "")

        return {
            "full_author": full_author,
            "authors":     full_author,
            "year":        year,
            "title":       title.strip(),
            "is_org":      is_org,
            "source":      "grobid",
        }

    @staticmethod
    def _to_initials(forename: str) -> str:
        if not forename:
            return ""
        parts = re.split(r"[\s\-]+", forename.strip())
        inits = []
        for p in parts:
            if not p:
                continue
            inits.append(p[0].upper() + ".")
        if "-" in forename:
            return "-".join(inits)
        return " ".join(inits)

# ── Bibliography Parser ───────────────────────────────────────────────────────
class BibliographyParser:
    _RE_APA   = re.compile(
        r"^(?P<authors>.+?)\s*\((?P<year>(?:19|20)\d{2}[a-z]?|n\.d\.|in\s+press)"
        r"(?:,\s*[^)]+)?\)\.",
        re.DOTALL | re.IGNORECASE)
    _RE_AMA_Y = re.compile(r"\b((?:19|20)\d{2})\b")

    _grobid_cache: Dict[str, Optional[Dict]] = {}

    @classmethod
    def batch_parse(cls, raw_strings: List[str],
                    grobid: Optional["GrobidClient"]) -> None:
        cls._grobid_cache.clear()
        if grobid is None or not raw_strings:
            return
        results = grobid.parse_references(raw_strings)
        if results is None:
            return
        for raw, parsed in zip(raw_strings, results):
            cls._grobid_cache[raw] = parsed

    @classmethod
    def parse_entry(cls, raw: str) -> Optional[Dict]:
        g = cls._grobid_cache.get(raw)
        if g is not None:
            full    = re.sub(r"\s+", " ", g["full_author"]).strip().rstrip(",.")
            year    = g.get("year") or "n.d."
            is_org  = g.get("is_org", _is_organization(full))
            disp    = cls._display(full, is_org)
            ac      = 1 if is_org else _count_authors(full)
            abm     = _RE_BIB_ABBREV.search(full)
            abbrev  = abm.group(1) if abm else None
            ca      = re.sub(r"\[.*?\]", "", full).strip()
            auto_ac = None
            if " " in ca and not re.search(r",\s*[A-Z]\.?", ca):
                a2 = _acronym_of(ca)
                if len(a2) >= 2:
                    auto_ac = a2
            return {
                "full_author": full, "display": disp, "year": year, "raw": raw,
                "cited": False, "author_count": ac, "abbrev": abbrev,
                "auto_acronym": auto_ac, "is_org": is_org,
                "sort_key": _sort_key_name(full.split(",")[0]),
                "grobid_title": g.get("title", ""),
            }

        m = cls._RE_APA.match(raw.strip())
        if m:
            authors_raw = m.group("authors").strip().rstrip(".")
            year        = m.group("year")
        else:
            ym = cls._RE_AMA_Y.search(raw)
            if not ym:
                return None
            year = ym.group(1)
            authors_raw = cls._extract_ama_authors(raw)

        full   = re.sub(r"\s+", " ", authors_raw).strip().rstrip(",.")
        is_org = _is_organization(full)
        disp   = cls._display(full, is_org)
        ac     = 1 if is_org else _count_authors(full)

        abm    = _RE_BIB_ABBREV.search(full)
        abbrev = abm.group(1) if abm else None
        ca     = re.sub(r"\[.*?\]", "", full).strip()
        auto_ac = None
        if " " in ca and not re.search(r",\s*[A-Z]\.?", ca):
            a2 = _acronym_of(ca)
            if len(a2) >= 2:
                auto_ac = a2

        return {
            "full_author": full, "display": disp, "year": year, "raw": raw,
            "cited": False, "author_count": ac, "abbrev": abbrev,
            "auto_acronym": auto_ac, "is_org": is_org,
            "sort_key": _sort_key_name(full.split(",")[0]),
        }

    @staticmethod
    def _extract_ama_authors(raw: str) -> str:
        m = re.match(r'^(.*?)\.\s+[A-Z"«\[]', raw)
        if m:
            return m.group(1)
        ym = re.search(r'\b((?:19|20)\d{2})\b', raw)
        if ym:
            return raw[:ym.start()].rstrip(' .,;')
        return raw.split(".", 1)[0].strip()

    @staticmethod
    def _display(full: str, is_org: bool = False) -> str:
        fc = re.sub(r'\[.*?\]', '', full).strip()
        if is_org:
            return fc

        parts = re.split(r'\s*&\s*|\s*,\s*', fc)
        sn = []
        has_et_al = False

        for part in parts:
            part = part.strip()
            if not part:
                continue
            if re.search(r'\bet\s+al\b', part, re.IGNORECASE):
                has_et_al = True
                part = re.sub(r'(?i)\bet\s+al\.?\b', '', part).strip()
                if not part:
                    continue
            if _INITIAL_ONLY_RE.match(part):
                continue
            sn.append(part)

        if not sn:
            sn = [w for w in fc.split()
                  if w[0].isupper() and not re.match(r'^[A-Z]\.$', w)]
        if not sn:
            return fc.split()[0] if fc.split() else fc

        if len(sn) == 1:
            return sn[0] if not has_et_al else f"{sn[0]} et al."
        if has_et_al or len(sn) >= ET_AL_MIN:
            return f"{sn[0]} et al."
        return f"{sn[0]} & {sn[1]}"

# ── Matcher ───────────────────────────────────────────────────────────────────
def _score_candidate(cite_norm: str, cite_author: str, cite_year: str,
                     ref_norm: str, ref: dict) -> float:
    """
    Composite match score: name similarity (0–1) + year bonus.

    FIX-53: Three-tier year bonus instead of binary:
      +2.0  full year match  (cite_year == ref_year)
      +0.5  base year match  (_strip_suffix equal but full years differ)
       0.0  different base year
       
    Includes penalty for year distance and bonus/penalty based on 'et al.' 
    and author counts.
    """
    name_sim   = difflib.SequenceMatcher(None, cite_norm, ref_norm).ratio()
    # Boost if the citation heavily matches the 'display' string (e.g. "Franck et al.")
    disp_sim   = difflib.SequenceMatcher(None, cite_norm, _norm(ref.get("display", ""))).ratio()
    name_sim   = max(name_sim, disp_sim)

    ref_year   = ref["year"]
    cb         = _strip_suffix(cite_year)
    rb         = _strip_suffix(ref_year)
    full_match = (cite_year == ref_year)
    base_match = (cb == rb) and not full_match
    year_bonus = 2.0 if full_match else (0.5 if base_match else 0.0)

    year_penalty = 0.0
    if not full_match and not base_match and cb.isdigit() and rb.isdigit():
        diff = abs(int(cb) - int(rb))
        year_penalty = min(diff * 0.1, 0.5)  # -0.1 per year, max -0.5

    et_al_modifier = 0.0
    has_et_al = bool(re.search(r'\bet\s+al\b', cite_author, re.IGNORECASE))
    ref_authors = ref.get("author_count", 1)
    
    if has_et_al and ref_authors >= 3:
        et_al_modifier = 0.5
    elif has_et_al and ref_authors < 3:
        et_al_modifier = -0.5
    elif not has_et_al and ref_authors >= 3 and not ref.get("is_org"):
        et_al_modifier = -0.5
    elif not has_et_al and ref_authors < 3:
        et_al_modifier = 0.2

    return name_sim + year_bonus - year_penalty + et_al_modifier


def match_citation(
    cite_author: str,
    cite_year: str,
    bibliography: Dict[str, dict],
) -> MatchResult:
    """
    Every candidate is scored; the highest-scoring candidate wins.
    Score = name_similarity + year_bonus (2.0 full, 0.5 base, 0.0 different).

    Returns MatchResult(key, match_type, score).
    """
    cn = _norm(cite_author)
    cf = _first_surname(cite_author)
    cb = _strip_suffix(cite_year)

    best_key:   Optional[str] = None
    best_mt:    str           = "not_found"
    best_score: float         = -1.0

    for key, ref in bibliography.items():
        rfull = ref["full_author"]
        rn    = _norm(rfull)
        ry    = ref["year"]
        rf    = _first_surname(rfull)
        rb    = _strip_suffix(ry)
        year_ok = (cite_year == ry)

        if cn == rn:
            mt    = "exact" if year_ok else ("suffix_mismatch" if cb == rb else "year_mismatch")
            score = _score_candidate(cn, cite_author, cite_year, rn, ref)
        elif (cf and rf and
              (cf == rf or
               (len(cf) > SHORT_NAME_MAX_LEN and
                difflib.SequenceMatcher(None, cf, rf).ratio() >= FUZZY_THRESHOLD))):
            mt    = "smart" if year_ok else ("suffix_mismatch" if cb == rb else "year_mismatch")
            score = _score_candidate(cn, cite_author, cite_year, rn, ref)
        elif _is_org_match(cite_author, rfull):
            mt    = "org_abbrev" if year_ok else "year_mismatch"
            score = _score_candidate(cn, cite_author, cite_year, rn, ref)
        else:
            ratio = difflib.SequenceMatcher(None, cn, rn).ratio()
            disp_ratio = difflib.SequenceMatcher(None, cn, _norm(ref.get("display", ""))).ratio()
            ratio = max(ratio, disp_ratio)
            
            if ratio >= FUZZY_THRESHOLD:
                mt    = "spelling_mismatch"
                
                # Penalty and bonuses just like in _score_candidate to match behavior
                year_penalty = 0.0
                if cb.isdigit() and rb.isdigit() and cb != rb:
                    year_penalty = min(abs(int(cb) - int(rb)) * 0.1, 0.5)
                
                et_al_modifier = 0.0
                has_et_al = bool(re.search(r'\bet\s+al\b', cite_author, re.IGNORECASE))
                ref_authors = ref.get("author_count", 1)
                if has_et_al and ref_authors >= 3:
                    et_al_modifier = 0.5
                elif has_et_al and ref_authors < 3:
                    et_al_modifier = -0.5
                    
                score = ratio + (2.0 if year_ok else (0.5 if cb == rb else 0.0)) - year_penalty + et_al_modifier
            else:
                continue

        if score > best_score:
            best_score = score
            best_key   = key
            best_mt    = mt

    if best_key is None:
        return MatchResult(None, "not_found", 0.0)

    # Reclassify suffix cases now that we have the single best candidate
    ref  = bibliography[best_key]
    ry   = ref["year"]
    rb   = _strip_suffix(ry)
    if best_mt in ("year_mismatch",) and cb == rb and cite_year != ry:
        others = [k for k, r in bibliography.items()
                  if _strip_suffix(r["year"]) == cb
                  and _norm(r["full_author"]) == _norm(ref["full_author"])
                  and k != best_key]
        best_mt = "suffix_ambiguous" if others else "suffix_mismatch"

    return MatchResult(best_key, best_mt, best_score)

# ── et al. checker ────────────────────────────────────────────────────────────
def check_etal_enforcement(cite_author: str, bib: Dict, et_al_min: int = ET_AL_MIN) -> Optional[str]:
    n   = bib.get("author_count", 1)
    has = bool(re.search(r"\bet\s+al\b", cite_author, re.IGNORECASE))

    if bib.get("is_org"):
        if has:
            return (f"et al. INCORRECT: '{bib['display']}' is an organization "
                    "— list full name, do not use et al.")
        return None

    if n >= et_al_min and not has:
        return (f"et al. REQUIRED: entry has {n} authors ('{bib['display']}') "
                f"— requires et al. (threshold: {et_al_min}+).")
    if n < et_al_min and has:
        return (f"et al. INCORRECT: entry has {n} author(s) ('{bib['display']}') "
                "— list all names.")
    if n >= et_al_min and has and not re.search(r"\bet\s+al\.", cite_author, re.IGNORECASE):
        return "et al. PUNCTUATION: missing period — must be 'et al.' (with period)."
    return None

# ── Bibliography structural checks ───────────────────────────────────────────
def check_bibliography_structure(entries: List[Dict]) -> List[Dict]:
    issues: List[Dict] = []
    seen_r:  Dict[str, int] = {}
    seen_ay: Dict[str, int] = {}

    for e in entries:
        raw  = e["raw"]
        pidx = e.get("para_idx", -1)
        ay   = f"{_norm(e['full_author'])}|{e['year']}"
        if raw in seen_r:
            issues.append({
                "type": "duplicate_entry", "para_idx": pidx,
                "raw": e["display"], "para": None,
                "message": (f"📋 DUPLICATE ENTRY: '{e['display']} ({e['year']})' "
                            f"at para {pidx + 1} (first at {seen_r[raw] + 1})."),
            })
        else:
            seen_r[raw] = pidx
        if ay in seen_ay and seen_ay[ay] != pidx:
            issues.append({
                "type": "duplicate_authyear", "para_idx": pidx,
                "raw": e["display"], "para": None,
                "message": (f"📋 SAME AUTHOR+YEAR: '{e['display']} ({e['year']})' "
                            "— add a/b/c suffix."),
            })
        seen_ay[ay] = pidx

    pk = ""; pd = ""; pp = -1
    for e in entries:
        sk = e.get("sort_key", "")
        if sk and pk and sk < pk:
            issues.append({
                "type": "order_error", "para_idx": e.get("para_idx", -1),
                "raw": e["display"], "para": None,
                "message": (f"🔤 ORDER: '{e['display']}' should precede '{pd}' "
                            "— bibliography must be alphabetical."),
            })
        pk = sk; pd = e["display"]; pp = e.get("para_idx", -1)

    sg: Dict[Tuple[str, str], List] = {}
    for e in entries:
        bare = _strip_suffix(e["year"])
        sfx  = e["year"][len(bare):]
        if sfx:
            sg.setdefault((_norm(e["full_author"]), bare), []).append(
                (sfx, e.get("para_idx", -1), e["display"])
            )
    for (auth, bare), items in sg.items():
        actual = [s for s, _, _ in items]
        exp    = sorted(actual)
        if actual != exp:
            for i, (s, pidx, disp) in enumerate(items):
                if actual[i] != exp[i]:
                    issues.append({
                        "type": "suffix_order_error", "para_idx": pidx,
                        "raw": disp, "para": None,
                        "message": (f"🔡 SUFFIX ORDER: '{disp} ({bare}{s})' out of order "
                                    f"— expected {', '.join(bare + x for x in exp)}."),
                    })

    rl = list({e["raw"] for e in entries})
    for i in range(len(rl)):
        for j in range(i + 1, len(rl)):
            r = difflib.SequenceMatcher(None, rl[i].lower(), rl[j].lower()).ratio()
            if r >= NEAR_DUP_THRESHOLD:
                for e in entries:
                    if e["raw"] == rl[j]:
                        issues.append({
                            "type": "near_duplicate", "para_idx": e.get("para_idx", -1),
                            "raw": e["display"], "para": None,
                            "message": (f"📋 NEAR-DUPLICATE ({r:.0%}): "
                                        f"'{rl[j][:80]}' may duplicate '{rl[i][:60]}'."),
                        })
                        break
    return issues

# ── Org tracker ───────────────────────────────────────────────────────────────
class OrgTracker:
    def __init__(self):
        self._introduced: Dict[str, Tuple[str, int]] = {}

    def record(self, abbrev: str, full: str, pidx: int):
        if abbrev not in self._introduced:
            self._introduced[abbrev] = (full, pidx)

    def check(self, abbrev: str) -> Optional[str]:
        if abbrev not in self._introduced:
            return (f"ORG ABBREV FIRST USE: '{abbrev}' used without prior introduction. "
                    "Spell out full name first: 'Full Name [{abbrev}]'.")
        return None

    def is_known(self, abbrev: str) -> bool:
        return abbrev in self._introduced

# ── Context filter ────────────────────────────────────────────────────────────
def _para_context(para) -> str:
    sn = getattr(getattr(para, "style", None), "name", "") or ""
    if _HEADING_RE.search(sn):  return "heading"
    if _FOOTNOTE_RE.search(sn): return "footnote"
    if _CAPTION_RE.search(sn):  return "caption"
    if _BQUOTE_RE.search(sn):   return "blockquote"
    return "body"

# ── Word XML helpers ──────────────────────────────────────────────────────────
def _ensure_style(doc):
    _log = logging.getLogger(__name__)
    if CITE_STYLE not in [s.name for s in doc.styles]:
        try:
            cs = doc.styles.add_style(CITE_STYLE, WD_STYLE_TYPE.CHARACTER)
            cs.font.bold = False
        except Exception as exc:
            _log.warning("Could not create '%s' character style: %s", CITE_STYLE, exc)
            return
    try:
        cs = doc.styles[CITE_STYLE]
        cs.hidden = False
        cs.quick_style = True
        cs.priority = 100
    except Exception as exc:
        _log.warning("Could not set visibility on '%s' style: %s", CITE_STYLE, exc)

    return doc.styles[CITE_STYLE].style_id


def apply_highlight_by_span(para, char_start: int, char_end: int,
                             highlight, style=None):
    if char_start >= char_end:
        return

    attempts = 0
    while attempts < 20:
        attempts += 1
        runs   = list(para.runs)
        runs_text = "".join(r.text for r in runs)
        if char_end > len(runs_text):
            char_end = len(runs_text)

        offset      = 0
        made_change = False

        for i, run in enumerate(runs):
            rlen    = len(run.text)
            run_end = offset + rlen

            if run_end <= char_start or offset >= char_end:
                offset = run_end
                continue

            overlap_start = max(char_start, offset)
            overlap_end   = min(char_end,   run_end)
            fully_inside  = (offset >= char_start and run_end <= char_end)

            if fully_inside:
                _tc_rpr_change(run._r, new_highlight=highlight, new_style_id=style)
                # Also set highlight_color property for direct access
                if highlight is not None:
                    run.font.highlight_color = highlight
            else:
                inner_txt = runs_text[overlap_start:overlap_end]
                if not inner_txt:
                    offset = run_end
                    continue
                result = safe_splice(para, overlap_start, overlap_end,
                                     inner_txt, highlight, style)
                if result is not None:
                    made_change = True
                    break

            offset = run_end

        if not made_change:
            break

def _full_text(p):
    from docx.oxml.ns import qn as _qn
    parts = []
    for child in p._p:
        tag = child.tag
        if tag == _qn('w:r'):
            for n in child:
                if n.tag == _qn('w:t') and n.text:
                    parts.append(n.text)
                elif n.tag == _qn('w:br') or n.tag == _qn('w:cr'):
                    parts.append('\n')
        elif tag == _qn('w:ins'):
            for r in child.findall(_qn('w:r')):
                for n in r:
                    if n.tag == _qn('w:t') and n.text:
                        parts.append(n.text)
                    elif n.tag == _qn('w:br') or n.tag == _qn('w:cr'):
                        parts.append('\n')
    return "".join(parts)


# ── Tracked-change helpers ────────────────────────────────────────────────────
import datetime as _datetime
_TC_AUTHOR = COMMENT_AUTHOR
_TC_DATE   = _datetime.datetime.now(
    _datetime.timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")
_tc_rev_id = 0


def _next_rev_id() -> int:
    global _tc_rev_id
    _tc_rev_id += 1
    return _tc_rev_id


def _tc_ins(text: str, base_rPr=None) -> "OxmlElement":
    import copy
    ins = OxmlElement('w:ins')
    ins.set(qn('w:id'),     str(_next_rev_id()))
    ins.set(qn('w:author'), _TC_AUTHOR)
    ins.set(qn('w:date'),   _TC_DATE)
    r = OxmlElement('w:r')
    if base_rPr is not None:
        r.append(copy.deepcopy(base_rPr))
    t = OxmlElement('w:t')
    t.text = text
    if text and (text[0] == ' ' or text[-1] == ' '):
        t.set(qn('xml:space'), 'preserve')
    r.append(t)
    ins.append(r)
    return ins


def _tc_del(text: str, base_rPr=None) -> "OxmlElement":
    import copy
    d = OxmlElement('w:del')
    d.set(qn('w:id'),     str(_next_rev_id()))
    d.set(qn('w:author'), _TC_AUTHOR)
    d.set(qn('w:date'),   _TC_DATE)
    r = OxmlElement('w:r')
    if base_rPr is not None:
        r.append(copy.deepcopy(base_rPr))
    dt = OxmlElement('w:delText')
    dt.text = text
    if text and (text[0] == ' ' or text[-1] == ' '):
        dt.set(qn('xml:space'), 'preserve')
    r.append(dt)
    d.append(r)
    return d


def _tc_rpr_change(run_el, new_highlight=None, new_style_id=None):
    import copy
    rpr = run_el.get_or_add_rPr()

    orig_rpr = OxmlElement('w:rPr')
    for child in list(rpr):
        if child.tag != qn('w:rPrChange'):
            orig_rpr.append(copy.deepcopy(child))

    if new_highlight is not None:
        hl = OxmlElement('w:highlight')
        colour_map = {
            WD_COLOR_INDEX.BRIGHT_GREEN: "green",
            WD_COLOR_INDEX.YELLOW:       "yellow",
        }
        hl.set(qn('w:val'), colour_map.get(new_highlight, "yellow"))
        for old in rpr.findall(qn('w:highlight')):
            rpr.remove(old)
        rpr.append(hl)

    if new_style_id is not None:
        rs = OxmlElement('w:rStyle')
        actual_id = new_style_id
        rs.set(qn('w:val'), actual_id)
        for old in rpr.findall(qn('w:rStyle')):
            rpr.remove(old)
        rpr.insert(0, rs)

    rpc = OxmlElement('w:rPrChange')
    rpc.set(qn('w:id'),     str(_next_rev_id()))
    rpc.set(qn('w:author'), _TC_AUTHOR)
    rpc.set(qn('w:date'),   _TC_DATE)
    rpc.append(orig_rpr)
    for old in rpr.findall(qn('w:rPrChange')):
        rpr.remove(old)
    rpr.append(rpc)


def _tracked_replace_one(para, old_text: str, new_text: str,
                         highlight, style_id: str = None) -> bool:
    new_text = _to_smart_quotes(new_text)
    if style_id is None:
        style_id = CITE_STYLE

    import copy

    runs_text = "".join(r.text for r in para.runs)
    pos = runs_text.find(old_text)
    if pos == -1:
        pos = runs_text.lower().find(old_text.lower())
    if pos == -1:
        return False

    end = pos + len(old_text)

    spliced = safe_splice(para, pos, end, old_text, None, None)
    if spliced is None:
        return False

    base_rPr = (copy.deepcopy(spliced.rPr)
                if hasattr(spliced, 'rPr') and spliced.rPr is not None
                else None)

    new_rPr = OxmlElement('w:rPr')
    if base_rPr is not None:
        for child in list(base_rPr):
            new_rPr.append(copy.deepcopy(child))
    colour_map = {WD_COLOR_INDEX.BRIGHT_GREEN: "green", WD_COLOR_INDEX.YELLOW: "yellow"}
    hl_el = OxmlElement('w:highlight')
    hl_el.set(qn('w:val'), colour_map.get(highlight, "yellow"))
    for old in new_rPr.findall(qn('w:highlight')):
        new_rPr.remove(old)
    new_rPr.append(hl_el)
    rs_el = OxmlElement('w:rStyle')
    actual_id = "citebib" if style_id == CITE_STYLE else style_id
    rs_el.set(qn('w:val'), actual_id)
    for old in new_rPr.findall(qn('w:rStyle')):
        new_rPr.remove(old)
    new_rPr.insert(0, rs_el)

    del_el = _tc_del(old_text, base_rPr)
    ins_el = _tc_ins(new_text, new_rPr)

    parent = spliced.getparent()
    idx    = list(parent).index(spliced)
    parent.remove(spliced)
    parent.insert(idx,     del_el)
    parent.insert(idx + 1, ins_el)
    return True


def safe_splice(para, start, end, new_text, highlight, style):
    """
    Isolate the character range [start, end) of para into a single run,
    then optionally apply highlight + style to that run.

    FIX-41b: Tightened empty-pre guard.  When start lands exactly at the
    beginning of a run (pre == ""), the original code still set
    runs[s_r_idx].text = pre (= ""), creating a zero-length run in the XML.
    That orphan run could be misidentified by a subsequent safe_splice call
    on the same paragraph, causing text like "et al." to be stranded outside
    the new run and then dropped when the zero-length run was removed.

    Fix: when pre is empty, zero out the original run's text FIRST, then
    insert the spliced run at the same XML position (ai), not ai+1.  This
    leaves no zero-length run artefact before the spliced content.
    """
    import copy
    runs = list(para.runs)
    if not runs:
        return None

    curr = 0; s_r_idx = -1; s_off = -1; e_r_idx = -1; e_off = -1
    for i, r in enumerate(runs):
        l = len(r.text)
        if s_r_idx == -1 and curr <= start < curr + l:
            s_r_idx, s_off = i, start - curr
        if e_r_idx == -1 and curr < end <= curr + l:
            e_r_idx, e_off = i, end - curr
        curr += l
    if e_r_idx == -1 and end == curr:
        e_r_idx = len(runs) - 1
        e_off   = len(runs[-1].text)
    if s_r_idx == -1 or e_r_idx == -1:
        return None

    st_text  = runs[s_r_idx].text
    end_text = runs[e_r_idx].text
    pre      = st_text[:s_off]
    post     = end_text[e_off:]

    for i in range(s_r_idx + 1, e_r_idx):
        runs[i].text = ""

    base_rPr = (copy.deepcopy(runs[s_r_idx]._r.rPr)
                if hasattr(runs[s_r_idx]._r, 'rPr')
                   and runs[s_r_idx]._r.rPr is not None
                else None)

    st = para.add_run(new_text)
    st._r.clear_content()
    if base_rPr is not None:
        st._r.append(copy.deepcopy(base_rPr))
    t = OxmlElement('w:t')
    t.text = new_text
    if new_text.startswith(' ') or new_text.endswith(' '):
        t.set(qn('xml:space'), 'preserve')
    st._r.append(t)
    if style is not None:
        try:
            actual_name = CITE_STYLE if style in ("citebib", "cite_bib") else style
            st.style = actual_name
        except Exception:
            pass
    if highlight is not None:
        st.font.highlight_color = highlight

    pe = runs[s_r_idx]._r.getparent()
    ai = list(pe).index(runs[s_r_idx]._r)
    pe.remove(st._r)

    if s_r_idx == e_r_idx:
        if pre:
            # Normal case: pre-text stays in original run, spliced run follows
            runs[s_r_idx].text = pre
            pe.insert(ai + 1, st._r)
            if post:
                post_r = OxmlElement('w:r')
                if base_rPr is not None:
                    post_r.append(copy.deepcopy(base_rPr))
                pt = OxmlElement('w:t')
                pt.text = post
                if post.startswith(' ') or post.endswith(' '):
                    pt.set(qn('xml:space'), 'preserve')
                post_r.append(pt)
                pe.insert(ai + 2, post_r)
        else:
            # FIX-41b: pre is empty — replace the original run IN-PLACE.
            # Zero it out and insert the spliced run at the SAME position (ai),
            # not ai+1, so no zero-length artefact precedes the new content.
            runs[s_r_idx].text = ""
            pe.insert(ai, st._r)   # <-- was ai+1; now ai to avoid orphan empty run
            if post:
                post_r = OxmlElement('w:r')
                if base_rPr is not None:
                    post_r.append(copy.deepcopy(base_rPr))
                pt = OxmlElement('w:t')
                pt.text = post
                if post.startswith(' ') or post.endswith(' '):
                    pt.set(qn('xml:space'), 'preserve')
                post_r.append(pt)
                pe.insert(ai + 1, post_r)
    else:
        runs[s_r_idx].text = pre
        runs[e_r_idx].text = post
        pe.insert(ai + 1, st._r)

    return st._r


def isolate_target_run(p, txt):
    if not txt:
        return None
    runs_text = "".join(r.text for r in p.runs)
    pos = runs_text.find(txt)
    if pos == -1:
        pos = runs_text.lower().find(txt.lower())
    if pos != -1:
        return safe_splice(p, pos, pos + len(txt),
                           runs_text[pos:pos + len(txt)], None, None)
    return None

def apply_style_to_text(p, txt, hl, style_id=None):
    if style_id is None:
        style_id = CITE_STYLE
    if not txt:
        return None
    res = None
    
    # Normalize the search text
    txt_normalized = re.sub(r'\s+', ' ', txt).strip()
    
    # For multi-citation blocks, try searching for citation without outer parens
    txt_inner = txt_normalized
    if txt_normalized.startswith('(') and txt_normalized.endswith(')'):
        txt_inner = txt_normalized[1:-1]
        
    runs_text = "".join(r.text for r in p.runs)
    
    # Helper to find exact offsets ignoring whitespace differences
    def get_offsets(original_text, search_text):
        if not search_text: return -1, -1
        # Build regex that allows any whitespace between non-whitespace characters
        pattern = r'\s*'.join(re.escape(char) for char in search_text if not char.isspace())
        match = re.search(pattern, original_text, re.IGNORECASE)
        if match:
            return match.start(), match.end()
        return -1, -1

    # 1. Try matching the exact full text (with parens)
    start, end = get_offsets(runs_text, txt_normalized)
    
    # 2. Try matching the inner text (without parens) if full match fails
    if start == -1 and txt_inner != txt_normalized:
        start, end = get_offsets(runs_text, txt_inner)
        
    if start != -1:
        # We found exactly where it is in runs_text!
        # First try applying with apply_highlight_by_span (which sets highlight and style cleanly)
        apply_highlight_by_span(p, start, end, hl, style_id)
        # Return None so that insert_comment falls back to text search rather than crashing on True
        res = None
def tracked_replace(p, old_text, new_text, highlight=None, style_id=None):
    if style_id is None:
        style_id = CITE_STYLE
    if not old_text:
        return None
    res = None
    full = _full_text(p)
    pos  = full.find(old_text)
    while pos != -1:
        ok = _tracked_replace_one(p, old_text, new_text, highlight, style_id)
        if ok:
            res = True
        full = _full_text(p)
        pos  = full.find(old_text, pos + max(len(new_text), 1))
    return res


def sort_citation_block(para, block_raw: str, highlight) -> tuple:
    if not block_raw:
        return False, False

    inner = block_raw.strip()
    if inner.startswith("(") and inner.endswith(")"):
        inner = inner[1:-1]

    raw_segs = [s.strip() for s in re.split(r'[;:]', inner) if s.strip()]
    if len(raw_segs) < 2:
        return False, False

    def _seg_sort_key(seg: str):
        um = _RE_CITE_UNIT.search(seg)
        if um:
            auth = um.group(1).strip().rstrip(",")
            yr   = um.group(2).strip()
            return (_sort_key_name(_first_surname(auth)), yr)
        return (seg, "")

    sorted_segs = sorted(raw_segs, key=_seg_sort_key)
    if sorted_segs == raw_segs:
        return False, False

    new_block = "(" + "; ".join(sorted_segs) + ")"

    if tracked_replace(para, block_raw, new_block, highlight, CITE_STYLE):
        return True, True

    runs_text = "".join(r.text for r in para.runs)
    first_seg = raw_segs[0]
    last_seg  = raw_segs[-1]
    idx_first = runs_text.find(first_seg)
    idx_last  = runs_text.rfind(last_seg)
    if idx_first == -1 or idx_last == -1:
        return False, False
    span_start = max(0, idx_first - 1)
    while span_start > 0 and runs_text[span_start] != '(':
        span_start -= 1
    span_end = idx_last + len(last_seg)
    while span_end < len(runs_text) and runs_text[span_end] != ')':
        span_end += 1
    span_end = min(span_end + 1, len(runs_text))

    old_block_found = runs_text[span_start:span_end]
    if old_block_found and tracked_replace(para, old_block_found, new_block, highlight, CITE_STYLE):
        return True, True

    result = safe_splice(para, span_start, span_end, new_block, highlight, CITE_STYLE)
    if result is not None:
        try:
            sid = para.part.document.styles[CITE_STYLE].style_id
        except KeyError:
            sid = CITE_STYLE
        _tc_rpr_change(result, new_highlight=highlight, new_style_id=sid)
        return True, False

    try:
        sid = para.part.document.styles[CITE_STYLE].style_id
    except KeyError:
        sid = CITE_STYLE
    apply_highlight_by_span(para, span_start, span_end, highlight, sid)
    return True, False

def _get_comments_part(doc):
    try:
        from docx.opc.part import Part
        from docx.opc.packuri import PackURI
        from lxml import etree as _lxml_etree

        REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships/comments"
        CT  = "application/vnd.openxmlformats-officedocument.wordprocessingml.comments+xml"

        for rel in doc.part.rels.values():
            if rel.reltype == REL:
                return rel.target_part

        for p in doc.part.package.iter_parts():
            if str(getattr(p, "partname", "")) == "/word/comments.xml":
                doc.part.relate_to(p, REL)
                return p

        _W_NS   = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
        _MC_NS  = "http://schemas.openxmlformats.org/markup-compatibility/2006"
        _W14_NS = "http://schemas.microsoft.com/office/word/2010/wordml"

        root = _lxml_etree.Element(
            f"{{{_W_NS}}}comments",
            nsmap={"w": _W_NS, "mc": _MC_NS, "w14": _W14_NS},
        )
        root.set(f"{{{_MC_NS}}}Ignorable", "w14 wp14")

        blob = _lxml_etree.tostring(
            root, xml_declaration=True, encoding="UTF-8", standalone=True
        )
        puri = PackURI("/word/comments.xml")
        part = Part(puri, CT, blob, doc.part.package)
        doc.part.relate_to(part, REL)
        return part

    except Exception as exc:
        logging.error(f"_get_comments_part failed: {exc}")
        return None

def insert_comment(doc, para, text, target_run=None, target_text=None):
    try:
        from docx.text.run import Run as _Run

        if not para.runs:
            return False

        def _find_in_ins(search_txt):
            if not search_txt:
                return None
            sl = search_txt.lower()
            for child in para._p:
                if child.tag == qn('w:ins'):
                    for r_el in child.findall(qn('w:r')):
                        t_el = r_el.find(qn('w:t'))
                        if t_el is not None and t_el.text and sl in t_el.text.lower():
                            return r_el
            return None

        if target_run is None and target_text is not None:
            norm_target = _norm_apos(target_text)
            target_run  = isolate_target_run(para, norm_target)
            if target_run is None:
                target_run = isolate_target_run(para, target_text)
            if target_run is None:
                stripped = target_text.strip("() \u201c\u201d\u2018\u2019")
                if stripped and stripped != target_text:
                    target_run = isolate_target_run(para, stripped)
            if target_run is None:
                first_token = (target_text.strip("() ").split(",")[0]).strip()
                if first_token and len(first_token) > 2:
                    target_run = isolate_target_run(para, first_token)
            if target_run is None:
                target_run = _find_in_ins(norm_target)
            if target_run is None:
                target_run = _find_in_ins(
                    target_text.strip("() ").split(",")[0].strip())

        if target_run is None:
            target_run = para.runs[0]._r if para.runs else None
        if target_run is None:
            return False

        if isinstance(target_run, _Run):
            run_obj = target_run
        else:
            run_obj = _Run(target_run, para)

        doc.add_comment(
            runs=run_obj,
            text=text,
            author=COMMENT_AUTHOR,
            initials=COMMENT_INITIALS,
        )
        return True

    except Exception as exc:
        logging.error(f"insert_comment failed: {exc}")
        return False


class ValidationReport:
    def __init__(self, issues, stats, total_refs, total_cites):
        self.issues      = issues
        self.stats       = stats
        self.total_refs  = total_refs
        self.total_cites = total_cites

    def summary(self) -> str:
        s = self.stats

        def R(lbl, k, w=44):
            return f"  {lbl:<{w}}: {s.get(k, 0)}"

        lines = [
            "=" * 68, "APA 7th CITATION VALIDATION REPORT", "=" * 68,
            f"  {'Total in-text citations':<44}: {self.total_cites}",
            f"  {'Total bibliography entries':<44}: {self.total_refs}",
            "-" * 68,
            R("Matched (green)",            "matched"),
            R("GROBID-parsed bib entries",  "grobid_parsed"),
            R("Missing references",         "missing"),
            R("Year mismatches",            "year_mismatch"),
            R("Suffix missing",             "suffix_mismatch"),
            R("Suffix ambiguous",           "suffix_ambiguous"),
            R("Spelling mismatches",        "spelling_mismatch"),
            R("et al. violations",          "etal_enforcement"),
            R("et al. auto-fixes",          "etal_autofixes"),
            R("Secondary citations",        "secondary"),
            R("Multi-citation blocks",      "multi_citation_blocks"),
            R("Duplicate citations in block","duplicate_citations"),
            R("Multi-year citations",       "multi_year"),
            R("Multi-year mismatches",      "multi_year_mismatches"),
            R("Block order violations",     "alphabetical_violations"),
            R("Org-author cases",           "organization_cases"),
            R("Org-abbrev first-use errors","org_abbrev_first_use"),
            R("Page locator errors",        "page_locator_errors"),
            R("Format auto-fixes",          "format_fixed"),
            R("Year-range (manual)",        "year_range"),
            R("Unused references",          "unused"),
            R("Duplicate bib entries",      "duplicate_entry"),
            R("Same author+year (needs suffix)","duplicate_authyear"),
            R("Near-duplicate bib entries", "near_duplicate"),
            R("Bib order errors",           "order_error"),
            R("Suffix order errors",        "suffix_order_error"),
            "-" * 68,
        ]

        _SEC = [
            ("missing",               "MISSING REFERENCES"),
            ("year_mismatch",         "YEAR MISMATCHES"),
            ("suffix_mismatch",       "SUFFIX MISSING"),
            ("suffix_ambiguous",      "SUFFIX AMBIGUOUS"),
            ("spelling_mismatch",     "SPELLING MISMATCHES"),
            ("etal_enforcement",      "et al. VIOLATIONS"),
            ("secondary",             "SECONDARY CITATIONS"),
            ("duplicate_citations",   "DUPLICATE WITHIN BLOCK"),
            ("multi_year_mismatches", "MULTI-YEAR MISMATCHES"),
            ("alphabetical_violations","BLOCK ORDER VIOLATIONS"),
            ("org_abbrev_first_use",  "ORG FIRST-USE ERRORS"),
            ("organization_cases",    "ORG ABBREV MATCHES"),
            ("page_locator_errors",   "PAGE LOCATOR ERRORS"),
            ("format_fixed",          "FORMAT AUTO-FIXES"),
            ("year_range",            "YEAR-RANGE CITATIONS"),
            ("unused",                "UNUSED REFERENCES"),
            ("duplicate_entry",       "DUPLICATE BIB ENTRIES"),
            ("duplicate_authyear",    "SAME AUTHOR+YEAR"),
            ("near_duplicate",        "NEAR-DUPLICATE BIB"),
            ("order_error",           "BIB ORDER ERRORS"),
            ("suffix_order_error",    "SUFFIX ORDER ERRORS"),
        ]
        for itype, heading in _SEC:
            items = [i for i in self.issues if i["type"] == itype]
            if items:
                lines += ["", heading, "-" * 52]
                for item in items:
                    pidx  = item.get('para_idx', -1)
                    label = _para_label(pidx)
                    lines.append(f"  Para {label:>8}: {item['message']}")
        lines += ["", "=" * 68, "END OF REPORT", "=" * 68]
        return "\n".join(lines)

    def to_dict(self) -> dict:
        issues_clean = [{k: v for k, v in i.items() if k != "para"}
                        for i in self.issues]

        def _f(t):
            return [i for i in issues_clean if i["type"] == t]

        return {
            "stats":                  dict(self.stats),
            "total_refs":             self.total_refs,
            "total_cites":            self.total_cites,
            "issues":                 issues_clean,
            "style":                  "APA",
            "et_al_issues":           _f("etal_enforcement"),
            "name_spelling_warnings": _f("spelling_mismatch"),
            "ordering_issues":        _f("alphabetical_violations") + _f("order_error"),
            "suffix_issues":          _f("suffix_mismatch") + _f("suffix_ambiguous") + _f("suffix_order_error"),
            "disambiguation_issues":  _f("disambiguation"),
            "personal_comm_citations":_f("personal_comm"),
            "secondary_citations":    _f("secondary"),
        }

# ── Main Processor ────────────────────────────────────────────────────────────
class CitationProcessor:
    """
    Usage:
        proc   = CitationProcessor("paper.docx", job_id="job-1",
                                   include_contexts={"body","footnote"})
        report = proc.run()
        proc.save("paper_out.docx")
        print(report.summary())
    """

    def __init__(
        self,
        doc_path: str,
        job_id: str = "default",
        include_contexts: Optional[Set[str]] = None,
        grobid: Optional["GrobidClient"] = None,
        et_al_threshold: Optional[int] = None,
    ):
        self.doc_path         = doc_path
        self.doc              = Document(doc_path)
        self.job_id           = job_id
        self.include_contexts = include_contexts or {"body"}
        self.log              = logging.getLogger(f"{__name__}.{job_id}")
        self.grobid           = grobid
        self._et_al_min       = et_al_threshold if et_al_threshold is not None else ET_AL_MIN
        self.cite_style_id    = _ensure_style(self.doc)
        self.bibliography:    Dict[str, dict]    = {}
        self._bib_ordered:    List[dict]         = []
        self._issues:         List[dict]         = []
        self._stats:          Dict[str, int]     = defaultdict(int)
        self._cited_keys:     set                = set()
        self._org_tracker:    OrgTracker         = OrgTracker()
        self._seen_blocks:    Set[str]           = set()
        self._seen_cites:     Dict[int, Set[str]] = {}
        self._lock                               = threading.Lock()
        self._citation_changes: List[Dict]       = []
        self._tracked_changes:  List[Dict]       = []
        self.enable_track_changes                = TRACK_CHANGES_AVAILABLE
        self._para_offset: int = 0
        self._pending_blocks: Dict[str, dict] = {}

    # ── Internal helpers ──────────────────────────────────────────────────────
    def _apply_tracked_fix(self, para, original, fixed, fix_type):
        if not track_changes or not self.enable_track_changes:
            tracked_replace(para, original, fixed, YELLOW, self.cite_style_id)
        else:
            try:
                full_text = _full_text(para)
                if full_text.find(original) != -1:
                    tracked_replace(para, original, fixed, YELLOW, self.cite_style_id)
                    self._tracked_changes.append({
                        "type": fix_type, "original": original, "fixed": fixed,
                        "para_idx": (list(self.doc.paragraphs).index(para)
                                     if para in self.doc.paragraphs else -1),
                        "author": COMMENT_AUTHOR,
                    })
                    self.log.info(f"Applied fix: '{original}' → '{fixed}' ({fix_type})")
                else:
                    self.log.warning(f"Could not find '{original}' in paragraph")
            except Exception as exc:
                self.log.error(f"Error applying tracked fix: {exc}")
                tracked_replace(para, original, fixed, YELLOW, self.cite_style_id)

    # ── Bibliography parsing ──────────────────────────────────────────────────
    def _parse_bibliography(self):
        raw_entries: List[Tuple[int, str]] = []
        in_bib = False
        for idx, para in enumerate(self.doc.paragraphs):
            txt = para.text.strip()
            if BIB_TAG_OPEN  in txt.lower(): in_bib = True;  continue
            if BIB_TAG_CLOSE in txt.lower(): in_bib = False; continue
            if in_bib and txt:
                raw_entries.append((idx, txt))

        if self.grobid is not None:
            try:
                BibliographyParser.batch_parse(
                    [txt for _, txt in raw_entries], self.grobid)
                self.log.info(
                    "GROBID batch parsed %d bib entries", len(raw_entries))
            except Exception as exc:
                self.log.warning("GROBID batch_parse error: %s — using regex fallback", exc)
                BibliographyParser._grobid_cache.clear()
        else:
            BibliographyParser._grobid_cache.clear()

        for idx, txt in raw_entries:
            e = BibliographyParser.parse_entry(txt)
            if not e:
                continue
            e["para_idx"] = idx
            self._bib_ordered.append(e)
            if e.get("abbrev"):
                self._org_tracker.record(e["abbrev"], e["full_author"], idx)
            if e.get("auto_acronym"):
                self._org_tracker.record(e["auto_acronym"], e["full_author"], idx)
            k1 = f"{_norm(e['display'])}|{e['year']}"
            k2 = f"{_norm(e['full_author'])}|{e['year']}"
            if k1 not in self.bibliography: self.bibliography[k1] = e
            if k2 not in self.bibliography: self.bibliography[k2] = e

        for iss in check_bibliography_structure(self._bib_ordered):
            pidx = iss.get("para_idx")
            iss["para"] = (self.doc.paragraphs[pidx]
                           if pidx is not None and pidx >= 0 else None)
            self._issues.append(iss)
            self._stats[iss["type"]] += 1

    # ── Body processing ───────────────────────────────────────────────────────
    def _process_body(self):
        self._para_offset = len(self.doc.paragraphs)

        for idx, para in enumerate(self.doc.paragraphs):
            txt = para.text
            if not txt.strip():
                continue
            if BIB_TAG_OPEN in txt.lower():
                break
            if _para_context(para) not in self.include_contexts:
                continue
            self._process_para(para, idx, txt)

        all_seen_tc_objs = []
        all_seen_tc = set()

        def _iter_unique_tcs(table):
            from docx.oxml.ns import qn as _qn
            for tr in table._tbl.iterchildren(_qn('w:tr')):
                for tc in tr.iterchildren(_qn('w:tc')):
                    yield tc

        def _scan_table(table, tbl_idx: int, cell_counter_ref: list):
            from docx.table import Table as DTable, _Cell
            for tc in _iter_unique_tcs(table):
                if id(tc) in all_seen_tc:
                    continue
                all_seen_tc_objs.append(tc) # Prevent GC to ensure id(tc) stays unique
                all_seen_tc.add(id(tc))

                cell = _Cell(tc, table)
                for cell_para in cell.paragraphs:
                    txt = cell_para.text
                    if not txt.strip():
                        continue
                    virtual_idx = (self._para_offset
                                   + tbl_idx * 1000
                                   + cell_counter_ref[0])
                    cell_counter_ref[0] += 1
                    self._process_para(
                        cell_para, virtual_idx, txt,
                        context_label=f"table {tbl_idx + 1}"
                    )
                for nested_table in cell.tables:
                    _scan_table(nested_table, tbl_idx, cell_counter_ref)

        for tbl_idx, table in enumerate(self.doc.tables):
            from docx.oxml.ns import qn as _qn
            first_tcs = {id(tc)
                         for tc in table._tbl.iterchildren(_qn('w:tr'))
                         for tc in tc.iterchildren(_qn('w:tc'))}
            if first_tcs and first_tcs.issubset(all_seen_tc):
                continue
            _scan_table(table, tbl_idx, [0])

        # Scan text boxes (w:txbxContent) — citations inside floating frames
        from docx.oxml.ns import qn as _qn
        from docx.text.paragraph import Paragraph as _Para
        txbx_counter = [0]
        for txbx in self.doc.element.body.iter(_qn('w:txbxContent')):
            for p_elem in txbx.iter(_qn('w:p')):
                try:
                    cell_para = _Para(p_elem, self.doc)
                    txt = _full_text(cell_para)
                    if not txt.strip():
                        continue
                    virtual_idx = self._para_offset + 9000 + txbx_counter[0]
                    txbx_counter[0] += 1
                    self._process_para(cell_para, virtual_idx, txt, context_label="text box")
                except Exception:
                    pass

    def _process_para(self, para, idx: int, txt: str,
                      context_label: str = "") -> None:
        # Always use full text to preserve line breaks, especially in tables
        txt = _full_text(para)
        
        if ApaFixer.needs_fix(txt):
            _, changes = ApaFixer.fix(txt)
            for ch in changes:
                ft = ch["fix_type"]
                loc_pfx = f"({context_label}) " if context_label else ""
                if ft == "year_range_not_apa":
                    self._add_issue(
                        "year_range", idx, para, ch["original"],
                        f"YEAR RANGE {loc_pfx}: '{ch['original']}' not valid APA.")
                    self._stats["year_range"] += 1
                else:
                    self._apply_tracked_fix(para, ch["original"], ch["fixed"], ft)
                    self._add_issue(
                        "format_fixed", idx, para, ch["fixed"],
                        f"FORMAT FIXED {loc_pfx}: '{ch['original']}' → '{ch['fixed']}'.")
                    self._stats["format_fixed"] += 1
            txt = _full_text(para)

        for om in re.finditer(r'([A-Z][A-Za-z\s]+?)\s*[\[(]([A-Z]{2,8})[\])]', txt):
            self._org_tracker.record(om.group(2), om.group(1).strip(), idx)

        for cite in CitationExtractor.extract(txt):
            self._validate_citation(cite, idx, para)
        self._apply_block_highlights()

    # ── Citation validation ───────────────────────────────────────────────────
    def _validate_citation(self, cite, para_idx, para):
        ct = cite.get("cite_type")

        if ct == "secondary":
            raw  = cite["raw"]
            oa   = cite.get("original_author", "?")
            oy   = cite.get("original_year", "?")
            auth = cite["author"]
            year = cite["year"]
            mr   = match_citation(auth, year, self.bibliography)
            if mr.key and mr.match_type in ("exact", "smart", "org_abbrev"):
                self.bibliography[mr.key]["cited"] = True
                self._cited_keys.add(mr.key)
            apply_style_to_text(para, raw, YELLOW, self.cite_style_id)
            self._add_issue(
                "secondary", para_idx, para, raw,
                f"SECONDARY: '({oa},{oy}, as cited in {auth},{year})' "
                f"— only '{auth},{year}' needs bib entry.")
            self._stats["secondary"] += 1
            return

        if ct == "multi_year":
            raw  = cite["raw"]
            auth = cite["author"]
            years = cite.get("years", [])
            self._stats["multi_year"] += 1
            any_miss = False
            for year in years:
                mr = match_citation(auth, year, self.bibliography)
                if mr.match_type in ("exact", "smart"):
                    self.bibliography[mr.key]["cited"] = True
                    self._cited_keys.add(mr.key)
                else:
                    any_miss = True
                    self._add_issue(
                        "multi_year_mismatches", para_idx, para, raw,
                        f"MULTI-YEAR: '{auth}, {year}' — no matching bib entry.")
                    self._stats["multi_year_mismatches"] += 1
            apply_style_to_text(para, raw, YELLOW if any_miss else GREEN, self.cite_style_id)
            return

        raw     = cite["raw"]
        auth    = cite["author"]
        year    = cite.get("year", "")
        blk_raw = cite.get("block_raw", raw)
        blk_sz  = cite.get("block_size", 1)

        # ── Multi-citation block bookkeeping ──────────────────────────────────
        is_new_block = blk_sz > 1 and blk_raw not in self._seen_blocks
        if is_new_block:
            self._seen_blocks.add(blk_raw)
            self._stats["multi_citation_blocks"] += 1
            
            has_colon_separator = ':' in blk_raw.strip("()")

            if not cite.get("block_order_ok", True) or has_colon_separator:
                sorted_ok, already_styled = sort_citation_block(para, blk_raw, YELLOW if not cite.get("block_order_ok", True) else None)
                inner = blk_raw.strip("()")
                segs  = [s.strip() for s in re.split(r'[;:]', inner) if s.strip()]
                def _sk(seg):
                    um = _RE_CITE_UNIT.search(seg)
                    if um:
                        return (_sort_key_name(_first_surname(
                                    um.group(1).strip().rstrip(","))),
                                um.group(2).strip())
                    return (seg, "")
                sorted_inner = "; ".join(sorted(segs, key=_sk))
                sorted_blk   = f"({sorted_inner})"
                if not cite.get("block_order_ok", True):
                    if not sorted_ok:
                        self._add_issue(
                            "alphabetical_violations", para_idx, para, blk_raw,
                            f"BLOCK ORDER: '{blk_raw}' not alphabetically sorted "
                            "— APA requires alphabetical order.")
                    self._stats["alphabetical_violations"] += 1
                elif has_colon_separator and sorted_ok:
                    self._add_issue(
                        "general_issues", para_idx, para, blk_raw,
                        f"CITATION SEPARATOR AUTO-FIXED: '{blk_raw}' → '{sorted_blk}' "
                        "— changed colon to semicolon per APA 7th.")
                        
                self._pending_blocks[sorted_blk] = {
                    "para": para, "para_idx": para_idx,
                    "all_green": True,
                    "orig_blk": blk_raw,
                    "sorted_blk": sorted_blk,
                    "was_sorted": True,
                    "already_styled": already_styled,
                }
            else:
                self._pending_blocks[blk_raw] = {
                    "para": para, "para_idx": para_idx,
                    "all_green": True,
                    "orig_blk": blk_raw,
                    "sorted_blk": blk_raw,
                    "was_sorted": False,
                }

        loc = cite.get("locator")
        if loc:
            for prob in check_locator(loc):
                self._add_issue(
                    "page_locator_errors", para_idx, para, raw,
                    f"LOCATOR '{loc}': {prob}.")
                self._stats["page_locator_errors"] += 1

        mr = match_citation(auth, year, self.bibliography)
        rk, mt = mr.key, mr.match_type

        self._citation_changes.append({
            "raw": raw, "author": auth, "year": year,
            "match_type": mt, "matched_ref": rk,
            "match_score": mr.score, "para_idx": para_idx,
        })
        self.log.debug(f"Citation: {auth} ({year}) → {mt} (ref={rk}, score={mr.score:.2f})")

        def _mark_block_yellow():
            if blk_sz > 1:
                for pb in self._pending_blocks.values():
                    if (pb.get("orig_blk") == blk_raw
                            or pb.get("sorted_blk", "") == blk_raw):
                        pb["all_green"] = False

        if mt in ("suffix_ambiguous", "suffix_mismatch"):
            _mark_block_yellow()
            if blk_sz == 1:
                apply_style_to_text(para, raw, YELLOW, self.cite_style_id)
            if mt == "suffix_ambiguous":
                self._add_issue(
                    "suffix_ambiguous", para_idx, para, raw,
                    f"SUFFIX AMBIGUOUS: '{raw}' — multiple entries; add a/b/c suffix.")
            else:
                ref = self.bibliography[rk]
                self._add_issue(
                    "suffix_mismatch", para_idx, para, raw,
                    f"SUFFIX MISSING: '{raw}' — entry is '{ref['display']} ({ref['year']})'. Add suffix.")
            self._stats[mt] += 1
            return

        if mt in ("exact", "smart", "org_abbrev"):
            ref = self.bibliography[rk]
            ref["cited"] = True
            self._cited_keys.add(rk)
            ew = check_etal_enforcement(auth, ref, self._et_al_min)
            if ew:
                _mark_block_yellow()
                ff = ApaFixer.fix_etal_expansion(auth, ref)
                if ff and ff != auth:
                    tracked_replace(para, auth, ff, YELLOW, self.cite_style_id)
                    self._add_issue(
                        "etal_enforcement", para_idx, para, raw,
                        f"👥 {ew} AUTO-FIXED → '{ff}'.",
                        original_text=auth,
                        corrected_text=ff,
                    )
                    self._stats["etal_autofixes"] += 1
                else:
                    if blk_sz == 1:
                        apply_style_to_text(para, raw, YELLOW, self.cite_style_id)
                    self._add_issue("etal_enforcement", para_idx, para, raw, f"👥 {ew}")
                self._stats["etal_enforcement"] += 1
            else:
                if blk_sz == 1:
                    apply_style_to_text(para, raw, GREEN, self.cite_style_id)
                self._stats["matched"] += 1

            if mt == "org_abbrev":
                abbr = auth.strip().upper()
                warn = self._org_tracker.check(abbr)
                if warn:
                    self._add_issue(
                        "org_abbrev_first_use", para_idx, para, raw, f"🏢 {warn}")
                    self._stats["org_abbrev_first_use"] += 1
                else:
                    self._add_issue(
                        "organization_cases", para_idx, para, raw,
                        f"ORG MATCHED: '{auth}' → '{ref['display']}' — verify.")
                    self._stats["organization_cases"] += 1

            sig = f"{_norm(auth)}|{year}"
            if sig in self._seen_cites.get(para_idx, set()):
                self._add_issue(
                    "duplicate_citations", para_idx, para, raw,
                    f"DUPLICATE: '{raw}' already cited in this paragraph/block.")
                self._stats["duplicate_citations"] += 1
            self._seen_cites.setdefault(para_idx, set()).add(sig)

        elif mt == "year_mismatch":
            ref = self.bibliography[rk]
            ref["cited"] = True
            self._cited_keys.add(rk)
            _mark_block_yellow()

            cb_cite = _strip_suffix(year)

            # ── FIX-53b: suppress auto-replace when any bib entry for this
            # author shares the same BASE year as the citation.
            # In that case the issue is a missing/wrong suffix, NOT a wrong
            # year.  Auto-replacing the year digit would be incorrect.
            # Downgrade to suffix_mismatch and flag without replacing. ────────
            same_base_year_exists = any(
                _strip_suffix(r["year"]) == cb_cite
                and _norm(r["full_author"]) == _norm(ref["full_author"])
                for r in self.bibliography.values()
            )
            if same_base_year_exists:
                if blk_sz == 1:
                    apply_style_to_text(para, raw, YELLOW, self.cite_style_id)
                self._add_issue(
                    "suffix_mismatch", para_idx, para, raw,
                    f"SUFFIX MISSING: '{raw}' — entry is "
                    f"'{ref['display']} ({ref['year']})'. "
                    "Please add the correct year suffix (a/b/c…).")
                self._stats["suffix_mismatch"] += 1
                return

            # ── FIX-41: surgical year replacement inside the citation string.
            # Instead of raw.replace(year, ref_year) which replaces ALL
            # occurrences of the four-digit string (can corrupt page numbers,
            # adjacent years, or "et al." context), use a targeted regex that
            # matches only the year token after the comma in the parenthetical
            # or narrative pattern. ─────────────────────────────────────────
            ref_year = ref["year"]

            def _replace_cite_year(s: str, old_yr: str, new_yr: str) -> str:
                """Replace the citation-year token only, not every occurrence."""
                # Pattern: comma + optional space + old_yr + optional locator + ) or ;
                def _sub(m):
                    return m.group(1) + new_yr + m.group(3)
                result = re.sub(
                    r'(,\s*)' + re.escape(old_yr) +
                    r'(\s*(?:,\s*pp?\.\s*\d[\d\-–]*)?(?:\s*\)|\s*;))',
                    lambda m: m.group(1) + new_yr + m.group(2),
                    s,
                    count=1,
                    flags=re.IGNORECASE,
                )
                if result == s:
                    # Fallback for narrative style: "Author (year)"
                    result = re.sub(
                        r'(\()' + re.escape(old_yr) + r'(\))',
                        lambda m: m.group(1) + new_yr + m.group(2),
                        s,
                        count=1,
                        flags=re.IGNORECASE,
                    )
                return result

            new_raw = _replace_cite_year(raw, year, ref_year)

            if blk_sz == 1:
                if new_raw != raw:
                    tracked_replace(para, raw, new_raw, YELLOW, self.cite_style_id)
                else:
                    apply_style_to_text(para, raw, YELLOW, self.cite_style_id)

            self._add_issue(
                "year_mismatch", para_idx, para, new_raw,
                f"AQ: Note that the citation of reference \"{auth}, {year}\" has been "
                f"changed to \"{auth}, {ref_year}\" to match with the reference list. "
                "Please confirm.",
                original_text=raw if raw != new_raw else None,
                corrected_text=new_raw if raw != new_raw else None,
            )
            self._stats["year_mismatch"] += 1

        elif mt == "spelling_mismatch":
            ref = self.bibliography[rk]
            _mark_block_yellow()
            if blk_sz == 1:
                apply_style_to_text(para, raw, YELLOW, self.cite_style_id)
            self._add_issue(
                "spelling_mismatch", para_idx, para, raw,
                f"SPELLING: cited '{auth}', bib has '{ref['display']}'.",
                original_text=auth,
                corrected_text=ref['display'],
            )
            self._stats["spelling_mismatch"] += 1

        else:
            _mark_block_yellow()
            if blk_sz == 1:
                apply_style_to_text(para, raw, YELLOW, self.cite_style_id)
            self._add_issue(
                "missing", para_idx, para, raw,
                f"AQ: The reference \"{raw}\" is cited in the text but not given in "
                "the list. Please provide complete publication details of this reference "
                "in the list or delete the citation from the text.")
            self._stats["missing"] += 1

    # ── Block highlight flush ─────────────────────────────────────────────────
    def _apply_block_highlights(self):
        from docx.oxml.ns import qn as _qn

        for blk_key, pb in self._pending_blocks.items():
            para     = pb["para"]
            blk_text = pb["sorted_blk"]
            colour   = GREEN if pb["all_green"] else YELLOW
            colour_str = "green" if pb["all_green"] else "yellow"

            if pb.get("already_styled"):
                for child in para._p:
                    if child.tag == _qn('w:ins'):
                        for r_el in child.findall(_qn('w:r')):
                            rpr = r_el.find(_qn('w:rPr'))
                            if rpr is not None:
                                hl_el = rpr.find(_qn('w:highlight'))
                                if hl_el is None:
                                    hl_el = OxmlElement('w:highlight')
                                    rpr.append(hl_el)
                                hl_el.set(_qn('w:val'), colour_str)
                                style_el = rpr.find(_qn('w:rStyle'))
                                if style_el is None:
                                    style_el = OxmlElement('w:rStyle')
                                    rpr.insert(0, style_el)
                                style_el.set(_qn('w:val'), self.cite_style_id)
                continue

            runs_text = "".join(r.text for r in para.runs)
            pos = runs_text.find(blk_text)
            if pos == -1:
                pos = runs_text.lower().find(blk_text.lower())
            if pos != -1:
                apply_highlight_by_span(para, pos, pos + len(blk_text),
                                        colour, self.cite_style_id)
            else:
                apply_style_to_text(para, pb["orig_blk"], colour, self.cite_style_id)

        self._pending_blocks.clear()

    # ── Unused references ─────────────────────────────────────────────────────
    def _flag_unused(self):
        seen: Set[Tuple[str, str]] = set()
        for k, ref in self.bibliography.items():
            key = (_norm(ref["full_author"]), ref["year"])
            if key in seen:
                continue
            seen.add(key)
            if ref.get("cited"):
                continue
            if any(self.bibliography.get(x, {}).get("raw") == ref["raw"]
                   for x in self._cited_keys):
                continue
            pidx = ref.get("para_idx")
            para = self.doc.paragraphs[pidx] if pidx is not None else None

            full_ref_text = ref.get("raw", ref["display"])
            msg = (
                f"AQ: The reference \"{ref['display']}, {ref['year']}\" is given in "
                "the list but not cited in the text. Please cite the reference in the "
                "text or delete from the list."
            )
            self._add_issue(
                "unused", pidx, para, ref["display"],
                msg,
                target_text=full_ref_text,
            )

            self._stats["unused"] += 1

    # ── Disambiguation detection ───────────────────────────────────────────────
    def _flag_disambiguation(self):
        by_surname: Dict[str, List[dict]] = {}
        for ref in self._bib_ordered:
            sn = _primary_surname(ref.get("full_author", ref["display"]))
            by_surname.setdefault(sn, []).append(ref)

        for sn, refs in by_surname.items():
            # Only flag when two+ *distinct* authors share the primary surname
            distinct = {_norm(r.get("full_author", r["display"])) for r in refs}
            if len(distinct) < 2:
                continue
            for ref in refs:
                pidx = ref.get("para_idx")
                para = self.doc.paragraphs[pidx] if pidx is not None else None
                self._add_issue(
                    "disambiguation", pidx, para, ref["display"],
                    f"DISAMBIGUATION (APA §8.20): Multiple authors share surname "
                    f"'{sn}' — include initials in all in-text citations "
                    f"(e.g. 'J. Smith, 2020' vs 'K. Smith, 2019').",
                )
                self._stats["disambiguation"] += 1

    # ── Comment insertion ─────────────────────────────────────────────────────
    def _insert_comments(self):
        _yr_re = re.compile(_YEAR_ANY, re.IGNORECASE)
        for iss in self._issues:
            p = iss.get("para")
            if p is None:
                continue

            search_src = iss.get("target_text") or iss.get("raw") or ""
            yr_m = _yr_re.search(search_src)
            year_token = yr_m.group(0) if yr_m else None

            if year_token:
                insert_comment(self.doc, p, iss["message"], target_text=year_token)
            else:
                hr = iss.get("highlight_run")
                if hr is not None:
                    insert_comment(self.doc, p, iss["message"], target_run=hr)
                else:
                    insert_comment(self.doc, p, iss["message"], target_text=search_src)

    # ── Issue tracker ─────────────────────────────────────────────────────────
    def _add_issue(self, itype, para_idx, para, raw, message, target_text=None,
                   original_text=None, corrected_text=None):
        if target_text is None:
            target_text = raw
        message = _to_smart_quotes(message)
        entry: Dict[str, Any] = {
            "type": itype,
            "para_idx": para_idx if para_idx is not None else -1,
            "para": para, "raw": raw, "message": message,
            "target_text": target_text,
        }
        if original_text is not None:
            entry["original_text"] = original_text
        if corrected_text is not None:
            entry["corrected_text"] = corrected_text
        with self._lock:
            self._issues.append(entry)

    # ── Public API ────────────────────────────────────────────────────────────
    def run(self) -> ValidationReport:
        self._parse_bibliography()
        self._process_body()
        self._flag_unused()
        self._flag_disambiguation()
        self._insert_comments()

        grobid_count = sum(
            1 for e in self._bib_ordered
            if e.get("grobid_title") is not None
        )
        if grobid_count:
            self._stats["grobid_parsed"] = grobid_count

        total_cites = sum(self._stats.get(k, 0) for k in (
            "matched", "missing", "year_mismatch", "suffix_mismatch",
            "suffix_ambiguous", "spelling_mismatch", "etal_enforcement",
            "secondary", "multi_year",
        ))
        return ValidationReport(
            issues=self._issues, stats=dict(self._stats),
            total_refs=len({r["raw"] for r in self.bibliography.values()}),
            total_cites=total_cites,
        )

    def save(self, output_path: str):
        self.doc.save(output_path)
        self.log.info("Saved: %s", output_path)

    def get_tracked_changes(self) -> List[Dict]:
        return self._tracked_changes

    def get_citation_changes(self) -> List[Dict]:
        return self._citation_changes

    def get_changes_summary(self) -> Dict:
        match_type_counts: Dict[str, int] = defaultdict(int)
        for change in self._citation_changes:
            match_type_counts[change["match_type"]] += 1

        return {
            "tracked_format_changes":     len(self._tracked_changes),
            "format_change_details":      self._tracked_changes,
            "total_citations_processed":  len(self._citation_changes),
            "citations_matched":          (match_type_counts.get("exact", 0)
                                           + match_type_counts.get("smart", 0)
                                           + match_type_counts.get("org_abbrev", 0)),
            "citations_not_found":        match_type_counts.get("not_found", 0),
            "citations_spelling_issues":  match_type_counts.get("spelling_mismatch", 0),
        }


def _find_runs_for_text(para, needle):
    """
    Return the list of runs in *para* whose combined text contains *needle*.
    Tries exact match first, then case-insensitive.
    """
    if not needle:
        return []
    full = "".join(r.text for r in para.runs)
    pos = full.find(needle)
    if pos == -1:
        pos = full.lower().find(needle.lower())
    if pos == -1:
        return []
    end = pos + len(needle)
    result, cur = [], 0
    for r in para.runs:
        rlen = len(r.text)
        if cur < end and cur + rlen > pos:
            result.append(r)
        cur += rlen
    return result


def _apply_style_to_runs(runs, style_name):
    """
    Apply a character style to a list of runs using XML-level insertion so the
    w:rStyle element is always present.
    """
    style_id = style_name.replace(" ", "")  # Word styleId strips spaces
    for r in runs:
        try:
            r_el = r._element
            rPr = r_el.find(qn("w:rPr"))
            if rPr is None:
                rPr = OxmlElement("w:rPr")
                r_el.insert(0, rPr)
            # Remove any existing rStyle to avoid duplicates
            for old in rPr.findall(qn("w:rStyle")):
                rPr.remove(old)
            rStyle = OxmlElement("w:rStyle")
            rStyle.set(qn("w:val"), style_id)
            rPr.insert(0, rStyle)
        except Exception:
            pass


def replace_text_in_para(para, target, replacement):
    """
    Replace target string with replacement inside a paragraph while trying to preserve
    run-level styling. If the target is fully inside a single run, replace it there.
    Otherwise, fall back to replacing it across combined runs.
    """
    if not target or target == replacement:
        return
    
    # Try single run match
    for run in para.runs:
        if target in run.text:
            run.text = run.text.replace(target, replacement)
            return
            
    # Try case-insensitive single run match
    for run in para.runs:
        if target.lower() in run.text.lower():
            pattern = re.compile(re.escape(target), re.IGNORECASE)
            run.text = pattern.sub(replacement, run.text)
            return

    # Fallback: target is split across runs
    runs = _find_runs_for_text(para, target)
    if runs:
        combined = "".join(r.text for r in runs)
        pattern = re.compile(re.escape(target), re.IGNORECASE)
        new_combined = pattern.sub(replacement, combined, count=1)
        runs[0].text = new_combined
        for r in runs[1:]:
            r.text = ""


def apply_apa_style_prep(doc):
    """
    Ensure 'cite_bib' character style exists, scan document paragraphs for citations,
    and apply 'cite_bib' character style to matched citations in-place.
    """
    # 1. Ensure 'cite_bib' character style exists
    try:
        doc.styles['cite_bib']
    except KeyError:
        try:
            cite_style = doc.styles.add_style('cite_bib', WD_STYLE_TYPE.CHARACTER)
            cite_style.font.superscript = False
        except Exception:
            pass

    # 2. Parse bibliography to build lookup mapping
    in_bib = False
    bibliography = {}
    for idx, para in enumerate(doc.paragraphs):
        txt = para.text.strip()
        if "<ref-open>" in txt.lower():
            in_bib = True
            continue
        if "<ref-close>" in txt.lower():
            in_bib = False
            continue
        if in_bib and txt:
            e = BibliographyParser.parse_entry(txt)
            if e:
                k1 = f"{_norm(e['display'])}|{e['year']}"
                k2 = f"{_norm(e['full_author'])}|{e['year']}"
                bibliography[k1] = e
                bibliography[k2] = e

    if not bibliography:
        # Fallback to REF-U paragraphs if no tags
        for idx, para in enumerate(doc.paragraphs):
            style_name = (para.style.name or "") if para.style else ""
            if "REF-U" in style_name and para.text.strip():
                e = BibliographyParser.parse_entry(para.text)
                if e:
                    k1 = f"{_norm(e['display'])}|{e['year']}"
                    k2 = f"{_norm(e['full_author'])}|{e['year']}"
                    bibliography[k1] = e
                    bibliography[k2] = e

    # 3. Apply style to runs for matched citations
    for para in doc.paragraphs:
        txt = para.text
        if "<ref-open>" in txt.lower():
            break
        style_name = (para.style.name or "") if para.style else ""
        if "REF-U" in style_name:
            continue
            
        citations = CitationExtractor.extract(txt)
        if not citations:
            continue
        for cite in citations:
            cite_type = cite["cite_type"]
            raw_text = cite["raw"]
            
            if cite_type == "parenthetical" and cite.get("block_size", 1) > 1:
                block_raw = cite["block_raw"]
                segs = [s.strip() for s in block_raw.replace("(", "").replace(")", "").split(";")]
                matched_all = True
                for seg in segs:
                    um = re.search(r"^(.*?)\s*,?\s*((?:19|20)\d{2}[a-z]?|n\.d\.|in\s+press)$", seg, re.IGNORECASE)
                    if um:
                        sa = um.group(1).strip().rstrip(",")
                        sy = um.group(2).strip()
                        mr = match_citation(sa, sy, bibliography)
                        rk = mr.key
                        if not (rk and mr.match_type in ("exact", "smart", "org_abbrev")):
                            matched_all = False
                    else:
                        matched_all = False
                if matched_all:
                    apply_style_to_text(para, block_raw, None, "cite_bib")
                    doc._dirty = True
            
            elif cite_type in ("parenthetical", "narrative"):
                auth = cite["author"]
                year = cite["year"]
                mr = match_citation(auth, year, bibliography)
                rk = mr.key
                if rk and mr.match_type in ("exact", "smart", "org_abbrev"):
                    apply_style_to_text(para, raw_text, None, "cite_bib")
                    doc._dirty = True


def process_document_apa(file_path):
    """
    Two-pass Name-Year (APA) reference validation, reordering, duplicate merging, and block sorting.
    Returns (doc, before_stats, after_stats, mapping, status_msg)
    """
    import tempfile
    import io
    import os

    doc = Document(file_path)
    apply_apa_style_prep(doc)

    # 1. Helper to parse bibliography from current doc state
    def parse_bib_from_doc(d):
        in_bib = False
        entries = []
        for idx, para in enumerate(d.paragraphs):
            txt = para.text.strip()
            if "<ref-open>" in txt.lower():
                in_bib = True
                continue
            if "<ref-close>" in txt.lower():
                in_bib = False
                continue
            if in_bib and txt:
                parsed = BibliographyParser.parse_entry(txt)
                if parsed:
                    entries.append({
                        "para": para,
                        "text": txt,
                        "parsed": parsed,
                        "idx": idx
                    })
        if not entries:
            for idx, para in enumerate(d.paragraphs):
                style_name = (para.style.name or "") if para.style else ""
                if "REF-U" in style_name and para.text.strip():
                    parsed = BibliographyParser.parse_entry(para.text)
                    if parsed:
                        entries.append({
                            "para": para,
                            "text": para.text,
                            "parsed": parsed,
                            "idx": idx
                        })
        return entries

    # Helper to compute sorting key
    def compute_sort_key(entry):
        parsed = entry["parsed"]
        full_author = parsed.get("full_author") or ""
        first_author_surname = parsed.get("sort_key") or ""
        year_str = _strip_suffix(parsed.get("year") or "")
        title = ""
        match = re.search(r"\((?:(?:19|20)\d{2}[a-z]?|n\.d\.|in\s+press)\)\.(.*)", entry["text"], re.IGNORECASE)
        if match:
            title = match.group(1).strip()
        else:
            title = entry["text"]
        title_norm = _norm(title)
        return (first_author_surname.lower(), year_str.lower(), title_norm.lower())

    # ── PASS 1: Sort and Suffix Resolution ──────────────────────────────────
    bib_entries = parse_bib_from_doc(doc)
    
    # Run a CitationProcessor initially to capture before stats
    cite_proc_before = CitationProcessor(file_path)
    report_before = cite_proc_before.run()
    before_stats = report_before.to_dict()

    # Reorder/Sort Bibliography Pass 1
    for entry in bib_entries:
        entry["sort_key"] = compute_sort_key(entry)
    bib_entries.sort(key=lambda x: x["sort_key"])

    # Suffix assignments
    groups = defaultdict(list)
    for entry in bib_entries:
        key = (_norm(entry["parsed"].get("full_author") or ""), _strip_suffix(entry["parsed"].get("year") or ""))
        groups[key].append(entry)

    suffix_mapping = {}
    for key, group in groups.items():
        if len(group) > 1:
            for i, entry in enumerate(group):
                suffix = chr(97 + i)
                bare_year = _strip_suffix(entry["parsed"]["year"])
                new_year = f"{bare_year}{suffix}"
                entry["new_year"] = new_year
                suffix_mapping[entry["text"]] = new_year
        else:
            entry = group[0]
            new_year = _strip_suffix(entry["parsed"]["year"])
            entry["new_year"] = new_year
            suffix_mapping[entry["text"]] = new_year

    # Update years in bibliography text runs
    for entry in bib_entries:
        old_year = entry["parsed"]["year"]
        new_year = entry["new_year"]
        if old_year != new_year:
            replace_text_in_para(entry["para"], f"({old_year})", f"({new_year})")

    # Update in-text citations for Pass 1
    # Build lookup for match_citation
    bibliography_lookup = {}
    for entry in bib_entries:
        e = entry["parsed"]
        k1 = f"{_norm(e['display'])}|{e['year']}"
        k2 = f"{_norm(e['full_author'])}|{e['year']}"
        entry_info = {
            "display": e["display"],
            "year": e["year"],
            "full_author": e["full_author"],
            "raw": entry["text"],
            "new_year": entry["new_year"]
        }
        bibliography_lookup[k1] = entry_info
        bibliography_lookup[k2] = entry_info

    # Track how we remap citations (old display -> new display)
    mapping = {}

    for para in doc.paragraphs:
        txt = para.text
        if "<ref-open>" in txt.lower():
            break
        citations = CitationExtractor.extract(txt)
        if not citations:
            continue
        citations.sort(key=lambda x: x["start"], reverse=True)

        for cite in citations:
            cite_type = cite["cite_type"]
            raw_text = cite["raw"]

            if cite_type == "parenthetical" and cite.get("block_size", 1) > 1:
                block_raw = cite["block_raw"]
                segs = [s.strip() for s in block_raw.replace("(", "").replace(")", "").split(";")]
                updated_units = []
                for seg in segs:
                    um = re.search(r"^(.*?)\s*,?\s*((?:19|20)\d{2}[a-z]?|n\.d\.|in\s+press)$", seg, re.IGNORECASE)
                    if um:
                        sa = um.group(1).strip().rstrip(",")
                        sy = um.group(2).strip()
                        mr = match_citation(sa, sy, bibliography_lookup)
                        rk = mr.key
                        if rk and mr.match_type in ("exact", "smart", "org_abbrev"):
                            ref_info = bibliography_lookup[rk]
                            new_y = ref_info["new_year"]
                            updated_units.append({
                                "surname": _first_surname(sa),
                                "text": f"{sa}, {new_y}"
                            })
                        else:
                            updated_units.append({
                                "surname": _first_surname(sa),
                                "text": seg
                            })
                    else:
                        updated_units.append({
                            "surname": _first_surname(seg),
                            "text": seg
                        })
                updated_units.sort(key=lambda x: x["surname"].lower())
                sorted_block_text = "(" + "; ".join(u["text"] for u in updated_units) + ")"
                if block_raw != sorted_block_text:
                    replace_text_in_para(para, block_raw, sorted_block_text)
                    mapping[block_raw] = sorted_block_text
                    
                # Style runs
                apply_style_to_text(para, sorted_block_text, None, "cite_bib")

            elif cite_type in ("parenthetical", "narrative"):
                auth = cite["author"]
                year = cite["year"]
                mr = match_citation(auth, year, bibliography_lookup)
                rk = mr.key
                if rk and mr.match_type in ("exact", "smart", "org_abbrev"):
                    ref_info = bibliography_lookup[rk]
                    new_y = ref_info["new_year"]
                    if year != new_y:
                        new_raw_text = raw_text.replace(year, new_y)
                        replace_text_in_para(para, raw_text, new_raw_text)
                        mapping[raw_text] = new_raw_text
                        
                        apply_style_to_text(para, new_raw_text, None, "cite_bib")
                    else:
                        apply_style_to_text(para, raw_text, None, "cite_bib")

    # Reorder bibliography paragraphs physically
    if bib_entries:
        body = doc._element.body
        indices = []
        for entry in bib_entries:
            try:
                idx = body.index(entry["para"]._element)
                indices.append(idx)
            except ValueError:
                pass
        if indices:
            anchor = min(indices)
            for entry in bib_entries:
                p = entry["para"]._element
                if p.getparent() == body:
                    body.remove(p)
            insert_idx = anchor
            for entry in bib_entries:
                body.insert(insert_idx, entry["para"]._element)
                insert_idx += 1


    # ── PASS 2: Find duplicates & Merge ──────────────────────────────────────
    bib_entries = parse_bib_from_doc(doc)
    
    # Perform fuzzy duplicate matching (similarity threshold 0.85)
    removed_paras = set()
    merged_count = 0
    n = len(bib_entries)

    for i in range(n):
        entry_a = bib_entries[i]
        if entry_a["para"] in removed_paras:
            continue
        for j in range(i + 1, n):
            entry_b = bib_entries[j]
            if entry_b["para"] in removed_paras:
                continue
                
            len_a = len(entry_a["text"])
            len_b = len(entry_b["text"])
            if len_a == 0 or len_b == 0:
                continue
            if min(len_a, len_b) / max(len_a, len_b) < 0.6:
                continue
                
            ratio = difflib.SequenceMatcher(None, entry_a["text"], entry_b["text"]).ratio()
            if ratio > 0.85:
                removed_paras.add(entry_b["para"])
                merged_count += 1
                
                body = doc._element.body
                p_element = entry_b["para"]._element
                if p_element.getparent() == body:
                    body.remove(p_element)

    # If any duplicates were removed, we re-run suffix resolution and ordering
    if merged_count > 0:
        bib_entries = [e for e in bib_entries if e["para"] not in removed_paras]
        
        # Suffix resolution again
        groups = defaultdict(list)
        for entry in bib_entries:
            key = (_norm(entry["parsed"].get("full_author") or ""), _strip_suffix(entry["parsed"].get("year") or ""))
            groups[key].append(entry)

        for key, group in groups.items():
            if len(group) > 1:
                for i, entry in enumerate(group):
                    suffix = chr(97 + i)
                    bare_year = _strip_suffix(entry["parsed"]["year"])
                    new_year = f"{bare_year}{suffix}"
                    entry["new_year"] = new_year
            else:
                entry = group[0]
                entry["new_year"] = _strip_suffix(entry["parsed"]["year"])

        # Update years in bibliography text runs
        for entry in bib_entries:
            old_year = entry["parsed"]["year"]
            new_year = entry["new_year"]
            if old_year != new_year:
                replace_text_in_para(entry["para"], f"({old_year})", f"({new_year})")

    # Determine status message
    if merged_count > 0:
        status_msg = f"Two-pass validation: Pass 1 renumbered, Pass 2 {merged_count} duplicate reference{'s' if merged_count > 1 else ''} removed and renumbered."
    elif len(mapping) > 0:
        status_msg = "Renumbering completed successfully."
    else:
        status_msg = "Validation completed."

    # Return doc and stats
    temp_buffer = io.BytesIO()
    doc.save(temp_buffer)
    temp_buffer.seek(0)
    
    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tf:
        tf.write(temp_buffer.read())
        temp_file_path = tf.name
        
    try:
        cite_proc_after = CitationProcessor(temp_file_path)
        report_after = cite_proc_after.run()
        after_stats = report_after.to_dict()
    finally:
        try:
            os.remove(temp_file_path)
        except Exception:
            pass

    return doc, before_stats, after_stats, mapping, status_msg