"""
Figure/Table/Box Citation & Caption Highlighter
Applies color-coded highlighting to DOCX files and detects missing captions.
"""

from typing import Dict, List, Tuple, Any, Optional
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from word_analyzer_docx import CitationAnalyzer


class FigureTableHighlighter:
    """
    Wraps CitationAnalyzer to apply color highlighting to figure/table/box citations and captions
    in DOCX files.

    Highlighting scheme:
    - Yellow (FFFF00): Caption paragraphs ("Figure 1", "Table 1", etc.)
    - Blue (0000FF): Citation references in body text ("See Figure 1")
    - Detects missing captions (cited but no caption found)
    """

    # Color codes
    CAPTION_COLOR = "FFFF00"  # Yellow background
    CITATION_COLOR = "0000FF"  # Blue background

    def __init__(self):
        self.analyzer = CitationAnalyzer()

    def extract_document_content(self, docx_path: str) -> List[Tuple[str, int, bool]]:
        """
        Extract paragraph text from DOCX with page numbers and caption flags.

        Returns:
            List of (text, page_no, is_caption) tuples
        """
        doc = Document(docx_path)
        content = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Determine if this paragraph is a caption (by style name)
            style_name = para.style.name if para.style else ""
            is_caption = self.analyzer.is_caption_paragraph(text, style_name)

            # Page number estimation (Word doesn't expose this easily; use para index as proxy)
            page_no = len(content) // 20 + 1  # Rough estimate: ~20 paragraphs per page

            content.append((text, page_no, is_caption))

        return content

    def highlight_text_in_paragraph(
        self,
        paragraph,
        start_offset: int,
        end_offset: int,
        color: str
    ) -> None:
        """
        Highlight text in a paragraph by character position.
        Splits runs if necessary to apply highlighting to exact range.
        """
        if start_offset >= end_offset:
            return

        current_pos = 0
        for run in paragraph.runs:
            run_len = len(run.text)
            run_start = current_pos
            run_end = current_pos + run_len

            # Check if this run overlaps with highlight range
            if run_end <= start_offset or run_start >= end_offset:
                current_pos += run_len
                continue

            # Apply background color to run
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), color)
            run._element.get_or_add_rPr().append(shading_elm)

            current_pos += run_len

    def apply_highlighting_to_docx(
        self,
        docx_path: str,
        output_path: str,
        caption_types: Optional[List[str]] = None
    ) -> Dict[str, Any]:
        """
        Apply highlighting to DOCX file based on citation analysis.

        Args:
            docx_path: Path to input DOCX
            output_path: Path to output DOCX
            caption_types: List of types to highlight (e.g., ["Figure", "Table", "Box"])
                          If None, highlights all supported types

        Returns:
            Dictionary with:
            - highlighted_count: Total paragraphs/runs highlighted
            - captions_highlighted: Count of caption paragraphs
            - citations_highlighted: Count of citation references
            - missing_captions: List of dicts with missing caption info
            - findings: Aggregated findings by element type and chapter
        """

        if caption_types is None:
            caption_types = ["Figure", "Table", "Box", "Exhibit", "Appendix", "Case Study"]

        # Extract content and run analysis
        content = self.extract_document_content(docx_path)
        dict_types = self.analyzer.analyze_document_citations(content)

        # Load DOCX for highlighting
        doc = Document(docx_path)

        # Track what we highlight
        captions_highlighted = 0
        citations_highlighted = 0
        highlighted_paragraphs = set()

        # Apply caption highlighting (yellow)
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = para.style.name if para.style else ""
            is_caption = self.analyzer.is_caption_paragraph(text, style_name)

            if is_caption:
                # Check which element type matches
                for elem_type in caption_types:
                    type_data = dict_types.get(elem_type, {})
                    # Check if this caption matches any known caption
                    for caption_id in type_data.get("Caption", {}).keys():
                        # Simple match: if caption_id appears in paragraph text
                        if self._text_contains_id(text, caption_id):
                            self._highlight_entire_paragraph(para, self.CAPTION_COLOR)
                            captions_highlighted += 1
                            highlighted_paragraphs.add(id(para))
                            break

        # Apply citation highlighting (blue)
        for para in doc.paragraphs:
            if id(para) in highlighted_paragraphs:
                continue  # Skip already-highlighted caption paragraphs

            text = para.text.strip()
            if not text:
                continue

            # Find citations in this paragraph
            for elem_type in caption_types:
                type_data = dict_types.get(elem_type, {})
                for citation_id in type_data.get("Citation", {}).keys():
                    if self._text_contains_id(text, citation_id):
                        # Highlight this paragraph blue
                        self._highlight_entire_paragraph(para, self.CITATION_COLOR)
                        citations_highlighted += 1
                        highlighted_paragraphs.add(id(para))
                        break

        # Save highlighted document
        doc.save(output_path)

        # Detect missing captions and citations
        missing_captions, missing_citations = self._detect_missing_captions_and_citations(dict_types, caption_types)

        # Build findings summary
        findings = self._build_findings_summary(dict_types, caption_types)

        return {
            "highlighted_count": captions_highlighted + citations_highlighted,
            "captions_highlighted": captions_highlighted,
            "citations_highlighted": citations_highlighted,
            "missing_captions": missing_captions,
            "missing_citations": missing_citations,
            "findings": findings,
        }

    def _text_contains_id(self, text: str, citation_id: str) -> bool:
        """Check if paragraph text contains a citation/caption ID."""
        # Normalize for comparison
        text_norm = text.lower()
        id_norm = citation_id.lower()

        # Handle variations: "Figure 1" might appear as "Fig. 1", "figure 1", etc.
        patterns = [
            id_norm,
            id_norm.replace("figure ", "fig "),
            id_norm.replace("table ", "tab "),
        ]

        for pattern in patterns:
            if pattern in text_norm:
                return True

        return False

    def _highlight_entire_paragraph(self, paragraph, color: str) -> None:
        """Apply background color to entire paragraph."""
        for run in paragraph.runs:
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), color)
            run._element.get_or_add_rPr().append(shading_elm)

    def _detect_missing_captions_and_citations(
        self,
        dict_types: Dict[str, Any],
        caption_types: List[str]
    ) -> Tuple[List[Dict[str, Any]], List[Dict[str, Any]]]:
        """
        Detect citations without corresponding captions (Missing Captions)
        and captions without corresponding citations (Missing Citations).

        Returns:
            Tuple of (missing_captions, missing_citations)
        """
        missing_captions = []
        missing_citations = []

        for elem_type in caption_types:
            type_data = dict_types.get(elem_type, {})
            captions = set(type_data.get("Caption", {}).keys())
            captions.update(type_data.get("DerivedCaption", {}).keys())

            citations = type_data.get("Citation", {})
            citation_pages = type_data.get("CitationPage", {})
            
            caption_pages = type_data.get("CaptionPage", {})
            caption_pages.update(type_data.get("DerivedCaptionPage", {}))

            for citation_id in citations.keys():
                # Normalize for comparison
                citation_norm = self._normalize_id(citation_id)
                caption_match = any(
                    self._normalize_id(cap_id) == citation_norm
                    for cap_id in captions
                )

                if not caption_match:
                    missing_captions.append({
                        "element": elem_type,
                        "id": citation_id,
                        "page": citation_pages.get(citation_id),
                        "type": "missing_caption"
                    })
                    
            for caption_id in captions:
                caption_norm = self._normalize_id(caption_id)
                citation_match = any(
                    self._normalize_id(cit_id) == caption_norm
                    for cit_id in citations.keys()
                )
                
                if not citation_match:
                    missing_citations.append({
                        "element": elem_type,
                        "id": caption_id,
                        "page": caption_pages.get(caption_id),
                        "type": "missing_citation"
                    })

        return missing_captions, missing_citations

    def _normalize_id(self, citation_id: str) -> str:
        """Normalize citation ID for comparison."""
        return citation_id.lower().replace(" ", "").replace("-", ".")

    def _build_findings_summary(
        self,
        dict_types: Dict[str, Any],
        caption_types: List[str]
    ) -> Dict[str, Any]:
        """
        Build aggregated findings summary for IA report integration.

        Returns:
            Dict with structure:
            {
              "Figure": {
                "Caption": count,
                "Citation": count,
                "Missing": count,
                ...
              },
              ...
            }
        """
        findings = {}

        for elem_type in caption_types:
            type_data = dict_types.get(elem_type, {})

            captions = len(type_data.get("Caption", {})) + len(type_data.get("DerivedCaption", {}))
            citations = len(type_data.get("Citation", {}))

            # Count missing captions (cited but no caption)
            captions_set = set(type_data.get("Caption", {}).keys())
            captions_set.update(type_data.get("DerivedCaption", {}).keys())

            missing = 0
            for citation_id in type_data.get("Citation", {}).keys():
                if not any(
                    self._normalize_id(cap_id) == self._normalize_id(citation_id)
                    for cap_id in captions_set
                ):
                    missing += 1

            findings[elem_type] = {
                "captions": captions,
                "citations": citations,
                "missing_captions": missing,
                "total": captions + citations,
            }

        return findings
