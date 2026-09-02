import logging
import os
import urllib.parse
from typing import Any, Dict, Optional

from fastapi import HTTPException
from sqlalchemy.orm import Session

logger = logging.getLogger("app.domains.review.service")

from app.models import File
from app.processing.docx_to_xhtml import DocxToXhtmlEngine
from app.processing.docx_to_xhtml_runs import DocxToXhtmlRunsEngine
from app.processing.xhtml_to_docx import XhtmlToDocxEngine
from app.processing.xhtml_to_docx_delta import XhtmlToDocxDeltaEngine
from app.utils.utils.structuring_lib.annotator import normalize_structural_tag_case
import re

def _build_body_to_allp_translator(doc):
    """Return `body_idx → all-w:p idx` translator for `doc`.

    Two indexing schemes co-exist in this codebase:

    - `enumerate(doc.paragraphs)` — body-level only, skips paragraphs inside
      tables/cells. Used by `CitationProcessor` (APA validator).
    - `doc.element.body.iter(w:p)` — every `<w:p>` in document order,
      table-cell paragraphs included. Used by `docx_to_xhtml_runs`, and
      therefore by every `data-para-idx` value the WYSIWYG editor sees.

    When there's a table above the References section the two schemes
    diverge (body-level index is smaller than all-w:p index by the number
    of cell paragraphs). Reference entries reported with a body-level index
    then map onto whatever table-cell paragraph happens to sit at the same
    all-w:p position in the editor — which is why REF{n} bookmarks stamped
    from that index land inside table cells.

    The returned function accepts a body-level index and returns the
    equivalent all-w:p index. `None`/negative inputs pass through as `-1`.
    "Virtual" encoded indices produced by `CitationProcessor` for
    table/textbox contexts (values >= len(doc.paragraphs)) are returned
    unchanged — they were never meant to be looked up in the frontend.
    """
    from docx.oxml.ns import qn as _qn_body
    body_paragraphs = list(doc.paragraphs)
    all_p_map = {
        p_elem: i
        for i, p_elem in enumerate(doc.element.body.iter(_qn_body("w:p")))
    }

    def _translate(body_idx):
        if body_idx is None:
            return -1
        if body_idx < 0 or body_idx >= len(body_paragraphs):
            return body_idx
        elem = body_paragraphs[body_idx]._element
        return all_p_map.get(elem, body_idx)

    return _translate


def _get_full_reference_text(para, doc_paragraphs, para_index_map) -> str:
    idx = para_index_map.get(para._element)
    if idx is None:
        return para.text.strip()
    
    text_parts = [para.text.strip()]
    curr_idx = idx + 1
    while curr_idx < len(doc_paragraphs):
        next_p = doc_paragraphs[curr_idx]
        next_text = next_p.text.strip()
        if not next_text:
            break
        
        if next_p.style and next_p.style.name == "REF-N":
            break
            
        is_doi = next_text.lower().startswith(("doi:", "https://doi.org/", "http://dx.doi.org/", "url:", "http://", "https://"))
        is_short_doi = len(next_text) < 100 and bool(re.search(r"\b10\.\d{4,9}/", next_text))
        
        if is_doi or is_short_doi:
            text_parts.append(next_text)
            curr_idx += 1
        else:
            break
            
    return " ".join(text_parts)


ADDITIONAL_REVIEW_STYLES = [
    "ACK1", "ACKTXT", "ANS-NL-FIRST", "ANS-NL-MID", "APX", "APX-TXT", "APX-TXT-FLUSH", "APXAU",
    "APXST", "APXT", "TXT", "TXT-FLUSH",
    "BIB", "BIBH1", "BIBH2", "BL-FIRST", "BL-LAST", "BL-MID", "BL2-MID", "BL3-MID", "BL4-MID", "BL5-MID", "BL6-MID",
    "BX1-BL-FIRST", "BX1-BL-LAST", "BX1-BL-MID", "BX1-BL2-MID", "BX1-EXT-ONLY", "BX1-EQ-FIRST", "BX1-EQ-MID", "BX1-EQ-LAST", "BX1-EQ-ONLY",
    "BX1-FN", "BX1-H1", "BX1-H2", "BX1-H3", "BX1-L1", "BX1-MCUL-FIRST", "BX1-MCUL-LAST", "BX1-MCUL-MID", "BX1-NL-FIRST", "BX1-NL-LAST", "BX1-NL-MID",
    "BX1-OUT1-FIRST", "BX1-OUT1-MID", "BX1-OUT2", "BX1-OUT2-LAST", "BX1-OUT3", "BX1-QUO", "BX1-QUO-AU", "BX1-TTL", "BX1-TXT", "BX1-TXT-DC", "BX1-TXT-FIRST", "BX1-TYPE", "BX1-UL-FIRST", "BX1-UL-LAST", "BX1-UL-MID",
    "CAU", "CHAP", "CN", "COQ", "COQA", "COUT-1", "COUT-2", "COUT-BL", "COUTH1", "COUT-NL-FIRST", "COUT-NL-MID", "CPAU", "CPT", "CST", "CT",
    "DIA-FIRST", "DIA-LAST", "DIA-MID", "EQ-FIRST", "EQ-LAST", "EQ-MID", "EQ-ONLY", "EQN-FIRST", "EQN-LAST", "EQN-MID", "EQN-ONLY",
    "EXT-FIRST", "EXT-LAST", "EXT-MID", "EXT-ONLY", "FIG-CRED", "FIG-LEG", "FN",
    "H1", "H2", "H3", "H4", "H5", "H6",
    "KP1", "KP-BL-FIRST", "KP-BL-LAST", "KP-BL-MID", "KP-NL-FIRST", "KP-NL-LAST", "KP-NL-MID", "KT-BL-FIRST", "KT-NL-FIRST",
]


def resolve_processed_target(db: Session, *, file_id: int):
    file_record = db.query(File).filter(File.id == file_id).first()
    if not file_record:
        raise HTTPException(status_code=404, detail="File not found")

    filename_lower = file_record.filename.lower()
    if not (filename_lower.endswith(".docx") or filename_lower.endswith(".doc")):
        raise HTTPException(
            status_code=400,
            detail=f"Technical or structuring review is only supported for Word documents. Got: {file_record.filename}"
        )

    original_path = file_record.path
    if original_path.endswith("_Processed.docx"):
        processed_path = original_path
        processed_filename = os.path.basename(original_path)
    else:
        dir_name = os.path.dirname(original_path)
        base_name = os.path.basename(original_path)
        name_only = os.path.splitext(base_name)[0]
        processed_filename = f"{name_only}_Processed.docx"
        processed_path = os.path.join(dir_name, processed_filename)

        # Fallback if _Processed.docx does not exist
        if not os.path.exists(processed_path) and os.path.exists(original_path):
            processed_path = original_path
            processed_filename = os.path.basename(original_path)

    return {
        "file_record": file_record,
        "processed_path": processed_path,
        "processed_filename": processed_filename,
    }


def build_review_page_state(
    db: Session,
    *,
    file_id: int,
    collabora_public_url: str,
    wopi_base_url: str,
    extract_document_structure_func,
    get_rules_loader_func,
):
    resolved = resolve_processed_target(db, file_id=file_id)
    processed_path = resolved["processed_path"]
    processed_filename = resolved["processed_filename"]

    if not os.path.exists(processed_path):
        return {
            "status": "error",
            "error_message": "Processed file not found. Please run Structuring process first.",
        }

    # Keep current side effect/validation behavior: structure extraction runs even
    # though the template does not currently consume the result directly.
    extract_document_structure_func(processed_path)

    # Build styles list based on client (Springer vs. Other Clients)
    from app.utils.client_styles import get_paragraph_styles_for_client
    from app.models import File

    file_record = db.query(File).filter(File.id == file_id).first()
    client_name = None
    if file_record and file_record.chapter:
        ch = file_record.chapter
        client_name = getattr(ch, "client", None)
        if not client_name or not isinstance(client_name, str):
            proj = getattr(ch, "project_rel", None)
            if proj and proj.client:
                client_name = (
                    getattr(proj.client, "company", None)
                    or getattr(proj.client, "name_company", None)
                    or getattr(proj.client, "division", None)
                )

    client_styles = get_paragraph_styles_for_client(client_name)
    styles = set(client_styles)

    try:
        from docx import Document

        doc = Document(processed_path)

        def _collect(paragraphs):
            for para in paragraphs:
                name = getattr(getattr(para, "style", None), "name", None)
                if name:
                    styles.add(name)

        _collect(doc.paragraphs)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    _collect(cell.paragraphs)

        char_styles = set()
        _CAPTION_CHAR_STYLES = {"FigureCitation", "TableCitation", "FIG-NUM", "TN"}
        from docx.enum.style import WD_STYLE_TYPE
        for style in doc.styles:
            try:
                if style.type == WD_STYLE_TYPE.CHARACTER and style.name != "Default Paragraph Font":
                    n = style.name
                    # Only include pipeline-created styles; skip Word built-ins
                    if (n.startswith("bib_") or n.startswith("cite_")
                            or (n.isalpha() and n.islower())
                            or n in _CAPTION_CHAR_STYLES):
                        char_styles.add(n)
            except Exception:
                pass
    except Exception as exc:  # pragma: no cover - defensive
        logger.warning("Could not read applied styles from %s: %s", processed_path, exc)
        char_styles = set()

    style_list = sorted(styles)
    char_style_list = sorted(char_styles)

    wopi_src = f"{wopi_base_url}/wopi/files/{file_id}/structuring"
    wopi_src_encoded = urllib.parse.quote(wopi_src, safe="")
    collabora_url = (
        f"{collabora_public_url}/browser/dist/cool.html"
        f"?WOPISrc={wopi_src_encoded}"
        f"&lang=en"
    )

    return {
        "status": "ok",
        "file": resolved["file_record"],
        "filename": processed_filename,
        "collabora_url": collabora_url,
        "styles": style_list,
        "char_styles": char_style_list,
    }


def save_changes(
    db: Session,
    *,
    file_id: int,
    changes: Dict[str, Any],
    update_document_structure_func,
    logger,
):
    from docx import Document

    resolved = resolve_processed_target(db, file_id=file_id)
    processed_path = resolved["processed_path"]
    processed_filename = resolved["processed_filename"]

    if not os.path.exists(processed_path):
        logger.warning(f"Save failed: Processed file not found at {processed_path}")
        raise HTTPException(status_code=404, detail="Processed file not found")

    modifications = changes.get("changes", {})

    # Handle paragraph_styles changes (from new WYSIWYG editor)
    if "paragraph_styles" in modifications:
        try:
            doc = Document(processed_path)
            style_changes = modifications["paragraph_styles"]

            block_idx = 0
            for elem in doc.element.body:
                from lxml import etree
                tag = etree.QName(elem.tag).localname

                if tag == "p":
                    # Find matching change for this paragraph block
                    for change in style_changes:
                        if change["para_index"] == block_idx:
                            new_style = normalize_structural_tag_case(change["style_name"])
                            para = doc.paragraphs[block_idx]
                            try:
                                para.style = new_style
                            except Exception as e:
                                logger.warning(f"Could not apply style '{new_style}' to paragraph: {e}")
                            break
                    block_idx += 1
                elif tag == "tbl":
                    # Tables count as a block index
                    for change in style_changes:
                        if change["para_index"] == block_idx:
                            # Table styling is limited; skip for now
                            break
                    block_idx += 1

            file_record = resolved["file_record"]
            try:
                from app.domains.files import version_service as vs
                vs.archive_existing_file(
                    db,
                    existing_file=file_record,
                    base_path=os.path.dirname(processed_path),
                    uploaded_by_id=None,
                    source_path=processed_path,
                )
                file_record.version += 1
            except Exception as _e:
                logger.warning(f"Version archive failed (non-fatal): {_e}")

            doc.save(processed_path)
            db.commit()
            logger.info(f"Successfully updated paragraph styles for {processed_filename}")
            return {"status": "success"}
        except Exception as exc:
            logger.error(f"Error saving paragraph styles: {exc}", exc_info=True)
            raise HTTPException(status_code=500, detail=f"Failed to save changes: {str(exc)}")

    # Handle legacy XHTML-based structure updates
    try:
        file_record = resolved["file_record"]
        try:
            from app.domains.files import version_service as vs
            vs.archive_existing_file(
                db,
                existing_file=file_record,
                base_path=os.path.dirname(processed_path),
                uploaded_by_id=None,
                source_path=processed_path,
            )
            file_record.version += 1
        except Exception as _e:
            logger.warning(f"Version archive failed (non-fatal): {_e}")

        success = update_document_structure_func(processed_path, processed_path, modifications)
        if success:
            db.commit()
            logger.info(f"Successfully updated structure for {processed_filename}")
            return {"status": "success"}
        raise Exception("update_document_structure returned False")
    except Exception as exc:
        logger.error(f"Error saving changes: {exc}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to save changes: {str(exc)}")


def import_word_comments_into_xhtml(
    *,
    db: Session,
    file_id: int,
    docx_path: str,
    xhtml_content: str,
    logger,
) -> str:
    """
    Read native Word comments from the DOCX and:
      1. Upsert each into the `comments` table for this file. Idempotent — keyed
         on a deterministic UUID derived from (file_id, Word comment id), so
         re-importing the same DOCX does not duplicate or overwrite rows.
      2. Wrap every paragraph covered by a Word comment's range in a
         <span data-comment-id="UUID"> in the XHTML so the in-browser editor
         renders the comment highlight and the panel finds it.

    Granularity is paragraph-level (matches the export side). Sub-paragraph
    ranges in Word are widened to the paragraph(s) they cover.
    """
    import uuid as _uuid
    from lxml import etree
    import lxml.html

    W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
    def qn(tag: str) -> str:
        return f"{{{W}}}{tag}"

    try:
        from docx import Document
        from docx.opc.constants import RELATIONSHIP_TYPE as RT
        doc = Document(docx_path)
    except Exception as e:
        logger.warning(f"Word-comment import: could not open {docx_path}: {e}")
        return xhtml_content

    comments_part = None
    for rel in doc.part.rels.values():
        if rel.reltype == RT.COMMENTS:
            comments_part = rel.target_part
            break
    if comments_part is None:
        return xhtml_content

    try:
        comments_root = etree.fromstring(comments_part.blob)
    except Exception as e:
        logger.warning(f"Word-comment import: could not parse comments.xml: {e}")
        return xhtml_content

    word_comments: Dict[str, Dict[str, str]] = {}
    for cmt in comments_root.findall(qn("comment")):
        cid = cmt.get(qn("id"))
        if cid is None:
            continue
        word_comments[cid] = {
            "author": cmt.get(qn("author")) or "Unknown",
            "text": "".join(t.text or "" for t in cmt.iter(qn("t"))).strip(),
        }
    if not word_comments:
        return xhtml_content

    # Walk paragraphs in document order, tracking which comment ranges are open.
    # A range can span multiple paragraphs — each one gets a span wrapper.
    body = doc.element.body
    open_ids: set = set()
    bookmark_for_comment_id: Dict[str, list] = {}

    def first_bookmark(p_el) -> Optional[str]:
        for bm in p_el.iter(qn("bookmarkStart")):
            name = bm.get(qn("name"))
            if name:
                return name
        return None

    for p_el in body.iter(qn("p")):
        opens_in = set()
        closes_in = set()
        for child in p_el.iter():
            if child.tag == qn("commentRangeStart"):
                cid = child.get(qn("id"))
                if cid is not None:
                    opens_in.add(cid)
                    open_ids.add(cid)
            elif child.tag == qn("commentRangeEnd"):
                cid = child.get(qn("id"))
                if cid is not None:
                    closes_in.add(cid)

        covering = open_ids | opens_in
        if covering:
            bm = first_bookmark(p_el)
            if bm:
                for cid in covering:
                    if cid in word_comments:
                        bookmark_for_comment_id.setdefault(cid, []).append(bm)

        for cid in closes_in:
            open_ids.discard(cid)

    if not bookmark_for_comment_id:
        return xhtml_content

    NAMESPACE = _uuid.uuid5(_uuid.NAMESPACE_DNS, "cms.comments")

    def uuid_for(word_comment_id: str) -> str:
        return str(_uuid.uuid5(NAMESPACE, f"file:{file_id}|wc:{word_comment_id}"))

    uuid_for_wid = {cid: uuid_for(cid) for cid in bookmark_for_comment_id.keys()}

    # Upsert: insert if missing, leave existing rows alone so user edits in the
    # in-browser editor are not clobbered by a re-import.
    from app.models import Comment as CommentModel
    try:
        existing_rows = (
            db.query(CommentModel)
            .filter(
                CommentModel.file_id == file_id,
                CommentModel.comment_uuid.in_(list(uuid_for_wid.values())),
            )
            .all()
        )
        existing_uuids = {r.comment_uuid for r in existing_rows}
        inserted = 0
        for cid, meta in word_comments.items():
            u = uuid_for_wid.get(cid)
            if not u or u in existing_uuids:
                continue
            db.add(
                CommentModel(
                    file_id=file_id,
                    comment_uuid=u,
                    text=meta["text"],
                    author_name=meta["author"],
                    resolved=False,
                )
            )
            inserted += 1
        if inserted:
            db.commit()
            logger.info(f"Imported {inserted} Word comment(s) into DB for file {file_id}")
    except Exception as e:
        logger.warning(f"Word-comment import: DB upsert failed: {e}", exc_info=True)
        try:
            db.rollback()
        except Exception:
            pass
        # Continue — we can still annotate the XHTML even if the DB write failed,
        # but realistically the panel won't show anything without the rows. Bail.
        return xhtml_content

    # Wrap matching XHTML paragraphs with <span data-comment-id="UUID">.
    try:
        root = lxml.html.fromstring(xhtml_content)
    except Exception as e:
        logger.warning(f"Word-comment import: could not parse XHTML: {e}")
        return xhtml_content

    block_by_bm: Dict[str, Any] = {}
    for el in root.iter():
        bm = el.get("data-bookmark")
        if bm and bm not in block_by_bm:
            block_by_bm[bm] = el

    wrapped = 0
    for cid, bms in bookmark_for_comment_id.items():
        u = uuid_for_wid[cid]
        for bm in bms:
            block = block_by_bm.get(bm)
            if block is None:
                continue
            # Idempotent: if an inner span already carries this UUID, skip.
            if block.xpath(f'.//span[@data-comment-id="{u}"]'):
                continue
            span = lxml.html.Element(
                "span",
                attrib={"data-comment-id": u, "class": "tc-comment"},
            )
            span.text = block.text
            block.text = None
            for child in list(block):
                span.append(child)
            block.append(span)
            wrapped += 1

    if not wrapped:
        return xhtml_content

    logger.info(f"Wrapped {wrapped} XHTML paragraph(s) with Word-comment spans for file {file_id}")
    return lxml.html.tostring(root, encoding="unicode")


def _build_export_docx_with_comments(
    *,
    processed_path: str,
    xhtml_path: str,
    comments: list,
    logger,
) -> Optional[str]:
    """
    Copy the processed DOCX to a temp path and inject native Word
    `<w:commentRangeStart>`/`<w:comment>` markers for every
    `span[data-comment-id]` in the saved XHTML that corresponds to a stored
    Comment record. Returns the temp file path, or None if there is nothing
    to inject.

    Granularity is paragraph-level for v1: each comment is pinned to the first
    XHTML block that contains the matching span, mapped to the DOCX via the
    existing `data-bookmark` ↔ Word-bookmark correspondence.
    """
    if not comments or not os.path.exists(xhtml_path):
        return None

    by_uuid = {c.comment_uuid: c for c in comments if c.comment_uuid}
    if not by_uuid:
        return None

    import lxml.html
    with open(xhtml_path, "r", encoding="utf-8") as f:
        html_content = f.read()
    try:
        root = lxml.html.fromstring(html_content)
    except Exception as e:
        logger.warning(f"Could not parse XHTML for comment injection: {e}")
        return None

    # Map UUID -> first bookmark seen in document order. Comments that span
    # multiple paragraphs anchor to the paragraph where the run starts.
    BLOCK_TAGS = {"p", "h1", "h2", "h3", "h4", "h5", "h6", "li", "td", "th"}
    bookmark_for_uuid: Dict[str, str] = {}
    for span in root.xpath("//span[@data-comment-id]"):
        uuid = span.get("data-comment-id")
        if not uuid or uuid in bookmark_for_uuid or uuid not in by_uuid:
            continue
        anc = span
        bm = None
        while anc is not None:
            if anc.tag in BLOCK_TAGS:
                bm = anc.get("data-bookmark")
                if bm:
                    break
            # Some blocks carry the bookmark on a wrapping <div>, so keep walking
            # up even after we've found a block-tag without a bookmark attribute.
            if anc.get("data-bookmark"):
                bm = anc.get("data-bookmark")
                break
            anc = anc.getparent()
        if bm:
            bookmark_for_uuid[uuid] = bm

    if not bookmark_for_uuid:
        return None

    import shutil
    import tempfile
    fd, tmp_path = tempfile.mkstemp(suffix=".docx", prefix="export_with_comments_")
    os.close(fd)
    shutil.copyfile(processed_path, tmp_path)

    from docx import Document
    from app.docx_pipeline.utils.docx_helpers import add_comment_to_paragraph
    from app.processing.xhtml_to_docx_delta import (
        _build_bookmark_para_index,
        _find_note_para_by_bookmark,
    )

    doc = Document(tmp_path)
    para_index = _build_bookmark_para_index(doc)

    injected = 0
    for uuid, bm in bookmark_for_uuid.items():
        comment = by_uuid[uuid]
        para = para_index.get(bm) or _find_note_para_by_bookmark(doc, bm)
        if para is None:
            logger.info(f"Comment {uuid}: bookmark {bm} unresolved in DOCX, skipping")
            continue
        try:
            add_comment_to_paragraph(
                doc,
                para,
                text=comment.text or "",
                author=comment.author_name or "Unknown",
            )
            injected += 1
        except Exception as e:
            logger.warning(f"Comment {uuid}: failed to inject: {e}", exc_info=True)

    if injected == 0:
        try:
            os.unlink(tmp_path)
        except Exception:
            pass
        return None

    try:
        doc.save(tmp_path)
    except Exception as e:
        logger.warning(f"Failed to save DOCX with injected comments: {e}", exc_info=True)
        try:
            os.unlink(tmp_path)
        except Exception:
            pass
        return None

    logger.info(f"Injected {injected} Word comment(s) into export copy at {tmp_path}")
    return tmp_path


def get_export_payload(db: Session, *, file_id: int, logger):
    resolved = resolve_processed_target(db, file_id=file_id)
    processed_path = resolved["processed_path"]
    processed_filename = resolved["processed_filename"]

    if not os.path.exists(processed_path):
        logger.warning(f"Export failed: Processed file not found at {processed_path}")
        raise HTTPException(status_code=404, detail="Processed file not found")

    # If any comments exist for this file, build a copy with native Word
    # comments injected. Source of truth lives in the DB; the on-disk DOCX
    # stays clean so re-saves don't accumulate stale comment XML.
    enriched_path = None
    try:
        from app.models import Comment
        stored_comments = db.query(Comment).filter(Comment.file_id == file_id).all()
        if stored_comments:
            xhtml_path = _get_xhtml_path(processed_path)
            enriched_path = _build_export_docx_with_comments(
                processed_path=processed_path,
                xhtml_path=xhtml_path,
                comments=stored_comments,
                logger=logger,
            )
    except Exception as e:
        logger.warning(f"Comment injection failed; falling back to plain DOCX: {e}", exc_info=True)
        enriched_path = None

    return {
        "path": enriched_path or processed_path,
        "filename": processed_filename,
        "is_temp": bool(enriched_path),
    }


def _get_xhtml_path(processed_path: str) -> str:
    """Derive XHTML path from processed DOCX path."""
    dir_name = os.path.dirname(processed_path)
    base_name = os.path.splitext(os.path.basename(processed_path))[0]
    xhtml_dir = os.path.join(dir_name, "xhtml")
    return os.path.join(xhtml_dir, f"{base_name}.html")


def _augment_with_word_comments(
    *,
    db: Session,
    file_id: int,
    docx_path: str,
    xhtml_path: str,
    xhtml_content: str,
    logger,
) -> str:
    """
    Run `import_word_comments_into_xhtml` and persist the result. Always safe
    to call — idempotent thanks to deterministic UUIDs and the
    INSERT-only DB strategy. If the XHTML didn't change, this is essentially a
    cheap DOCX peek (early-exits when no `word/comments.xml` part is present).
    """
    try:
        new_content = import_word_comments_into_xhtml(
            db=db,
            file_id=file_id,
            docx_path=docx_path,
            xhtml_content=xhtml_content,
            logger=logger,
        )
    except Exception as e:
        logger.warning(f"Word-comment import skipped: {e}", exc_info=True)
        return xhtml_content

    if new_content == xhtml_content:
        return xhtml_content

    try:
        os.makedirs(os.path.dirname(xhtml_path), exist_ok=True)
        with open(xhtml_path, "w", encoding="utf-8") as f:
            f.write(new_content)
        # Keep the XHTML cache "newer than" the DOCX so the next read still
        # serves the cached, span-augmented content rather than regenerating.
        docx_mtime = os.path.getmtime(docx_path)
        os.utime(xhtml_path, (docx_mtime + 1, docx_mtime + 1))
    except Exception as e:
        logger.warning(f"Failed to persist XHTML with imported comment spans: {e}")
    return new_content


def get_xhtml_content(db: Session, *, file_id: int) -> Dict[str, Any]:
    """
    Get XHTML content for a file.

    Returns: {"content": str, "exists": bool, "mtime": float, "xhtml_path": str}
    """
    resolved = resolve_processed_target(db, file_id=file_id)
    processed_path = resolved["processed_path"]

    if not os.path.exists(processed_path):
        raise HTTPException(status_code=404, detail="Processed DOCX file not found")

    xhtml_path = _get_xhtml_path(processed_path)

    # Use cached XHTML if the processed docx hasn't changed since the cached XHTML was written
    docx_mtime = os.path.getmtime(processed_path)
    if os.path.exists(xhtml_path) and os.path.getmtime(xhtml_path) >= docx_mtime:
        logger.info(f"Serving cached XHTML for file {file_id}")
    else:
        try:
            os.makedirs(os.path.dirname(xhtml_path), exist_ok=True)
            engine = DocxToXhtmlRunsEngine()
            content = engine.convert(processed_path, file_id=file_id)
            with open(xhtml_path, "w", encoding="utf-8") as f:
                f.write(content)
        except Exception as e:
            import logging
            logging.getLogger(__name__).warning(f"Failed to force-convert XHTML in get_xhtml_content: {e}")

    if not os.path.exists(xhtml_path):
        return {
            "content": "",
            "exists": False,
            "mtime": None,
            "xhtml_path": xhtml_path,
        }

    try:
        with open(xhtml_path, "r", encoding="utf-8") as f:
            content = f.read()
        # Pull any native Word comments (added externally in Word) into the DB
        # and decorate the XHTML — runs on every request, but is a near no-op
        # for DOCXs without a comments part, and idempotent otherwise.
        content = _augment_with_word_comments(
            db=db, file_id=file_id, docx_path=processed_path,
            xhtml_path=xhtml_path, xhtml_content=content, logger=logger,
        )
        mtime = os.path.getmtime(xhtml_path)
        return {
            "content": content,
            "exists": True,
            "mtime": mtime,
            "xhtml_path": xhtml_path,
        }
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to read XHTML: {str(e)}")


def generate_xhtml(db: Session, *, file_id: int, logger) -> Dict[str, Any]:
    """
    Generate XHTML from processed DOCX.

    Returns: {"status": str, "xhtml_path": str}
    """
    resolved = resolve_processed_target(db, file_id=file_id)
    processed_path = resolved["processed_path"]

    if not os.path.exists(processed_path):
        raise HTTPException(status_code=404, detail="Processed DOCX file not found")

    xhtml_path = _get_xhtml_path(processed_path)

    try:
        os.makedirs(os.path.dirname(xhtml_path), exist_ok=True)
        engine = DocxToXhtmlRunsEngine()
        content = engine.convert(processed_path, file_id=file_id)
        with open(xhtml_path, "w", encoding="utf-8") as f:
            f.write(content)
        # Pull native Word comments into DB + XHTML in a single idempotent pass.
        _augment_with_word_comments(
            db=db, file_id=file_id, docx_path=processed_path,
            xhtml_path=xhtml_path, xhtml_content=content, logger=logger,
        )
        logger.info(f"Generated XHTML runs: {xhtml_path}")
        return {
            "status": "ok",
            "xhtml_path": xhtml_path,
        }
    except Exception as e:
        logger.error(f"XHTML generation failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"XHTML generation failed: {str(e)}")


def save_xhtml_and_convert(
    db: Session,
    *,
    file_id: int,
    html_content: str,
    username: str = "WYSIWYG Editor",
    logger,
) -> Dict[str, Any]:
    """
    Save edited XHTML and apply paragraph style and format changes back to the
    existing processed DOCX in-place, incrementing version and archiving the old file.

    Returns: {"status": str, "file_id": int}
    """
    resolved = resolve_processed_target(db, file_id=file_id)
    file_record = resolved["file_record"]
    processed_path = resolved["processed_path"]

    if not os.path.exists(processed_path):
        raise HTTPException(status_code=404, detail="Processed DOCX file not found")

    # Check if there are actual changes made to the XHTML content
    xhtml_path_existing = _get_xhtml_path(processed_path)
    if os.path.exists(xhtml_path_existing):
        try:
            with open(xhtml_path_existing, "r", encoding="utf-8") as f:
                current_xhtml = f.read()
            if current_xhtml.strip() == html_content.strip():
                logger.info(f"No changes detected in edited HTML for file {file_id}. Version bump skipped.")
                return {"status": "ok", "file_id": file_id}
        except Exception as e:
            logger.warning(f"Failed to read existing XHTML to check for changes: {e}")

    # Get user id for archiving
    from app.domains.auth.models import User
    user_record = db.query(User).filter(User.username == username).first()
    user_id = user_record.id if user_record else None

    # 1. Archive the existing/old file before modifying it on disk
    from app.domains.files import version_service
    base_path = os.path.dirname(processed_path)
    try:
        version_service.archive_existing_file(
            db,
            existing_file=file_record,
            base_path=base_path,
            uploaded_by_id=user_id,
        )
    except Exception as e:
        logger.error(f"Failed to archive existing file: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to archive the file: {str(e)}")

    xhtml_path = _get_xhtml_path(processed_path)
    tmp_docx = processed_path + ".tmp"

    try:
        # 2. Write XHTML to disk
        os.makedirs(os.path.dirname(xhtml_path), exist_ok=True)

        # Strip existing XML declarations, DOCTYPEs, head, html, and body wrappers
        cleaned_content = re.sub(r'<!--\?xml.*?\?-->', '', html_content, flags=re.DOTALL)
        cleaned_content = re.sub(r'<\?xml.*?\?>', '', cleaned_content, flags=re.DOTALL)
        cleaned_content = re.sub(r'<!DOCTYPE.*?>', '', cleaned_content, flags=re.DOTALL | re.IGNORECASE)

        m_body = re.search(r'<body[^>]*>(.*?)</body>', cleaned_content, flags=re.DOTALL | re.IGNORECASE)
        if m_body:
            inner_body = m_body.group(1).strip()
        else:
            m_html = re.search(r'<html[^>]*>(.*?)</html>', cleaned_content, flags=re.DOTALL | re.IGNORECASE)
            inner_body = m_html.group(1).strip() if m_html else cleaned_content.strip()

        # Remove stray head/html tags if left over
        inner_body = re.sub(r'<head[^>]*>.*?</head>', '', inner_body, flags=re.DOTALL | re.IGNORECASE)
        inner_body = re.sub(r'</?(?:html|head|body)[^>]*>', '', inner_body, flags=re.IGNORECASE)

        # Convert inline XML citation tags from <div> to <span> to prevent paragraph breaking in browsers/editors
        inline_tags = (
            "edition|publisher-name|publisher-loc|month|year|surname|given-names|collab|"
            "comment|volume|issue|fpage|lpage|ext-link|uri|xref|named-content|label|"
            "string-name|person-group|mixed-citation|bold|italic|sub|sup"
        )
        div_inline_pattern = re.compile(rf'<div(\s+[^>]*)data-xml-tag="({inline_tags})"([^>]*)>(.*?)</div>', re.DOTALL | re.IGNORECASE)
        inner_body = div_inline_pattern.sub(r'<span\1data-xml-tag="\2"\3>\4</span>', inner_body)

        # Merge adjacent consecutive track changes insertion spans into a single span
        prev_body = ""
        merge_pattern = re.compile(r'(<(?:span|ins)[^>]*class="[^"]*tc-insert[^"]*"[^>]*>)(.*?)</(?:span|ins)>\s*<(?:span|ins)[^>]*class="[^"]*tc-insert[^"]*"[^>]*>(.*?)</(?:span|ins)>', re.DOTALL | re.IGNORECASE)
        while prev_body != inner_body:
            prev_body = inner_body
            inner_body = merge_pattern.sub(r'\1\2\3</span>', inner_body)

        # Escape non-standard named HTML entities for strict XML validation in web browsers
        entity_map = {
            "&nbsp;": "&#160;",
            "&rdquo;": "&#8221;",
            "&ldquo;": "&#8220;",
            "&rsquo;": "&#8217;",
            "&lsquo;": "&#8216;",
            "&ndash;": "&#8211;",
            "&mdash;": "&#8212;",
            "&hellip;": "&#8230;",
            "&thinsp;": "&#8201;",
            "&ensp;": "&#8194;",
            "&emsp;": "&#8195;",
        }
        for ent, num in entity_map.items():
            inner_body = inner_body.replace(ent, num)

        content_to_save = (
            '<?xml version="1.0" encoding="UTF-8"?>\n'
            '<!DOCTYPE html PUBLIC "-//W3C//DTD XHTML 1.0 Strict//EN" "http://www.w3.org/TR/xhtml1/DTD/xhtml1-strict.dtd">\n'
            '<html xmlns="http://www.w3.org/1999/xhtml" xmlns:xlink="http://www.w3.org/1999/xlink">\n'
            '  <head>\n'
            '    <meta http-equiv="Content-Type" content="text/html; charset=UTF-8"/>\n'
            '    <title>XHTML Document</title>\n'
            '    <link rel="stylesheet" type="text/css" href="idGeneratedStyles.css"/>\n'
            '  </head>\n'
            '  <body>\n'
            f'    {inner_body}\n'
            '  </body>\n'
            '</html>'
        )
            
        with open(xhtml_path, "w", encoding="utf-8") as f:
            f.write(content_to_save)
        logger.info(f"Saved XHTML: {xhtml_path}")

        # 3. Patch in-place directly on the existing file path!
        engine = XhtmlToDocxDeltaEngine()
        engine.convert(xhtml_path, processed_path, username=username)
        logger.info(f"Style-patched DOCX version {file_record.version + 1} in-place: {processed_path}")

        # Update the mtime of xhtml_path to match or be newer than processed_path's mtime to ensure cache hit on reload
        try:
            docx_mtime = os.path.getmtime(processed_path)
            os.utime(xhtml_path, (docx_mtime + 1, docx_mtime + 1))
        except Exception as utime_err:
            logger.warning(f"Failed to update XHTML cache mtime: {utime_err}")

        # 4. Update the existing file record version and timestamp in-place
        from app.utils.timezone import now_ist_naive
        file_record.version += 1
        file_record.uploaded_at = now_ist_naive()
        db.commit()

        # Invalidate results cache for the file
        from app.domains.processing.technical_editor_service import RESULTS_DIR
        cache_path = RESULTS_DIR / f"{file_id}_scan.json"
        if cache_path.exists():
            try:
                cache_path.unlink()
            except Exception:
                pass

        return {
            "status": "ok",
            "file_id": file_record.id,
        }

    except Exception as e:
        logger.error(f"XHTML save/convert failed: {e}", exc_info=True)
        if os.path.exists(tmp_docx):
            try:
                os.remove(tmp_docx)
            except Exception:
                pass
        raise HTTPException(
            status_code=500, detail=f"XHTML save/convert failed: {str(e)}"
        )


def get_file_xhtml_runs(db: Session, *, file_id: int, logger) -> Dict[str, Any]:
    """
    Build run-anchored XHTML (every run carries its original w:rPr as data-rpr) for the
    formatting-preserving WYSIWYG editor.

    Returns: {"content": str, "filename": str}
    """
    resolved = resolve_processed_target(db, file_id=file_id)
    file_record = resolved["file_record"]
    processed_path = resolved["processed_path"]

    # Fall back to the original upload if no processed copy exists yet.
    source_path = processed_path if os.path.exists(processed_path) else file_record.path
    if not os.path.exists(source_path):
        raise HTTPException(status_code=404, detail="Physical file missing on disk")

    xhtml_path = _get_xhtml_path(source_path)
    docx_mtime = os.path.getmtime(source_path)

    # 1. Try XHTML cache first
    if os.path.exists(xhtml_path) and os.path.getmtime(xhtml_path) >= docx_mtime:
        logger.info(f"Serving cached runs XHTML from xhtml file for file {file_id}")
        try:
            with open(xhtml_path, "r", encoding="utf-8") as f:
                content = f.read()
            content = _augment_with_word_comments(
                db=db, file_id=file_id, docx_path=source_path,
                xhtml_path=xhtml_path, xhtml_content=content, logger=logger,
            )
            return {"content": content, "filename": file_record.filename}
        except Exception as e:
            logger.warning(f"Failed to read runs XHTML cache from {xhtml_path}: {e}")

    # 2. Try reference review cache as secondary fallback
    cache_path = _ref_review_cache_path(source_path)
    if os.path.exists(cache_path):
        try:
            import json
            with open(cache_path, "r", encoding="utf-8") as cf:
                cached = json.load(cf)
            if cached.get("mtime") == docx_mtime and cached.get("cache_version") == REF_REVIEW_CACHE_VERSION and cached.get("content"):
                logger.info(f"Serving cached runs XHTML from ref cache for file {file_id}")
                # Save to xhtml_path for future fast loads
                try:
                    os.makedirs(os.path.dirname(xhtml_path), exist_ok=True)
                    with open(xhtml_path, "w", encoding="utf-8") as f:
                        f.write(cached["content"])
                except Exception as ce:
                    logger.warning(f"Failed to write ref cache content to {xhtml_path}: {ce}")
                content = _augment_with_word_comments(
                    db=db, file_id=file_id, docx_path=source_path,
                    xhtml_path=xhtml_path, xhtml_content=cached["content"], logger=logger,
                )
                return {"content": content, "filename": file_record.filename}
        except Exception as e:
            logger.warning(f"Failed to read runs XHTML cache: {e}")

    # 3. Generate from scratch
    try:
        content = DocxToXhtmlRunsEngine().convert(source_path, file_id=file_id)
        # Write to cache
        try:
            os.makedirs(os.path.dirname(xhtml_path), exist_ok=True)
            with open(xhtml_path, "w", encoding="utf-8") as f:
                f.write(content)
        except Exception as ce:
            logger.warning(f"Failed to write runs XHTML cache to {xhtml_path}: {ce}")
        content = _augment_with_word_comments(
            db=db, file_id=file_id, docx_path=source_path,
            xhtml_path=xhtml_path, xhtml_content=content, logger=logger,
        )
    except Exception as e:
        logger.error(f"Run-anchored XHTML generation failed for file {file_id}: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Failed to generate editor content: {str(e)}")

    return {"content": content, "filename": file_record.filename}


def save_xhtml_delta_and_convert(
    db: Session,
    *,
    file_id: int,
    html_content: str,
    username: str = "WYSIWYG Editor",
    logger,
) -> Dict[str, Any]:
    """
    Save edited run-anchored XHTML and apply ONLY the changed runs/marks back to the
    existing processed DOCX in-place, incrementing version and archiving the old file.

    Returns: {"status": str, "file_id": int}
    """
    resolved = resolve_processed_target(db, file_id=file_id)
    file_record = resolved["file_record"]
    processed_path = resolved["processed_path"]

    # Delta patch requires an existing DOCX to patch; fall back to the original upload.
    source_path = processed_path if os.path.exists(processed_path) else file_record.path
    if not os.path.exists(source_path):
        raise HTTPException(status_code=404, detail="Processed DOCX file not found")

    # Get user id for archiving
    from app.domains.auth.models import User
    user_record = db.query(User).filter(User.username == username).first()
    user_id = user_record.id if user_record else None

    # 1. Archive the existing/old file before modifying it on disk
    from app.domains.files import version_service
    base_path = os.path.dirname(source_path)
    try:
        version_service.archive_existing_file(
            db,
            existing_file=file_record,
            base_path=base_path,
            uploaded_by_id=user_id,
        )
    except Exception as e:
        logger.error(f"Failed to archive existing file: {e}")
        raise HTTPException(status_code=500, detail=f"Failed to archive the file: {str(e)}")

    xhtml_path = _get_xhtml_path(source_path)

    try:
        os.makedirs(os.path.dirname(xhtml_path), exist_ok=True)
        with open(xhtml_path, "w", encoding="utf-8") as f:
            f.write(html_content)

        # 2. Patch in-place directly on the existing file path!
        XhtmlToDocxDeltaEngine().convert(xhtml_path, source_path, username=username)
        logger.info(f"Delta-patched DOCX version {file_record.version + 1} in-place: {source_path}")

        # Update the mtime of xhtml_path to match or be newer than source_path's mtime to ensure cache hit on reload
        try:
            docx_mtime = os.path.getmtime(source_path)
            os.utime(xhtml_path, (docx_mtime + 1, docx_mtime + 1))
        except Exception as utime_err:
            logger.warning(f"Failed to update XHTML cache mtime: {utime_err}")

        # 3. Update the existing file record version and timestamp in-place
        from app.utils.timezone import now_ist_naive
        file_record.version += 1
        file_record.uploaded_at = now_ist_naive()
        db.commit()

        # Invalidate cache for the file
        from app.domains.processing.technical_editor_service import RESULTS_DIR
        cache_path = RESULTS_DIR / f"{file_id}_scan.json"
        if cache_path.exists():
            try:
                cache_path.unlink()
            except Exception:
                pass

        return {"status": "ok", "file_id": file_record.id}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"XHTML delta save/convert failed: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"XHTML delta save/convert failed: {str(e)}")


REF_REVIEW_CACHE_VERSION = 5


def _ref_review_cache_path(processed_path: str) -> str:
    """Derive a sidecar cache file path for the reference review."""
    dir_name = os.path.dirname(processed_path)
    base_name = os.path.splitext(os.path.basename(processed_path))[0]
    return os.path.join(dir_name, f".{base_name}.refcache.json")


def _run_validation_on_doc(
    doc,
    processed_path: str,
    style: Optional[str] = None,
    citation_format: Optional[str] = None,
    logger=None,
    et_al_threshold: Optional[int] = None,
) -> Dict[str, Any]:
    validation_logs = {
        "stats": {},
        "total_refs": 0,
        "total_cites": 0,
        "issues": [],
        "renumbering_map": {},
        "duplicates": [],
        "sequence_issues": [],
        "missing_references": [],
        "unused_references": [],
        "pipeline_log": [],
    }

    detected_style = style
    if not detected_style:
        # Detect style: REF-N -> AMA, REF-U or <ref-open> -> APA
        is_ama = False
        is_apa = False
        for p in doc.paragraphs:
            style_name = (p.style.name or "") if p.style else ""
            if "REF-N" in style_name:
                is_ama = True
            if "REF-U" in style_name or "<ref-open>" in p.text:
                is_apa = True

        detected_style = "APA" if is_apa else "AMA"

    validation_logs["detected_style"] = detected_style
    validation_logs["citation_format"] = citation_format or "auto"

    if detected_style == "AMA":
        # A. Numerical/Sequence Validation (AMA 11th Edition / Vancouver)
        from app.processing.legacy.Referencenumvalidation import ReferenceProcessor, iter_document_paragraphs
        ref_proc = ReferenceProcessor(doc, citation_format=validation_logs["citation_format"])
        num_stats = ref_proc.get_validation_stats()

        validation_logs["total_refs"] = num_stats.get("total_references", 0)
        validation_logs["total_cites"] = num_stats.get("total_citations", 0)
        validation_logs["duplicates"] = num_stats.get("duplicate_references", [])
        validation_logs["sequence_issues"] = num_stats.get("sequence_issues", [])
        validation_logs["broken_ranges"] = num_stats.get("broken_ranges", [])
        validation_logs["invalid_numbers"] = num_stats.get("invalid_numbers", [])
        validation_logs["mixed_citation_style"] = num_stats.get("mixed_citation_style")
        validation_logs["inline_text_citations"] = num_stats.get("inline_text_citations", [])
        validation_logs["roman_numeral_citations"] = num_stats.get("roman_numeral_citations", [])
        validation_logs["summary"] = num_stats.get("summary", {})
        validation_logs["tagged_cites"] = sum(
            1 for para in iter_document_paragraphs(doc)
            for run in para.runs
            if ref_proc.is_citation_run(run)
        )
        validation_logs["autonum_converted"] = 0

        # Map numerical missing/unused references to issues list
        for missing_num in num_stats.get("missing_references", []):
            msg = f"Numerical citation [{missing_num}] has no matching reference entry."
            issue_item = {
                "type": "missing",
                "para_idx": -1,
                "message": msg,
                "citation": f"[{missing_num}]"
            }
            validation_logs["issues"].append(issue_item)
            validation_logs["missing_references"].append(issue_item)

        for unused_num in num_stats.get("unused_references", []):
            msg = f"Reference [{unused_num}] is in bibliography but never cited."
            issue_item = {
                "type": "unused",
                "para_idx": -1,
                "message": msg,
                "citation": f"[{unused_num}]"
            }
            validation_logs["issues"].append(issue_item)
            validation_logs["unused_references"].append(issue_item)

        # Build citation_pairs and reference_entries
        refs_found, ref_objects = ref_proc.get_references_in_bibliography()
        all_cited_ids, appearance_order = ref_proc.get_citations_in_text()
        cited_set = set(all_cited_ids)

        from docx.oxml.ns import qn
        import docx.text.paragraph
        doc_paragraphs = [docx.text.paragraph.Paragraph(p_elem, doc) for p_elem in doc.element.body.iter(qn("w:p"))]
        para_index_map = {p._element: idx for idx, p in enumerate(doc_paragraphs)}

        # Pre-compute reference lookups once (was O(N²) via nested "for obj in ref_objects" scans).
        ref_text_map = {obj["id"]: _get_full_reference_text(obj["para"], doc_paragraphs, para_index_map) for obj in ref_objects}
        ref_para_by_id = {obj["id"]: obj["para"] for obj in ref_objects}
        ref_id_by_element = {obj["para"]._element: obj["id"] for obj in ref_objects}

        # Build citation_pairs: one entry per unique cited number, in appearance order
        citation_pairs = []
        for num in appearance_order:
            ref_para = ref_para_by_id.get(num)
            citation_pairs.append({
                "citation": str(num),
                "ref_number": num,
                "ref_text": ref_text_map.get(num, ""),
                "status": "ok" if num in refs_found else "missing",
                "para_idx": para_index_map.get(ref_para._element, -1) if ref_para else -1,
            })
        # Add unused entries (in bibliography but never cited)
        for obj in ref_objects:
            if obj["id"] not in cited_set:
                citation_pairs.append({
                    "citation": None,
                    "ref_number": obj["id"],
                    "ref_text": ref_text_map.get(obj["id"], ""),
                    "status": "unused",
                    "para_idx": para_index_map.get(obj["para"]._element, -1),
                })

        # Build reference_entries with global para_idx
        reference_entries = []
        for p in doc_paragraphs:
            style_name = (p.style.name or "") if p.style else ""
            if "REF-N" in style_name:
                ref_num = ref_id_by_element.get(p._element)
                if ref_num is not None:
                    reference_entries.append({
                        "number": ref_num,
                        "text": ref_text_map.get(ref_num, ""),
                        "style": "REF-N",
                        "is_cited": ref_num in cited_set,
                        "para_idx": para_index_map.get(p._element, -1),
                    })

        validation_logs["citation_pairs"] = citation_pairs
        validation_logs["reference_entries"] = reference_entries
    else:
        # B. APA Name & Year Validation (APA 7th Edition)
        from app.processing.legacy.validation_core import CitationProcessor
        cite_proc = CitationProcessor(processed_path, et_al_threshold=et_al_threshold)
        report = cite_proc.run()
        report_dict = report.to_dict()

        _body_to_allp_idx = _build_body_to_allp_translator(cite_proc.doc)

        validation_logs["stats"] = dict(report.stats)
        validation_logs["total_refs"] = report.total_refs
        validation_logs["total_cites"] = report.total_cites

        # Forward grouped APA issue categories
        for key in ("et_al_issues", "name_spelling_warnings", "ordering_issues",
                    "suffix_issues", "disambiguation_issues",
                    "personal_comm_citations", "secondary_citations"):
            validation_logs[key] = report_dict.get(key, [])

        # Collect Name & Year issues
        for issue in report.issues:
            issue_copy = {k: v for k, v in issue.items() if k != "para"}
            validation_logs["issues"].append(issue_copy)

            # Map into categories
            itype = issue_copy.get("type")
            if itype == "missing":
                validation_logs["missing_references"].append(issue_copy)
            elif itype == "unused":
                validation_logs["unused_references"].append(issue_copy)

        # Build match_score lookup keyed by raw citation string
        score_by_raw: Dict[str, float] = {
            c["raw"]: c["match_score"]
            for c in cite_proc.get_citation_changes()
            if "raw" in c and "match_score" in c
        }

        # Build citation_pairs and reference_entries
        bib_items = list(getattr(cite_proc, "_bib_ordered", []) or cite_proc.bibliography.values())

        citation_pairs = []
        # Iterate bib in document order to build pairs
        for entry in bib_items:
            status = "ok" if entry.get("cited") else "unused"
            raw_cite = f"({entry.get('display', '')}, {entry.get('year', '')})"
            pair: Dict[str, Any] = {
                "citation": raw_cite,
                "author": entry.get("display", ""),
                "year": entry.get("year", ""),
                "ref_text": entry.get("raw", ""),
                "status": status,
                "para_idx": _body_to_allp_idx(entry.get("para_idx", -1)),
            }
            score = score_by_raw.get(raw_cite)
            if score is not None:
                pair["match_score"] = round(score, 2)
            citation_pairs.append(pair)

        # Add missing citations (cited in text but no bib entry)
        for issue in report.issues:
            if issue.get("type") == "missing":
                citation_pairs.append({
                    "citation": issue.get("raw", issue.get("citation", "")),
                    "author": "",
                    "year": "",
                    "ref_text": "",
                    "status": "missing",
                    "para_idx": _body_to_allp_idx(issue.get("para_idx", -1)),
                })

        reference_entries = []
        for i, entry in enumerate(bib_items):
            reference_entries.append({
                "number": None,
                "text": entry.get("raw", ""),
                "style": "REF-U",
                "is_cited": entry.get("cited", False),
                "para_idx": _body_to_allp_idx(entry.get("para_idx", -1)),
            })

        validation_logs["citation_pairs"] = citation_pairs
        validation_logs["reference_entries"] = reference_entries

    return validation_logs


def build_reference_review_page_state(db: Session, *, file_id: int, style: Optional[str] = None, citation_format: Optional[str] = None, et_al_threshold: Optional[int] = None, logger) -> Dict[str, Any]:
    import json
    resolved = resolve_processed_target(db, file_id=file_id)
    file_record = resolved["file_record"]
    processed_path = resolved["processed_path"]

    if not os.path.exists(processed_path):
        raise HTTPException(
            status_code=404,
            detail="Processed reference file not found on disk."
        )

    docx_mtime = os.path.getmtime(processed_path)
    cache_path = _ref_review_cache_path(processed_path)

    # ── Fast path: serve from cache if the DOCX hasn't been modified ──────────
    if os.path.exists(cache_path):
        try:
            with open(cache_path, "r", encoding="utf-8") as cf:
                cached = json.load(cf)
            if cached.get("mtime") == docx_mtime and cached.get("cache_version") == REF_REVIEW_CACHE_VERSION:
                cached_style = cached.get("validation_logs", {}).get("detected_style")
                cached_format = cached.get("validation_logs", {}).get("citation_format", "auto")
                req_format = citation_format or "auto"
                if (style is None or cached_style == style) and cached_format == req_format:
                    logger.info(f"Reference review cache HIT for file {file_id}")
                    # Styles list (static, cheap to rebuild)
                    styles = _ref_review_styles()
                    cached_logs = cached["validation_logs"]
                    merge_manual_links_into_logs(
                        cached_logs,
                        read_manual_links(processed_path, logger=logger),
                    )
                    return {
                        "status": "ok",
                        "file": file_record,
                        "filename": cached["filename"],
                        "content": cached["content"],
                        "styles": styles,
                        "validation_logs": cached_logs,
                    }
        except Exception as e:
            logger.warning(f"Reference review cache read failed (will regenerate): {e}")

    logger.info(f"Reference review cache MISS for file {file_id} — regenerating…")

    # 1. Load run-anchored XHTML
    try:
        xhtml_data = get_file_xhtml_runs(db, file_id=file_id, logger=logger)
        content = xhtml_data["content"]
        filename = xhtml_data["filename"]
    except Exception as e:
        logger.error(f"Failed to convert references doc to XHTML: {e}")
        raise HTTPException(
            status_code=500,
            detail=f"Failed to load document content: {str(e)}"
        )

    # 2. Extract live stats and validation issues using refactored helper
    try:
        from docx import Document
        doc = Document(processed_path)

        detected_style = style
        if not detected_style:
            is_ama = False
            is_apa = False
            for p in doc.paragraphs:
                style_name = (p.style.name or "") if p.style else ""
                if "REF-N" in style_name:
                    is_ama = True
                if "REF-U" in style_name or "<ref-open>" in p.text:
                    is_apa = True
            detected_style = "APA" if is_apa else "AMA"

        validation_logs = _run_validation_on_doc(
            doc,
            processed_path,
            style=detected_style,
            citation_format=citation_format,
            logger=logger,
            et_al_threshold=et_al_threshold,
        )

    except Exception as e:
        logger.error(f"Error computing live validation stats: {e}", exc_info=True)
        validation_logs = {
            "stats": {},
            "total_refs": 0,
            "total_cites": 0,
            "issues": [],
            "renumbering_map": {},
            "duplicates": [],
            "sequence_issues": [],
            "missing_references": [],
            "unused_references": [],
            "pipeline_log": [],
        }

    # 3. Read log file if it exists
    try:
        base_dir = os.path.dirname(file_record.path)
        filename_base = os.path.splitext(file_record.filename)[0]
        clean_base = filename_base.replace("_Processed", "").replace("_Structured", "")
        log_filename = f"{clean_base}_log.txt"
        log_path = os.path.join(base_dir, log_filename)

        if os.path.exists(log_path):
            with open(log_path, "r", encoding="utf-8") as lf:
                raw_log = lf.read()
            validation_logs["raw_log"] = raw_log
    except Exception as e:
        logger.warning(f"Could not read raw log file: {e}")

    styles = _ref_review_styles()

    # ── Write cache so the next page load is instant ──────────────────────────
    # Cache the raw validation_logs (before manual-links merge) so the merge
    # always reflects the current sidecar without stale linked-status bleeding
    # into the cache.
    try:
        import json
        docx_mtime = os.path.getmtime(processed_path)
        with open(cache_path, "w", encoding="utf-8") as cf:
            json.dump(
                {
                    "mtime": docx_mtime,
                    "filename": filename,
                    "content": content,
                    "validation_logs": validation_logs,
                    "cache_version": REF_REVIEW_CACHE_VERSION
                },
                cf,
                ensure_ascii=False,
            )
        logger.info(f"Reference review cache written for file {file_id}")
    except Exception as e:
        logger.warning(f"Reference review cache write failed (non-fatal): {e}")

    merge_manual_links_into_logs(
        validation_logs,
        read_manual_links(processed_path, logger=logger),
    )

    return {
        "status": "ok",
        "file": file_record,
        "filename": filename,
        "content": content,
        "styles": styles,
        "validation_logs": validation_logs
    }


def run_validation_only(db: Session, *, file_id: int, style: Optional[str] = None, citation_format: Optional[str] = None, et_al_threshold: Optional[int] = None, logger=None) -> Dict[str, Any]:
    """
    Run validation-only (no XHTML conversion, no re-structuring).
    Returns only validation_logs and detected_style.
    This is fast and doesn't reload editor content.
    """
    resolved = resolve_processed_target(db, file_id=file_id)
    file_record = resolved["file_record"]
    processed_path = resolved["processed_path"]

    if not os.path.exists(processed_path):
        raise HTTPException(status_code=404, detail="Processed reference file not found.")

    try:
        from docx import Document
        doc = Document(processed_path)

        detected_style = style
        if not detected_style:
            is_ama = False
            is_apa = False
            for p in doc.paragraphs:
                style_name = (p.style.name or "") if p.style else ""
                if "REF-N" in style_name:
                    is_ama = True
                if "REF-U" in style_name or "<ref-open>" in p.text:
                    is_apa = True
            detected_style = "APA" if is_apa else "AMA"

        pipeline_log = []
        mapping = {}

        if detected_style == "AMA":
            from app.processing.legacy.Referencenumvalidation import process_document
            pipeline_log.append("Started numeric reference validation pipeline.")
            pipeline_log.append("Applied punctuation swapping (placed citations after punctuation).")

            # Run validation report on the original state BEFORE modifying the document
            validation_logs = _run_validation_on_doc(
                doc,
                processed_path,
                style="AMA",
                citation_format=citation_format,
                logger=logger,
                et_al_threshold=et_al_threshold,
            )

            # Run renumbering/formatting pipeline in Referencenumvalidation
            doc, before_stats, after_stats, mapping, status_msg = process_document(processed_path, citation_format=citation_format)

            pipeline_log.append(f"Initial check: {len(before_stats.get('missing_references', []))} missing, {len(before_stats.get('unused_references', []))} unused reference(s).")

            if before_stats.get("unused_references"):
                pipeline_log.append("Aborted renumbering: Unused references detected in the bibliography.")
            elif before_stats.get("missing_references"):
                pipeline_log.append("Aborted renumbering: Missing references detected in the text.")
            elif before_stats.get("is_perfect"):
                pipeline_log.append("No sequence or formatting issues detected. Document is already perfect.")
            else:
                pipeline_log.append("Pass 1: Reordered bibliography and renumbered citations to resolve sequence issues.")
                dup_before = len(before_stats.get("duplicate_references", []))
                dup_after = len(after_stats.get("duplicate_references", []))
                duplicates_resolved = dup_before - dup_after
                if duplicates_resolved > 0:
                    pipeline_log.append(f"Pass 2: Found and merged {duplicates_resolved} duplicate reference(s) and renumbered remaining citations.")
                else:
                    pipeline_log.append("Pass 2: No duplicate references to merge.")
                pipeline_log.append("Successfully saved renumbered document.")

            if not status_msg.startswith("Aborted:"):
                # Save the updated doc, archive previous version, and bump file version
                from app.domains.files import version_service as vs
                try:
                    vs.archive_existing_file(
                        db,
                        existing_file=file_record,
                        base_path=os.path.dirname(processed_path),
                        uploaded_by_id=None,
                        source_path=processed_path,
                    )
                    file_record.version += 1
                except Exception as _e:
                    if logger:
                        logger.warning(f"Version archive failed (non-fatal): {_e}")

                doc.save(processed_path)
                db.commit()

                # Regenerate validation logs against the updated document — but
                # only when renumbering actually changed something. When
                # process_document reports "Validation completed." the doc was
                # already perfect (only bookmarks were emitted), so the first
                # validation_logs are still authoritative. Skipping this second
                # pass avoids a redundant full doc walk + O(N²) duplicate scan.
                if status_msg != "Validation completed.":
                    validation_logs = _run_validation_on_doc(
                        doc,
                        processed_path,
                        style="AMA",
                        citation_format=citation_format,
                        logger=logger,
                        et_al_threshold=et_al_threshold,
                    )

                # Invalidate caches
                invalidate_ref_review_cache(processed_path, logger=logger)

                from app.domains.processing.technical_editor_service import RESULTS_DIR
                cache_path = RESULTS_DIR / f"{file_id}_scan.json"
                if cache_path.exists():
                    try:
                        cache_path.unlink()
                    except Exception:
                        pass

                xhtml_path = _get_xhtml_path(processed_path)
                if os.path.exists(xhtml_path):
                    try:
                        os.remove(xhtml_path)
                    except Exception as e:
                        if logger:
                            logger.warning(f"Could not remove stale XHTML cache: {e}")

            validation_logs["pipeline_log"] = pipeline_log
            validation_logs["renumbering_map"] = mapping
        else:
            from app.processing.legacy.validation_core import process_document_apa
            pipeline_log.append("Started Name-Year (APA) reference validation pipeline.")
            
            # Run the two-pass APA validation/renumbering pipeline
            doc, before_stats, after_stats, mapping, status_msg = process_document_apa(processed_path)
            
            pipeline_log.append(f"Pass 1: Reordered bibliography alphabetically and resolved suffix additions.")
            dup_before = len(before_stats.get("duplicate_references", []))
            dup_after = len(after_stats.get("duplicate_references", []))
            duplicates_resolved = dup_before - dup_after
            if duplicates_resolved > 0:
                pipeline_log.append(f"Pass 2: Found and merged {duplicates_resolved} duplicate reference(s) and re-sorted remaining citations.")
            else:
                pipeline_log.append("Pass 2: No duplicate references to merge.")
            pipeline_log.append("Successfully saved sorted and merged document.")
            
            # Save the updated doc, archive previous version, and bump file version
            from app.domains.files import version_service as vs
            try:
                vs.archive_existing_file(
                    db,
                    existing_file=file_record,
                    base_path=os.path.dirname(processed_path),
                    uploaded_by_id=None,
                    source_path=processed_path,
                )
                file_record.version += 1
            except Exception as _e:
                if logger:
                    logger.warning(f"Version archive failed (non-fatal): {_e}")

            doc.save(processed_path)
            db.commit()

            # Regenerate validation logs using the updated/saved document state!
            validation_logs = _run_validation_on_doc(
                doc,
                processed_path,
                style="APA",
                citation_format=citation_format,
                logger=logger,
                et_al_threshold=et_al_threshold,
            )

            # Invalidate caches
            invalidate_ref_review_cache(processed_path, logger=logger)

            from app.domains.processing.technical_editor_service import RESULTS_DIR
            cache_path = RESULTS_DIR / f"{file_id}_scan.json"
            if cache_path.exists():
                try:
                    cache_path.unlink()
                except Exception:
                    pass

            xhtml_path = _get_xhtml_path(processed_path)
            if os.path.exists(xhtml_path):
                try:
                    os.remove(xhtml_path)
                except Exception as e:
                    if logger:
                        logger.warning(f"Could not remove stale XHTML cache: {e}")

            validation_logs["pipeline_log"] = pipeline_log
            validation_logs["renumbering_map"] = mapping

    except Exception as e:
        if logger:
            logger.error(f"Error computing validation stats: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Validation failed: {str(e)}")

    merge_manual_links_into_logs(
        validation_logs,
        read_manual_links(processed_path, logger=logger),
    )

    return {
        "validation_logs": validation_logs,
        "detected_style": validation_logs.get("detected_style", "AMA"),
    }


def _ref_review_styles() -> list:
    """Static list of bibliography and citation character styles."""
    return [
        # bibliography character styles
        "bib_alt-year", "bib_article", "bib_base", "bib_book", "bib_chapterno", "bib_chaptertitle",
        "bib_comment", "bib_confacronym", "bib_confdate", "bib_conference", "bib_conflocation",
        "bib_confpaper", "bib_confproceedings", "bib_day", "bib_deg", "bib_doi", "bib_ed-etal",
        "bib_ed-fname", "bib_ed-organization", "bib_ed-suffix", "bib_ed-surname", "bib_editionno",
        "bib_etal", "bib_extlink", "bib_fname", "bib_fpage", "bib_institution", "bib_isbn",
        "bib_issue", "bib_journal", "bib_location", "bib_lpage", "bib_medline", "bib_month",
        "bib_number", "bib_organization", "bib_pagecount", "bib_papernumber", "bib_patent",
        "bib_publisher", "bib_reportnum", "bib_school", "bib_season", "bib_series", "bib_seriesno",
        "bib_suffix", "bib_suppl", "bib_surname", "bib_title", "bib_trans", "bib_unpubl",
        "bib_url", "bib_volcount", "bib_volume", "bib_year",
        # citation styles
        "cite_app", "cite_bib", "cite_eq", "cite_fig", "cite_fn", "cite_sec", "cite_tbl",
        "cite_tfn", "cite_base", "cite_box"
    ]


def run_structuring_pipeline(db: Session, *, file_id: int, logger, on_event=None) -> dict:
    """
    Run the 10-step docx_pipeline on the current processed DOCX for a file.
    Returns {"status": "ok", "content": str, "file_id": int}.

    on_event: optional callable invoked for each step + final summary event,
    enabling streaming progress updates to the client.
    """
    import shutil
    import sys
    import tempfile
    from pathlib import Path

    # Make docx_pipeline importable as a top-level package (it uses bare imports)
    _pipeline_base = str(Path(__file__).parent.parent.parent)
    if _pipeline_base not in sys.path:
        sys.path.insert(0, _pipeline_base)

    from docx_pipeline.pipeline.runner import process_file

    resolved = resolve_processed_target(db, file_id=file_id)
    file_record = resolved["file_record"]
    processed_path = resolved["processed_path"]

    if not os.path.exists(processed_path):
        raise HTTPException(status_code=404, detail="Processed DOCX file not found")

    with tempfile.TemporaryDirectory(prefix="pipeline_run_") as tmp_dir:
        input_path = Path(processed_path)
        output_dir = Path(tmp_dir)

        logger.info(f"Running docx_pipeline on file {file_id}: {processed_path}")
        result = process_file(input_path, output_dir, on_event=on_event)

        if result["status"] == "error":
            issues = "; ".join(
                i["message"] for i in result.get("issues", []) if i.get("level") == "ERROR"
            ) or "Unknown pipeline error"
            logger.error(f"Pipeline failed for file {file_id}: {issues}")
            raise HTTPException(status_code=500, detail=f"Pipeline failed: {issues}")

        pipeline_output = output_dir / input_path.name
        if not pipeline_output.exists():
            raise HTTPException(status_code=500, detail="Pipeline produced no output file")

        shutil.copy2(str(pipeline_output), processed_path)
        logger.info(f"Pipeline output copied back to {processed_path}")

    # Invalidate stale XHTML cache
    xhtml_path = _get_xhtml_path(processed_path)
    if os.path.exists(xhtml_path):
        try:
            os.remove(xhtml_path)
        except Exception as e:
            logger.warning(f"Could not remove stale XHTML cache: {e}")

    # Bump version
    file_record.version = (file_record.version or 0) + 1
    db.commit()

    # Regenerate XHTML from the pipeline output
    try:
        os.makedirs(os.path.dirname(xhtml_path), exist_ok=True)
        engine = DocxToXhtmlRunsEngine()
        content = engine.convert(processed_path, file_id=file_id)
        with open(xhtml_path, "w", encoding="utf-8") as f:
            f.write(content)
        logger.info(f"XHTML regenerated for file {file_id} after pipeline run")
        return {"status": "ok", "content": content, "file_id": file_record.id}
    except Exception as e:
        logger.error(f"XHTML regeneration failed after pipeline: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"Pipeline ran but XHTML generation failed: {str(e)}")


def xslt_convert_to_xhtml(db: Session, *, file_id: int) -> Dict[str, Any]:
    """
    Convert the processed DOCX to XHTML using the XSLT 2.0 pipeline.
    Writes result to a separate _xslt.xhtml cache — does NOT overwrite the
    Python-converted cache used by the existing load flow.
    Returns: {"status": "ok", "content": str}
    """
    from app.processing.xslt_docx_to_xhtml import XsltDocxToXhtmlEngine
    from pathlib import Path as _Path

    resolved = resolve_processed_target(db, file_id=file_id)
    processed_path = resolved["processed_path"]
    source_path = processed_path if os.path.exists(processed_path) else resolved["file_record"].path

    if not os.path.exists(source_path):
        raise HTTPException(status_code=404, detail="Physical file missing on disk")

    try:
        full_html = XsltDocxToXhtmlEngine().convert(source_path)
    except Exception as e:
        logger.error(f"XSLT conversion failed for file {file_id}: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"XSLT conversion failed: {str(e)}")

    # Extract only <body> content — TipTap's setContent expects a fragment, not a
    # full HTML document with <html>/<head>/<body> wrappers.
    body_match = re.search(r'<body[^>]*>(.*?)</body>', full_html, re.DOTALL | re.IGNORECASE)
    content = body_match.group(1).strip() if body_match else full_html

    xslt_cache = _Path(_get_xhtml_path(source_path)).with_suffix("._xslt.xhtml")
    try:
        os.makedirs(str(xslt_cache.parent), exist_ok=True)
        xslt_cache.write_text(content, encoding="utf-8")
    except Exception as e:
        logger.warning(f"Could not write XSLT cache: {e}")

    return {"status": "ok", "content": content}


def xslt_save_to_docx(db: Session, *, file_id: int, xhtml: str) -> Dict[str, Any]:
    """
    Convert XHTML back to DOCX using the XSLT 2.0 pipeline.
    Writes to a separate _xslt_output.docx — does NOT overwrite the original
    processed DOCX until the pipeline is fully validated.
    Returns: {"status": "ok", "output_path": str}
    """
    from app.processing.xslt_xhtml_to_docx import XsltXhtmlToDocxEngine
    from pathlib import Path as _Path

    resolved = resolve_processed_target(db, file_id=file_id)
    processed_path = resolved["processed_path"]

    docx_path = _Path(processed_path)
    output_path = docx_path.with_name(docx_path.stem + "_xslt_output.docx")

    try:
        XsltXhtmlToDocxEngine().convert(xhtml, str(output_path))
    except Exception as e:
        logger.error(f"XSLT DOCX save failed for file {file_id}: {e}", exc_info=True)
        raise HTTPException(status_code=500, detail=f"XSLT save failed: {str(e)}")

    return {"status": "ok", "output_path": str(output_path)}


def invalidate_ref_review_cache(processed_path: str, logger=None) -> None:
    """Delete the reference review sidecar cache so the next load regenerates it."""
    cache_path = _ref_review_cache_path(processed_path)
    if os.path.exists(cache_path):
        try:
            os.remove(cache_path)
            if logger:
                logger.info(f"Reference review cache invalidated: {cache_path}")
        except Exception as e:
            if logger:
                logger.warning(f"Failed to remove ref review cache: {e}")


# ── Manual bookmark → citation → reference links ──────────────────────────────
# Sidecar JSON next to the processed DOCX. Persists user-created mappings for
# bookmarks that the auto-linker couldn't resolve (Word-native bookmarks the
# validator didn't recognize, or manual bookmarks the user added). Kept out of
# the DB on purpose: matches the existing "derive at read time" pattern and
# needs no migrations. Promote to a table later if audit trail is required.

MANUAL_LINKS_VERSION = 1


def _manual_links_path(processed_path: str) -> str:
    dir_name = os.path.dirname(processed_path)
    base_name = os.path.splitext(os.path.basename(processed_path))[0]
    return os.path.join(dir_name, f".{base_name}.manual_links.json")


def _empty_manual_links_doc() -> Dict[str, Any]:
    return {"version": MANUAL_LINKS_VERSION, "links": []}


def read_manual_links(processed_path: str, logger=None) -> Dict[str, Any]:
    """Load the manual-links sidecar. Returns an empty doc if missing/corrupt."""
    import json
    path = _manual_links_path(processed_path)
    if not os.path.exists(path):
        return _empty_manual_links_doc()
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        if not isinstance(data, dict) or not isinstance(data.get("links"), list):
            return _empty_manual_links_doc()
        data.setdefault("version", MANUAL_LINKS_VERSION)
        return data
    except Exception as e:
        if logger:
            logger.warning(f"manual_links read failed at {path}: {e}")
        return _empty_manual_links_doc()


def _write_manual_links(processed_path: str, doc: Dict[str, Any], logger=None) -> None:
    import json
    path = _manual_links_path(processed_path)
    try:
        with open(path, "w", encoding="utf-8") as f:
            json.dump(doc, f, ensure_ascii=False, indent=2)
    except Exception as e:
        if logger:
            logger.error(f"manual_links write failed at {path}: {e}")
        raise


def upsert_manual_link(
    processed_path: str,
    *,
    bookmark_name: str,
    ref_number: Optional[int],
    ref_text: str,
    citation_text: Optional[str],
    linked_by: Optional[str],
    logger=None,
) -> Dict[str, Any]:
    """Insert or replace the manual link for `bookmark_name`. Returns the stored link."""
    from datetime import datetime, timezone
    doc = read_manual_links(processed_path, logger=logger)
    now = datetime.now(timezone.utc).isoformat()
    entry = {
        "bookmark_name": bookmark_name,
        "ref_number": ref_number,
        "ref_text": ref_text,
        "citation_text": citation_text,
        "linked_by": linked_by,
        "linked_at": now,
    }
    links = [lnk for lnk in doc["links"] if lnk.get("bookmark_name") != bookmark_name]
    links.append(entry)
    doc["links"] = links
    _write_manual_links(processed_path, doc, logger=logger)
    return entry


def delete_manual_link(
    processed_path: str,
    *,
    bookmark_name: str,
    logger=None,
) -> bool:
    """Remove a manual link. Returns True if something was deleted."""
    doc = read_manual_links(processed_path, logger=logger)
    before = len(doc["links"])
    doc["links"] = [lnk for lnk in doc["links"] if lnk.get("bookmark_name") != bookmark_name]
    if len(doc["links"]) == before:
        return False
    _write_manual_links(processed_path, doc, logger=logger)
    return True


def merge_manual_links_into_logs(
    validation_logs: Dict[str, Any],
    manual_links_doc: Dict[str, Any],
) -> None:
    """
    Fold persisted manual links into validation_logs so the panel's stats/statuses
    reflect them. Mutates `validation_logs` in place.

    - Adds `manual_links` array (raw list, for the frontend).
    - For each manual link with a `ref_number`, marks matching `citation_pairs`
      as "ok" and matching `reference_entries.is_cited = True`, so counts
      shift from Missing/Unused into Matched exactly like an auto-link would.
    """
    links = list(manual_links_doc.get("links") or [])
    validation_logs["manual_links"] = links
    if not links:
        return

    linked_ref_numbers = {int(lnk["ref_number"]) for lnk in links if lnk.get("ref_number") is not None}
    linked_ref_texts = {
        (lnk.get("ref_text") or "").strip().lower()
        for lnk in links
        if lnk.get("ref_text")
    }

    for pair in validation_logs.get("citation_pairs") or []:
        rn = pair.get("ref_number")
        rt = (pair.get("ref_text") or "").strip().lower()
        if (rn is not None and rn in linked_ref_numbers) or (rt and rt in linked_ref_texts):
            if pair.get("status") != "ok":
                pair["status"] = "ok"
                pair["manual_linked"] = True

    for entry in validation_logs.get("reference_entries") or []:
        num = entry.get("number")
        txt = (entry.get("text") or "").strip().lower()
        if (num is not None and num in linked_ref_numbers) or (txt and txt in linked_ref_texts):
            entry["is_cited"] = True
            entry["manual_linked"] = True
